"""
Exporters de relatorio: MD, DOCX, PDF.

Recebem o dict do endpoint /api/testes/<id>/relatorio (mesma estrutura
que o frontend consome) e retornam bytes pra send_file.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Any

_PROJECT = Path(__file__).resolve().parents[3]


def _resumir_acao(a: dict) -> str:
    """Gera linha humana de uma acao do tutorial.

    Ex: {"tipo":"click","seletor":"button:has-text('Salvar')"} ->
        'CLICK em "button:has-text(...)"'
    """
    if not isinstance(a, dict):
        return str(a)
    tipo = (a.get("tipo") or "?").upper()
    seletor = a.get("seletor")
    alt = a.get("alternativa")
    valor = a.get("valor_literal")
    valor_ds = a.get("valor_from_dataset")
    valor_ctx = a.get("valor_from_contexto")
    url = a.get("url") or a.get("destino")
    timeout = a.get("timeout")
    sub = a.get("sequencia")

    partes = [tipo]
    if sub and isinstance(sub, list):
        # Acao composta: lista cada subacao
        partes.append(f"(sequencia de {len(sub)} sub-acoes)")
        return " ".join(partes)

    if tipo in ("CLICK", "WAIT_FOR") and seletor:
        partes.append(f'em "{seletor[:120]}"')
    elif tipo == "FILL" and seletor:
        v_str = valor if valor is not None else (
            f"<dataset:{valor_ds}>" if valor_ds else
            f"<contexto:{valor_ctx}>" if valor_ctx else "?"
        )
        partes.append(f'"{seletor[:80]}" com "{v_str}"')
    elif tipo == "SELECT" and seletor:
        v_str = valor if valor is not None else (
            f"<dataset:{valor_ds}>" if valor_ds else "?"
        )
        partes.append(f'"{seletor[:80]}" -> "{v_str}"')
    elif tipo in ("NAVIGATE", "NAVEGACAO", "GOTO") and url:
        partes.append(f"url={url}")
    elif tipo == "WAIT":
        partes.append(f"{valor or '?'} ms")
    elif tipo == "EVALUATE":
        # JS arbitrario - so mostra primeiras linhas
        codigo = valor or ""
        primeira = next((l.strip() for l in codigo.splitlines() if l.strip()), "")
        partes.append(f"JS: {primeira[:100]}")
    elif tipo == "UPLOAD_ARQUIVO" and seletor:
        partes.append(f'em "{seletor[:80]}"')
    elif tipo == "CHAMAR_API" and url:
        metodo = a.get("metodo", "POST")
        partes.append(f"{metodo} {url}")
    elif seletor:
        partes.append(f'em "{seletor[:120]}"')
    if alt and seletor:
        partes.append(f"(alternativa: \"{alt[:60]}\")")
    if timeout:
        partes.append(f"[t={timeout}ms]")
    return " ".join(partes)


def _limpar_descricao_painel(s: str | None) -> str:
    """Remove o cabecalho '## Passo NN — titulo' (redundante com h4) e mantem o resto."""
    if not s:
        return ""
    linhas = s.strip().splitlines()
    # Pula primeira linha se for "## Passo ..."
    while linhas and (linhas[0].strip().startswith("## ") or not linhas[0].strip()):
        linhas.pop(0)
    return "\n".join(linhas).strip()


# ============================================================
# Markdown
# ============================================================

def gerar_md(rel: dict, screenshot_url_fn=None) -> bytes:
    """Gera relatorio em Markdown (texto puro).

    screenshot_url_fn: opcional; se passado, transforma o path relativo do
    screenshot em URL absoluta. Caso contrario embute o caminho relativo
    como ![alt](path).
    """
    teste = rel["teste"]
    execucoes = rel["execucoes"]

    lines = []
    lines.append(f"# Relatorio do Teste — {teste['titulo']}")
    lines.append("")
    lines.append(f"- **ID:** `{teste['id']}`")
    lines.append(f"- **Sprint:** {teste.get('sprint_nome','-')}")
    lines.append(f"- **Tester:** {teste.get('tester','-')}")
    lines.append(f"- **Ciclo:** `{teste.get('ciclo_id','-')}`")
    lines.append(f"- **Estado:** **{teste.get('estado','-')}**")
    lines.append(f"- **Criado em:** {teste.get('criado_em') or '-'}")
    lines.append(f"- **Iniciado em:** {teste.get('iniciado_em') or '-'}")
    lines.append(f"- **Ultima atividade:** {teste.get('atualizado_em') or '-'}")
    lines.append(f"- **Concluido em:** {teste.get('concluido_em') or '-'}")
    lines.append("")

    # Sumario por CT
    lines.append("## Sumario")
    lines.append("")
    lines.append("| # | CT | UC | Estado | Passos | Aprov | Reprov | Obs | Duracao |")
    lines.append("|---|---|---|---|---|---|---|---|---|")
    total_passos = total_aprov = total_reprov = total_obs = 0
    total_dur = 0.0
    for e in execucoes:
        passos = e.get("passos") or []
        n = len(passos)
        ap = sum(1 for p in passos if p.get("veredicto_po") == "APROVADO")
        rp = sum(1 for p in passos if p.get("veredicto_po") == "REPROVADO")
        obs = sum(len(p.get("observacoes") or []) for p in passos)
        dur = (e.get("duracao_ms") or 0) / 1000.0
        lines.append(
            f"| {e['ordem']} | `{e['ct_id']}` | {e.get('uc_id','-')} — {e.get('uc_nome','-')} | "
            f"{e.get('estado','-')} | {n} | {ap} | {rp} | {obs} | {dur:.1f}s |"
        )
        total_passos += n; total_aprov += ap; total_reprov += rp
        total_obs += obs; total_dur += dur
    lines.append(f"| **TOTAL** | | | | **{total_passos}** | **{total_aprov}** | **{total_reprov}** | **{total_obs}** | **{total_dur:.1f}s** |")
    lines.append("")

    # Detalhes
    lines.append("## Detalhes")
    lines.append("")
    for e in execucoes:
        lines.append(f"### {e['ordem']}. `{e['ct_id']}` — {e.get('uc_id','-')} {e.get('uc_nome','')}")
        lines.append("")
        if e.get("ct_descricao"):
            lines.append(f"*{e['ct_descricao']}*")
            lines.append("")
        lines.append(f"Estado: **{e.get('estado','-')}** | Veredicto auto: {e.get('veredito_automatico','-')} | Veredicto PO: {e.get('veredicto_po','-')}")
        lines.append("")
        for p in (e.get("passos") or []):
            lines.append(f"#### Passo {p['ordem']} — `{p['passo_id']}`")
            if p.get("passo_titulo"):
                lines.append(f"**{p['passo_titulo']}**")
            lines.append("")
            lines.append(
                f"- Veredicto auto: **{p.get('veredito_automatico','-')}** | PO: **{p.get('veredicto_po','-')}** | "
                f"Duracao: {(p.get('duracao_ms') or 0)/1000:.1f}s"
            )
            if p.get("correcao_necessaria"):
                lines.append(f"- ⚠ Correcao necessaria: {p.get('correcao_descricao','(sem descricao)')}")
            lines.append("")

            # === Instrucao do passo (descricao_painel) ===
            descricao = _limpar_descricao_painel(p.get("descricao_painel"))
            if descricao:
                lines.append("**Instrucao do passo (o que o tester vai ver):**")
                lines.append("")
                lines.append("> " + descricao.replace("\n", "\n> "))
                lines.append("")

            # === O que foi feito (acoes do tutorial) ===
            acoes = p.get("acoes") or []
            if acoes:
                lines.append("**O que foi clicado/digitado neste passo:**")
                lines.append("")
                for i, a in enumerate(acoes, 1):
                    if isinstance(a, dict) and isinstance(a.get("sequencia"), list):
                        lines.append(f"{i}. **{(a.get('tipo') or 'sequencia').upper()}** — sequencia de {len(a['sequencia'])} sub-acoes:")
                        for j, sub in enumerate(a["sequencia"], 1):
                            lines.append(f"   {i}.{j}. {_resumir_acao(sub)}")
                    else:
                        lines.append(f"{i}. {_resumir_acao(a)}")
                lines.append("")

            # === Pontos a observar na tela ===
            pontos = p.get("pontos_observacao") or []
            if pontos:
                lines.append("**Pontos a observar na tela:**")
                lines.append("")
                for ponto in pontos:
                    lines.append(f"- {ponto}")
                lines.append("")

            # Screenshots
            for label, key in [("Antes", "screenshot_antes_path"), ("Depois", "screenshot_depois_path")]:
                ss = p.get(key)
                if ss:
                    if screenshot_url_fn:
                        url = screenshot_url_fn(ss)
                    else:
                        url = ss
                    lines.append(f"**{label}:**")
                    lines.append("")
                    lines.append(f"![{label}]({url})")
                    lines.append("")

            # Observacoes
            obs_list = p.get("observacoes") or []
            if obs_list:
                lines.append("**Observacoes:**")
                lines.append("")
                for o in obs_list:
                    txt = (o.get("texto") or "").replace("\n", " ")
                    when = o.get("criado_em", "")[:19].replace("T", " ")
                    lines.append(f"- *{when}* — {txt}")
                lines.append("")

    return "\n".join(lines).encode("utf-8")


# ============================================================
# DOCX
# ============================================================

def gerar_docx(rel: dict) -> bytes:
    """Gera relatorio DOCX com screenshots embutidos."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    teste = rel["teste"]
    execucoes = rel["execucoes"]

    doc = Document()

    # Titulo
    h = doc.add_heading(f"Relatorio — {teste['titulo']}", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Sumario do teste
    p = doc.add_paragraph()
    p.add_run("ID: ").bold = True
    p.add_run(teste.get("id", "-"))
    p = doc.add_paragraph()
    p.add_run("Sprint: ").bold = True
    p.add_run(teste.get("sprint_nome", "-"))
    p = doc.add_paragraph()
    p.add_run("Tester: ").bold = True
    p.add_run(teste.get("tester", "-"))
    p = doc.add_paragraph()
    p.add_run("Ciclo: ").bold = True
    p.add_run(teste.get("ciclo_id", "-") or "-")
    p = doc.add_paragraph()
    p.add_run("Estado: ").bold = True
    p.add_run(teste.get("estado", "-"))
    p = doc.add_paragraph()
    p.add_run("Criado em: ").bold = True
    p.add_run(teste.get("criado_em") or "-")
    p = doc.add_paragraph()
    p.add_run("Iniciado em: ").bold = True
    p.add_run(teste.get("iniciado_em") or "-")
    p = doc.add_paragraph()
    p.add_run("Ultima atividade: ").bold = True
    p.add_run(teste.get("atualizado_em") or "-")
    p = doc.add_paragraph()
    p.add_run("Concluido em: ").bold = True
    p.add_run(teste.get("concluido_em") or "-")

    # Sumario por CT
    doc.add_heading("Sumario", level=1)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["#", "CT", "UC", "Estado", "Passos", "Aprov", "Reprov", "Obs"]):
        hdr[i].text = h
    total_passos = total_aprov = total_reprov = total_obs = 0
    for e in execucoes:
        passos = e.get("passos") or []
        n = len(passos)
        ap = sum(1 for pp in passos if pp.get("veredicto_po") == "APROVADO")
        rp = sum(1 for pp in passos if pp.get("veredicto_po") == "REPROVADO")
        obs = sum(len(pp.get("observacoes") or []) for pp in passos)
        row = table.add_row().cells
        row[0].text = str(e.get("ordem", ""))
        row[1].text = e.get("ct_id", "")
        row[2].text = f"{e.get('uc_id','')} — {e.get('uc_nome','')}"
        row[3].text = e.get("estado", "") or ""
        row[4].text = str(n)
        row[5].text = str(ap)
        row[6].text = str(rp)
        row[7].text = str(obs)
        total_passos += n; total_aprov += ap; total_reprov += rp; total_obs += obs
    row = table.add_row().cells
    row[0].text = "TOTAL"
    row[4].text = str(total_passos)
    row[5].text = str(total_aprov)
    row[6].text = str(total_reprov)
    row[7].text = str(total_obs)
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True

    # Detalhes
    doc.add_heading("Detalhes", level=1)
    for e in execucoes:
        doc.add_heading(f"{e.get('ordem','')}. {e.get('ct_id','')} — {e.get('uc_id','')} {e.get('uc_nome','')}", level=2)
        if e.get("ct_descricao"):
            doc.add_paragraph(e["ct_descricao"], style="Intense Quote")
        p = doc.add_paragraph()
        p.add_run(f"Estado: {e.get('estado','-')} | Veredicto auto: {e.get('veredito_automatico','-')} | Veredicto PO: {e.get('veredicto_po','-')}")

        for pp in (e.get("passos") or []):
            doc.add_heading(f"Passo {pp.get('ordem','')} — {pp.get('passo_id','')}", level=3)
            if pp.get("passo_titulo"):
                p = doc.add_paragraph()
                p.add_run(pp["passo_titulo"]).bold = True
            p = doc.add_paragraph()
            p.add_run(
                f"Auto: {pp.get('veredito_automatico','-')} | PO: {pp.get('veredicto_po','-')} | "
                f"Duracao: {(pp.get('duracao_ms') or 0)/1000:.1f}s"
            )
            if pp.get("correcao_necessaria"):
                p = doc.add_paragraph()
                p.add_run(f"⚠ Correcao: {pp.get('correcao_descricao','-')}").bold = True

            # === Instrucao do passo ===
            descricao = _limpar_descricao_painel(pp.get("descricao_painel"))
            if descricao:
                p = doc.add_paragraph()
                p.add_run("Instrucao do passo:").bold = True
                doc.add_paragraph(descricao, style="Intense Quote")

            # === O que foi feito (acoes) ===
            acoes = pp.get("acoes") or []
            if acoes:
                p = doc.add_paragraph()
                p.add_run("O que foi clicado/digitado:").bold = True
                for i, a in enumerate(acoes, 1):
                    if isinstance(a, dict) and isinstance(a.get("sequencia"), list):
                        doc.add_paragraph(f"{i}. {(a.get('tipo') or 'sequencia').upper()} — {len(a['sequencia'])} sub-acoes:", style="List Number")
                        for j, sub in enumerate(a["sequencia"], 1):
                            doc.add_paragraph(f"  {i}.{j}. {_resumir_acao(sub)}", style="List Bullet 2")
                    else:
                        doc.add_paragraph(f"{i}. {_resumir_acao(a)}", style="List Number")

            # === Pontos a observar ===
            pontos = pp.get("pontos_observacao") or []
            if pontos:
                p = doc.add_paragraph()
                p.add_run("Pontos a observar na tela:").bold = True
                for ponto in pontos:
                    doc.add_paragraph(ponto, style="List Bullet")

            # Screenshots — embutidos
            for label, key in [("Antes", "screenshot_antes_path"), ("Depois", "screenshot_depois_path")]:
                ss = pp.get(key)
                if not ss:
                    continue
                full = _PROJECT / ss
                if not full.exists() or not full.is_file():
                    doc.add_paragraph(f"[{label}: arquivo nao encontrado: {ss}]")
                    continue
                p = doc.add_paragraph()
                p.add_run(f"{label}:").bold = True
                try:
                    doc.add_picture(str(full), width=Inches(6.0))
                except Exception as ex:
                    doc.add_paragraph(f"[falha embutir {label}: {ex}]")

            # Observacoes
            obs_list = pp.get("observacoes") or []
            if obs_list:
                p = doc.add_paragraph()
                p.add_run("Observacoes:").bold = True
                for o in obs_list:
                    when = (o.get("criado_em") or "")[:19].replace("T", " ")
                    txt = o.get("texto") or ""
                    doc.add_paragraph(f"[{when}] {txt}", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ============================================================
# PDF (via WeasyPrint a partir de HTML)
# ============================================================

def gerar_pdf(rel: dict) -> bytes:
    """Gera PDF a partir de HTML renderizado com WeasyPrint."""
    import weasyprint

    teste = rel["teste"]
    execucoes = rel["execucoes"]

    # Calcula totais
    total_passos = total_aprov = total_reprov = total_obs = 0
    for e in execucoes:
        passos = e.get("passos") or []
        total_passos += len(passos)
        total_aprov += sum(1 for p in passos if p.get("veredicto_po") == "APROVADO")
        total_reprov += sum(1 for p in passos if p.get("veredicto_po") == "REPROVADO")
        total_obs += sum(len(p.get("observacoes") or []) for p in passos)

    def _esc(s: Any) -> str:
        if s is None: return "-"
        return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _img_src(path: str | None) -> str:
        """Inline de imagem como file:// (WeasyPrint le do disco direto)."""
        if not path:
            return ""
        full = _PROJECT / path
        if not full.exists():
            return ""
        return f"file://{full.as_posix()}"

    html_parts = ["""<!DOCTYPE html>
<html lang="pt-br">
<head>
<meta charset="utf-8">
<title>Relatorio de Teste</title>
<style>
  @page { size: A4; margin: 1.5cm; }
  body { font-family: 'Helvetica', sans-serif; font-size: 10pt; color: #222; }
  h1 { font-size: 18pt; color: #1a1a40; border-bottom: 2px solid #1a1a40; padding-bottom: 4px; }
  h2 { font-size: 14pt; color: #1a1a40; margin-top: 1.5em; border-bottom: 1px solid #ccc; }
  h3 { font-size: 12pt; color: #2a4a8a; margin-top: 1.2em; }
  h4 { font-size: 11pt; color: #444; margin-top: 0.8em; }
  table { width: 100%; border-collapse: collapse; margin: 0.5em 0; font-size: 9pt; }
  th, td { border: 1px solid #ccc; padding: 4px 6px; text-align: left; }
  th { background: #f0f0f0; font-weight: bold; }
  .meta { background: #f8f8f8; padding: 0.5em 1em; border-left: 4px solid #1a1a40; margin: 0.5em 0; font-size: 9pt; }
  .meta strong { color: #1a1a40; }
  .veredicto { display: inline-block; padding: 1px 6px; border-radius: 3px; font-size: 8pt; font-weight: bold; }
  .v-aprovado { background: #1a3a1a; color: #4a8a4a; }
  .v-reprovado { background: #3a1a1a; color: #8a4a4a; }
  .v-inconclusivo { background: #3a3a1a; color: #8a8a4a; }
  .v-pendente { background: #333; color: #888; }
  .ss-pair { display: flex; gap: 8px; margin: 6px 0; flex-wrap: wrap; page-break-inside: avoid; }
  .ss-pair > div { flex: 1; min-width: 45%; }
  .ss-pair img { width: 100%; max-width: 360px; border: 1px solid #999; }
  .ss-label { font-size: 8pt; color: #666; font-weight: bold; }
  .obs { background: #f0f4ff; border-left: 3px solid #4a6aaa; padding: 4px 8px; margin: 4px 0; font-size: 9pt; }
  .obs-when { color: #888; font-size: 8pt; margin-right: 6px; }
  .bloco-label { font-weight: bold; color: #1a1a40; margin: 6px 0 2px 0; font-size: 9pt; }
  .instrucao { background: #fafafa; border-left: 3px solid #888; padding: 4px 10px; margin: 6px 0; font-size: 9pt; }
  .instrucao p { margin: 4px 0; }
  .acoes { background: #f4f7f4; border-left: 3px solid #4a8a4a; padding: 4px 10px; margin: 6px 0; font-size: 8.5pt; }
  .acoes ol { padding-left: 20px; margin: 4px 0; }
  .acoes code { background: #fff; border: 1px solid #ddd; padding: 0 4px; font-size: 8pt; }
  .pontos { background: #fff8f0; border-left: 3px solid #c08040; padding: 4px 10px; margin: 6px 0; font-size: 9pt; }
  .pontos ul { padding-left: 20px; margin: 4px 0; }
  .passo-block { page-break-inside: avoid; margin-bottom: 1em; padding: 6px; border-left: 3px solid #ccc; }
  .passo-aprovado { border-left-color: #4a8a4a; }
  .passo-reprovado { border-left-color: #8a4a4a; }
  .ct-block { margin-top: 1.5em; page-break-before: auto; }
  .total-row { background: #e0e0e0; font-weight: bold; }
  code { background: #f0f0f0; padding: 1px 4px; border-radius: 2px; font-size: 9pt; }
</style>
</head>
<body>"""]

    html_parts.append(f"<h1>Relatorio — {_esc(teste.get('titulo'))}</h1>")
    html_parts.append('<div class="meta">')
    html_parts.append(f"<p><strong>ID:</strong> <code>{_esc(teste.get('id'))}</code></p>")
    html_parts.append(f"<p><strong>Sprint:</strong> {_esc(teste.get('sprint_nome'))}</p>")
    html_parts.append(f"<p><strong>Tester:</strong> {_esc(teste.get('tester'))}</p>")
    html_parts.append(f"<p><strong>Ciclo:</strong> <code>{_esc(teste.get('ciclo_id'))}</code></p>")
    html_parts.append(f"<p><strong>Estado:</strong> {_esc(teste.get('estado'))}</p>")
    html_parts.append(f"<p><strong>Criado em:</strong> {_esc(teste.get('criado_em'))}</p>")
    html_parts.append(f"<p><strong>Iniciado em:</strong> {_esc(teste.get('iniciado_em'))}</p>")
    html_parts.append(f"<p><strong>Ultima atividade:</strong> {_esc(teste.get('atualizado_em'))}</p>")
    html_parts.append(f"<p><strong>Concluido em:</strong> {_esc(teste.get('concluido_em'))}</p>")
    html_parts.append("</div>")

    # Sumario
    html_parts.append("<h2>Sumario</h2>")
    html_parts.append("<table><thead><tr><th>#</th><th>CT</th><th>UC</th><th>Estado</th>"
                      "<th>Passos</th><th>Aprov</th><th>Reprov</th><th>Obs</th></tr></thead><tbody>")
    for e in execucoes:
        passos = e.get("passos") or []
        n = len(passos)
        ap = sum(1 for p in passos if p.get("veredicto_po") == "APROVADO")
        rp = sum(1 for p in passos if p.get("veredicto_po") == "REPROVADO")
        obs = sum(len(p.get("observacoes") or []) for p in passos)
        html_parts.append(
            f"<tr><td>{e.get('ordem','')}</td><td><code>{_esc(e.get('ct_id'))}</code></td>"
            f"<td>{_esc(e.get('uc_id'))} — {_esc(e.get('uc_nome'))}</td>"
            f"<td>{_esc(e.get('estado'))}</td>"
            f"<td>{n}</td><td>{ap}</td><td>{rp}</td><td>{obs}</td></tr>"
        )
    html_parts.append(
        f'<tr class="total-row"><td colspan="4">TOTAL</td>'
        f'<td>{total_passos}</td><td>{total_aprov}</td><td>{total_reprov}</td><td>{total_obs}</td></tr>'
    )
    html_parts.append("</tbody></table>")

    # Detalhes
    html_parts.append("<h2>Detalhes</h2>")
    for e in execucoes:
        html_parts.append('<div class="ct-block">')
        html_parts.append(f"<h3>{e.get('ordem','')}. <code>{_esc(e.get('ct_id'))}</code> — "
                          f"{_esc(e.get('uc_id'))} {_esc(e.get('uc_nome'))}</h3>")
        if e.get("ct_descricao"):
            html_parts.append(f"<p><em>{_esc(e['ct_descricao'])}</em></p>")
        html_parts.append(f"<p>Estado: <strong>{_esc(e.get('estado'))}</strong></p>")
        for pp in (e.get("passos") or []):
            v_po = (pp.get("veredicto_po") or "pendente").lower()
            v_auto = (pp.get("veredito_automatico") or "pendente").lower()
            html_parts.append(f'<div class="passo-block passo-{v_po}">')
            html_parts.append(f"<h4>Passo {pp.get('ordem','')} — <code>{_esc(pp.get('passo_id'))}</code></h4>")
            if pp.get("passo_titulo"):
                html_parts.append(f"<p><strong>{_esc(pp['passo_titulo'])}</strong></p>")
            html_parts.append(
                f'<p>Auto: <span class="veredicto v-{v_auto}">{_esc(pp.get("veredito_automatico"))}</span> | '
                f'PO: <span class="veredicto v-{v_po}">{_esc(pp.get("veredicto_po"))}</span> | '
                f'Duracao: {(pp.get("duracao_ms") or 0)/1000:.1f}s</p>'
            )
            if pp.get("correcao_necessaria"):
                html_parts.append(f"<p><strong>⚠ Correcao:</strong> {_esc(pp.get('correcao_descricao'))}</p>")

            # === Instrucao do passo ===
            descricao = _limpar_descricao_painel(pp.get("descricao_painel"))
            if descricao:
                html_parts.append('<div class="instrucao">')
                html_parts.append('<div class="bloco-label">Instrucao do passo:</div>')
                # Converte \n em <br> e preserva markdown leve
                desc_html = _esc(descricao).replace("\n\n", "</p><p>").replace("\n", "<br>")
                html_parts.append(f"<p>{desc_html}</p>")
                html_parts.append("</div>")

            # === O que foi feito ===
            acoes = pp.get("acoes") or []
            if acoes:
                html_parts.append('<div class="acoes">')
                html_parts.append('<div class="bloco-label">O que foi clicado/digitado:</div>')
                html_parts.append("<ol>")
                for a in acoes:
                    if isinstance(a, dict) and isinstance(a.get("sequencia"), list):
                        html_parts.append(f"<li><strong>{_esc((a.get('tipo') or 'sequencia').upper())}</strong> — sequencia de {len(a['sequencia'])} sub-acoes:")
                        html_parts.append("<ol>")
                        for sub in a["sequencia"]:
                            html_parts.append(f"<li><code>{_esc(_resumir_acao(sub))}</code></li>")
                        html_parts.append("</ol></li>")
                    else:
                        html_parts.append(f"<li><code>{_esc(_resumir_acao(a))}</code></li>")
                html_parts.append("</ol></div>")

            # === Pontos a observar ===
            pontos = pp.get("pontos_observacao") or []
            if pontos:
                html_parts.append('<div class="pontos">')
                html_parts.append('<div class="bloco-label">Pontos a observar na tela:</div>')
                html_parts.append("<ul>")
                for ponto in pontos:
                    html_parts.append(f"<li>{_esc(ponto)}</li>")
                html_parts.append("</ul></div>")

            # Screenshots
            ss_a = _img_src(pp.get("screenshot_antes_path"))
            ss_d = _img_src(pp.get("screenshot_depois_path"))
            if ss_a or ss_d:
                html_parts.append('<div class="ss-pair">')
                if ss_a:
                    html_parts.append(f'<div><div class="ss-label">ANTES</div><img src="{ss_a}"/></div>')
                if ss_d:
                    html_parts.append(f'<div><div class="ss-label">DEPOIS</div><img src="{ss_d}"/></div>')
                html_parts.append("</div>")

            # Observacoes
            for o in (pp.get("observacoes") or []):
                when = (o.get("criado_em") or "")[:19].replace("T", " ")
                html_parts.append(
                    f'<div class="obs"><span class="obs-when">[{_esc(when)}]</span>{_esc(o.get("texto",""))}</div>'
                )
            html_parts.append("</div>")
        html_parts.append("</div>")

    html_parts.append("</body></html>")
    html = "\n".join(html_parts)

    pdf_bytes = weasyprint.HTML(string=html, base_url=str(_PROJECT)).write_pdf()
    return pdf_bytes
