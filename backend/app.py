"""
Agente de Editais - Backend Flask
MVP com 9 ações via Select + Prompt
"""
import os
import json
from flask import Flask, request, jsonify, send_file, Response
from flask_cors import CORS

from models import init_db, get_db, User, Session, Message, RefreshToken, Produto, ProdutoEspecificacao, Edital, Analise, Proposta, FonteEdital, Contrato, EstrategiaEdital, Monitoramento, ModalidadeLicitacao, OrigemOrgao, AreaProduto, ClasseProdutoV2, SubclasseProduto, Empresa
from llm import call_deepseek
from tools import (
    _web_search, tool_web_search, tool_download_arquivo, tool_processar_upload,
    tool_extrair_especificacoes, tool_cadastrar_fonte, tool_listar_fontes,
    tool_buscar_editais_fonte, tool_buscar_editais_scraper, tool_extrair_requisitos,
    tool_listar_editais, tool_listar_produtos, tool_calcular_aderencia, tool_gerar_proposta,
    tool_calcular_score_aderencia, tool_salvar_editais_selecionados,
    tool_reprocessar_produto, tool_atualizar_produto,
    tool_buscar_links_editais,
    execute_tool, _extrair_info_produto, PROMPT_EXTRAIR_SPECS,
    tool_calcular_scores_validacao, tool_atualizar_status_proposta
)
from config import UPLOAD_FOLDER, MAX_HISTORY_MESSAGES

import bcrypt
import jwt
import uuid
from datetime import datetime, timedelta
from functools import wraps

app = Flask(__name__)
CORS(app, origins=["http://localhost:5173", "http://localhost:5175", "http://localhost:3000"])

# Registrar blueprint CRUD genérico
from crud_routes import crud_bp
app.register_blueprint(crud_bp)

# JWT Config
JWT_SECRET = "editais-ia-secret-key-change-in-production-2024"
JWT_EXPIRY_HOURS = 24

# Prompts prontos para o dropdown
PROMPTS_PRONTOS = [
    # === CONSULTAS ANALÍTICAS (MindsDB) ===
    {"id": "mindsdb_totais", "nome": "📊 Totais (produtos/editais)", "prompt": "Quantos produtos e editais existem no banco?"},
    {"id": "mindsdb_editais_novos", "nome": "📊 Editais com status novo", "prompt": "Quais editais estão com status novo?"},
    {"id": "mindsdb_editais_orgao", "nome": "📊 Editais por órgão", "prompt": "Liste editais do Ministério da Saúde"},
    {"id": "mindsdb_editais_mes", "nome": "📊 Editais do mês", "prompt": "Quais editais têm data de abertura em fevereiro de 2026?"},
    {"id": "mindsdb_score_medio", "nome": "📊 Score médio de aderência", "prompt": "Qual é o score médio de aderência das análises?"},
    {"id": "mindsdb_produtos_categoria", "nome": "📊 Produtos por categoria", "prompt": "Quantos produtos temos em cada categoria?"},
    {"id": "mindsdb_alta_aderencia", "nome": "📊 Produtos c/ alta aderência", "prompt": "Quais produtos têm aderência acima de 70% em algum edital?"},
    {"id": "mindsdb_propostas", "nome": "📊 Total de propostas", "prompt": "Quantas propostas foram geradas?"},
    {"id": "mindsdb_editais_semana", "nome": "📊 Editais da semana", "prompt": "Quais editais vencem esta semana?"},
    {"id": "mindsdb_melhor_produto", "nome": "📊 Produto c/ melhor score", "prompt": "Qual produto tem o melhor score de aderência?"},
    {"id": "mindsdb_editais_uf", "nome": "📊 Editais por UF", "prompt": "Quantos editais temos por estado (UF)?"},
    {"id": "mindsdb_resumo", "nome": "📊 Resumo geral do banco", "prompt": "Faça um resumo do banco: total de produtos, editais, análises e propostas"},
    # === AÇÕES DO SISTEMA ===
    {"id": "listar_produtos", "nome": "Listar meus produtos", "prompt": "Liste todos os meus produtos cadastrados"},
    {"id": "listar_editais", "nome": "Listar editais abertos", "prompt": "Quais editais estão abertos?"},
    {"id": "calcular_aderencia", "nome": "Calcular aderência", "prompt": "Calcule a aderência do produto [NOME_PRODUTO] ao edital [NUMERO_EDITAL]"},
    {"id": "gerar_proposta", "nome": "Gerar proposta", "prompt": "Gere uma proposta do produto [NOME_PRODUTO] para o edital [NUMERO_EDITAL] com preço R$ [VALOR]"},
    {"id": "buscar_editais", "nome": "Buscar editais", "prompt": "Busque editais de [TERMO] no PNCP"},
    {"id": "cadastrar_fonte", "nome": "Cadastrar fonte", "prompt": "Cadastre a fonte [NOME], tipo [api/scraper], URL [URL]"},
    {"id": "listar_fontes", "nome": "Listar fontes", "prompt": "Quais são as fontes de editais cadastradas?"},
    {"id": "ajuda", "nome": "O que posso fazer?", "prompt": "O que você pode fazer? Quais são suas capacidades?"},
    # === REGISTRO DE RESULTADOS (Sprint 1) ===
    {"id": "registrar_derrota", "nome": "📉 Registrar derrota", "prompt": "Perdemos o edital [NUMERO] para [EMPRESA] com R$ [VALOR_VENCEDOR], nosso preço foi R$ [NOSSO_VALOR]"},
    {"id": "registrar_vitoria", "nome": "🏆 Registrar vitória", "prompt": "Ganhamos o edital [NUMERO] com R$ [VALOR]"},
    {"id": "registrar_cancelado", "nome": "⛔ Edital cancelado", "prompt": "O edital [NUMERO] foi cancelado"},
    {"id": "consultar_resultado", "nome": "🔎 Consultar resultado", "prompt": "Qual o resultado do edital [NUMERO]?"},
    # === BUSCA COM FILTROS ===
    {"id": "buscar_modalidade", "nome": "🔍 Buscar por modalidade", "prompt": "Busque editais de [TERMO] com modalidade [MODALIDADE]"},
    {"id": "buscar_tipo_produto", "nome": "🔍 Buscar por tipo produto", "prompt": "Busque editais de [TERMO] do tipo [TIPO_PRODUTO]"},
    {"id": "buscar_origem", "nome": "🔍 Buscar por origem", "prompt": "Busque editais de [TERMO] de origem [ORIGEM]"},
    {"id": "buscar_filtros", "nome": "🔍 Buscar com filtros", "prompt": "Busque editais de [TERMO] em [UF], modalidade [MODALIDADE], origem [ORIGEM]"},
    # === BUSCA WEB E ANVISA ===
    {"id": "buscar_anvisa", "nome": "🛡️ Buscar registro ANVISA", "prompt": "Busque registros ANVISA para o produto [NOME_PRODUTO]"},
    {"id": "buscar_web_manual", "nome": "🌐 Buscar manual na web", "prompt": "Busque o manual do produto [NOME_PRODUTO] na web e cadastre"},
]


PROMPT_CLASSIFICAR_INTENCAO = """Você é um agente classificador de intenções para um sistema de gestão de editais e licitações.

Analise a mensagem do usuário e classifique em UMA das categorias abaixo:

## CATEGORIAS DO SISTEMA:

### AÇÕES COM ARQUIVOS (quando usuário envia um PDF/documento):
1. **arquivo_cadastrar**: Cadastrar o arquivo como produto no sistema (PADRÃO se não especificar outra ação)
   Exemplos: "cadastre", "salve como produto", "registre", "" (vazio), "cadastre como Analisador X"

2. **arquivo_mostrar**: Mostrar/exibir o conteúdo do arquivo
   Exemplos: "mostre o conteúdo", "exiba o texto", "o que tem nesse PDF?", "leia o documento", "mostra"

3. **arquivo_specs**: Extrair e listar especificações técnicas (sem cadastrar)
   Exemplos: "quais especificações?", "extraia as specs", "liste as características técnicas"

4. **arquivo_resumir**: Fazer um resumo do documento
   Exemplos: "resuma", "faça um resumo", "sintetize", "resuma o documento"

5. **arquivo_analisar**: Fazer análise detalhada do documento
   Exemplos: "analise", "faça uma análise", "avalie o documento", "o que você acha desse manual?"

6. **extrair_ata**: Extrair resultados de uma ata de sessão de pregão/licitação
   Exemplos: "extraia os resultados desta ata", "quem ganhou este pregão?", "registre os resultados da ata", "extraia vencedores", "resultado da licitação"
   Palavras-chave: ata, resultados da ata, vencedor do pregão, extrair resultados, ata de sessão
   IMPORTANTE: Use quando o arquivo é uma ATA de sessão (não um manual de produto)

### AÇÕES DE BUSCA:
7. **buscar_web**: Buscar MATERIAIS/MANUAIS/DATASHEETS na WEB (não editais!)
   Exemplos: "busque na web o manual do equipamento X", "encontre o datasheet do Y"

8. **download_url**: Baixar arquivo de uma URL específica
   Exemplos: "baixe o arquivo da URL: http://...", "baixe https://..."
   IMPORTANTE: Se contém URL (http:// ou https://), classifique como download_url!

9. **buscar_editais**: Buscar EDITAIS/LICITAÇÕES em portais (PNCP, BEC) por TERMO/ÁREA COM cálculo de score de aderência
   Exemplos: "busque editais de tecnologia", "editais da área médica", "busque editais de hematologia"
   Também cobre score rápido, híbrido e profundo: "busque editais de hematologia com score rápido", "busque editais com score profundo nos 5 melhores", "busque editais com score híbrido", "busque editais com análise profunda dos 10 melhores"
   IMPORTANTE: Use quando buscar por TERMO genérico (área, categoria, produto) E calcular score de aderência (qualquer tipo de score: rápido, híbrido ou profundo)

9b. **buscar_editais_simples**: Buscar EDITAIS SEM calcular score - apenas listar os editais encontrados
   Exemplos: "busque editais de tecnologia sem score", "liste editais de hematologia sem calcular aderência", "busque editais de informática apenas listando"
   Palavras-chave: sem score, sem calcular, sem aderência, apenas listar, só listar, listar editais
   IMPORTANTE: Use quando o usuário quer apenas ver os editais sem NENHUMA análise de aderência

10. **buscar_edital_numero**: Buscar UM edital específico pelo NÚMERO
   Exemplos: "busque o edital PE-001/2026", "encontre o edital 90186", "busque edital número 123/2025"
   Palavras-chave: busque o edital, encontre o edital, buscar edital número, edital PE-, edital nº
   IMPORTANTE: Use quando o usuário menciona um NÚMERO específico de edital

### AÇÕES DE LISTAGEM:
9. **listar_editais**: Ver editais JÁ SALVOS no sistema
   Exemplos: "liste meus editais", "editais salvos"

10. **listar_produtos**: Ver produtos cadastrados
    Exemplos: "liste meus produtos", "quais produtos tenho"

11. **listar_fontes**: Ver fontes de editais cadastradas
    Exemplos: "quais fontes?", "liste fontes"

12. **listar_propostas**: Ver propostas técnicas geradas
    Exemplos: "liste minhas propostas", "quais propostas tenho", "propostas geradas", "ver propostas"

### AÇÕES DE PROCESSAMENTO:
13. **calcular_aderencia**: Calcular score produto vs edital
    Exemplos: "calcule aderência do produto X com edital Y"

14. **gerar_proposta**: Gerar proposta técnica
    Exemplos: "gere proposta para o edital X"

15. **cadastrar_fonte**: Cadastrar nova fonte de editais
    Exemplos: "cadastre a fonte BEC-SP"

16. **salvar_editais**: Salvar editais da última busca (um específico ou todos)
    Exemplos: "salve os editais", "salvar recomendados", "salvar todos", "salvar edital 02223/2025", "salvar edital PE-001/2026"
    IMPORTANTE: Use quando o usuário quer SALVAR editais que vieram de uma BUSCA anterior. Diferente de cadastrar_edital que é para criar um edital MANUALMENTE com dados informados.

17. **reprocessar_produto**: Reprocessar/atualizar especificações de um produto
    Exemplos: "reprocesse o produto X", "atualize specs do produto X", "extraia novamente as especificações"

18. **consulta_mindsdb**: Consultas analíticas complexas sobre editais e produtos via linguagem natural
    Exemplos: "qual o score médio de aderência?", "quantos editais por estado?", "qual produto tem melhor desempenho?", "estatísticas dos editais", "análise dos dados", "relatório de editais"
    Use quando: perguntas analíticas, estatísticas, agregações, comparações, rankings, tendências

19. **registrar_resultado**: Registrar resultado de certame (vitória ou derrota) - AFIRMAÇÕES
    Exemplos: "perdemos o edital PE-001", "ganhamos o pregão", "vencedor foi empresa X com R$ 100k", "registre derrota no PE-002", "perdemos por preço para MedLab"
    Palavras-chave: perdemos, ganhamos, vencedor, derrota, vitória, segundo lugar
    IMPORTANTE: Use apenas quando o usuário está AFIRMANDO um resultado, não perguntando.

20. **consultar_resultado**: Consultar/perguntar sobre resultado de certames - PERGUNTAS
    Exemplos: "qual o resultado do edital PE-001?", "quem ganhou o pregão?", "como foi o edital?", "mostre os resultados de todos os editais", "ver resultados dos editais", "listar resultados", "resultados dos certames"
    Palavras-chave: qual o resultado, quem ganhou, quem venceu, como foi, resultados de todos, ver resultados, listar resultados
    IMPORTANTE: Use quando o usuário está PERGUNTANDO sobre resultados (um edital ou todos).

21. **buscar_atas_pncp**: Buscar atas de sessão/registro de preço no PNCP para download
    Exemplos: "busque atas de hematologia", "encontre atas de pregão de equipamentos", "baixe atas de registro de preço", "atas de sessão de pregão"
    Palavras-chave: buscar atas, encontrar atas, baixar atas, atas de registro, atas de sessão, atas pncp
    IMPORTANTE: Use quando o usuário quer BUSCAR atas no portal PNCP (não quando já tem um arquivo)

22. **buscar_precos_pncp**: Buscar preços históricos de contratos no PNCP
    Exemplos: "busque preços de hematologia", "qual o preço de mercado para analisador?", "preços de contratos de equipamentos", "quanto custa um equipamento X no PNCP?"
    Palavras-chave: buscar preços, preço de mercado, preços pncp, quanto custa, preço médio, valores de contrato
    Use quando: usuário quer saber preços praticados em licitações anteriores

23. **historico_precos**: Consultar histórico de preços registrados no sistema
    Exemplos: "mostre histórico de preços de hematologia", "histórico de preços do produto X", "quais preços já registramos?"
    Palavras-chave: histórico de preços, preços registrados, preços salvos, histórico preço

24. **listar_concorrentes**: Listar todos os concorrentes conhecidos
    Exemplos: "liste os concorrentes", "quais concorrentes conhecemos?", "mostre os concorrentes"
    Palavras-chave: listar concorrentes, concorrentes conhecidos, nossos concorrentes

25. **analisar_concorrente**: Analisar um concorrente específico
    Exemplos: "analise o concorrente MedLab", "como está a empresa TechSaúde?", "histórico do concorrente X"
    Palavras-chave: analisar concorrente, análise concorrente, histórico concorrente

26. **recomendar_preco**: Recomendar preço para um produto/edital
    Exemplos: "qual preço sugerir para hematologia?", "recomende preço para analisador", "que preço colocar?"
    Palavras-chave: recomendar preço, sugerir preço, que preço, qual preço colocar

27. **classificar_edital**: Classificar tipo de edital (comodato, venda, aluguel)
    Exemplos: "classifique este edital", "que tipo de edital é este?", "é comodato ou venda?"
    Palavras-chave: classificar edital, tipo de edital, qual modalidade, é comodato

28. **verificar_completude**: Verificar se produto tem todas informações necessárias
    Exemplos: "produto X está completo?", "verifique completude do produto", "falta algo no produto?"
    Palavras-chave: verificar completude, produto completo, falta informação, completude produto

29. **cadastrar_edital**: Cadastrar/registrar manualmente um edital no sistema COM DADOS INFORMADOS PELO USUÁRIO
    Exemplos: "cadastre o edital PE-001/2026, órgão Ministério da Saúde, objeto: aquisição de equipamentos", "registre este edital com os dados...", "adicione o edital número X do órgão Y"
    Palavras-chave: cadastre edital, registre edital, adicione edital, cadastrar edital manualmente, inserir edital
    IMPORTANTE: Use APENAS quando o usuário quer CRIAR um edital MANUALMENTE informando dados (órgão, objeto, etc).
    NÃO USE para "salvar edital NUMERO" que veio de uma busca - isso é salvar_editais!

### SPRINT 2 - ALERTAS E MONITORAMENTO:
30. **configurar_alertas**: Configurar alertas de prazo para um edital
    Exemplos: "configure alertas para PE-001", "avise-me 24h antes da abertura", "quero alerta de impugnação"
    Palavras-chave: configurar alerta, avise-me, lembre-me, alertar antes, alertas para edital

31. **listar_alertas**: Ver alertas configurados / próximos pregões
    Exemplos: "quais alertas tenho?", "meus alertas", "próximos pregões", "alertas configurados"
    Palavras-chave: listar alertas, meus alertas, alertas configurados, próximos pregões

32. **dashboard_prazos**: Ver dashboard de prazos e contagem regressiva
    Exemplos: "mostre dashboard de prazos", "quais editais abrem esta semana?", "timer dos editais"
    Palavras-chave: dashboard prazos, editais abrem, contagem regressiva, timer editais

33. **calendario_editais**: Ver calendário de editais
    Exemplos: "calendário de fevereiro", "calendário de editais", "editais do mês"
    Palavras-chave: calendário editais, calendário mês, ver calendário

34. **configurar_monitoramento**: Configurar monitoramento automático de editais
    Exemplos: "monitore editais de hematologia", "configure busca automática", "avise novos editais de X"
    Palavras-chave: monitorar editais, monitoramento automático, busca automática, avisar novos

35. **listar_monitoramentos**: Ver monitoramentos configurados
    Exemplos: "quais monitoramentos tenho?", "monitoramentos ativos", "ver minhas buscas automáticas"
    Palavras-chave: listar monitoramentos, monitoramentos ativos, minhas buscas

36. **desativar_monitoramento**: Desativar um monitoramento
    Exemplos: "desative monitoramento de hematologia", "pare de monitorar X", "cancele busca automática"
    Palavras-chave: desativar monitoramento, parar monitorar, cancelar busca

37. **configurar_notificacoes**: Configurar preferências de notificação
    Exemplos: "configure meu email de notificação", "quero alertas das 8h às 18h", "preferências de alerta"
    Palavras-chave: configurar notificação, email notificação, preferências alerta

38. **historico_notificacoes**: Ver histórico de notificações
    Exemplos: "histórico de notificações", "notificações não lidas", "ver notificações"
    Palavras-chave: histórico notificações, notificações não lidas, ver notificações

39. **extrair_datas_edital**: Extrair datas importantes de um edital (PDF)
    Exemplos: "extraia as datas deste edital", "quando abre o edital?", "prazo de impugnação"
    Palavras-chave: extrair datas, datas edital, quando abre, prazo impugnação

40. **cancelar_alerta**: Cancelar alertas configurados
    Exemplos: "cancele alertas do PE-001", "remova meus alertas", "desative alertas"
    Palavras-chave: cancelar alerta, remover alerta, desativar alerta

### FASE 1 — PRECIFICAÇÃO:
41b. **precif_organizar_lotes**: Organizar itens do edital em lotes
    Exemplos: "organize os lotes do edital PE-001", "importe itens do PNCP e crie lotes", "crie lotes para o edital"
    Palavras-chave: organizar lotes, criar lotes, importar itens, lotes do edital

42b. **precif_selecao_portfolio**: Seleção inteligente de produto do portfolio para item do edital
    Exemplos: "selecione produto para o item 1", "qual produto atende o item?", "seleção inteligente de portfolio"
    Palavras-chave: selecionar produto, seleção portfolio, match produto item, vincular produto

43b. **precif_calcular_volumetria**: Calcular volumetria técnica (kits necessários)
    Exemplos: "calcule a volumetria", "quantos kits preciso?", "calcule kits para 50000 testes"
    Palavras-chave: calcular volumetria, quantos kits, volume ajustado, rendimento

44b. **precif_configurar_custos**: Configurar base de custos (Camada A)
    Exemplos: "configure os custos do item", "custo base do produto", "configure camada A de custos"
    Palavras-chave: configurar custos, custo base, camada A, base de custos, NCM, ICMS

45b. **precif_preco_base**: Montar preço base (Camada B)
    Exemplos: "defina preço base com markup 76%", "preço base manual R$ 150", "configure camada B"
    Palavras-chave: preço base, markup, camada B, montar preço

46b. **precif_valor_referencia**: Definir valor de referência/target (Camada C)
    Exemplos: "defina valor de referência", "target de preço", "configure camada C", "referência do edital"
    Palavras-chave: valor referência, target, camada C, referência edital

47b. **precif_estruturar_lances**: Estruturar lances — valor inicial e mínimo (Camadas D e E)
    Exemplos: "configure lances: inicial R$ 145, mínimo R$ 95", "estruture os lances", "defina faixa de disputa"
    Palavras-chave: estruturar lances, lance inicial, lance mínimo, faixa de disputa, camada D, camada E

48b. **precif_estrategia**: Definir estratégia competitiva e simular cenários
    Exemplos: "quero ganhar este edital", "simule estratégia competitiva", "estratégia agressiva"
    Palavras-chave: estratégia competitiva, quero ganhar, simular disputa, cenários

49b. **precif_historico_camada_f**: Consultar histórico de preços para Camada F
    Exemplos: "busque histórico de preços do produto", "consulte preços históricos", "camada F do item"
    Palavras-chave: histórico preços camada, camada F, preços passados

50b. **precif_comodato**: Gestão de comodato
    Exemplos: "cadastre comodato para o edital", "configure equipamento em comodato", "amortização do equipamento"
    Palavras-chave: comodato, amortização, equipamento comodato, gestão comodato

41. **chat_livre**: Dúvidas gerais, conversas
    Exemplos: "o que é pregão?", "olá", "obrigado"

### CADASTRO DE PRODUTOS:
47. **cadastrar_produto**: Cadastrar um produto manualmente com dados do formulário
    Exemplos: "cadastre manualmente o produto: Nome=...", "registre o produto X", "cadastre o produto Y"
    Palavras-chave: cadastre manualmente, cadastre o produto, registre o produto, criar produto

48. **buscar_anvisa**: Buscar registros de produtos na ANVISA
    Exemplos: "busque o registro ANVISA numero 12345", "busque registros ANVISA para o produto X", "consulte ANVISA do produto Y"
    Palavras-chave: anvisa, registro anvisa, buscar anvisa, consultar anvisa

### ANÁLISE DE EDITAIS:
42. **resumir_edital**: Fazer um resumo de um edital cadastrado
    Exemplos: "resuma o edital PE-001/2026", "faça um resumo do edital", "resumo do edital PE-001", "sintetize o edital"
    Palavras-chave: resumir edital, resumo do edital, sintetize edital, resumo edital
    IMPORTANTE: O usuário quer um resumo executivo do edital (objeto, valor, prazos, requisitos principais)

43. **perguntar_edital**: Responder dúvidas/perguntas sobre um edital específico
    Exemplos: "qual o prazo de entrega do edital PE-001?", "o edital PE-001 exige garantia?", "quais documentos são exigidos no PE-001?", "pergunte ao edital PE-001 sobre [DÚVIDA]"
    Palavras-chave: perguntar ao edital, dúvida sobre edital, o edital exige, o edital pede, prazo do edital, requisitos do edital
    IMPORTANTE: Use quando o usuário tem uma dúvida específica sobre um edital cadastrado

44. **baixar_pdf_edital**: Baixar o PDF de um edital já cadastrado (a partir da URL salva)
    Exemplos: "baixe o PDF do edital PE-001/2026", "faça download do edital PE-001", "baixar edital PE-001", "download do pdf do edital"
    Palavras-chave: baixar pdf edital, download edital, baixar edital, baixe o edital, download pdf edital
    IMPORTANTE: Use quando o usuário quer BAIXAR o arquivo PDF de um edital que já está cadastrado no sistema

45. **atualizar_url_edital**: Atualizar a URL de um edital cadastrado
    Exemplos: "atualize o edital PE-001 com URL: https://...", "mude a URL do edital PE-001 para https://...", "corrija a URL do edital PE-001", "atualize URL do edital"
    Palavras-chave: atualizar url, atualize edital com url, mude url, corrija url, atualizar link edital
    IMPORTANTE: Use quando o usuário quer ATUALIZAR/CORRIGIR a URL de download de um edital já cadastrado

46. **buscar_links_editais**: Retornar links de editais em uma área/categoria específica
    Exemplos: "retorne os links para os editais na área de hematologia", "links de editais de equipamentos médicos", "mostre links de editais de TI", "links para editais de laboratório"
    Palavras-chave: links de editais, links para editais, retorne os links, mostre links editais
    IMPORTANTE: Use quando o usuário quer VER LINKS clicáveis de editais, não calcular score

## CONTEXTO IMPORTANTE:
- **tem_arquivo**: {tem_arquivo} (true se usuário enviou um arquivo junto com a mensagem)
- Se tem_arquivo=true E mensagem vazia ou genérica → **arquivo_cadastrar**
- Se tem_arquivo=true E pede para mostrar/ler → **arquivo_mostrar**

## PARÂMETROS EXTRAS (extraia se mencionados):
- "termo_busca": termo de busca OTIMIZADO para APIs de licitação
- "nome_produto": nome do produto
- "url": URL completa se houver
- "produto": nome do produto para aderência/proposta
- "edital": número/identificador do edital
- "nome_fonte": nome da fonte de editais (ex: "ComprasNet", "BEC-SP")
- "tipo_fonte": tipo da fonte ("api" ou "scraper")

## IMPORTANTE - TERMO DE BUSCA:
Se a intenção for **buscar_editais** ou **buscar_editais_simples**, extraia o termo_busca da mensagem.
REGRA CRÍTICA: Use o termo EXATO que o usuário digitou. NÃO substitua, NÃO "otimize", NÃO troque por sinônimos.
REMOVA do termo_busca: palavras de comando (busque, editais, no, pncp), modificadores de busca (incluindo encerrados, sem calcular score, últimos X dias, sem score, apenas listando, com score rápido, com score profundo, com score híbrido, análise profunda, nos X melhores), nomes de fontes (PNCP, ComprasNet, BEC-SP, Licitacoes-e).
MANTENHA APENAS as palavras-chave do ASSUNTO/PRODUTO.
Exemplos:
- "busque editais de equipamento medico" → termo_busca: "equipamento medico"
- "busque editais de hematologia" → termo_busca: "hematologia"
- "busque editais de material hospitalar" → termo_busca: "material hospitalar"
- "busque editais de informática" → termo_busca: "informática"
- "busque editais de medicamentos incluindo encerrados sem calcular score" → termo_busca: "medicamentos"
- "busque editais de equipamento medico no PNCP últimos 60 dias" → termo_busca: "equipamento medico"
- "busque editais de medicamentos incluindo encerrados últimos 30 dias sem calcular score" → termo_busca: "medicamentos"
- "busque editais de reagentes com score profundo nos 5 melhores" → termo_busca: "reagentes"
- "busque editais de hematologia com score híbrido nos 10 melhores" → termo_busca: "hematologia"
- "busque editais de computadores com análise profunda" → termo_busca: "computadores"

## MENSAGEM DO USUÁRIO:
"{mensagem}"

## RESPOSTA:
Retorne APENAS um JSON:
{{"intencao": "<categoria>", "termo_busca": null, "nome_produto": null, "url": null, "produto": null, "edital": null}}"""


def detectar_intencao_ia(message: str, tem_arquivo: bool = False) -> dict:
    """
    Usa DeepSeek-chat para classificar a intenção do usuário.
    Retorna dict com 'intencao' e parâmetros extraídos.

    Args:
        message: Mensagem do usuário
        tem_arquivo: True se o usuário enviou um arquivo junto
    """
    import json
    import re

    prompt = PROMPT_CLASSIFICAR_INTENCAO.format(
        mensagem=message or "(mensagem vazia)",
        tem_arquivo="true" if tem_arquivo else "false"
    )

    try:
        resposta = call_deepseek(
            [{"role": "user", "content": prompt}],
            max_tokens=150,
            model_override="deepseek-chat"  # Rápido para classificação
        )

        # Extrair JSON da resposta
        json_match = re.search(r'\{[\s\S]*?\}', resposta)
        if json_match:
            resultado = json.loads(json_match.group())
            print(f"[AGENTE] Intenção detectada: {resultado.get('intencao')} | Termo: {resultado.get('termo_busca')}")
            return resultado
    except Exception as e:
        print(f"[AGENTE] Erro na classificação: {e}")

    # Fallback para detecção por palavras-chave
    return {"intencao": detectar_intencao_fallback(message), "termo_busca": None}


def detectar_intencao_fallback(message: str) -> str:
    """Fallback: detecção por palavras-chave (usado se IA falhar)."""
    msg = message.lower()

    # 1. Buscar na WEB (manuais, datasheets) - ANTES de buscar editais!
    if any(p in msg for p in ["busque na web", "buscar na web", "pesquise na web", "datasheet", "manual do"]):
        return "buscar_web"
    if any(p in msg for p in ["especificações do", "especificacoes do"]) and "edital" not in msg:
        return "buscar_web"

    # 2. Upload de manual
    if any(p in msg for p in ["upload", "enviei", "arquivo que", "processe o manual", "processe o pdf"]):
        return "upload_manual"

    # 2.1 Extrair ata de sessão (ANTES de outras ações com arquivo)
    if any(p in msg for p in ["extraia os resultados", "extrair resultados", "resultados da ata",
                               "ata de sessão", "ata de sessao", "vencedor do pregão", "vencedor do pregao",
                               "quem ganhou o pregão", "quem ganhou o pregao", "extraia da ata",
                               "registre os resultados da ata", "resultado da licitação", "resultado da licitacao"]):
        return "extrair_ata"

    # 2.5. Download de URL - ANTES de outras ações
    if "http://" in msg or "https://" in msg:
        if any(p in msg for p in ["baixe", "baixar", "download", "faça download"]):
            return "download_url"
        # Se tem URL e fala de PDF/manual/arquivo, também é download
        if any(p in msg for p in [".pdf", "manual", "arquivo", "documento"]):
            return "download_url"

    # 3. Salvar editais (da busca)
    # Detecta: "salvar edital", "salvar editais", "salvar todos", "salvar recomendados"
    # Também detecta "salvar edital NUMERO" (quando tem número de edital)
    if any(p in msg for p in ["salvar edital", "salvar editais", "salvar todos", "salvar recomendados",
                               "guardar edital", "guardar editais"]):
        return "salvar_editais"
    # "salve" sozinho ou com número de edital
    if "salve" in msg and ("edital" in msg or "editais" in msg or re.search(r'\d{2,}[/]\d{4}', msg)):
        return "salvar_editais"

    # 4. Listar produtos
    if any(p in msg for p in ["meus produtos", "listar produtos", "produtos cadastrados", "ver produtos"]):
        return "listar_produtos"

    # 5. Listar editais salvos
    if any(p in msg for p in ["meus editais", "editais salvos", "editais cadastrados", "ver editais"]):
        return "listar_editais"

    # 5.1 Listar propostas
    if any(p in msg for p in ["minhas propostas", "listar propostas", "propostas geradas", "ver propostas", "propostas cadastradas"]):
        return "listar_propostas"

    # 5.2 Consultar resultado de certame (perguntas sobre resultado)
    # IMPORTANTE: Deve vir ANTES de buscar_editais para ter prioridade
    if any(p in msg for p in ["qual o resultado", "qual foi o resultado", "resultado do edital",
                               "resultado dos editais", "resultados dos editais", "resultado existente",
                               "resultados existentes", "busque o resultado", "buscar resultado",
                               "mostre os resultados", "ver resultados", "listar resultados",
                               "quem ganhou", "quem venceu", "como foi o edital",
                               "resultado do certame", "resultados dos certames"]):
        return "consultar_resultado"

    # 5.3 Registrar resultado de certame (afirmações de vitória/derrota)
    if any(p in msg for p in ["perdemos", "ganhamos", "vencedor foi", "vencedora foi",
                               "derrota", "vitória", "vitoria", "segundo lugar", "terceiro lugar",
                               "registre resultado", "registrar resultado", "perdemos o", "ganhamos o",
                               "foi cancelado", "ficou deserto", "foi revogado", "edital cancelado",
                               "edital deserto", "edital revogado"]):
        return "registrar_resultado"

    # 5.4 Buscar atas no PNCP
    if any(p in msg for p in ["buscar atas", "busque atas", "encontrar atas", "encontre atas",
                               "baixar atas", "baixe atas", "atas de registro", "atas de sessão",
                               "atas de sessao", "atas pncp", "atas do pncp"]):
        return "buscar_atas_pncp"

    # 5.4.1 Buscar preços no PNCP (Funcionalidade 4 Sprint 1)
    if any(p in msg for p in ["buscar preços", "busque preços", "buscar precos", "busque precos",
                               "preço de mercado", "preco de mercado", "preços pncp", "precos pncp",
                               "quanto custa", "preço médio", "preco medio", "valores de contrato",
                               "preços de contrato", "precos de contrato", "preço praticado",
                               "preco praticado", "preços praticados", "precos praticados"]):
        return "buscar_precos_pncp"

    # 5.4.2 Histórico de preços (Funcionalidade 5 Sprint 1)
    if any(p in msg for p in ["histórico de preços", "historico de precos", "preços registrados",
                               "precos registrados", "preços salvos", "precos salvos",
                               "histórico preço", "historico preco"]):
        return "historico_precos"

    # 5.4.3 Listar concorrentes (Funcionalidade 6 Sprint 1)
    if any(p in msg for p in ["listar concorrentes", "liste concorrentes", "concorrentes conhecidos",
                               "nossos concorrentes", "quais concorrentes", "ver concorrentes"]):
        return "listar_concorrentes"

    # 5.4.4 Analisar concorrente (Funcionalidade 6 Sprint 1)
    if any(p in msg for p in ["analisar concorrente", "analise concorrente", "análise concorrente",
                               "analise o concorrente", "histórico concorrente", "historico concorrente"]):
        return "analisar_concorrente"

    # 5.4.5 Recomendar preço (Funcionalidade 7 Sprint 1)
    if any(p in msg for p in ["recomendar preço", "recomendar preco", "sugerir preço", "sugerir preco",
                               "que preço colocar", "que preco colocar", "qual preço sugerir",
                               "qual preco sugerir", "recomende preço", "recomende preco"]):
        return "recomendar_preco"

    # 5.4.6 Classificar edital (Funcionalidade 8 Sprint 1)
    if any(p in msg for p in ["classificar edital", "classifique edital", "tipo de edital",
                               "que tipo de edital", "é comodato", "e comodato", "é venda",
                               "é aluguel", "qual modalidade"]):
        return "classificar_edital"

    # 5.4.7 Verificar completude produto (Funcionalidade 9 Sprint 1)
    if any(p in msg for p in ["verificar completude", "produto completo", "falta informação",
                               "falta informacao", "completude produto", "está completo",
                               "esta completo", "informações faltando"]):
        return "verificar_completude"

    # 5.4.8 Resumir edital
    if any(p in msg for p in ["resumir edital", "resuma o edital", "resumo do edital", "resuma edital",
                               "sintetize o edital", "sintetize edital", "resumo edital"]):
        return "resumir_edital"

    # 5.4.9 Perguntar ao edital
    if any(p in msg for p in ["perguntar ao edital", "pergunte ao edital", "dúvida sobre edital",
                               "duvida sobre edital", "o edital exige", "o edital pede",
                               "prazo do edital", "requisitos do edital", "no edital pe-",
                               "do edital pe-", "edital pe-"]) and "?" in msg:
        return "perguntar_edital"

    # 5.4.10 Baixar PDF do edital
    if any(p in msg for p in ["baixar pdf edital", "baixe o pdf do edital", "download do edital",
                               "baixar edital", "baixe o edital", "download pdf edital",
                               "faça download do edital", "baixe edital"]):
        return "baixar_pdf_edital"

    # 5.4.11 Atualizar URL do edital
    if any(p in msg for p in ["atualize o edital", "atualizar url", "atualize url", "mude a url",
                               "corrija a url", "corrija url", "atualizar link", "atualize link"]):
        if "url" in msg or "http" in msg:
            return "atualizar_url_edital"

    # 5.5 Reprocessar produto
    if any(p in msg for p in ["reprocess", "atualize specs", "atualizar specs", "extraia novamente"]):
        return "reprocessar_produto"

    # 5.6 Excluir edital
    if any(p in msg for p in ["excluir edital", "excluir editais", "deletar edital", "remover edital", "apagar edital"]):
        return "excluir_edital"

    # 5.7 Excluir produto
    if any(p in msg for p in ["excluir produto", "deletar produto", "remover produto", "apagar produto"]):
        return "excluir_produto"

    # 5.8 Atualizar/Editar edital
    if any(p in msg for p in ["editar edital", "atualizar edital", "modificar edital", "alterar edital"]):
        return "atualizar_edital"

    # 5.9 Atualizar/Editar produto
    if any(p in msg for p in ["editar produto", "atualizar produto", "modificar produto", "alterar produto"]):
        return "atualizar_produto"

    # 6. Calcular aderência
    if any(p in msg for p in ["aderência", "aderencia", "score", "compatível", "compatibilidade"]):
        return "calcular_aderencia"

    # 7. Gerar proposta
    if any(p in msg for p in ["proposta", "gerar proposta", "elaborar proposta"]):
        return "gerar_proposta"

    # 8. Fontes
    if any(p in msg for p in ["fonte"]):
        if any(p in msg for p in ["cadastr", "adicion", "nova fonte"]):
            return "cadastrar_fonte"
        return "listar_fontes"

    # 8.5 Cadastrar produto manualmente (via formulário do Portfolio)
    if any(p in msg for p in ["cadastre manualmente o produto", "cadastre o produto", "cadastrar produto",
                               "registre o produto", "criar produto", "crie o produto", "adicione o produto"]):
        return "cadastrar_produto"

    # 8.6 Buscar registros ANVISA
    if any(p in msg for p in ["anvisa", "registro anvisa", "busque anvisa", "consulte anvisa",
                               "buscar anvisa", "consultar anvisa", "registros anvisa"]):
        return "buscar_anvisa"

    # 9. Cadastrar edital manualmente - ANTES de buscar editais
    if any(p in msg for p in ["cadastre o edital", "cadastrar edital", "registre o edital", "adicione o edital", "inserir edital"]):
        return "cadastrar_edital"

    # 10. Detectar se é busca no BANCO ou na WEB
    # Palavras que indicam BANCO LOCAL: "no banco", "cadastrado", "salvo", "no sistema", "banco de dados"
    busca_local = any(p in msg for p in ["no banco", "cadastrado", "salvo", "no sistema", "banco de dados",
                                          "tenho o edital", "tenho edital", "já tenho", "ja tenho"])
    # Palavras que indicam WEB: "na web", "no pncp", "internet", "online", "portal"
    busca_web = any(p in msg for p in ["na web", "no pncp", "pncp", "internet", "online", "portal", "bec"])

    import re

    # 10.1 Buscar edital específico por número
    tem_numero_edital = re.search(r'(PE[-]?\d+|[Pp]reg[aã]o\s*n?[ºo°]?\s*\d+|\d{4,}[/]\d{4}|n[ºo°]\s*\d+)', msg, re.IGNORECASE)
    if any(p in msg for p in ["busque o edital", "encontre o edital", "buscar edital"]) or tem_numero_edital:
        # Sempre usa buscar_edital_numero - a função internamente decide banco/web
        return "buscar_edital_numero"

    # 10.2 Buscar editais por termo
    if any(p in msg for p in ["busque editais", "buscar editais", "encontre editais", "encontrar editais"]):
        if busca_local:
            return "listar_editais"  # Lista do banco
        # Verificar se quer sem score
        sem_score = any(p in msg for p in ["sem score", "sem calcular", "sem aderência", "sem aderencia",
                                            "apenas listar", "só listar", "so listar", "apenas liste",
                                            "só liste", "so liste", "sem análise", "sem analise"])
        if sem_score:
            return "buscar_editais_simples"  # Busca sem calcular score
        else:
            return "buscar_editais"  # Busca na web com score (padrão)

    # 10.3 Buscar produtos
    if any(p in msg for p in ["busque produto", "buscar produto", "encontre produto", "encontrar produto"]):
        if busca_web:
            return "buscar_web"  # Busca manual na web
        else:
            return "listar_produtos"  # Lista do banco (padrão)

    # 11. Consultas analíticas via MindsDB - ANTES de buscar_editais genérico!
    # Inclui consultas com filtros de status, agregações, estatísticas
    palavras_mindsdb = [
        "estatística", "estatistica", "score médio", "score medio", "média de", "media de",
        "quantos editais", "quantos produtos", "análise dos dados", "analise dos dados",
        "relatório", "relatorio", "ranking", "tendência", "tendencia", "comparar",
        "por estado", "por uf", "por categoria", "desempenho", "performance",
        # Consultas de status/resultado
        "status perdido", "status ganho", "status novo", "status cancelado",
        "resultado perdido", "resultado ganho", "editais perdidos", "editais ganhos",
        "editais com status", "editais que estão", "editais que estao",
        "quais editais têm", "quais editais tem", "liste editais com",
        # Agregações
        "total de", "soma de", "contagem de", "quantidade de"
    ]
    if any(p in msg for p in palavras_mindsdb):
        return "consulta_mindsdb"

    # 12. FALLBACK INTELIGENTE: Se parece ser consulta sobre dados do banco → MindsDB
    # Palavras que indicam que é uma pergunta sobre dados armazenados
    palavras_dados_banco = [
        # Entidades do banco
        "edital", "editais", "produto", "produtos", "proposta", "propostas",
        "análise", "analise", "análises", "analises", "ata", "atas",
        "resultado", "resultados", "fonte", "fontes", "concorrente", "concorrentes",
        # Verbos de consulta
        "liste", "listar", "mostre", "mostrar", "exiba", "exibir",
        "quais", "qual", "quantos", "quantas", "onde", "quando",
        # Filtros
        "com valor", "acima de", "abaixo de", "maior que", "menor que",
        "entre", "desde", "até", "depois de", "antes de",
        "do mês", "da semana", "do ano", "de hoje", "de ontem",
        "em são paulo", "em sp", "em minas", "em mg", "no rio",
        # Ordenação
        "ordenado", "ordenados", "mais recente", "mais antigo", "últimos", "ultimos"
    ]

    # Se contém palavras de dados E parece ser uma pergunta/consulta
    eh_pergunta = any(p in msg for p in ["?", "quais", "qual", "quantos", "quantas",
                                          "liste", "mostre", "exiba", "me diga", "me fale"])
    menciona_dados = any(p in msg for p in palavras_dados_banco)

    if menciona_dados and eh_pergunta:
        return "consulta_mindsdb"

    # =============================================================================
    # SPRINT 2: ALERTAS E AUTOMAÇÃO
    # =============================================================================

    # 13.1 Configurar alertas de prazo
    if any(p in msg for p in ["configurar alerta", "configure alerta", "criar alerta", "crie alerta",
                               "avise-me", "lembre-me antes", "alerta para o edital", "alertar sobre",
                               "quero ser avisado", "me avise quando", "notifique-me"]):
        return "configurar_alertas"

    # 13.2 Listar alertas
    if any(p in msg for p in ["meus alertas", "listar alertas", "alertas configurados", "ver alertas",
                               "quais alertas", "alertas ativos", "próximos pregões", "proximos pregoes"]):
        return "listar_alertas"

    # 13.3 Dashboard de prazos
    if any(p in msg for p in ["dashboard de prazo", "dashboard prazos", "contagem regressiva",
                               "prazos dos editais", "próximos prazos", "proximos prazos",
                               "ver prazos", "mostre os prazos", "quais prazos"]):
        return "dashboard_prazos"

    # 13.4 Calendário de editais
    if any(p in msg for p in ["calendário", "calendario", "calendário de editais", "calendario de editais",
                               "editais do mês", "editais do mes", "editais da semana", "agenda de editais",
                               "datas importantes", "próximas datas", "proximas datas"]):
        return "calendario_editais"

    # 13.5 Configurar monitoramento
    if any(p in msg for p in ["configurar monitoramento", "configure monitoramento", "criar monitoramento",
                               "monitorar editais", "monitore editais", "quero monitorar",
                               "acompanhar editais", "busca automática", "busca automatica"]):
        return "configurar_monitoramento"

    # 13.6 Listar monitoramentos
    if any(p in msg for p in ["meus monitoramentos", "listar monitoramentos", "monitoramentos ativos",
                               "ver monitoramentos", "quais monitoramentos", "monitoramentos configurados"]):
        return "listar_monitoramentos"

    # 13.7 Desativar monitoramento
    if any(p in msg for p in ["desativar monitoramento", "parar monitoramento", "cancelar monitoramento",
                               "desative o monitoramento", "pare de monitorar", "remover monitoramento"]):
        return "desativar_monitoramento"

    # 13.8 Configurar notificações
    if any(p in msg for p in ["configurar notificações", "configurar notificacoes", "preferências de notificação",
                               "preferencias de notificacao", "email de notificação", "configurar email",
                               "configurar preferências", "configurar preferencias"]):
        return "configurar_notificacoes"

    # 13.9 Histórico de notificações
    if any(p in msg for p in ["histórico de notificações", "historico de notificacoes", "notificações recebidas",
                               "notificacoes recebidas", "ver notificações", "ver notificacoes",
                               "minhas notificações", "minhas notificacoes", "notificações não lidas",
                               "notificacoes nao lidas"]):
        return "historico_notificacoes"

    # 13.10 Extrair datas de edital
    if any(p in msg for p in ["extrair datas", "extraia as datas", "datas do edital", "prazos do edital",
                               "quais são as datas", "quais sao as datas", "identifique as datas",
                               "encontre as datas"]):
        return "extrair_datas_edital"

    # 13.11 Cancelar alerta
    if any(p in msg for p in ["cancelar alerta", "cancele o alerta", "remover alerta", "remova o alerta",
                               "excluir alerta", "desativar alerta", "não me avise mais",
                               "nao me avise mais"]):
        return "cancelar_alerta"

    return "chat_livre"


def detectar_intencao(message: str) -> str:
    """Wrapper para compatibilidade - retorna apenas a intenção."""
    resultado = detectar_intencao_ia(message)
    return resultado.get("intencao", "chat_livre")


# =============================================================================
# Auth Helpers
# =============================================================================

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')


def verify_password(password: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))
    except Exception:
        return False


def create_access_token(user_id: str, email: str, empresa_id: str = None) -> str:
    payload = {
        "user_id": user_id,
        "email": email,
        "empresa_id": empresa_id,
        "exp": datetime.utcnow() + timedelta(hours=JWT_EXPIRY_HOURS)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm="HS256")


def require_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        auth_header = request.headers.get('Authorization')

        if auth_header and auth_header.startswith('Bearer '):
            token = auth_header.split(' ')[1]

        if not token:
            return jsonify({"error": "Token não fornecido"}), 401

        try:
            payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
            request.user_id = payload["user_id"]
            request.user_email = payload["email"]
            request.empresa_id = payload.get("empresa_id")
            # Fallback: se JWT sem empresa_id, buscar primeira empresa do user
            if not request.empresa_id:
                db = get_db()
                try:
                    empresa = db.query(Empresa).filter(Empresa.user_id == request.user_id).first()
                    request.empresa_id = empresa.id if empresa else None
                finally:
                    db.close()
        except jwt.ExpiredSignatureError:
            return jsonify({"error": "Token expirado"}), 401
        except jwt.InvalidTokenError:
            return jsonify({"error": "Token inválido"}), 401

        return f(*args, **kwargs)
    return decorated


def get_current_user_id():
    return getattr(request, 'user_id', None)


def get_current_empresa_id():
    return getattr(request, 'empresa_id', None)


# =============================================================================
# Auth Routes
# =============================================================================

@app.route("/api/auth/login", methods=["POST"])
def login():
    data = request.json or {}
    email = data.get("email", "").strip().lower()
    password = data.get("password", "")

    if not email or not password:
        return jsonify({"error": "Email e senha são obrigatórios"}), 400

    db = get_db()
    try:
        user = db.query(User).filter(User.email == email).first()
        if not user or not user.password_hash:
            return jsonify({"error": "Email ou senha inválidos"}), 401

        if not verify_password(password, user.password_hash):
            return jsonify({"error": "Email ou senha inválidos"}), 401

        user.last_login_at = datetime.now()

        # Buscar empresa do user
        empresa = db.query(Empresa).filter(Empresa.user_id == user.id).first()
        empresa_id = empresa.id if empresa else None

        # Create refresh token
        refresh_token_value = str(uuid.uuid4())
        refresh_token = RefreshToken(
            token=refresh_token_value,
            user_id=user.id,
            expires_at=datetime.now() + timedelta(days=30)
        )
        db.add(refresh_token)
        db.commit()

        access_token = create_access_token(user.id, user.email, empresa_id=empresa_id)

        resp = {
            "access_token": access_token,
            "refresh_token": refresh_token_value,
            "user": user.to_dict(),
            "has_empresa": empresa is not None,
        }
        if empresa:
            resp["empresa"] = empresa.to_dict()

        return jsonify(resp)
    finally:
        db.close()


@app.route("/api/auth/register", methods=["POST"])
def register():
    data = request.json or {}
    name = data.get("name", "").strip()
    email = data.get("email", "").strip().lower()
    password = data.get("password", "")

    if not name or not email or not password:
        return jsonify({"error": "Nome, email e senha são obrigatórios"}), 400

    if len(password) < 6:
        return jsonify({"error": "A senha deve ter pelo menos 6 caracteres"}), 400

    db = get_db()
    try:
        existing = db.query(User).filter(User.email == email).first()
        if existing:
            return jsonify({"error": "Este email já está cadastrado"}), 409

        user = User(
            email=email,
            name=name,
            password_hash=hash_password(password)
        )
        db.add(user)
        db.commit()

        # Create refresh token
        refresh_token_value = str(uuid.uuid4())
        refresh_token = RefreshToken(
            token=refresh_token_value,
            user_id=user.id,
            expires_at=datetime.now() + timedelta(days=30)
        )
        db.add(refresh_token)
        db.commit()

        access_token = create_access_token(user.id, user.email)

        return jsonify({
            "access_token": access_token,
            "refresh_token": refresh_token_value,
            "user": user.to_dict(),
            "has_empresa": False,
        }), 201
    finally:
        db.close()


@app.route("/api/auth/user", methods=["GET"])
@require_auth
def get_current_user():
    db = get_db()
    try:
        user = db.query(User).filter(User.id == get_current_user_id()).first()
        if not user:
            return jsonify({"error": "Usuário não encontrado"}), 404
        empresa = db.query(Empresa).filter(Empresa.user_id == user.id).first()
        resp = user.to_dict()
        resp["has_empresa"] = empresa is not None
        if empresa:
            resp["empresa"] = empresa.to_dict()
        return jsonify(resp)
    finally:
        db.close()


@app.route("/api/auth/refresh", methods=["POST"])
def refresh():
    data = request.json or {}
    refresh_token_value = data.get("refresh_token", "")

    if not refresh_token_value:
        return jsonify({"error": "Refresh token não fornecido"}), 400

    db = get_db()
    try:
        token_record = db.query(RefreshToken).filter(
            RefreshToken.token == refresh_token_value,
            RefreshToken.revoked == False,
            RefreshToken.expires_at > datetime.now()
        ).first()

        if not token_record:
            return jsonify({"error": "Refresh token inválido ou expirado"}), 401

        user = db.query(User).filter(User.id == token_record.user_id).first()
        if not user:
            return jsonify({"error": "Usuário não encontrado"}), 404

        empresa = db.query(Empresa).filter(Empresa.user_id == user.id).first()
        empresa_id = empresa.id if empresa else None
        access_token = create_access_token(user.id, user.email, empresa_id=empresa_id)

        resp = {
            "access_token": access_token,
            "user": user.to_dict(),
            "has_empresa": empresa is not None,
        }
        if empresa:
            resp["empresa"] = empresa.to_dict()
        return jsonify(resp)
    finally:
        db.close()


@app.route("/api/auth/logout", methods=["POST"])
@require_auth
def logout():
    data = request.json or {}
    refresh_token_value = data.get("refresh_token", "")

    if refresh_token_value:
        db = get_db()
        try:
            token_record = db.query(RefreshToken).filter(
                RefreshToken.token == refresh_token_value
            ).first()
            if token_record:
                token_record.revoked = True
                db.commit()
        finally:
            db.close()

    return jsonify({"message": "Logout realizado com sucesso"})


@app.route("/api/auth/switch-empresa", methods=["POST"])
@require_auth
def switch_empresa():
    """Troca a empresa ativa e gera novo token JWT."""
    data = request.json or {}
    empresa_id = data.get("empresa_id", "").strip()
    if not empresa_id:
        return jsonify({"error": "empresa_id é obrigatório"}), 400

    db = get_db()
    try:
        empresa = db.query(Empresa).filter(
            Empresa.id == empresa_id,
            Empresa.user_id == request.user_id
        ).first()
        if not empresa:
            return jsonify({"error": "Empresa não encontrada ou sem permissão"}), 404

        access_token = create_access_token(request.user_id, request.user_email, empresa_id=empresa.id)
        return jsonify({
            "access_token": access_token,
            "empresa": empresa.to_dict(),
        })
    finally:
        db.close()


# =============================================================================
# Ações Routes
# =============================================================================

@app.route("/api/acoes", methods=["GET"])
def listar_acoes():
    """Lista prompts prontos para o dropdown."""
    return jsonify({"prompts": PROMPTS_PRONTOS})


# =============================================================================
# Auto-rename session
# =============================================================================

def generate_session_title(first_question: str) -> str:
    """Generate a short title for the session based on the first question."""
    prompt = f"""Crie um título curto (3-5 palavras) que resuma esta pergunta sobre licitações/editais:
"{first_question}"
Responda apenas com o título, sem aspas ou pontuação final."""

    try:
        messages = [{"role": "user", "content": prompt}]
        # Usar deepseek-chat para tarefas simples (reasoner retorna vazio para prompts curtos)
        title = call_deepseek(messages, max_tokens=50, model_override="deepseek-chat")
        # Clean up the title
        title = title.strip().strip('"\'').strip()
        # Limit length
        if len(title) > 50:
            title = title[:47] + "..."
        return title if title else None
    except Exception as e:
        print(f"Erro ao gerar título: {e}")
        return None


def count_session_messages(session_id: str, db) -> int:
    """Count messages in a session."""
    return db.query(Message).filter(Message.session_id == session_id).count()


# =============================================================================
# Chat Routes (com suporte a ações)
# =============================================================================

@app.route("/api/chat", methods=["POST"])
@require_auth
def chat():
    """
    Endpoint principal do chat.
    Detecta automaticamente a intenção do usuário.
    """
    data = request.json or {}
    session_id = data.get("session_id")
    message = data.get("message", "").strip()
    user_id = get_current_user_id()
    empresa_id = get_current_empresa_id()

    if not session_id or not message:
        return jsonify({"error": "session_id e message são obrigatórios"}), 400

    db = get_db()
    try:
        # Verificar sessão
        session = db.query(Session).filter(
            Session.id == session_id,
            Session.user_id == user_id
        ).first()

        if not session:
            return jsonify({"error": "Sessão não encontrada"}), 404

        # Check if this is the first user message (for auto-rename)
        is_first_message = count_session_messages(session_id, db) == 0

        # Detectar intenção usando agente IA
        print(f"[CHAT] Detectando intenção para: {message[:50]}...")
        intencao_resultado = detectar_intencao_ia(message)
        action_type = intencao_resultado.get("intencao", "chat_livre")
        termo_busca_ia = intencao_resultado.get("termo_busca")
        print(f"[CHAT] Intenção: {action_type} | Termo: {termo_busca_ia}")

        # Salvar mensagem do usuário
        user_msg = Message(
            session_id=session_id,
            role='user',
            content=message,
            action_type=action_type
        )
        db.add(user_msg)

        # Processar de acordo com a ação detectada
        response_text = ""
        resultado = None

        if action_type == "buscar_web":
            response_text, resultado = processar_buscar_web(message, user_id, intencao_resultado)

        elif action_type == "upload_manual":
            response_text, resultado = processar_upload_manual(message, user_id, intencao_resultado)

        elif action_type == "download_url":
            response_text, resultado = processar_download_url(message, user_id, intencao_resultado)

        elif action_type == "cadastrar_fonte":
            response_text, resultado = processar_cadastrar_fonte(message, user_id, intencao_resultado)

        elif action_type == "buscar_editais":
            response_text, resultado = processar_buscar_editais(message, user_id, termo_ia=termo_busca_ia, empresa_id=empresa_id)

        elif action_type == "buscar_editais_simples":
            response_text, resultado = processar_buscar_editais(message, user_id, termo_ia=termo_busca_ia, calcular_score=False, empresa_id=empresa_id)

        elif action_type == "buscar_edital_numero":
            response_text, resultado = processar_buscar_edital_numero(message, user_id)

        elif action_type == "listar_editais":
            response_text, resultado = processar_listar_editais(message, user_id)

        elif action_type == "listar_produtos":
            response_text, resultado = processar_listar_produtos(message, user_id)

        elif action_type == "listar_fontes":
            response_text, resultado = processar_listar_fontes(message)

        elif action_type == "listar_propostas":
            response_text, resultado = processar_listar_propostas(message, user_id)

        elif action_type == "calcular_aderencia":
            response_text, resultado = processar_calcular_aderencia(message, user_id)

        elif action_type == "gerar_proposta":
            response_text, resultado = processar_gerar_proposta(message, user_id)

        elif action_type == "salvar_editais":
            response_text, resultado = processar_salvar_editais(message, user_id, session_id, db)

        elif action_type == "reprocessar_produto":
            response_text, resultado = processar_reprocessar_produto(message, user_id)

        elif action_type == "excluir_edital":
            response_text, resultado = processar_excluir_edital(message, user_id)

        elif action_type == "excluir_produto":
            response_text, resultado = processar_excluir_produto(message, user_id)

        elif action_type == "atualizar_edital":
            response_text, resultado = processar_atualizar_edital(message, user_id)

        elif action_type == "atualizar_produto":
            response_text, resultado = processar_atualizar_produto(message, user_id)

        elif action_type == "consulta_mindsdb":
            response_text, resultado = processar_consulta_mindsdb(message, user_id)

        elif action_type == "registrar_resultado":
            response_text, resultado = processar_registrar_resultado(message, user_id)

        elif action_type == "consultar_resultado":
            response_text, resultado = processar_consultar_resultado(message, user_id)

        elif action_type == "buscar_atas_pncp":
            response_text, resultado = processar_buscar_atas_pncp(message, user_id)

        elif action_type == "buscar_precos_pncp":
            response_text, resultado = processar_buscar_precos_pncp(message, user_id)

        elif action_type == "historico_precos":
            response_text, resultado = processar_historico_precos(message, user_id)

        elif action_type == "listar_concorrentes":
            response_text, resultado = processar_listar_concorrentes(user_id)

        elif action_type == "analisar_concorrente":
            response_text, resultado = processar_analisar_concorrente(message, user_id)

        elif action_type == "recomendar_preco":
            response_text, resultado = processar_recomendar_preco(message, user_id)

        elif action_type == "classificar_edital":
            response_text, resultado = processar_classificar_edital(message, user_id)

        elif action_type == "verificar_completude":
            response_text, resultado = processar_verificar_completude(message, user_id)

        elif action_type == "cadastrar_produto":
            response_text, resultado = processar_cadastrar_produto(message, user_id)

        elif action_type == "buscar_anvisa":
            response_text, resultado = processar_buscar_anvisa(message, user_id)

        elif action_type == "cadastrar_edital":
            response_text, resultado = processar_cadastrar_edital(message, user_id, intencao_resultado)

        # =============================================================================
        # SPRINT 2: ALERTAS E AUTOMAÇÃO
        # =============================================================================
        elif action_type == "configurar_alertas":
            response_text = processar_configurar_alertas(message, user_id)

        elif action_type == "listar_alertas":
            response_text = processar_listar_alertas(message, user_id)

        elif action_type == "dashboard_prazos":
            response_text = processar_dashboard_prazos(message, user_id)

        elif action_type == "calendario_editais":
            response_text = processar_calendario_editais(message, user_id)

        elif action_type == "configurar_monitoramento":
            response_text = processar_configurar_monitoramento(message, user_id)

        elif action_type == "listar_monitoramentos":
            response_text = processar_listar_monitoramentos(message, user_id)

        elif action_type == "desativar_monitoramento":
            response_text = processar_desativar_monitoramento(message, user_id)

        elif action_type == "configurar_notificacoes":
            response_text = processar_configurar_notificacoes(message, user_id)

        elif action_type == "historico_notificacoes":
            response_text = processar_historico_notificacoes(message, user_id)

        elif action_type == "extrair_datas_edital":
            response_text = processar_extrair_datas_edital(message, user_id)

        elif action_type == "cancelar_alerta":
            response_text = processar_cancelar_alerta(message, user_id)

        # =============================================================================
        # ANÁLISE DE EDITAIS (Resumir e Perguntar)
        # =============================================================================
        elif action_type == "resumir_edital":
            response_text, resultado = processar_resumir_edital(message, user_id, intencao_resultado)

        elif action_type == "perguntar_edital":
            response_text, resultado = processar_perguntar_edital(message, user_id, intencao_resultado)

        elif action_type == "baixar_pdf_edital":
            response_text, resultado = processar_baixar_pdf_edital(message, user_id, intencao_resultado)

        elif action_type == "atualizar_url_edital":
            response_text, resultado = processar_atualizar_url_edital(message, user_id, intencao_resultado)

        elif action_type == "buscar_links_editais":
            response_text, resultado = processar_buscar_links_editais(message, user_id)

        # =============================================================================
        # FASE 1: PRECIFICAÇÃO (UC-P01 a UC-P10)
        # =============================================================================
        elif action_type == "precif_organizar_lotes":
            response_text, resultado = processar_precif_organizar_lotes(message, user_id, empresa_id)

        elif action_type == "precif_selecao_portfolio":
            response_text, resultado = processar_precif_selecao_portfolio(message, user_id, empresa_id)

        elif action_type == "precif_calcular_volumetria":
            response_text, resultado = processar_precif_calcular_volumetria(message, user_id)

        elif action_type == "precif_configurar_custos":
            response_text, resultado = processar_precif_configurar_custos(message, user_id, empresa_id)

        elif action_type == "precif_preco_base":
            response_text, resultado = processar_precif_preco_base(message, user_id)

        elif action_type == "precif_valor_referencia":
            response_text, resultado = processar_precif_valor_referencia(message, user_id)

        elif action_type == "precif_estruturar_lances":
            response_text, resultado = processar_precif_estruturar_lances(message, user_id)

        elif action_type == "precif_estrategia":
            response_text, resultado = processar_precif_estrategia(message, user_id, empresa_id)

        elif action_type == "precif_historico_camada_f":
            response_text, resultado = processar_precif_historico_camada_f(message, user_id)

        elif action_type == "precif_comodato":
            response_text, resultado = processar_precif_comodato(message, user_id, empresa_id)

        else:  # chat_livre
            response_text = processar_chat_livre(message, user_id, session_id, db)

        # Salvar resposta do assistente
        # Se foi busca de editais, salvar os editais no sources_json para recuperar depois
        sources_data = None
        if action_type in ["buscar_editais", "buscar_editais_simples"] and resultado:
            # Salvar editais para uso posterior (salvar_editais) — banco é LONGTEXT, aguenta tudo
            sources_data = {
                "editais": resultado.get("editais", []),
                "editais_com_score": resultado.get("editais_com_score", []),
                "editais_recomendados": resultado.get("editais_recomendados", []),
                "editais_participar": resultado.get("editais_participar", []),
                "termo": resultado.get("termo")
            }

        assistant_msg = Message(
            session_id=session_id,
            role='assistant',
            content=response_text,
            action_type=action_type,
            sources_json=sources_data
        )
        db.add(assistant_msg)

        # Auto-rename session if first message
        new_session_name = None
        print(f"DEBUG: is_first_message={is_first_message}, session.name='{session.name}'")
        if is_first_message and session.name == "Nova conversa":
            try:
                print(f"DEBUG: Gerando título para: {message[:50]}...")
                new_session_name = generate_session_title(message)
                print(f"DEBUG: Título gerado: {new_session_name}")
                if new_session_name:
                    session.name = new_session_name
            except Exception as e:
                print(f"DEBUG: Erro ao gerar título: {e}")
                pass  # Don't fail the request if rename fails

        # Atualizar sessão
        session.updated_at = datetime.now()
        db.commit()

        response_data = {
            "response": response_text,
            "session_id": session_id,
            "action_type": action_type,
            "resultado": resultado
        }

        if new_session_name:
            response_data["session_name"] = new_session_name

        return jsonify(response_data)

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


# =============================================================================
# Processadores de Ações
# =============================================================================

def processar_buscar_web(message: str, user_id: str, intencao_resultado: dict):
    """
    Processa ação: Buscar material/manuais/datasheets na web (DuckDuckGo/Serper/SerpAPI).

    Diferente de buscar_editais - aqui buscamos MANUAIS e ESPECIFICAÇÕES de produtos,
    não licitações/editais.
    """
    # Extrair termo de busca da IA ou usar mensagem
    termo = intencao_resultado.get("termo_busca") or message

    resultado = tool_web_search(termo, user_id)

    if resultado.get("success"):
        response = f"""## 🔍 Busca de Material na Web

**Termo pesquisado:** {termo}
**Total de resultados:** {resultado.get('total_resultados', 0)}
**PDFs encontrados:** {resultado.get('pdfs_encontrados', 0)}

"""
        # Mostrar PDFs encontrados
        pdfs = resultado.get('resultados_pdf', [])
        if pdfs:
            response += "### 📄 PDFs Encontrados\n\n"
            for i, pdf in enumerate(pdfs, 1):
                response += f"**{i}. {pdf['titulo']}**\n"
                response += f"   {pdf['descricao'][:150]}...\n" if len(pdf.get('descricao', '')) > 150 else f"   {pdf.get('descricao', '')}\n"
                response += f"   🔗 [Baixar PDF]({pdf['link']})\n\n"

        # Mostrar outros resultados
        outros = resultado.get('outros_resultados', [])
        if outros:
            response += "### 🌐 Outros Resultados\n\n"
            for i, item in enumerate(outros, 1):
                response += f"**{i}. {item['titulo']}**\n"
                response += f"   🔗 {item['link']}\n\n"

        response += """---
### Próximos passos:
Para baixar um PDF e cadastrar como produto, envie:
`Baixe o arquivo da URL: <cole_a_url_do_pdf>`

O sistema irá:
1. Baixar o arquivo
2. Extrair texto e especificações
3. Cadastrar como produto no sistema"""

    else:
        response = f"❌ Erro na busca: {resultado.get('error', 'Erro desconhecido')}"

    return response, resultado


def processar_upload_manual(message: str, user_id: str, intencao_resultado: dict):
    """
    Processa ação: Upload de manual/PDF para cadastrar produto.

    Nota: O upload físico do arquivo é feito via /api/upload.
    Esta função processa a intenção quando o usuário menciona que quer processar um arquivo.
    """
    nome_produto = intencao_resultado.get("nome_produto")

    if nome_produto:
        response = f"""## 📄 Upload de Manual

Para cadastrar o produto **{nome_produto}**, faça o seguinte:

1. Clique no botão **📎** ao lado do campo de mensagem
2. Selecione o arquivo PDF do manual
3. Após o upload, envie uma mensagem confirmando: "Processe como {nome_produto}"

O sistema irá:
- Extrair o texto do PDF
- Identificar especificações técnicas
- Cadastrar o produto com todas as specs"""
    else:
        response = """## 📄 Upload de Manual

Para cadastrar um produto a partir de um manual PDF:

1. Clique no botão **📎** ao lado do campo de mensagem
2. Selecione o arquivo PDF do manual
3. Após o upload, informe o nome do produto

Exemplo: "Processe o manual que enviei e cadastre como Analisador Bioquímico BS-240"

O sistema extrairá automaticamente as especificações técnicas do manual."""

    return response, {"status": "aguardando_upload", "nome_produto": nome_produto}


def processar_download_url(message: str, user_id: str, intencao_resultado: dict):
    """
    Processa ação: Baixar arquivo de URL, extrair especificações e cadastrar produto.

    Fluxo completo:
    1. Baixa o arquivo da URL
    2. Extrai texto do PDF
    3. Usa IA para extrair especificações técnicas
    4. Cadastra produto no banco
    """
    import re

    intencao_resultado = intencao_resultado or {}

    # Extrair URL da mensagem ou do resultado da IA
    url = intencao_resultado.get("url")
    nome_produto = intencao_resultado.get("nome_produto")

    # Se IA não extraiu a URL, tentar extrair via regex
    if not url:
        url_match = re.search(r'https?://[^\s<>"\']+', message)
        if url_match:
            url = url_match.group()

    if not url:
        return "❌ Não encontrei uma URL na mensagem. Envie no formato:\n`Baixe o arquivo da URL: https://exemplo.com/manual.pdf`", {"error": "URL não encontrada"}

    # Se não tem nome do produto, tentar extrair do nome do arquivo ou pedir
    if not nome_produto:
        # Tentar extrair do nome do arquivo na URL
        filename = url.split('/')[-1].split('?')[0]
        if filename and len(filename) > 5:
            nome_produto = filename.replace('.pdf', '').replace('_', ' ').replace('-', ' ')[:50]

    response = f"## ⏳ Baixando arquivo...\n\n**URL:** {url}\n\n"

    # 1. Baixar o arquivo
    resultado_download = tool_download_arquivo(url, user_id, nome_produto)

    if not resultado_download.get("success"):
        return f"❌ Erro ao baixar arquivo: {resultado_download.get('error')}", resultado_download

    filepath = resultado_download.get("filepath")
    filesize = resultado_download.get("size", 0)
    response += f"✅ Arquivo baixado: {resultado_download.get('filename')} ({filesize/1024:.1f} KB)\n\n"

    # 2. Se não tem nome do produto, pedir ao usuário
    if not nome_produto or nome_produto == "documento":
        response += """## ⚠️ Nome do produto não identificado

Envie o nome do produto para cadastrar. Exemplo:
`Cadastre como Analisador Bioquímico BS-240 da Mindray`

Ou informe mais detalhes:
`Cadastre como [nome], fabricante [fabricante], categoria [categoria]`"""
        return response, {
            "success": True,
            "status": "aguardando_nome_produto",
            "filepath": filepath,
            "filesize": filesize
        }

    # 3. Processar o arquivo e cadastrar produto
    response += f"## ⏳ Processando PDF e extraindo especificações...\n\n"

    # Determinar categoria automaticamente
    categoria = "equipamento"  # Padrão
    nome_lower = nome_produto.lower()
    if any(t in nome_lower for t in ["analisador", "bioquímic", "laborat"]):
        categoria = "equipamento"
    elif any(t in nome_lower for t in ["centrifuga", "microscop"]):
        categoria = "equipamento"
    elif any(t in nome_lower for t in ["cama", "maca", "cadeira"]):
        categoria = "mobiliario"
    elif any(t in nome_lower for t in ["monitor", "desfibrilador", "eletrocard"]):
        categoria = "equipamento"

    resultado_processo = tool_processar_upload(
        filepath=filepath,
        user_id=user_id,
        nome_produto=nome_produto,
        categoria=categoria,
        fabricante="Mindray" if "mindray" in message.lower() else None,
        modelo=None
    )

    if resultado_processo.get("success"):
        produto = resultado_processo.get("produto", {})
        specs = resultado_processo.get("especificacoes", [])

        response += f"""## ✅ Produto Cadastrado com Sucesso!

**Nome:** {produto.get('nome', nome_produto)}
**Categoria:** {categoria}
**Fabricante:** {produto.get('fabricante', 'Não informado')}

### Especificações Extraídas ({len(specs)} encontradas):
"""
        for spec in specs[:10]:  # Mostrar até 10 specs
            response += f"- **{spec.get('nome', 'N/A')}:** {spec.get('valor', 'N/A')}\n"

        if len(specs) > 10:
            response += f"\n... e mais {len(specs) - 10} especificações.\n"

        response += f"\n---\n✅ Produto pronto para calcular aderência com editais!"
    else:
        response += f"❌ Erro ao processar: {resultado_processo.get('error')}"

    return response, resultado_processo


def processar_cadastrar_fonte(message: str, user_id: str, intencao_resultado: dict = None):
    """
    Processa ação: Cadastrar fonte de editais.
    Se tiver todos os dados, cadastra direto.
    Se faltar URL ou tipo, busca na web automaticamente.
    """
    import re

    intencao_resultado = intencao_resultado or {}

    # Verificar se a IA já extraiu os dados (aceitar vários nomes de campo)
    nome_fonte = intencao_resultado.get("nome_fonte") or intencao_resultado.get("nome")
    tipo_fonte = intencao_resultado.get("tipo_fonte") or intencao_resultado.get("tipo")
    url_fonte = intencao_resultado.get("url_fonte") or intencao_resultado.get("url")

    # Se não tem nome_fonte, tentar extrair da mensagem com regex
    if not nome_fonte:
        # Padrão: "fonte NOME" ou "fonte: NOME" ou "cadastre a fonte NOME"
        # Inclui caracteres acentuados (À-ú)
        match = re.search(r'fonte[:\s]+([A-Za-zÀ-ú0-9\-_\s]+?)(?:,|\s+tipo|\s+url|$)', message, re.IGNORECASE)
        if match:
            nome_fonte = match.group(1).strip()

    # Se não tem tipo_fonte, tentar extrair da mensagem
    if not tipo_fonte:
        if 'tipo api' in message.lower() or ', api,' in message.lower() or ' api ' in message.lower():
            tipo_fonte = 'api'
        elif 'tipo scraper' in message.lower() or ', scraper,' in message.lower() or ' scraper ' in message.lower():
            tipo_fonte = 'scraper'

    # Se não tem URL, tentar extrair da mensagem
    if not url_fonte:
        url_match = re.search(r'https?://[^\s,]+', message)
        if url_match:
            url_fonte = url_match.group(0).strip()

    print(f"[FONTE] Dados extraídos: nome={nome_fonte}, tipo={tipo_fonte}, url={url_fonte}")

    # Se tem nome mas falta URL, buscar na web
    if nome_fonte and not url_fonte:
        print(f"[FONTE] URL não informada, buscando na web...")

        # Buscar na web
        resultado_busca = tool_web_search(f"{nome_fonte} portal licitações governo site oficial", user_id, num_results=5)

        # Combinar todos os resultados (PDFs + outros)
        todos_resultados = resultado_busca.get("resultados_pdf", []) + resultado_busca.get("outros_resultados", [])

        if resultado_busca.get("success") and todos_resultados:
            # Usar IA para extrair a URL correta dos resultados
            resultados_texto = "\n".join([
                f"- {r.get('titulo')}: {r.get('link')}"
                for r in todos_resultados[:5]
            ])

            prompt_extrair = f"""Analise os resultados de busca abaixo e identifique a URL oficial do portal de licitações "{nome_fonte}".

Resultados:
{resultados_texto}

Retorne APENAS um JSON com:
- url: URL oficial do portal (a mais provável)
- tipo: "api" se for portal do governo federal ou tiver API conhecida, "scraper" caso contrário
- nome_completo: nome completo/oficial da fonte

JSON:"""

            try:
                resposta_ia = call_deepseek([{"role": "user", "content": prompt_extrair}], max_tokens=300, model_override="deepseek-chat")
                json_match = re.search(r'\{[\s\S]*?\}', resposta_ia)
                if json_match:
                    dados_web = json.loads(json_match.group())
                    url_fonte = dados_web.get("url")
                    if not tipo_fonte:
                        tipo_fonte = dados_web.get("tipo", "scraper")
                    nome_completo = dados_web.get("nome_completo")
                    if nome_completo:
                        nome_fonte = nome_completo
                    print(f"[FONTE] Dados da web: url={url_fonte}, tipo={tipo_fonte}, nome={nome_fonte}")
            except Exception as e:
                print(f"[FONTE] Erro ao extrair dados da web: {e}")
                # Fallback: usar primeiro resultado
                if todos_resultados:
                    primeiro = todos_resultados[0]
                    url_fonte = primeiro.get("link")
                    if not tipo_fonte:
                        tipo_fonte = "scraper"

    # Se ainda não tem tipo, usar padrão
    if not tipo_fonte:
        tipo_fonte = "scraper"

    print(f"[FONTE] Dados finais: nome={nome_fonte}, tipo={tipo_fonte}, url={url_fonte}")

    if nome_fonte and url_fonte:
        resultado = tool_cadastrar_fonte(
            nome=nome_fonte,
            tipo=tipo_fonte,
            url_base=url_fonte,
            descricao=f"Fonte cadastrada via chat: {nome_fonte}"
        )
        if resultado.get("success"):
            response = f"""✅ Fonte cadastrada com sucesso!

**Nome:** {nome_fonte}
**Tipo:** {tipo_fonte}
**URL:** {url_fonte}"""
        elif resultado.get("duplicada"):
            fonte_exist = resultado.get('fonte_existente', {})
            response = f"""⚠️ Fonte já existe!

**Nome:** {fonte_exist.get('nome')}
**URL:** {fonte_exist.get('url')}"""
        else:
            response = f"❌ Erro ao cadastrar: {resultado.get('error')}"
        return response, resultado

    # Se não conseguiu extrair nem da web, pedir mais informações
    response = f"""Não consegui encontrar informações sobre a fonte "{nome_fonte or 'informada'}".

Por favor, forneça os dados completos:
- **Nome**: Nome da fonte
- **Tipo**: api ou scraper
- **URL**: URL base da fonte

Exemplo: `cadastre a fonte BEC-SP, tipo scraper, url https://bec.sp.gov.br`"""
    return response, {"status": "aguardando_dados"}


def _buscar_editais_multifonte(termo: str, user_id: str, uf: str = None,
                                incluir_encerrados: bool = False,
                                buscar_detalhes: bool = True,
                                dias_busca: int = 90,
                                fonte: str = None,
                                modalidade: str = None,
                                tipo_produto: str = None,
                                origem: str = None) -> dict:
    """
    Busca editais em múltiplas fontes (PNCP API + scraper web) EM PARALELO,
    deduplica e retorna resultado consolidado.
    Reutilizada pelo chat (processar_buscar_editais) e pelo endpoint REST (/api/editais/buscar).

    Args:
        fonte: Se especificado, busca APENAS nessa fonte. Valores: 'PNCP', nome de fonte scraper, ou None (todas).
        modalidade: Filtro pós-busca por modalidade (ex: "Pregão Eletrônico").
        tipo_produto: Filtro pós-busca por tipo de produto inferido do objeto (ex: "Reagentes").
        origem: Filtro pós-busca por origem do órgão inferida (ex: "Hospital").
    """
    import re
    from datetime import datetime as _dt
    hoje = _dt.now()
    editais = []
    fontes_consultadas = []
    erros_fontes = []

    # Resolver fonte: pode vir como UUID, nome ou URL
    fonte_nome = None
    if fonte:
        fonte_str = fonte.strip()
        # Se parece UUID, buscar nome no banco
        if len(fonte_str) >= 32 and '-' in fonte_str:
            try:
                _db = get_db()
                fonte_obj = _db.query(FonteEdital).filter(FonteEdital.id == fonte_str).first()
                if fonte_obj:
                    fonte_nome = fonte_obj.nome
                    print(f"[BUSCA-MULTI] Fonte UUID '{fonte_str}' → '{fonte_nome}'")
                _db.close()
            except Exception:
                pass
        if not fonte_nome:
            fonte_nome = fonte_str

    # Normalizar fonte para comparação
    fonte_lower = fonte_nome.lower().strip() if fonte_nome else None
    _fontes_com_api = ('pncp', 'bec', 'bec-sp', 'bec sp', 'comprasnet')
    buscar_pncp = fonte_lower is None or 'pncp' in fonte_lower
    buscar_bec = fonte_lower is None or fonte_lower in ('bec', 'bec-sp', 'bec sp')
    buscar_comprasnet = fonte_lower is None or fonte_lower in ('comprasnet',)

    # Scraper: apenas quando busca geral ou fonte sem API própria
    if fonte_lower is None:
        buscar_scraper = True
    elif fonte_lower in _fontes_com_api:
        buscar_scraper = False
    else:
        buscar_scraper = True
        buscar_pncp = False

    # Se fonte é BEC ou ComprasNet, não buscar PNCP
    if fonte_lower and 'pncp' not in fonte_lower:
        buscar_pncp = False

    if fonte_lower:
        print(f"[BUSCA-MULTI] Buscando '{termo}' na fonte '{fonte_nome}'...")
    else:
        print(f"[BUSCA-MULTI] Buscando '{termo}' em todas as fontes (PNCP + BEC + ComprasNet + Scraper)...")

    # 1. Buscar no PNCP (API oficial) — com tratamento de erro
    if buscar_pncp:
        try:
            resultado_pncp = tool_buscar_editais_fonte("PNCP", termo, user_id, uf=uf,
                                                        buscar_detalhes=buscar_detalhes,
                                                        incluir_encerrados=incluir_encerrados,
                                                        dias_busca=dias_busca)
            if resultado_pncp.get("success"):
                editais_pncp = resultado_pncp.get("editais", [])
                for ed in editais_pncp:
                    ed['fonte'] = 'PNCP (API)'
                if not incluir_encerrados:
                    antes = len(editais_pncp)
                    editais_pncp = [e for e in editais_pncp if not e.get('encerrado')]
                    filtrados = antes - len(editais_pncp)
                    if filtrados:
                        print(f"[BUSCA-MULTI] PNCP: {filtrados} encerrados removidos")
                editais.extend(editais_pncp)
                fontes_consultadas.append("PNCP (API)")
                print(f"[BUSCA-MULTI] PNCP: {len(editais_pncp)} editais encontrados")
            else:
                erros_fontes.append(f"PNCP: {resultado_pncp.get('error', 'sem resultados')}")
        except Exception as e:
            print(f"[BUSCA-MULTI] Erro PNCP: {e}")
            erros_fontes.append(f"PNCP: {e}")

    # 2. Buscar na BEC-SP (API direta)
    if buscar_bec:
        try:
            from tools import tool_buscar_editais_bec
            resultado_bec = tool_buscar_editais_bec(
                termo, user_id,
                incluir_encerrados=incluir_encerrados,
                dias_busca=dias_busca
            )
            if resultado_bec.get("success"):
                editais_bec = resultado_bec.get("editais", [])
                if not incluir_encerrados:
                    antes = len(editais_bec)
                    editais_bec = [e for e in editais_bec if not e.get('encerrado')]
                    filtrados = antes - len(editais_bec)
                    if filtrados:
                        print(f"[BUSCA-MULTI] BEC-SP: {filtrados} encerrados removidos")
                editais.extend(editais_bec)
                fontes_consultadas.append("BEC-SP (API)")
                print(f"[BUSCA-MULTI] BEC-SP: {len(editais_bec)} editais encontrados")
            else:
                erros_fontes.append(f"BEC-SP: {resultado_bec.get('error', 'sem resultados')}")
        except Exception as e:
            print(f"[BUSCA-MULTI] Erro BEC-SP: {e}")
            erros_fontes.append(f"BEC-SP: {e}")

    # 3. Buscar no ComprasNet (API direta com fallback para scraper)
    _comprasnet_precisa_scraper = False
    if buscar_comprasnet:
        try:
            from tools import tool_buscar_editais_comprasnet
            resultado_cn = tool_buscar_editais_comprasnet(
                termo, user_id,
                incluir_encerrados=incluir_encerrados,
                dias_busca=dias_busca
            )
            if resultado_cn.get("success"):
                editais_cn = resultado_cn.get("editais", [])
                if editais_cn:
                    if not incluir_encerrados:
                        antes = len(editais_cn)
                        editais_cn = [e for e in editais_cn if not e.get('encerrado')]
                        filtrados = antes - len(editais_cn)
                        if filtrados:
                            print(f"[BUSCA-MULTI] ComprasNet: {filtrados} encerrados removidos")
                    editais.extend(editais_cn)
                    fontes_consultadas.append("ComprasNet (API)")
                    print(f"[BUSCA-MULTI] ComprasNet API: {len(editais_cn)} editais")
                elif not resultado_cn.get("api_disponivel", True):
                    _comprasnet_precisa_scraper = True
                    print(f"[BUSCA-MULTI] ComprasNet API indisponível, ativando scraper fallback")
                    erros_fontes.append(f"ComprasNet API: {resultado_cn.get('mensagem', 'indisponível')}")
        except Exception as e:
            print(f"[BUSCA-MULTI] Erro ComprasNet: {e}")
            _comprasnet_precisa_scraper = True
            erros_fontes.append(f"ComprasNet: {e}")

    # Mapear nomes amigáveis de fonte para domínios de URL
    _fonte_to_dominios = {
        'bec-sp': ['bec.sp.gov.br'],
        'bec sp': ['bec.sp.gov.br'],
        'bec': ['bec.sp.gov.br'],
        'comprasnet': ['comprasnet.gov.br'],
        'licitacoes-e': ['licitacoes-e.com.br'],
        'compras.rs': ['compras.rs.gov.br'],
        'compras.mg': ['compras.mg.gov.br'],
        'gov.br': ['gov.br'],
        'licitanet': ['licitanet.com.br'],
        'portal de compras': ['portaldecompraspublicas.com.br'],
    }

    # 4. Buscar em outras fontes via scraper web (Brave/Serper/SerpAPI) — com tratamento de erro
    resultado_scraper = None
    if buscar_scraper or _comprasnet_precisa_scraper:
        try:
            # Determinar domínios para o scraper
            fontes_scraper = None  # None = todas
            if _comprasnet_precisa_scraper and not buscar_scraper:
                # Scraper APENAS para ComprasNet (fallback da API indisponível)
                fontes_scraper = ['comprasnet.gov.br', 'gov.br/compras']
                print(f"[BUSCA-MULTI] Scraper fallback restrito a ComprasNet: {fontes_scraper}")
            elif fonte_lower and 'pncp' not in fonte_lower:
                # Se fonte específica (não PNCP/BEC/ComprasNet), buscar SÓ naquela fonte
                dominios = _fonte_to_dominios.get(fonte_lower, [])
                if not dominios:
                    # Tentar buscar domínio da fonte no banco
                    try:
                        _db = get_db()
                        _fobj = _db.query(FonteEdital).filter(
                            FonteEdital.nome.ilike(f"%{fonte_nome}%")
                        ).first()
                        if _fobj and _fobj.url_base:
                            dom = _fobj.url_base.replace("https://", "").replace("http://", "").split("/")[0]
                            dominios = [dom]
                        _db.close()
                    except Exception:
                        pass
                if dominios:
                    fontes_scraper = dominios
                    print(f"[BUSCA-MULTI] Scraper restrito a: {fontes_scraper}")
            resultado_scraper = tool_buscar_editais_scraper(termo, fontes=fontes_scraper, user_id=user_id)
        except Exception as e:
            print(f"[BUSCA-MULTI] Erro scraper: {e}")
            resultado_scraper = None
            erros_fontes.append(f"Scraper: {e}")

    if resultado_scraper and resultado_scraper.get("success"):
        editais_scraper = resultado_scraper.get("editais", [])
        links_pncp = {ed.get('url', '') for ed in editais}

        palavras_excluir_objeto = [
            'prestação de serviço', 'mão de obra', 'dedicação exclusiva',
            'terceirização', 'lavanderia', 'limpeza', 'manutenção preventiva',
            'manutenção corretiva', 'prorrogação da ata', 'prorrogação parcial',
            'termo aditivo', 'credenciadas no sistema', 'poderão participar'
        ]

        editais_novos = []
        editais_para_enriquecer = []  # (indice, numero_edital, fonte_nome)

        for ed in editais_scraper:
            if ed.get('link') not in links_pncp and ed.get('link'):
                # Filtrar resultados gov.br/pncp.gov.br/compras.gov.br do scraper
                # quando já temos busca PNCP API (são duplicatas em formato inferior)
                ed_link = (ed.get('link', '') or '').lower()
                ed_fonte = (ed.get('fonte', '') or '').lower()
                if buscar_pncp and any(d in ed_link or d in ed_fonte for d in ['pncp.gov.br', 'compras.gov.br', 'www.gov.br']):
                    continue

                # Filtrar por fonte específica se solicitado
                if fonte_lower and 'pncp' not in fonte_lower:
                    ed_fonte = (ed.get('fonte', '') or '').lower()
                    ed_link = (ed.get('link', '') or '').lower()
                    # Verificar se a fonte do edital corresponde à fonte pedida
                    # Tentar match direto e por domínio
                    dominios_alvo = _fonte_to_dominios.get(fonte_lower, [fonte_lower])
                    fonte_match = any(d in ed_fonte or d in ed_link for d in dominios_alvo)
                    if not fonte_match and fonte_lower not in ed_fonte and fonte_lower not in ed_link:
                        continue

                texto = (ed.get('descricao', '') + ' ' + ed.get('titulo', '')).lower()
                eh_servico_ou_prorrogacao = any(p in texto for p in palavras_excluir_objeto)
                if eh_servico_ou_prorrogacao:
                    print(f"[BUSCA-MULTI] Filtrando (serviço/prorrogação): {ed.get('numero', ed.get('titulo', '')[:30])}")
                    continue

                numero_edital = ed.get('numero', ed.get('titulo', '')[:50])

                # Filtrar encerrados pelo ano no número (ex: "04/2025" → ano 2025 < atual)
                if not incluir_encerrados and numero_edital:
                    ano_match = re.search(r'/(\d{4})', numero_edital)
                    if ano_match:
                        ano_edital = int(ano_match.group(1))
                        ano_atual = hoje.year
                        if ano_edital < ano_atual:
                            print(f"[BUSCA-MULTI] Filtrando (ano {ano_edital}): {numero_edital}")
                            continue

                ed_normalizado = {
                    'numero': numero_edital,
                    'orgao': ed.get('orgao', 'Não identificado'),
                    'objeto': ed.get('descricao', ed.get('titulo', '')),
                    'url': ed.get('link'),
                    'fonte': f"{ed.get('fonte', 'Web')} (Scraper)",
                    'fonte_tipo': 'scraper',
                    'modalidade': 'Identificar no portal',
                    'valor_referencia': None,
                    'data_abertura': 'Ver no portal'
                }

                idx = len(editais_novos)
                editais_novos.append(ed_normalizado)

                # Marcar para enriquecimento via PNCP
                # Se NÃO incluir encerrados: enriquecer TODOS de 2026+ (precisamos checar data real)
                # Se incluir encerrados: enriquecer até 5 (só para dados extras)
                if numero_edital and '/' in numero_edital:
                    if not incluir_encerrados:
                        # Precisa verificar data de TODOS os editais do ano atual ou futuro
                        editais_para_enriquecer.append((idx, numero_edital, ed.get('fonte', 'Web')))
                    elif len(editais_para_enriquecer) < 5:
                        editais_para_enriquecer.append((idx, numero_edital, ed.get('fonte', 'Web')))

        # Enriquecer editais do scraper em paralelo via PNCP
        if editais_para_enriquecer:
            from concurrent.futures import ThreadPoolExecutor, as_completed
            from tools import _buscar_edital_pncp_por_numero

            def _enriquecer(args):
                idx, numero, fonte = args
                try:
                    return idx, numero, fonte, _buscar_edital_pncp_por_numero(numero)
                except Exception as e:
                    print(f"[BUSCA-MULTI] Erro ao enriquecer {numero}: {e}")
                    return idx, numero, fonte, None

            n_enriquecer = len(editais_para_enriquecer)
            max_workers = min(n_enriquecer, 8)
            timeout_global = max(30, n_enriquecer * 5)  # 5s por edital, mínimo 30s
            print(f"[BUSCA-MULTI] Enriquecendo {n_enriquecer} editais em paralelo ({max_workers} workers, timeout {timeout_global}s)...")
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = {executor.submit(_enriquecer, args): args for args in editais_para_enriquecer}
                for future in as_completed(futures, timeout=timeout_global):
                    try:
                        idx, numero, fonte, dados_pncp = future.result(timeout=10)
                        if dados_pncp:
                            # Filtrar encerrado por data de encerramento/fim de vigência
                            data_enc = dados_pncp.get('data_encerramento') or dados_pncp.get('data_fim_vigencia') or ''
                            data_ab = dados_pncp.get('data_abertura', '')
                            if not incluir_encerrados and data_enc:
                                try:
                                    dt_enc = datetime.fromisoformat(str(data_enc).replace('Z', '')[:19])
                                    if dt_enc < hoje:
                                        editais_novos[idx]['_remover'] = True
                                        print(f"[BUSCA-MULTI] Edital {numero} encerrado (fim: {data_enc[:10]}), removendo")
                                        continue
                                except (ValueError, TypeError):
                                    pass
                            editais_novos[idx].update({
                                'cnpj_orgao': dados_pncp.get('cnpj_orgao'),
                                'ano_compra': dados_pncp.get('ano_compra'),
                                'seq_compra': dados_pncp.get('seq_compra'),
                                'numero_pncp': dados_pncp.get('numero_pncp'),
                                'valor_referencia': dados_pncp.get('valor_referencia'),
                                'data_abertura': data_ab,
                                'uf': dados_pncp.get('uf'),
                                'cidade': dados_pncp.get('cidade'),
                                'orgao': dados_pncp.get('orgao') or editais_novos[idx]['orgao'],
                                'objeto': dados_pncp.get('objeto') or editais_novos[idx]['objeto'],
                                'url': dados_pncp.get('url') or editais_novos[idx]['url'],
                                'fonte': f"{fonte} → PNCP",
                                'fonte_tipo': 'api',
                                'dados_completos': True,
                            })
                            print(f"[BUSCA-MULTI] Edital {numero} enriquecido com dados PNCP")
                    except Exception as e:
                        print(f"[BUSCA-MULTI] Timeout/erro no enriquecimento: {e}")
        # Se NÃO incluir encerrados, remover editais do scraper que não foram verificados
        # (data continua "Ver no portal" = não sabemos se está aberto)
        if not incluir_encerrados:
            indices_enriquecidos = {args[0] for args in editais_para_enriquecer}
            removidos_nao_verificados = 0
            for i, ed in enumerate(editais_novos):
                if ed.get('_remover'):
                    continue
                if i in indices_enriquecidos and not ed.get('dados_completos'):
                    # Foi enviado para enriquecimento mas PNCP não encontrou → não verificável
                    ed['_remover'] = True
                    removidos_nao_verificados += 1
                    print(f"[BUSCA-MULTI] Removendo (não verificado no PNCP): {ed.get('numero', '?')}")
            if removidos_nao_verificados:
                print(f"[BUSCA-MULTI] {removidos_nao_verificados} editais removidos por data não verificável")

        editais.extend([e for e in editais_novos if not e.get('_remover')])
        fontes_scraper = resultado_scraper.get('fontes_consultadas', [])
        fontes_consultadas.extend([f"{f} (Scraper)" for f in fontes_scraper if 'pncp' not in f.lower()])
        total_scraper_final = len([e for e in editais_novos if not e.get('_remover')])
        print(f"[BUSCA-MULTI] Scraper: {total_scraper_final} editais adicionais após filtragem")
        if resultado_scraper.get('erros'):
            for err in resultado_scraper.get('erros', []):
                erros_fontes.append(f"{err.get('fonte')}: {err.get('erro')}")

    # 3. Deduplicar por número de edital + fallback orgão+valor (priorizar PNCP)
    editais_unicos = []
    numeros_vistos = set()
    orgao_valor_vistos = set()
    for ed in editais:
        numero = ed.get('numero', '')
        chave = numero if numero and numero not in ['N/A', 'None', ''] else ed.get('url', '')
        # Chave secundária: orgão + valor (pega duplicatas com números diferentes)
        orgao = (ed.get('orgao') or '').strip().lower()
        valor = ed.get('valor_estimado') or ed.get('valor_referencia') or 0
        chave2 = f"{orgao}|{valor}" if orgao and valor else ""
        if chave and chave in numeros_vistos:
            continue
        if chave2 and chave2 in orgao_valor_vistos:
            continue
        if chave:
            numeros_vistos.add(chave)
        if chave2:
            orgao_valor_vistos.add(chave2)
        editais_unicos.append(ed)

    if len(editais) != len(editais_unicos):
        print(f"[BUSCA-MULTI] Removidas {len(editais) - len(editais_unicos)} duplicatas")
    editais = editais_unicos

    # 4. Filtros pós-busca: modalidade, tipo_produto, origem
    if modalidade:
        # Normalizar: "Pregão Eletrônico" → "pregao eletronico", "pregao_eletronico" → "pregao eletronico"
        import unicodedata
        def _norm_mod(s):
            s = (s or '').lower().replace("_", " ")
            # Remover acentos
            return ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')
        mod_norm = _norm_mod(modalidade)
        editais_pre = len(editais)
        editais = [e for e in editais if mod_norm in _norm_mod(e.get('modalidade') or '')]
        print(f"[BUSCA-MULTI] Filtro modalidade='{modalidade}': {editais_pre} → {len(editais)}")

    if tipo_produto:
        tp_lower = tipo_produto.lower()
        editais_pre = len(editais)
        editais = [e for e in editais if tp_lower in inferir_tipo_produto(e.get('objeto', '')).lower()]
        print(f"[BUSCA-MULTI] Filtro tipo_produto='{tipo_produto}': {editais_pre} → {len(editais)}")

    if origem:
        orig_lower = origem.lower()
        editais_pre = len(editais)
        editais = [e for e in editais if orig_lower in inferir_origem_orgao(e.get('orgao', ''), e.get('orgao_tipo', '')).lower()]
        print(f"[BUSCA-MULTI] Filtro origem='{origem}': {editais_pre} → {len(editais)}")

    return {
        "success": len(editais) > 0,
        "termo": termo,
        "fontes_consultadas": fontes_consultadas,
        "total_resultados": len(editais),
        "editais": editais,
        "erros": erros_fontes if erros_fontes else None
    }


def processar_buscar_editais(message: str, user_id: str, termo_ia: str = None, calcular_score: bool = True, incluir_encerrados: bool = None, empresa_id: str = None):
    """
    Processa ação: Buscar editais

    Novo fluxo:
    1. Busca editais (sem salvar)
    2. Calcula score de aderência para cada edital vs produtos do usuário
       - tipo_score "rapido" (padrão): score rápido em batch via deepseek-chat
       - tipo_score "hibrido": rápido em todos + profundo (6 dimensões) nos top N
       - tipo_score "profundo": profundo (6 dimensões) direto nos N primeiros
    3. Ordena por score
    4. Mostra recomendações (PARTICIPAR/AVALIAR/NÃO PARTICIPAR) com justificativas
    5. Oferece opção de salvar os recomendados

    Args:
        message: Mensagem original do usuário
        user_id: ID do usuário
        termo_ia: Termo de busca já extraído pelo agente classificador (opcional)
        calcular_score: Se True, calcula score de aderência. Se False, apenas lista os editais.
        incluir_encerrados: Se True, inclui editais já encerrados. Se None, detecta da mensagem.
    """
    import json
    import re

    fonte = None  # None = todas as fontes
    uf = None

    # Detectar fonte da mensagem
    msg_lower_fonte = message.lower()
    if 'no pncp' in msg_lower_fonte or 'do pncp' in msg_lower_fonte or 'pelo pncp' in msg_lower_fonte:
        fonte = "PNCP"
    elif 'comprasnet' in msg_lower_fonte:
        fonte = "comprasnet"
    elif 'licitacoes-e' in msg_lower_fonte or 'licitações-e' in msg_lower_fonte or 'licitanet' in msg_lower_fonte:
        fonte = "licitacoes-e"
    elif 'bec' in msg_lower_fonte and 'sp' in msg_lower_fonte:
        fonte = "bec-sp"
    elif 'compras.rs' in msg_lower_fonte or 'compras rs' in msg_lower_fonte:
        fonte = "compras.rs"
    elif 'portal de compras' in msg_lower_fonte and ('mg' in msg_lower_fonte or 'minas' in msg_lower_fonte):
        fonte = "compras.mg"
    elif 'gov.br' in msg_lower_fonte:
        fonte = "gov.br"
    # Se não especificou fonte → None (busca todas)

    # Detectar se deve incluir editais encerrados
    if incluir_encerrados is None:
        msg_lower = message.lower()
        # Negações explícitas → NÃO incluir encerrados
        negar = any(p in msg_lower for p in [
            'sem incluir encerrados', 'sem encerrados', 'não incluir encerrados',
            'nao incluir encerrados', 'excluir encerrados', 'apenas abertos',
            'somente abertos', 'só abertos', 'so abertos',
        ])
        if negar:
            incluir_encerrados = False
        else:
            # Afirmações explícitas → incluir encerrados
            incluir_encerrados = any(p in msg_lower for p in [
                'incluindo encerrados', 'incluir encerrados', 'com encerrados',
                'mostrar encerrados', 'ver encerrados',
            ])

    # Detectar janela de busca (dias) da mensagem
    dias_busca = 90  # default
    msg_lower_dias = message.lower()
    dias_match = re.search(r'janela\s+de\s+(\d+)\s*dias?', msg_lower_dias)
    if not dias_match:
        dias_match = re.search(r'(?:ultimos|últimos)\s+(\d+)\s*dias?', msg_lower_dias)
    if dias_match:
        dias_busca = int(dias_match.group(1))
        if dias_busca != 0:
            dias_busca = max(7, min(730, dias_busca))
    elif 'sem limite' in msg_lower_dias or 'periodo indefinido' in msg_lower_dias:
        dias_busca = 0

    # Detectar tipo de score da mensagem (rapido/hibrido/profundo)
    tipo_score = "rapido"  # padrão quando calcular_score=True
    limite_score_profundo = 10  # padrão
    if not calcular_score:
        tipo_score = "nenhum"
    else:
        msg_lower_score = message.lower()
        if any(p in msg_lower_score for p in ['score profundo', 'análise profunda', 'analise profunda',
                                               'score detalhado', 'análise detalhada', 'analise detalhada',
                                               'com profundo', '6 dimensões', '6 dimensoes', 'score completo']):
            tipo_score = "profundo"
        elif any(p in msg_lower_score for p in ['score híbrido', 'score hibrido', 'com híbrido', 'com hibrido',
                                                 'modo híbrido', 'modo hibrido']):
            tipo_score = "hibrido"
        elif any(p in msg_lower_score for p in ['score rápido', 'score rapido', 'com rápido', 'com rapido']):
            tipo_score = "rapido"

        # Detectar limite para profundo/híbrido: "nos 5 melhores", "top 10", "nos 20 primeiros", "todos"
        if tipo_score in ("profundo", "hibrido"):
            limite_match = re.search(r'(?:nos|top|primeiros|melhores)\s+(\d+)', msg_lower_score)
            if not limite_match:
                limite_match = re.search(r'(\d+)\s+(?:melhores|primeiros|editais)', msg_lower_score)
            if limite_match:
                limite_score_profundo = int(limite_match.group(1))
                limite_score_profundo = max(1, min(100, limite_score_profundo))
            elif 'todos' in msg_lower_score and ('profundo' in msg_lower_score or 'híbrido' in msg_lower_score or 'hibrido' in msg_lower_score):
                limite_score_profundo = 9999  # Todos

    # Usar termo da IA se disponível, senão extrair da mensagem
    if termo_ia:
        termo = termo_ia
        # Limpar modificadores que a IA pode ter incluído no termo
        import re as _re
        for mod in ['incluindo encerrados', 'sem calcular score', 'sem score', 'sem calcular',
                     'apenas listando', 'no pncp', 'no comprasnet', 'na web',
                     'com modalidade', 'modalidade', 'do tipo', 'de origem',
                     'pregão eletrônico', 'pregao eletronico', 'pregão presencial', 'pregao presencial',
                     'concorrência', 'concorrencia', 'tomada de preços', 'tomada de precos',
                     'convite', 'dispensa', 'inexigibilidade',
                     'origem hospital', 'origem municipal', 'origem estadual', 'origem federal',
                     'origem universidade', 'origem autarquia', 'origem lacen',
                     'tipo reagente', 'tipo equipamento', 'tipo comodato', 'tipo aluguel',
                     'tipo informática', 'tipo informatica', 'tipo insumo',
                     'com score profundo', 'com score rápido', 'com score rapido',
                     'com score híbrido', 'com score hibrido', 'score profundo', 'score rápido',
                     'score rapido', 'score híbrido', 'score hibrido', 'score detalhado',
                     'score completo', 'análise profunda', 'analise profunda',
                     'análise detalhada', 'analise detalhada', 'com profundo', 'com rápido',
                     'com rapido', 'com híbrido', 'com hibrido', 'modo híbrido', 'modo hibrido',
                     '6 dimensões', '6 dimensoes']:
            termo = _re.sub(r'\b' + _re.escape(mod) + r'\b', '', termo, flags=_re.IGNORECASE)
        # Limpar "últimos X dias"
        termo = _re.sub(r'(?:ultimos|últimos)\s+\d+\s*dias?', '', termo, flags=_re.IGNORECASE)
        # Limpar "nos N melhores/primeiros", "top N"
        termo = _re.sub(r'(?:nos|top)\s+\d+\s*(?:melhores|primeiros|editais)?', '', termo, flags=_re.IGNORECASE)
        termo = _re.sub(r'\d+\s+(?:melhores|primeiros)', '', termo, flags=_re.IGNORECASE)
        # Limpar espaços duplicados
        termo = _re.sub(r'\s+', ' ', termo).strip()
        print(f"[BUSCA] Usando termo da IA: '{termo}'")
    else:
        termo = None
        # Tentar extrair parâmetros com LLM (usar deepseek-chat para rapidez)
        prompt = f"""Extraia os parâmetros de busca de editais da mensagem.
Retorne APENAS um JSON válido com: fonte (PNCP, ComprasNet, BEC-SP ou null), termo (palavras-chave da busca), uf (sigla do estado com 2 letras ou null)

Mensagem: "{message}"

JSON:"""

        try:
            resposta = call_deepseek([{"role": "user", "content": prompt}], max_tokens=200, model_override="deepseek-chat")
            json_match = re.search(r'\{[\s\S]*?\}', resposta)
            if json_match:
                dados = json.loads(json_match.group())
                # Só usar fonte do DeepSeek se não foi detectada por palavras-chave
                if fonte is None:
                    fonte_ds = dados.get('fonte')
                    if fonte_ds:
                        fonte = fonte_ds
                termo = dados.get('termo')
                uf = dados.get('uf')
        except Exception as e:
            print(f"Erro ao extrair parâmetros com LLM: {e}")

    # Fallback: extrair termos da própria mensagem
    if not termo:
        # Remover palavras comuns de comando e modificadores de busca
        palavras_ignorar = ['busque', 'buscar', 'procure', 'procurar', 'editais', 'edital', 'de', 'do', 'da',
                           'no', 'na', 'em', 'para', 'pncp', 'comprasnet', 'bec', 'sp', 'são', 'paulo',
                           'retorne', 'mostre', 'liste', 'quero', 'ver', 'todos', 'área', 'area',
                           'incluindo', 'encerrados', 'sem', 'calcular', 'score', 'aderência', 'aderencia',
                           'apenas', 'listando', 'últimos', 'ultimos', 'dias', 'com', 'web',
                           'licitacoes-e', 'licitacoes', 'portal', 'compras', 'fonte',
                           'rápido', 'rapido', 'híbrido', 'hibrido', 'profundo', 'detalhado',
                           'completo', 'análise', 'analise', 'profunda', 'detalhada', 'modo',
                           'melhores', 'primeiros', 'top', 'nos', 'dimensões', 'dimensoes']
        palavras = message.lower().split()
        termos = [p for p in palavras if p not in palavras_ignorar and len(p) > 2]
        termo = ' '.join(termos) if termos else message

    # Limpar frases de filtro do termo (modalidade, tipo, origem)
    _filtros_limpar = [
        r'com\s+modalidade\s+\S+(?:\s+\S+)?', r'modalidade\s+\S+(?:\s+\S+)?',
        r'do\s+tipo\s+\S+', r'tipo\s+\S+',
        r'de\s+origem\s+\S+', r'origem\s+\S+',
    ]
    for padrao in _filtros_limpar:
        termo = re.sub(padrao, '', termo, flags=re.IGNORECASE)
    termo = re.sub(r'\s+', ' ', termo).strip()

    # Detectar UF na mensagem
    ufs = ['AC','AL','AP','AM','BA','CE','DF','ES','GO','MA','MT','MS','MG','PA','PB','PR','PE','PI','RJ','RN','RS','RO','RR','SC','SP','SE','TO']
    msg_upper = message.upper()
    for sigla in ufs:
        if f" {sigla} " in f" {msg_upper} " or msg_upper.endswith(f" {sigla}"):
            uf = sigla
            break
    # Detectar por nome do estado
    if "SÃO PAULO" in msg_upper or "SAO PAULO" in msg_upper:
        uf = "SP"
    elif "RIO DE JANEIRO" in msg_upper:
        uf = "RJ"
    elif "MINAS GERAIS" in msg_upper:
        uf = "MG"

    # Detectar modalidade da mensagem
    modalidade = None
    _modalidades_map = {
        'pregão eletrônico': 'Pregão Eletrônico', 'pregao eletronico': 'Pregão Eletrônico',
        'pregão presencial': 'Pregão Presencial', 'pregao presencial': 'Pregão Presencial',
        'concorrência': 'Concorrência', 'concorrencia': 'Concorrência',
        'tomada de preços': 'Tomada de Preços', 'tomada de precos': 'Tomada de Preços',
        'convite': 'Convite', 'dispensa': 'Dispensa', 'inexigibilidade': 'Inexigibilidade',
    }
    msg_lower_mod = message.lower()
    for chave, valor in _modalidades_map.items():
        if chave in msg_lower_mod:
            modalidade = valor
            break

    # Detectar tipo de produto da mensagem
    tipo_produto = None
    _tipos_map = {
        'reagente': 'Reagentes', 'equipamento': 'Equipamentos', 'comodato': 'Comodato',
        'aluguel': 'Aluguel', 'locação': 'Aluguel', 'locacao': 'Aluguel',
        'insumo': 'Insumos Hospitalares', 'informática': 'Informática', 'informatica': 'Informática',
        'rede': 'Redes', 'mobiliário': 'Mobiliário', 'mobiliario': 'Mobiliário',
    }
    for chave, valor in _tipos_map.items():
        if f'tipo {chave}' in msg_lower_mod or f'do tipo {chave}' in msg_lower_mod:
            tipo_produto = valor
            break

    # Detectar origem da mensagem
    origem = None
    _origens_map = {
        'hospital': 'Hospital', 'universidade': 'Universidade', 'lacen': 'LACEN',
        'força armada': 'Força Armada', 'forca armada': 'Força Armada',
        'autarquia': 'Autarquia', 'municipal': 'Municipal', 'estadual': 'Estadual',
        'federal': 'Federal', 'centro de pesquisa': 'Centros de Pesquisas',
        'fundação de pesquisa': 'Fundações de Pesquisa', 'fundacao de pesquisa': 'Fundações de Pesquisa',
    }
    for chave, valor in _origens_map.items():
        if f'origem {chave}' in msg_lower_mod or f'de origem {chave}' in msg_lower_mod:
            origem = valor
            break

    print(f"[BUSCA] Termo final: '{termo}', Fonte: '{fonte}', UF: '{uf}', Modalidade: '{modalidade}', Tipo: '{tipo_produto}', Origem: '{origem}', Score: '{tipo_score}' (limite={limite_score_profundo})")

    # ========== PASSO 1: Buscar editais em MÚLTIPLAS FONTES ==========
    resultado = _buscar_editais_multifonte(
        termo=termo, user_id=user_id, uf=uf, incluir_encerrados=incluir_encerrados,
        dias_busca=dias_busca, fonte=fonte,
        modalidade=modalidade, tipo_produto=tipo_produto, origem=origem,
    )
    editais = resultado.get("editais", [])
    fontes_consultadas = resultado.get("fontes_consultadas", [])
    erros_fontes = resultado.get("erros") or []

    if not editais:
        fontes_str = ', '.join(fontes_consultadas) if fontes_consultadas else 'nenhuma fonte disponível'
        response = f"""**Busca realizada:** {termo}
**Fontes consultadas:** {fontes_str}

⚠️ Nenhum edital encontrado para '{termo}'.

**Sugestões:**
- Tente termos mais específicos (ex: "monitor LCD 24 polegadas")
- Verifique se há editais salvos: "liste meus editais"
- Cadastre mais fontes de editais: "cadastre a fonte BEC-SP"
"""
        if erros_fontes:
            response += f"\n**Erros nas fontes:** {'; '.join(erros_fontes)}\n"
        return response, resultado

    # ========== PASSO 2: Calcular score de aderência (se solicitado) ==========
    aviso_produtos = None
    if tipo_score == "rapido" and editais:
        print(f"[APP] Score RAPIDO para {len(editais)} editais...")
        resultado_score = tool_calcular_score_aderencia(editais, user_id, empresa_id=empresa_id)
        if resultado_score.get("success"):
            editais_com_score = resultado_score.get("editais_com_score", editais)
            aviso_produtos = resultado_score.get("aviso")
        else:
            editais_com_score = editais

    elif tipo_score == "hibrido" and editais:
        print(f"[APP] Score HIBRIDO para {len(editais)} editais (profundo nos top {limite_score_profundo})...")
        # Passo 1: Score rápido em todos
        resultado_score = tool_calcular_score_aderencia(editais, user_id, empresa_id=empresa_id)
        if resultado_score.get("success"):
            editais_com_score = resultado_score.get("editais_com_score", editais)
            aviso_produtos = resultado_score.get("aviso")
        else:
            editais_com_score = editais
        # Passo 2: Ordenar por score, pegar top N
        editais_com_score.sort(key=lambda x: x.get("score_tecnico", 0), reverse=True)
        top_n = editais_com_score[:limite_score_profundo]
        rest = editais_com_score[limite_score_profundo:]
        # Passo 3: Score profundo nos top N
        top_n = _calcular_score_profundo(top_n, user_id, empresa_id)
        editais_com_score = top_n

    elif tipo_score == "profundo" and editais:
        print(f"[APP] Score PROFUNDO direto em {min(len(editais), limite_score_profundo)} editais...")
        target = editais[:limite_score_profundo]
        target = _calcular_score_profundo(target, user_id, empresa_id)
        editais_com_score = target

    elif tipo_score == "nenhum":
        print(f"[APP] Busca SIMPLES (sem score) - {len(editais)} editais encontrados")
        editais_com_score = editais
    else:
        editais_com_score = editais

    # ========== PASSO 3: Formatar resposta ==========
    total_encontrados = len(editais_com_score)

    fontes_str = ', '.join(fontes_consultadas) if fontes_consultadas else fonte
    _modos_label = {
        "nenhum": "listagem simples (sem score)",
        "rapido": "com score rápido de aderência",
        "hibrido": f"com score híbrido (rápido + profundo nos top {limite_score_profundo})",
        "profundo": f"com score profundo (6 dimensões) em {min(total_encontrados, limite_score_profundo)} editais",
    }
    modo_busca = _modos_label.get(tipo_score, "com análise de aderência")
    response = f"""**Busca realizada:** {termo}
**Fontes consultadas:** {fontes_str}
**Modo:** {modo_busca}
**Resultados:** {total_encontrados} edital(is) encontrado(s)

"""

    if aviso_produtos:
        response += f"⚠️ {aviso_produtos}\n\n"

    # Ordenar: abertos primeiro (por encerramento mais próximo), encerrados no final
    from datetime import datetime as _dt_chat
    def _sort_key_chat(e):
        encerrado = e.get("encerrado", False)
        d = e.get("data_encerramento") or e.get("data_abertura") or ""
        dt = _dt_chat(2099, 12, 31)
        for fmt in ("%d/%m/%Y %H:%M", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d"):
            try:
                dt = _dt_chat.strptime(d[:19], fmt)
                break
            except (ValueError, TypeError):
                continue
        # Encerrados vão para o final (1), abertos primeiro (0)
        return (1 if encerrado else 0, dt)
    editais_com_score.sort(key=_sort_key_chat)

    # Separar por recomendação — todos os editais, sem limite
    participar = [e for e in editais_com_score if e.get('recomendacao') == 'PARTICIPAR']
    avaliar = [e for e in editais_com_score if e.get('recomendacao') == 'AVALIAR']
    nao_participar = [e for e in editais_com_score if e.get('recomendacao') == 'NÃO PARTICIPAR']
    sem_score = [e for e in editais_com_score if not e.get('recomendacao')]

    def formatar_edital(ed, i):
        """Formata um edital para exibição com botões de ação"""
        numero = ed.get('numero', 'N/A')
        orgao = ed.get('orgao', 'N/A')
        uf_ed = ed.get('uf', '')
        cidade = ed.get('cidade', '')
        local = f"{cidade}/{uf_ed}" if cidade and uf_ed else (uf_ed or cidade or 'Brasil')
        objeto = ed.get('objeto', '')[:200]
        valor = ed.get('valor_referencia')
        valor_str = f"R$ {valor:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') if valor else "Não informado"
        data_abertura = ed.get('data_abertura', 'Não informada')
        modalidade = ed.get('modalidade', 'N/A')
        url = ed.get('url', '')
        score = ed.get('score_tecnico')
        justificativa = ed.get('justificativa', '')
        fonte_edital = ed.get('fonte', '')
        pdf_url = ed.get('pdf_url', '')
        total_itens = ed.get('total_itens', 0)
        cnpj = ed.get('cnpj_orgao', '')
        ano = ed.get('ano_compra', '')
        seq = ed.get('seq_compra', '')
        dados_completos = ed.get('dados_completos', False)
        encerrado = ed.get('encerrado', False)

        # Badge de fonte com cor
        fonte_badge = ""
        if 'PNCP' in fonte_edital:
            fonte_badge = f"🟢 {fonte_edital}"
        elif 'ComprasNet' in fonte_edital:
            fonte_badge = f"🔵 {fonte_edital}"
        elif 'BEC' in fonte_edital:
            fonte_badge = f"🟡 {fonte_edital}"
        elif 'Scraper' in fonte_edital:
            fonte_badge = f"🟠 {fonte_edital}"
        else:
            fonte_badge = f"⚪ {fonte_edital}" if fonte_edital else ""

        # Badge de status (encerrado ou aberto)
        status_badge = "🔴 ENCERRADO" if encerrado else ""

        texto = f"---\n"
        texto += f"### {i}. [{fonte_badge}] {numero}"
        if score is not None:
            texto += f" | Score: **{score:.0f}%**"
        if status_badge:
            texto += f" | {status_badge}"
        texto += "\n"
        texto += f"**Órgão:** {orgao} ({local})\n"
        texto += f"**Modalidade:** {modalidade}\n"
        texto += f"**Valor estimado:** {valor_str}\n"
        texto += f"**Data abertura:** {data_abertura}\n"
        if total_itens > 0:
            texto += f"**Itens:** {total_itens} item(ns)\n"
        if dados_completos:
            texto += f"**Dados:** ✅ Completos (PNCP)\n"
        texto += f"**Objeto:** {objeto}\n"
        if justificativa:
            texto += f"\n**Análise:** {justificativa}\n"

        # Score profundo (6 dimensões) se disponível
        score_profundo = ed.get('score_profundo')
        if score_profundo:
            sp_scores = score_profundo.get('scores', {})
            decisao = score_profundo.get('decisao', 'AVALIAR')
            decisao_icon = "✅" if decisao == "GO" else ("❌" if decisao == "NO-GO" else "⚠️")
            texto += f"\n**Score Profundo:** {decisao_icon} **{decisao}** (Geral: {score_profundo.get('score_geral', 0)}%)\n"
            texto += f"| Dimensão | Score |\n|---|---|\n"
            for dim_key, dim_label in [('tecnico','Técnico'), ('documental','Documental'), ('complexidade','Complexidade'),
                                       ('juridico','Jurídico'), ('logistico','Logístico'), ('comercial','Comercial')]:
                val = sp_scores.get(dim_key, 0)
                texto += f"| {dim_label} | {val}% |\n"
            sp_just = score_profundo.get('justificativa', '')
            if sp_just:
                texto += f"\n{sp_just}\n"
            positivos = score_profundo.get('pontos_positivos', [])
            atencao = score_profundo.get('pontos_atencao', [])
            if positivos:
                texto += f"\n**Pontos positivos:** " + "; ".join(positivos[:3]) + "\n"
            if atencao:
                texto += f"**Pontos de atenção:** " + "; ".join(atencao[:3]) + "\n"

        # Botões de ação
        texto += f"\n"
        if url:
            texto += f"🔗 [Acessar Portal]({url}) "

        # Botão PDF - usar proxy local para visualização inline
        # URL absoluta do backend (porta 5007)
        backend_url = "http://localhost:5007"
        if cnpj and ano and seq:
            # URL do proxy local que permite visualização no navegador
            proxy_url = f"{backend_url}/api/proxy/pdf/pncp/{cnpj}/{ano}/{seq}"
            texto += f"| 📄 [Ver PDF]({proxy_url}) "
            texto += f"| ⬇️ [Baixar PDF]({proxy_url}?download=true) "
        elif pdf_url:
            # Se tem pdf_url mas não tem dados PNCP, usa direto (com fallback)
            texto += f"| 📄 [Ver PDF]({pdf_url}) "
            texto += f"| ⬇️ [Baixar PDF]({pdf_url}?download=true) "

        texto += "\n\n"
        return texto

    contador = 1

    # Editais recomendados (PARTICIPAR)
    if participar:
        response += "## ✅ RECOMENDADOS PARA PARTICIPAR\n\n"
        for ed in participar:
            response += formatar_edital(ed, contador)
            contador += 1

    # Editais para avaliar
    if avaliar:
        response += "## ⚠️ AVALIAR PARTICIPAÇÃO\n\n"
        for ed in avaliar:
            response += formatar_edital(ed, contador)
            contador += 1

    # Editais não recomendados
    if nao_participar:
        response += "## ❌ NÃO RECOMENDADOS\n\n"
        for ed in nao_participar:
            response += formatar_edital(ed, contador)
            contador += 1

    # Sem score (sem produtos cadastrados)
    if sem_score:
        response += "## 📋 EDITAIS ENCONTRADOS\n\n"
        for ed in sem_score:
            response += formatar_edital(ed, contador)
            contador += 1

    # ========== PASSO 4: Oferecer salvamento ==========
    qtd_participar = len(participar)
    qtd_avaliar = len(avaliar)
    qtd_recomendados = qtd_participar + qtd_avaliar

    if tipo_score != "nenhum" and qtd_recomendados > 0:
        # Busca COM score - mostrar opções por recomendação
        response += f"\n---\n"
        response += f"## 💾 Deseja salvar os editais?\n\n"
        response += f"Encontrei **{qtd_recomendados} edital(is)** recomendado(s):\n"
        if qtd_participar > 0:
            response += f"- ✅ {qtd_participar} para PARTICIPAR\n"
        if qtd_avaliar > 0:
            response += f"- ⚠️ {qtd_avaliar} para AVALIAR\n"
        response += f"\n"

        # Botões de ação (marcação especial para o frontend)
        response += f"<!-- BOTOES_SALVAR -->\n"
        response += f"[[btn:salvar_recomendados:💾 Salvar Recomendados ({qtd_recomendados})]]\n"
        if qtd_participar > 0 and qtd_avaliar > 0:
            response += f"[[btn:salvar_participar:✅ Salvar só PARTICIPAR ({qtd_participar})]]\n"
        response += f"[[btn:salvar_todos:📋 Salvar Todos ({len(editais_com_score)})]]\n"
        response += f"<!-- /BOTOES_SALVAR -->\n\n"

        response += f"*Ou digite: \"salvar editais\", \"salvar recomendados\", \"salvar edital PE-2026/001\"*\n"

    elif tipo_score == "nenhum" and len(editais_com_score) > 0:
        # Busca SEM score - oferecer salvar todos
        response += f"\n---\n"
        response += f"## 💾 Deseja salvar os editais?\n\n"
        response += f"Encontrei **{len(editais_com_score)} edital(is)**.\n\n"

        # Botões de ação
        response += f"<!-- BOTOES_SALVAR -->\n"
        response += f"[[btn:salvar_todos:💾 Salvar Todos ({len(editais_com_score)})]]\n"
        response += f"<!-- /BOTOES_SALVAR -->\n\n"

        response += f"*Ou digite: \"salvar editais\", \"salvar todos\", \"salvar edital [NÚMERO]\"*\n"

    # Adicionar editais ao resultado para possível salvamento posterior
    resultado["editais_com_score"] = editais_com_score
    resultado["editais_recomendados"] = participar + avaliar
    resultado["editais_participar"] = participar
    resultado["editais_avaliar"] = avaliar

    return response, resultado


def processar_buscar_links_editais(message: str, user_id: str):
    """
    Processa ação: Retornar links de editais em uma área.
    Retorna links formatados para o usuário clicar.
    """
    import re

    # Extrair termo/área da mensagem
    termo = None

    # Padrões comuns
    padroes = [
        r'links?\s+(?:para\s+)?(?:os\s+)?editais?\s+(?:na\s+|da\s+|de\s+)?(?:área|area)\s+(.+)',
        r'links?\s+(?:de\s+)?editais?\s+(?:de\s+|para\s+|em\s+)?(.+)',
        r'editais?\s+(?:de\s+|para\s+|em\s+|na\s+área\s+)?(.+)',
        r'busca.+links?\s+(.+)',
        r'(?:retorne|mostre|liste)\s+(?:os\s+)?links?\s+(?:para\s+|de\s+)?(.+)',
    ]

    msg_lower = message.lower()
    for padrao in padroes:
        match = re.search(padrao, msg_lower, re.IGNORECASE)
        if match:
            termo = match.group(1).strip()
            # Limpar palavras desnecessárias
            palavras_remover = ['por favor', 'obrigado', 'pncp', 'web', 'internet']
            for p in palavras_remover:
                termo = termo.replace(p, '').strip()
            break

    # Fallback: usar toda a mensagem após limpeza
    if not termo:
        palavras_ignorar = ['retorne', 'mostre', 'liste', 'busque', 'links', 'link', 'editais',
                           'edital', 'para', 'os', 'de', 'da', 'do', 'na', 'no', 'área', 'area']
        palavras = msg_lower.split()
        termos = [p for p in palavras if p not in palavras_ignorar and len(p) > 2]
        termo = ' '.join(termos) if termos else "equipamentos"

    print(f"[LINKS] Buscando links para área: '{termo}'")

    # Chamar a função de busca de links
    resultado = tool_buscar_links_editais(termo, user_id=user_id)

    if resultado.get("success"):
        return resultado.get("texto", "Nenhum resultado"), resultado
    else:
        return f"❌ Erro ao buscar links: {resultado.get('error', 'Erro desconhecido')}", resultado


def processar_buscar_edital_numero(message: str, user_id: str, buscar_apenas_banco: bool = False):
    """
    Processa busca de um edital específico pelo número.

    Args:
        message: Mensagem do usuário contendo o número do edital
        user_id: ID do usuário
        buscar_apenas_banco: Se True, busca APENAS no banco local. Se False, busca no banco e depois na web.

    Returns:
        Tuple (response_text, resultado)
    """
    import re
    import requests

    # Detectar se é busca no banco ou na web
    msg_lower = message.lower()
    busca_local = any(p in msg_lower for p in ["no banco", "cadastrado", "salvo", "no sistema", "banco de dados",
                                                "tenho o edital", "tenho edital", "já tenho", "ja tenho"])

    # Se especificou banco na mensagem, força busca local
    if busca_local:
        buscar_apenas_banco = True

    # Extrair número do edital da mensagem
    # Padrões: PE-001/2026, PE0013/2025, PE 050/2025, 90186/2025, nº 123, número 456
    padroes = [
        r'PE[-\s]?\d+[/\-]?\d*',  # PE-001/2026, PE0013/2025, PE 050/2025
        r'[Pp]reg[aã]o\s*(?:n[ºo°]?\s*)?\d+[/\-]?\d*',  # Pregão nº 123/2025
        r'[Ee]dital\s*(?:n[ºo°]?\s*)?\d+[/\-]?\d*',  # Edital nº 123/2025
        r'\d{3,}[/\-]\d{4}',  # 90186/2025, 050/2025
        r'n[ºo°]\s*\d+[/\-]?\d*',  # nº 123/2025
    ]

    numero_edital = None
    for padrao in padroes:
        match = re.search(padrao, message, re.IGNORECASE)
        if match:
            numero_edital = match.group().strip()
            # Limpar prefixos comuns que não fazem parte do número
            numero_edital = re.sub(r'^(edital|pregão|pregao|pe|nº|no|n°)\s*', '', numero_edital, flags=re.IGNORECASE).strip()
            break

    if not numero_edital:
        return """❌ **Não consegui identificar o número do edital.**

Por favor, informe o número no formato:
- "Busque o edital PE-001/2026 no banco"
- "Busque o edital PE-001/2026 no PNCP"
- "Tenho o edital PE-001/2026 cadastrado?"
""", None

    print(f"[BUSCA-EDITAL] Buscando edital: {numero_edital} | Apenas banco: {buscar_apenas_banco}")

    # 1. Verificar se está salvo no sistema
    from models import Edital
    from database import SessionLocal

    db = SessionLocal()
    try:
        edital_local = db.query(Edital).filter(
            Edital.numero.ilike(f"%{numero_edital}%"),
            Edital.user_id == user_id
        ).first()

        if edital_local:
            from models import EditalItem, EditalDocumento

            valor_ref = f"R$ {edital_local.valor_referencia:,.2f}" if edital_local.valor_referencia else '-'
            data_ab = edital_local.data_abertura.strftime('%d/%m/%Y %H:%M') if edital_local.data_abertura else '-'
            data_pub = edital_local.data_publicacao.strftime('%d/%m/%Y') if edital_local.data_publicacao else '-'
            objeto_texto = (edital_local.objeto or '')[:200]
            objeto_sufixo = '...' if len(edital_local.objeto or '') > 200 else ''

            response = f"""## ✅ Edital Encontrado no Sistema

### Dados Gerais
| Campo | Valor |
|-------|-------|
| **Número** | {edital_local.numero} |
| **Órgão** | {edital_local.orgao} |
| **UF/Cidade** | {edital_local.uf or '-'} / {edital_local.cidade or '-'} |
| **Modalidade** | {edital_local.modalidade or '-'} |
| **Status** | {edital_local.status or '-'} |
| **Valor Referência** | {valor_ref} |
| **Data Publicação** | {data_pub} |
| **Data Abertura** | {data_ab} |

### Objeto
{objeto_texto}{objeto_sufixo}

"""
            # Buscar itens do edital
            itens = db.query(EditalItem).filter(EditalItem.edital_id == edital_local.id).order_by(EditalItem.numero_item).all()
            if itens:
                response += f"### Itens ({len(itens)})\n"
                response += "| Item | Descrição | Qtd | Valor Total |\n"
                response += "|------|-----------|-----|-------------|\n"
                for item in itens:
                    desc = (item.descricao or '')[:50]
                    desc_sufixo = '...' if len(item.descricao or '') > 50 else ''
                    qtd = f"{item.quantidade:,.0f} {item.unidade_medida or ''}" if item.quantidade else '-'
                    valor = f"R$ {item.valor_total_estimado:,.2f}" if item.valor_total_estimado else '-'
                    response += f"| {item.numero_item or '-'} | {desc}{desc_sufixo} | {qtd} | {valor} |\n"
                response += "\n"

            # Verificar se tem PDF
            doc = db.query(EditalDocumento).filter(EditalDocumento.edital_id == edital_local.id).first()
            if doc and doc.path_arquivo:
                import os
                if os.path.exists(doc.path_arquivo):
                    response += f"### Documento\n"
                    response += f"📄 **{doc.nome_arquivo}** ({len(doc.texto_extraido or ''):,} caracteres extraídos)\n\n"
                else:
                    response += f"### Documento\n"
                    response += f"⚠️ PDF não disponível (arquivo removido)\n\n"

            # Dados PNCP
            if edital_local.cnpj_orgao:
                response += f"### Dados PNCP\n"
                response += f"- **CNPJ Órgão:** {edital_local.cnpj_orgao}\n"
                response += f"- **Nº PNCP:** {edital_local.numero_pncp or '-'}\n"
                response += f"- **Situação:** {edital_local.situacao_pncp or '-'}\n"
                response += f"- **SRP:** {'Sim' if edital_local.srp else 'Não'}\n\n"

            # URL
            if edital_local.url:
                response += f"### Link\n🔗 {edital_local.url}\n\n"

            response += f"""---
**Ações disponíveis:**
- Baixar PDF: "Baixe o PDF do edital {edital_local.numero}"
- Fazer perguntas: "Qual o prazo de entrega do edital {edital_local.numero}?"
- Calcular aderência: "Calcule aderência do produto X ao edital {edital_local.numero}"
"""
            return response, {"edital": edital_local.numero, "encontrado_local": True, "id": edital_local.id}

        # Se não encontrou no banco e é busca apenas local
        if buscar_apenas_banco:
            return f"""## ❌ Edital não encontrado no banco

O edital **{numero_edital}** não está cadastrado no sistema.

**Opções:**
- Buscar na web: "Busque o edital {numero_edital} no PNCP"
- Cadastrar manualmente: "Cadastre o edital {numero_edital}, órgão [ORGAO], objeto: [OBJETO]"
- Buscar por termo: "Busque editais de [TERMO] no PNCP"
""", {"numero": numero_edital, "encontrado": False}

    finally:
        db.close()

    # 2. Buscar no PNCP (apenas se não for busca exclusiva no banco)
    try:
        # Limpar número para busca
        numero_limpo = re.sub(r'[^\d/]', '', numero_edital)

        url = f"https://pncp.gov.br/api/consulta/v1/contratacoes/publicacao"
        params = {
            "q": numero_edital,
            "pagina": 1,
            "tamanhoPagina": 10
        }

        response_api = requests.get(url, params=params, timeout=30)

        if response_api.status_code == 200:
            dados = response_api.json()
            resultados = dados.get("data", []) or dados.get("items", []) or []

            if resultados:
                response = f"## 🌐 Resultados da Web para: {numero_edital}\n\n"
                response += f"Encontrados: **{len(resultados)}** edital(is)\n\n"

                for i, ed in enumerate(resultados[:5], 1):
                    orgao = ed.get("orgaoEntidade", {}).get("razaoSocial", ed.get("nomeOrgao", "N/A"))
                    numero = ed.get("numeroControlePNCP", ed.get("numero", "N/A"))
                    objeto = ed.get("objetoCompra", ed.get("objeto", "N/A"))[:150]

                    response += f"""### {i}. {numero}
- **Órgão:** {orgao}
- **Objeto:** {objeto}...
- **Modalidade:** {ed.get("modalidadeNome", "N/A")}

"""

                response += "\n---\n*Para salvar, use: \"Salve os editais encontrados\"*"
                return response, {"editais": resultados, "termo": numero_edital}

        # Se não encontrou na API principal, tentar busca genérica
        return processar_buscar_editais(f"edital {numero_edital}", user_id, calcular_score=False)

    except Exception as e:
        print(f"[BUSCA-EDITAL] Erro na API PNCP: {e}")
        # Fallback para busca genérica
        return processar_buscar_editais(f"edital {numero_edital}", user_id, calcular_score=False)


def processar_listar_produtos(message: str, user_id: str):
    """Processa ação: Listar produtos do usuário"""
    resultado = tool_listar_produtos(user_id)

    if resultado.get("success"):
        produtos = resultado.get("produtos", [])
        if produtos:
            response = f"**Seus produtos cadastrados:** {len(produtos)}\n\n"

            # Agrupar por categoria
            por_categoria = {}
            for p in produtos:
                cat = p.get("categoria", "outro")
                if cat not in por_categoria:
                    por_categoria[cat] = []
                por_categoria[cat].append(p)

            for cat, prods in sorted(por_categoria.items()):
                response += f"**[{cat.upper()}]**\n"
                for p in prods:
                    response += f"- {p['nome']} ({p.get('fabricante', 'N/A')} - {p.get('modelo', 'N/A')})\n"
                response += "\n"
        else:
            response = "Você não tem produtos cadastrados ainda. Faça upload de um manual PDF para cadastrar."
    else:
        response = f"Erro ao listar produtos: {resultado.get('error')}"

    return response, resultado


def processar_reprocessar_produto(message: str, user_id: str):
    """
    Reprocessa um produto para extrair especificações novamente.
    Útil quando a extração inicial falhou ou foi incompleta.
    """
    # Tentar identificar o produto na mensagem
    # Primeiro listar produtos do usuário
    produtos_resultado = tool_listar_produtos(user_id)

    if not produtos_resultado.get("success"):
        return "Erro ao buscar seus produtos.", produtos_resultado

    produtos = produtos_resultado.get("produtos", [])
    if not produtos:
        return "Você não tem produtos cadastrados para reprocessar.", {"success": False}

    # Tentar encontrar o produto mencionado na mensagem
    msg_lower = message.lower()
    produto_id = None
    produto_nome = None

    for p in produtos:
        nome_lower = p.get("nome", "").lower()
        modelo_lower = (p.get("modelo") or "").lower()

        # Verificar se nome ou modelo está na mensagem
        if nome_lower and any(parte in msg_lower for parte in nome_lower.split()[:3]):
            produto_id = p.get("id")
            produto_nome = p.get("nome")
            break
        if modelo_lower and modelo_lower in msg_lower:
            produto_id = p.get("id")
            produto_nome = p.get("nome")
            break

    # Se não encontrou, usar o último produto cadastrado
    if not produto_id:
        ultimo = produtos[-1]
        produto_id = ultimo.get("id")
        produto_nome = ultimo.get("nome")

    # Reprocessar
    print(f"[APP] Reprocessando produto: {produto_nome} ({produto_id})")
    resultado = tool_reprocessar_produto(produto_id, user_id)

    if resultado.get("success"):
        specs = resultado.get("specs", [])
        response = f"""## 🔄 Produto Reprocessado!

**Produto:** {resultado.get('produto_nome', produto_nome)}
**ID:** {produto_id}

### Especificações Extraídas ({resultado.get('specs_extraidas', 0)} encontradas):

"""
        for spec in specs[:30]:
            nome = spec.get('nome_especificacao', 'N/A')
            valor = spec.get('valor', 'N/A')
            unidade = spec.get('unidade', '')
            response += f"- **{nome}:** {valor} {unidade}\n"

        if len(specs) > 30:
            response += f"\n... e mais {len(specs) - 30} especificações.\n"

        response += "\n✅ Produto atualizado e pronto para calcular aderência!"
    else:
        response = f"❌ Erro ao reprocessar: {resultado.get('error')}"

    return response, resultado


def processar_excluir_edital(message: str, user_id: str):
    """
    Processa ação: Excluir edital(is).
    Identifica editais por número, ID ou palavras-chave na mensagem.
    """
    from tools import tool_excluir_edital, tool_excluir_editais_multiplos, tool_listar_editais

    msg_lower = message.lower()

    # Verificar se é exclusão de todos
    if "todos" in msg_lower:
        editais_resultado = tool_listar_editais(user_id)
        if not editais_resultado.get("success") or not editais_resultado.get("editais"):
            return "Você não tem editais salvos para excluir.", {"success": False}

        edital_ids = [e["id"] for e in editais_resultado.get("editais", [])]
        resultado = tool_excluir_editais_multiplos(edital_ids, user_id)

        if resultado.get("success"):
            return f"✅ {resultado.get('excluidos', 0)} edital(is) excluído(s) com sucesso!", resultado
        else:
            return f"❌ Erro ao excluir editais: {resultado.get('error')}", resultado

    # Listar editais para identificar qual excluir
    editais_resultado = tool_listar_editais(user_id)
    if not editais_resultado.get("success"):
        return "Erro ao buscar seus editais.", editais_resultado

    editais = editais_resultado.get("editais", [])
    if not editais:
        return "Você não tem editais salvos para excluir.", {"success": False}

    # Tentar encontrar edital por número ou ID na mensagem
    edital_a_excluir = None
    for ed in editais:
        numero = ed.get("numero", "").lower()
        edital_id = ed.get("id", "")

        if numero and numero in msg_lower:
            edital_a_excluir = ed
            break
        if edital_id[:8].lower() in msg_lower:
            edital_a_excluir = ed
            break

    if not edital_a_excluir:
        # Mostrar lista de editais para o usuário escolher
        response = "**Qual edital você deseja excluir?**\n\nEditais salvos:\n"
        for i, ed in enumerate(editais[:10], 1):
            response += f"{i}. **{ed.get('numero')}** - {ed.get('orgao', 'N/A')[:40]}\n"
        response += "\nDigite: 'excluir edital [número]' para confirmar."
        return response, {"success": False, "editais": editais}

    # Excluir o edital encontrado
    resultado = tool_excluir_edital(edital_a_excluir["id"], user_id)

    if resultado.get("success"):
        return f"✅ Edital **{edital_a_excluir.get('numero')}** excluído com sucesso!", resultado
    else:
        return f"❌ Erro ao excluir edital: {resultado.get('error')}", resultado


def processar_excluir_produto(message: str, user_id: str):
    """
    Processa ação: Excluir produto.
    Identifica produto por nome ou ID na mensagem.
    """
    from tools import tool_excluir_produto, tool_listar_produtos

    msg_lower = message.lower()

    # Listar produtos para identificar qual excluir
    produtos_resultado = tool_listar_produtos(user_id)
    if not produtos_resultado.get("success"):
        return "Erro ao buscar seus produtos.", produtos_resultado

    produtos = produtos_resultado.get("produtos", [])
    if not produtos:
        return "Você não tem produtos cadastrados para excluir.", {"success": False}

    # Verificar se é exclusão de todos
    if "todos" in msg_lower:
        excluidos = 0
        erros = 0
        for p in produtos:
            resultado = tool_excluir_produto(p["id"], user_id)
            if resultado.get("success"):
                excluidos += 1
            else:
                erros += 1
        return f"✅ {excluidos} produto(s) excluído(s)!" + (f" ({erros} erros)" if erros else ""), {"success": True, "excluidos": excluidos}

    # Tentar encontrar produto por nome na mensagem
    produto_a_excluir = None
    for p in produtos:
        nome = p.get("nome", "").lower()
        modelo = (p.get("modelo") or "").lower()
        produto_id = p.get("id", "")

        # Verificar se nome, modelo ou ID está na mensagem
        if nome and any(parte in msg_lower for parte in nome.split()[:3]):
            produto_a_excluir = p
            break
        if modelo and modelo in msg_lower:
            produto_a_excluir = p
            break
        if produto_id[:8].lower() in msg_lower:
            produto_a_excluir = p
            break

    if not produto_a_excluir:
        # Mostrar lista de produtos para o usuário escolher
        response = "**Qual produto você deseja excluir?**\n\nProdutos cadastrados:\n"
        for i, p in enumerate(produtos[:10], 1):
            response += f"{i}. **{p.get('nome')}** ({p.get('fabricante', 'N/A')})\n"
        response += "\nDigite: 'excluir produto [nome]' para confirmar."
        return response, {"success": False, "produtos": produtos}

    # Excluir o produto encontrado
    resultado = tool_excluir_produto(produto_a_excluir["id"], user_id)

    if resultado.get("success"):
        return f"✅ Produto **{produto_a_excluir.get('nome')}** excluído com sucesso!", resultado
    else:
        return f"❌ Erro ao excluir produto: {resultado.get('error')}", resultado


def processar_atualizar_edital(message: str, user_id: str):
    """
    Processa ação: Atualizar/Editar edital.
    Usa IA para extrair o que o usuário quer alterar.
    """
    from tools import tool_atualizar_edital, tool_listar_editais

    # Listar editais para identificar qual atualizar
    editais_resultado = tool_listar_editais(user_id)
    if not editais_resultado.get("success"):
        return "Erro ao buscar seus editais.", editais_resultado

    editais = editais_resultado.get("editais", [])
    if not editais:
        return "Você não tem editais salvos para editar.", {"success": False}

    msg_lower = message.lower()

    # Tentar encontrar edital por número na mensagem
    edital_a_editar = None
    for ed in editais:
        numero = ed.get("numero", "").lower()
        if numero and numero in msg_lower:
            edital_a_editar = ed
            break

    if not edital_a_editar:
        # Usar o último edital
        edital_a_editar = editais[0]

    # Extrair campos a atualizar usando IA
    prompt = f"""Analise a mensagem do usuário e extraia quais campos do edital ele quer alterar.

Mensagem: "{message}"

Campos possíveis: numero, orgao, objeto, modalidade, status, valor_referencia, data_abertura, url

Status possíveis: novo, analisando, participar, nao_participar, proposta_enviada, ganho, perdido, cancelado
Modalidades: pregao_eletronico, pregao_presencial, concorrencia, tomada_precos, convite, dispensa, inexigibilidade

IMPORTANTE: Se a mensagem contém uma URL (http:// ou https://), extraia como campo "url".

Retorne JSON com apenas os campos a alterar:
{{"campo1": "novo_valor", "campo2": "novo_valor"}}

Se não identificar campos claros, retorne {{}}
"""

    try:
        resposta_ia = call_deepseek([{"role": "user", "content": prompt}], max_tokens=100, model_override="deepseek-chat")
        import json
        import re
        json_match = re.search(r'\{[\s\S]*?\}', resposta_ia)
        if json_match:
            campos = json.loads(json_match.group())
        else:
            campos = {}
    except:
        campos = {}

    if not campos:
        # Mostrar edital atual e pedir para especificar
        response = f"""**Editar Edital: {edital_a_editar.get('numero')}**

Dados atuais:
- **Número:** {edital_a_editar.get('numero')}
- **Órgão:** {edital_a_editar.get('orgao')}
- **Status:** {edital_a_editar.get('status')}
- **Modalidade:** {edital_a_editar.get('modalidade')}
- **URL:** {edital_a_editar.get('url') or 'Não cadastrada'}

Por favor, especifique o que deseja alterar. Exemplos:
- "alterar status para participar"
- "mudar órgão para Prefeitura de SP"
- "atualizar URL para https://exemplo.com/edital.pdf"
"""
        return response, {"success": False, "edital": edital_a_editar}

    # Aplicar atualizações
    resultado = tool_atualizar_edital(
        edital_id=edital_a_editar["id"],
        user_id=user_id,
        **campos
    )

    if resultado.get("success"):
        edital_atualizado = resultado.get("edital", {})
        response = f"""✅ Edital **{edital_atualizado.get('numero')}** atualizado!

Novos dados:
- **Número:** {edital_atualizado.get('numero')}
- **Órgão:** {edital_atualizado.get('orgao')}
- **Status:** {edital_atualizado.get('status')}
- **Modalidade:** {edital_atualizado.get('modalidade')}
"""
        return response, resultado
    else:
        return f"❌ Erro ao atualizar edital: {resultado.get('error')}", resultado


def processar_cadastrar_produto(message: str, user_id: str):
    """
    Cadastra produto manualmente a partir dos dados enviados pelo formulário do Portfolio.
    Formato: Cadastre manualmente o produto: Nome="X", Fabricante="Y", Modelo="Z",
             Categoria="W", SubclasseId="uuid", NCM="N". Especificacoes: campo=valor[unidade]{tipo}, ...
    """
    import re
    from decimal import Decimal, InvalidOperation
    db = get_db()
    try:
        empresa_id = get_current_empresa_id()

        # Parsear campos do prompt
        def extrair(campo):
            m = re.search(rf'{campo}="([^"]*)"', message)
            return m.group(1).strip() if m else None

        nome = extrair("Nome")
        fabricante = extrair("Fabricante")
        modelo = extrair("Modelo")
        categoria = extrair("Categoria")
        subclasse_id = extrair("SubclasseId")
        ncm = extrair("NCM")

        if not nome:
            return "Informe ao menos o nome do produto. Ex: Cadastre o produto: Nome=\"Analisador X\"", {"error": "nome_obrigatorio"}

        # Validar categoria
        categorias_validas = ['equipamento', 'reagente', 'insumo_hospitalar', 'insumo_laboratorial',
                             'informatica', 'redes', 'mobiliario', 'eletronico', 'comodato', 'outro']
        if not categoria or categoria not in categorias_validas:
            categoria = "outro"

        # Validar subclasse_id se fornecido
        if subclasse_id:
            sub_exists = db.query(SubclasseProduto).filter(SubclasseProduto.id == subclasse_id).first()
            if not sub_exists:
                subclasse_id = None  # Ignorar ID inválido

        # Verificar duplicidade
        existente = db.query(Produto).filter(
            Produto.user_id == user_id,
            Produto.nome == nome
        ).first()
        if existente:
            return f"Produto **{nome}** ja existe (ID: {existente.id}). Use 'atualizar produto' para modificar.", {"error": "duplicado", "produto_id": existente.id}

        # Criar produto
        produto = Produto(
            user_id=user_id,
            empresa_id=empresa_id,
            nome=nome,
            fabricante=fabricante,
            modelo=modelo,
            categoria=categoria,
            subclasse_id=subclasse_id,
            ncm=ncm,
            status_pipeline="cadastrado",
        )
        db.add(produto)
        db.flush()
        produto_id = produto.id

        # Parsear especificações com formato enriquecido: campo=valor[unidade]{tipo}
        specs_criadas = []
        specs_match = re.search(r'Especificacoes:\s*(.+)$', message)
        if specs_match:
            specs_str = specs_match.group(1)
            for par in specs_str.split(","):
                par = par.strip()
                if "=" in par:
                    campo, resto = par.split("=", 1)
                    campo = campo.strip()
                    resto = resto.strip()

                    # Extrair unidade [xxx] e tipo {xxx} do resto
                    unidade = None
                    tipo = "texto"
                    un_match = re.search(r'\[([^\]]+)\]', resto)
                    if un_match:
                        unidade = un_match.group(1)
                        resto = resto[:un_match.start()] + resto[un_match.end():]
                    tp_match = re.search(r'\{([^\}]+)\}', resto)
                    if tp_match:
                        tipo = tp_match.group(1)
                        resto = resto[:tp_match.start()] + resto[tp_match.end():]

                    valor = resto.strip()
                    if campo and valor:
                        # Converter valor numérico se tipo é numero/decimal
                        valor_numerico = None
                        if tipo in ("numero", "decimal"):
                            try:
                                valor_numerico = Decimal(valor.replace(",", "."))
                            except (InvalidOperation, ValueError):
                                pass

                        spec = ProdutoEspecificacao(
                            produto_id=produto_id,
                            nome_especificacao=campo,
                            valor=valor,
                            unidade=unidade,
                            valor_numerico=valor_numerico,
                        )
                        db.add(spec)
                        specs_criadas.append(f"{campo}: {valor}{f' {unidade}' if unidade else ''}")

        db.commit()

        specs_text = ""
        if specs_criadas:
            specs_list = "\n".join([f"  - {s}" for s in specs_criadas])
            specs_text = f"\n\n**Especificacoes cadastradas ({len(specs_criadas)}):**\n{specs_list}"

        subclasse_text = ""
        if subclasse_id:
            sub = db.query(SubclasseProduto).filter(SubclasseProduto.id == subclasse_id).first()
            subclasse_text = f"\n- **Subclasse:** {sub.nome if sub else subclasse_id}"

        response = f"""## Produto cadastrado com sucesso

- **Nome:** {nome}
- **Categoria:** {categoria}{subclasse_text}
- **Fabricante:** {fabricante or 'N/I'}
- **Modelo:** {modelo or 'N/I'}
- **NCM:** {ncm or 'N/I'}
- **Status:** cadastrado{specs_text}

O produto foi salvo no banco. Voce pode ver na aba "Meus Produtos" do Portfolio."""

        return response, {"success": True, "produto_id": produto_id, "nome": nome}

    except Exception as e:
        db.rollback()
        print(f"[ERRO] processar_cadastrar_produto: {e}")
        return f"Erro ao cadastrar produto: {str(e)}", {"error": str(e)}
    finally:
        db.close()


def processar_buscar_anvisa(message: str, user_id: str):
    """
    Busca registros de produtos na ANVISA via web search.
    Aceita número de registro ou nome do produto.
    Atualiza campo registro_anvisa e anvisa_status do produto se encontrar.
    """
    import re
    db = get_db()
    try:
        # Extrair número de registro ou nome do produto
        registro_match = re.search(r'(?:numero|nº|n°|registro)\s*[:\s]*(\d{8,15})', message, re.IGNORECASE)
        registro = registro_match.group(1) if registro_match else None

        # Extrair nome do produto
        nome_match = re.search(r'(?:produto|para o|para)\s+(.+?)(?:\s+na\s+anvisa|\s*$)', message, re.IGNORECASE)
        nome_produto = nome_match.group(1).strip() if nome_match else None

        if not registro and not nome_produto:
            # Tentar extrair qualquer número grande
            num_match = re.search(r'\b(\d{8,15})\b', message)
            if num_match:
                registro = num_match.group(1)
            else:
                # Extrair o que sobra após remover palavras de comando
                cleaned = re.sub(r'(?:busque?|registros?|anvisa|consulte?|do|da|de|o|a|na|no|para)\s*', '', message, flags=re.IGNORECASE).strip()
                if cleaned:
                    nome_produto = cleaned

        if not registro and not nome_produto:
            return "Informe o número do registro ANVISA ou o nome do produto para buscar.", {"error": "dados_insuficientes"}

        # Buscar via _web_search direto (sem filetype:pdf que tool_web_search adiciona)
        if registro:
            search_label = f"registro **{registro}**"
        else:
            search_label = f"produto **{nome_produto}**"

        # Query 1: busca direta com ANVISA
        query1 = f"{registro or nome_produto} registro ANVISA produto saude"
        print(f"[ANVISA] Buscando: {query1}")
        results = _web_search(query1, num_results=5)

        # Query 2: complementar
        query2 = f"ANVISA consulta produtos {registro or nome_produto} situação registro"
        print(f"[ANVISA] Buscando: {query2}")
        results2 = _web_search(query2, num_results=5)

        # Deduplicar por link
        seen_links = {r.get("link") for r in results}
        for r in results2:
            if r.get("link") not in seen_links:
                results.append(r)
                seen_links.add(r.get("link"))

        if not results:
            return f"Nenhum resultado encontrado na ANVISA para {search_label}. Tente com outro termo ou verifique o número do registro.", {"error": "sem_resultados"}

        # Tentar extrair número de registro dos resultados
        import re as re2
        registro_encontrado = registro
        if not registro_encontrado:
            for r in results:
                title_snippet = f"{r.get('title', '')} {r.get('snippet', '')}"
                match = re2.search(r'(?:registro|nº|n°)\s*(?:ANVISA\s*)?[:\s]*(\d{8,15})', title_snippet, re2.IGNORECASE)
                if not match:
                    match = re2.search(r'\b(\d{11,15})\b', title_snippet)
                if match:
                    registro_encontrado = match.group(1)
                    break

        # Tentar atualizar produto no banco
        produto_atualizado = None
        if registro_encontrado and nome_produto:
            produto = db.query(Produto).filter(
                Produto.user_id == user_id,
                Produto.nome.ilike(f"%{nome_produto}%")
            ).first()
            if produto:
                produto.registro_anvisa = registro_encontrado
                produto.anvisa_status = "ativo"
                db.commit()
                produto_atualizado = produto.nome
        elif registro_encontrado:
            produto = db.query(Produto).filter(
                Produto.user_id == user_id,
                Produto.registro_anvisa == None  # noqa: E711
            ).first()
            if produto:
                produto.registro_anvisa = registro_encontrado
                produto.anvisa_status = "ativo"
                db.commit()
                produto_atualizado = produto.nome

        # Formatar resposta
        links_text = ""
        for i, r in enumerate(results[:8]):
            title = r.get("title", "Sem título")
            url = r.get("link", "")
            snippet = r.get("snippet", "")[:120]
            links_text += f"\n{i+1}. [{title}]({url})\n   {snippet}\n"

        update_text = ""
        if produto_atualizado:
            update_text = f"\n\n✅ **Produto atualizado:** {produto_atualizado} → Registro ANVISA: {registro_encontrado}, Status: Ativo"
        elif registro_encontrado:
            update_text = f"\n\n📋 **Registro encontrado:** {registro_encontrado}"

        response = f"""## Resultados ANVISA para {search_label}

{links_text}
{update_text}

**Dica:** Acesse os links acima para verificar a situação do registro no portal da ANVISA."""

        return response, {"success": True, "results_count": len(results), "registro": registro_encontrado}

    except Exception as e:
        print(f"[ERRO] processar_buscar_anvisa: {e}")
        return f"Erro ao buscar ANVISA: {str(e)}", {"error": str(e)}
    finally:
        db.close()


def processar_atualizar_produto(message: str, user_id: str):
    """
    Processa ação: Atualizar/Editar produto.
    Usa IA para extrair o que o usuário quer alterar.
    """
    from tools import tool_atualizar_produto, tool_listar_produtos

    # Listar produtos para identificar qual atualizar
    produtos_resultado = tool_listar_produtos(user_id)
    if not produtos_resultado.get("success"):
        return "Erro ao buscar seus produtos.", produtos_resultado

    produtos = produtos_resultado.get("produtos", [])
    if not produtos:
        return "Você não tem produtos cadastrados para editar.", {"success": False}

    msg_lower = message.lower()

    # Tentar encontrar produto por nome na mensagem
    produto_a_editar = None
    for p in produtos:
        nome = p.get("nome", "").lower()
        if nome and any(parte in msg_lower for parte in nome.split()[:3]):
            produto_a_editar = p
            break

    if not produto_a_editar:
        # Usar o último produto
        produto_a_editar = produtos[-1]

    # Extrair campos a atualizar usando IA
    prompt = f"""Analise a mensagem do usuário e extraia quais campos do produto ele quer alterar.

Mensagem: "{message}"

Campos possíveis: nome, fabricante, modelo, categoria

Categorias: equipamento, reagente, insumo_hospitalar, insumo_laboratorial, informatica, redes, mobiliario, eletronico, outro

Retorne JSON com apenas os campos a alterar:
{{"campo1": "novo_valor", "campo2": "novo_valor"}}

Se não identificar campos claros, retorne {{}}
"""

    try:
        resposta_ia = call_deepseek([{"role": "user", "content": prompt}], max_tokens=100, model_override="deepseek-chat")
        import json
        import re
        json_match = re.search(r'\{[\s\S]*?\}', resposta_ia)
        if json_match:
            campos = json.loads(json_match.group())
        else:
            campos = {}
    except:
        campos = {}

    if not campos:
        # Mostrar produto atual e pedir para especificar
        response = f"""**Editar Produto: {produto_a_editar.get('nome')}**

Dados atuais:
- **Nome:** {produto_a_editar.get('nome')}
- **Fabricante:** {produto_a_editar.get('fabricante', 'N/A')}
- **Modelo:** {produto_a_editar.get('modelo', 'N/A')}
- **Categoria:** {produto_a_editar.get('categoria', 'N/A')}

Por favor, especifique o que deseja alterar. Exemplos:
- "alterar fabricante para Philips"
- "mudar categoria para equipamento"
"""
        return response, {"success": False, "produto": produto_a_editar}

    # Aplicar atualizações
    resultado = tool_atualizar_produto(
        produto_id=produto_a_editar["id"],
        user_id=user_id,
        **campos
    )

    if resultado.get("success"):
        produto_atualizado = resultado.get("produto", {})
        response = f"""✅ Produto **{produto_atualizado.get('nome')}** atualizado!

Novos dados:
- **Nome:** {produto_atualizado.get('nome')}
- **Fabricante:** {produto_atualizado.get('fabricante', 'N/A')}
- **Modelo:** {produto_atualizado.get('modelo', 'N/A')}
- **Categoria:** {produto_atualizado.get('categoria', 'N/A')}
"""
        return response, resultado
    else:
        return f"❌ Erro ao atualizar produto: {resultado.get('error')}", resultado


def processar_listar_fontes(message: str):
    """Processa ação: Listar fontes de editais"""
    resultado = tool_listar_fontes()

    if resultado.get("success"):
        fontes = resultado.get("fontes", [])
        if fontes:
            response = f"**Fontes de editais cadastradas:** {len(fontes)}\n\n"
            for f in fontes:
                status = "✅ Ativa" if f.get("ativo") else "❌ Inativa"
                response += f"- **{f['nome']}** ({f['tipo']}) {status}\n"
                response += f"  URL: {f.get('url_base', 'N/A')}\n"
                if f.get('descricao'):
                    response += f"  {f['descricao'][:100]}\n"
                response += "\n"
        else:
            response = "Nenhuma fonte de editais cadastrada."
    else:
        response = f"Erro ao listar fontes: {resultado.get('error')}"

    return response, resultado


def processar_listar_propostas(message: str, user_id: str):
    """Processa ação: Listar propostas geradas pelo usuário"""
    db = get_db()
    try:
        propostas = db.query(Proposta).filter(
            Proposta.user_id == user_id
        ).order_by(Proposta.created_at.desc()).limit(20).all()

        if propostas:
            response = f"## 📝 Minhas Propostas ({len(propostas)})\n\n"
            for p in propostas:
                # Buscar edital e produto relacionados
                edital = db.query(Edital).filter(Edital.id == p.edital_id).first()
                produto = db.query(Produto).filter(Produto.id == p.produto_id).first()

                edital_num = edital.numero if edital else "N/A"
                produto_nome = produto.nome[:40] if produto else "N/A"
                preco = f"R$ {p.preco_total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") if p.preco_total else "N/A"
                data = p.created_at.strftime("%d/%m/%Y %H:%M") if p.created_at else "N/A"

                status_emoji = {
                    "rascunho": "📋",
                    "enviada": "📤",
                    "aceita": "✅",
                    "rejeitada": "❌"
                }.get(p.status, "📋")

                response += f"### {status_emoji} Proposta - Edital {edital_num}\n"
                response += f"- **Produto:** {produto_nome}\n"
                response += f"- **Preço:** {preco}\n"
                response += f"- **Status:** {p.status}\n"
                response += f"- **Data:** {data}\n"
                response += f"- **ID:** `{p.id[:8]}...`\n\n"

            resultado = {"success": True, "propostas": [p.to_dict() for p in propostas], "total": len(propostas)}
        else:
            response = "Você ainda não tem propostas geradas.\n\nPara gerar uma proposta, use:\n`Gere uma proposta do produto [NOME] para o edital [NUMERO] com preço R$ [VALOR]`"
            resultado = {"success": True, "propostas": [], "total": 0}

        return response, resultado

    except Exception as e:
        return f"Erro ao listar propostas: {str(e)}", {"success": False, "error": str(e)}
    finally:
        db.close()


def processar_buscar_editais_score(message: str, user_id: str):
    """Processa ação: Buscar editais + calcular score"""
    # Primeiro buscar editais (sem score — score é calculado abaixo separadamente)
    response_busca, resultado_busca = processar_buscar_editais(message, user_id, calcular_score=False)

    if not resultado_busca.get("success"):
        return response_busca, resultado_busca

    # Depois calcular score para cada edital com os produtos do usuário
    produtos = tool_listar_produtos(user_id)

    if not produtos.get("produtos"):
        return response_busca + "\n\n⚠️ Você não tem produtos cadastrados para calcular aderência.", resultado_busca

    response = response_busca + "\n\n**Análise de Aderência:**\n"
    analises = []

    for edital in resultado_busca.get("editais", [])[:3]:
        for produto in produtos.get("produtos", [])[:2]:
            analise = tool_calcular_aderencia(produto["id"], edital["id"], user_id)
            if analise.get("success"):
                analises.append(analise)
                response += f"\n- {produto['nome']} x {edital['numero']}: **{analise.get('score_tecnico', 0):.0f}%** - {analise.get('recomendacao', '')}"

    resultado_busca["analises"] = analises
    return response, resultado_busca


def processar_listar_editais(message: str, user_id: str):
    """Processa ação: Listar editais salvos"""
    # Extrair filtros da mensagem
    uf = None
    status = None

    message_lower = message.lower()
    if " sp" in message_lower or "são paulo" in message_lower:
        uf = "SP"
    elif " rj" in message_lower or "rio de janeiro" in message_lower:
        uf = "RJ"
    elif " mg" in message_lower or "minas gerais" in message_lower:
        uf = "MG"

    if "novo" in message_lower:
        status = "novo"
    elif "analisando" in message_lower:
        status = "analisando"

    # Verificar se usuário quer ver todos
    mostrar_todos = any(p in message_lower for p in ["todos", "all", "completo", "completa"])
    limite = 100 if mostrar_todos else 20  # Default 20, ou 100 se pedir todos

    resultado = tool_listar_editais(user_id, status=status, uf=uf)

    if resultado.get("success"):
        editais = resultado.get("editais", [])
        if editais:
            total = len(editais)
            editais_mostrar = editais[:limite]

            response = f"**Editais salvos:** {total}"
            if total > limite:
                response += f" (mostrando {limite})"
            response += "\n\n"

            for i, ed in enumerate(editais_mostrar, 1):
                fonte = ed.get('fonte', '')
                fonte_badge = ""
                if 'PNCP' in fonte:
                    fonte_badge = "🟢"
                elif 'ComprasNet' in fonte:
                    fonte_badge = "🔵"
                elif 'BEC' in fonte:
                    fonte_badge = "🟡"
                else:
                    fonte_badge = "⚪"

                response += f"{i}. **{ed['numero']}** ({ed['status']}) {fonte_badge} {fonte}\n"
                response += f"   {ed['orgao']} - {ed['uf'] or 'N/A'}\n"
                response += f"   {ed['objeto'][:80]}...\n"

                # Botões de ação
                botoes = []
                if ed.get('url'):
                    botoes.append(f"[🔗 Portal]({ed['url']})")

                # PDF - verificar se tem pdf_url ou dados PNCP
                pdf_url = ed.get('pdf_url')
                cnpj = ed.get('cnpj_orgao')
                ano = ed.get('ano_compra')
                seq = ed.get('seq_compra')

                if pdf_url:
                    botoes.append(f"[📄 Ver PDF]({pdf_url})")
                    botoes.append(f"[⬇️ Baixar]({pdf_url}?download=true)")
                elif cnpj and ano and seq:
                    pdf_api = f"https://pncp.gov.br/api/pncp/v1/orgaos/{cnpj}/compras/{ano}/{seq}/arquivos/1"
                    botoes.append(f"[📄 Ver PDF]({pdf_api})")
                    botoes.append(f"[⬇️ Baixar]({pdf_api})")

                if botoes:
                    response += f"   {' | '.join(botoes)}\n"
                response += "\n"

            if total > limite:
                response += f"\n📋 *Mostrando {limite} de {total} editais. Digite 'listar todos editais' para ver todos.*"
        else:
            response = "Você não tem editais salvos ainda. Use 'Buscar editais' para encontrar oportunidades."
    else:
        response = f"Erro ao listar: {resultado.get('error')}"

    return response, resultado


def _encontrar_produto(produtos: list, message_lower: str):
    """Helper para encontrar produto por nome, modelo ou palavras-chave"""
    for p in produtos:
        nome_lower = (p.get("nome") or "").lower()
        modelo_lower = (p.get("modelo") or "").lower()
        fabricante_lower = (p.get("fabricante") or "").lower()

        # Match exato do nome
        if nome_lower in message_lower:
            return p
        # Match por modelo completo
        if modelo_lower and modelo_lower in message_lower:
            return p
        # Match por parte principal do modelo (primeira palavra significativa)
        if modelo_lower:
            modelo_parts = modelo_lower.split()
            for part in modelo_parts:
                if len(part) >= 5 and part in message_lower:
                    return p
        # Match por fabricante + qualquer parte do modelo
        if fabricante_lower in message_lower:
            if modelo_lower:
                for part in modelo_lower.split():
                    if len(part) >= 3 and part in message_lower:
                        return p
        # Match parcial: qualquer palavra significativa do nome (>5 chars)
        palavras = nome_lower.split()
        for palavra in palavras:
            if len(palavra) > 5 and palavra in message_lower:
                return p
    return None


def processar_calcular_aderencia(message: str, user_id: str):
    """Processa ação: Calcular aderência"""
    # Listar produtos e editais disponíveis
    produtos = tool_listar_produtos(user_id)
    editais = tool_listar_editais(user_id)

    if not produtos.get("produtos"):
        return "Você não tem produtos cadastrados. Faça upload de um manual primeiro.", {"status": "sem_produtos"}

    if not editais.get("editais"):
        return "Você não tem editais salvos. Busque editais primeiro.", {"status": "sem_editais"}

    # Tentar identificar produto e edital na mensagem
    produto_encontrado = None
    edital_encontrado = None

    message_lower = message.lower()

    # Buscar produto - várias estratégias de match
    produto_encontrado = _encontrar_produto(produtos.get("produtos", []), message_lower)

    for e in editais.get("editais", []):
        numero = e.get("numero") or ""
        if numero and numero.lower() in message_lower:
            edital_encontrado = e
            break

    if produto_encontrado and edital_encontrado:
        resultado = tool_calcular_aderencia(
            produto_encontrado["id"],
            edital_encontrado["id"],
            user_id
        )

        if resultado.get("success"):
            response = f"""## Análise de Aderência

**Produto:** {resultado.get('produto')}
**Edital:** {resultado.get('edital')}

### Score Técnico: {resultado.get('score_tecnico', 0):.1f}%

"""
            # Se tem requisitos cadastrados
            if resultado.get('requisitos_total', 0) > 0:
                response += f"""**Requisitos:**
- Total: {resultado.get('requisitos_total', 0)}
- Atendidos: {resultado.get('requisitos_atendidos', 0)}
- Parciais: {resultado.get('requisitos_parciais', 0)}
- Não atendidos: {resultado.get('requisitos_nao_atendidos', 0)}

"""
            # Justificativa (da análise via IA)
            if resultado.get('justificativa'):
                response += f"""**Análise:** {resultado.get('justificativa')}

"""
            # Recomendação com emoji
            recomendacao = resultado.get('recomendacao', '')
            if 'RECOMENDADO' in recomendacao and 'NAO' not in recomendacao:
                emoji = "✅"
            elif 'AVALIAR' in recomendacao:
                emoji = "⚠️"
            else:
                emoji = "❌"

            response += f"""### {emoji} Recomendação: {recomendacao}
"""
            return response, resultado

    # Se não identificou, mostrar opções
    response = "Para calcular aderência, informe o produto e o edital.\n\n"
    response += "**Seus produtos:**\n"
    for p in produtos.get("produtos", [])[:5]:
        response += f"- {p['nome']}\n"
    response += "\n**Seus editais:**\n"
    for e in editais.get("editais", [])[:5]:
        response += f"- {e['numero']} ({e['orgao']})\n"
    response += "\nExemplo: 'Analise o Mindray BS-240 para o edital PE-2024-001'"

    return response, {"status": "aguardando_selecao", "produtos": produtos.get("produtos"), "editais": editais.get("editais")}


def processar_gerar_proposta(message: str, user_id: str):
    """Processa ação: Gerar proposta"""
    # Similar ao calcular aderência, precisa identificar produto e edital
    produtos = tool_listar_produtos(user_id)
    editais = tool_listar_editais(user_id)

    if not produtos.get("produtos") or not editais.get("editais"):
        return "Você precisa ter produtos e editais cadastrados para gerar uma proposta.", {"status": "incompleto"}

    # Tentar identificar e extrair preço
    produto_encontrado = None
    edital_encontrado = None
    preco = None

    message_lower = message.lower()

    # Usar helper para encontrar produto
    produto_encontrado = _encontrar_produto(produtos.get("produtos", []), message_lower)

    for e in editais.get("editais", []):
        numero = e.get("numero") or ""
        if numero and numero.lower() in message_lower:
            edital_encontrado = e
            break

    # Extrair preço - buscar padrão "R$ X" ou "preço X" ou "valor X"
    import re
    # Primeiro tenta R$ seguido de número
    preco_match = re.search(r'R\$\s*([\d.,]+)', message)
    if not preco_match:
        # Tenta "preço" ou "valor" seguido de número
        preco_match = re.search(r'(?:preço|preco|valor)\s*(?:de\s*)?R?\$?\s*([\d.,]+)', message, re.IGNORECASE)
    if preco_match:
        try:
            valor_str = preco_match.group(1)
            # Remove pontos de milhar e converte vírgula decimal
            preco = float(valor_str.replace('.', '').replace(',', '.'))
        except:
            pass

    if produto_encontrado and edital_encontrado:
        resultado = tool_gerar_proposta(
            edital_encontrado["id"],
            produto_encontrado["id"],
            user_id,
            preco=preco
        )

        if resultado.get("success"):
            response = f"""**Proposta Gerada com Sucesso!**

**Edital:** {resultado.get('edital')}
**Produto:** {resultado.get('produto')}
**Status:** {resultado.get('status')}

---

{resultado.get('texto_proposta', '')}

---

*Proposta salva com ID: {resultado.get('proposta_id')}*"""
            return response, resultado

    # Se não identificou, mostrar opções
    response = "Para gerar proposta, informe:\n- Produto\n- Edital\n- Preço (opcional)\n\n"
    response += "Exemplo: 'Gere proposta do Mindray BS-240 para edital PE-2024-001 com preço R$ 50.000'"

    return response, {"status": "aguardando_dados"}


def processar_salvar_editais(message: str, user_id: str, session_id: str, db):
    """
    Processa ação: Salvar editais

    Busca no histórico da sessão os editais da última busca e salva os recomendados
    ou os especificados pelo usuário.
    """
    import json
    import re

    msg_lower = message.lower()

    # Determinar o que salvar
    # - "salvar recomendados" ou "salvar editais recomendados" → PARTICIPAR + AVALIAR
    # - "salvar para participar" ou "salvar participar" → só PARTICIPAR
    # - "salvar todos" → todos os editais
    # - "salvar edital NUMERO" → edital específico

    # Primeiro verificar se tem número de edital específico na mensagem
    numero_especifico = None
    numero_match = re.search(r'edital\s+(\S+)', msg_lower)
    if numero_match:
        numero_especifico = numero_match.group(1).upper()
        # Limpar caracteres especiais do número
        numero_especifico = numero_especifico.strip('.,;:')

    # Determinar tipo de salvamento
    if numero_especifico:
        salvar_tipo = "especifico"
    elif "todos" in msg_lower:
        salvar_tipo = "todos"
    elif "participar" in msg_lower:
        salvar_tipo = "participar"
    elif "recomendados" in msg_lower or "recomendado" in msg_lower:
        salvar_tipo = "recomendados"
    else:
        salvar_tipo = "todos"  # Padrão para busca sem score

    # Buscar última mensagem de busca no histórico (com editais salvos em sources_json)
    # Aceita tanto buscar_editais (com score) quanto buscar_editais_simples (sem score)
    ultima_busca = db.query(Message).filter(
        Message.session_id == session_id,
        Message.action_type.in_(["buscar_editais", "buscar_editais_simples"]),
        Message.role == "assistant"
    ).order_by(Message.created_at.desc()).first()

    if not ultima_busca:
        return "Não encontrei uma busca de editais recente. Execute primeiro: **buscar editais de [tema]**", {"status": "sem_busca"}

    # Tentar recuperar editais do sources_json (salvo na busca)
    editais_para_salvar = []
    editais_com_score = []
    editais_participar = []
    editais_recomendados = []

    if ultima_busca.sources_json:
        # Recuperar editais salvos - SEM re-buscar!
        print(f"[SALVAR] Recuperando editais do sources_json...")
        sources = ultima_busca.sources_json
        editais_com_score = sources.get("editais_com_score", sources.get("editais", []))
        editais_participar = sources.get("editais_participar", [])
        editais_recomendados = sources.get("editais_recomendados", [])
        print(f"[SALVAR] Encontrados {len(editais_com_score)} editais salvos na sessão")
    else:
        # Fallback: sources_json vazio, precisa re-buscar (compatibilidade com buscas antigas)
        print(f"[SALVAR] sources_json vazio, buscando mensagem do usuário...")
        ultima_busca_user = db.query(Message).filter(
            Message.session_id == session_id,
            Message.action_type == "buscar_editais",
            Message.role == "user"
        ).order_by(Message.created_at.desc()).first()

        if ultima_busca_user:
            print(f"[SALVAR] Re-executando busca (fallback): {ultima_busca_user.content[:50]}...")
            classificacao = detectar_intencao_ia(ultima_busca_user.content, tem_arquivo=False)
            termo_ia = classificacao.get("termo_busca")
            _, resultado_busca = processar_buscar_editais(ultima_busca_user.content, user_id, termo_ia=termo_ia, calcular_score=False)

            if resultado_busca.get("success"):
                editais_com_score = resultado_busca.get("editais_com_score", [])
                editais_participar = resultado_busca.get("editais_participar", [])
                editais_recomendados = resultado_busca.get("editais_recomendados", [])

    if not editais_com_score:
        return "Não há editais para salvar. Execute uma busca primeiro: **buscar editais de [tema]**", {"status": "sem_editais"}

    print(f"[SALVAR] Tipo: {salvar_tipo}")
    print(f"[SALVAR] editais_com_score: {len(editais_com_score)}")
    print(f"[SALVAR] editais_participar: {len(editais_participar)}")
    print(f"[SALVAR] editais_recomendados: {len(editais_recomendados)}")

    if salvar_tipo == "especifico" and numero_especifico:
        # Salvar edital específico pelo número
        print(f"[SALVAR] Buscando edital específico: {numero_especifico}")
        for ed in editais_com_score:
            numero_edital = (ed.get("numero") or "").upper()
            # Tentar match exato ou parcial
            if numero_especifico in numero_edital or numero_edital in numero_especifico:
                editais_para_salvar.append(ed)
                print(f"[SALVAR] Encontrado edital: {numero_edital}")
                break
            # Tentar match só com números
            nums_busca = re.sub(r'[^\d]', '', numero_especifico)
            nums_edital = re.sub(r'[^\d]', '', numero_edital)
            if nums_busca and nums_edital and (nums_busca in nums_edital or nums_edital in nums_busca):
                editais_para_salvar.append(ed)
                print(f"[SALVAR] Encontrado edital (match numérico): {numero_edital}")
                break
    elif salvar_tipo == "todos":
        # Salvar TODOS os editais encontrados
        editais_para_salvar = editais_com_score
    elif salvar_tipo == "participar":
        # Salvar só os PARTICIPAR (score >= 80 ou recomendação PARTICIPAR)
        editais_para_salvar = editais_participar
        if not editais_para_salvar:
            # Fallback: pegar os com score >= 75 (margem para variação)
            editais_para_salvar = [e for e in editais_com_score if e.get("score_tecnico", 0) >= 75]
            print(f"[SALVAR] Fallback participar: {len(editais_para_salvar)} com score >= 75")
    elif salvar_tipo == "recomendados":
        # Salvar PARTICIPAR + AVALIAR
        editais_para_salvar = editais_recomendados
        if not editais_para_salvar:
            # Fallback: pegar os com score >= 50 ou todos se busca foi sem score
            editais_para_salvar = [e for e in editais_com_score if e.get("score_tecnico", 0) >= 50]
            if not editais_para_salvar:
                # Se ainda não tem (busca sem score), pegar todos
                editais_para_salvar = editais_com_score
            print(f"[SALVAR] Fallback recomendados: {len(editais_para_salvar)} editais")

    print(f"[SALVAR] editais_para_salvar: {len(editais_para_salvar)}")

    if not editais_para_salvar:
        return """Não encontrei editais para salvar.

**Opções:**
- Digite: **salvar editais recomendados** para salvar todos os recomendados
- Digite: **salvar edital [número]** para salvar um específico
- Execute uma nova busca: **buscar editais de [tema]**
""", {"status": "sem_editais"}

    # Salvar os editais selecionados (com verificação de duplicatas)
    resultado_salvar = tool_salvar_editais_selecionados(editais_para_salvar, user_id)

    if resultado_salvar.get("success"):
        salvos = resultado_salvar.get("salvos", [])
        duplicados = resultado_salvar.get("duplicados", [])
        erros = resultado_salvar.get("erros", [])
        incompletos = resultado_salvar.get("incompletos", [])

        response = "## 💾 Resultado do Salvamento\n\n"

        if salvos:
            response += f"**✅ Salvos com sucesso:** {len(salvos)} edital(is)\n"
            for num in salvos[:5]:
                response += f"- {num}\n"
            if len(salvos) > 5:
                response += f"- ... e mais {len(salvos) - 5}\n"
            response += "\n"

        if incompletos:
            response += f"**⚠️ Salvos com dados incompletos:** {len(incompletos)} edital(is)\n"
            response += "Estes editais não foram encontrados no PNCP e têm informações limitadas:\n"
            for num in incompletos[:3]:
                response += f"- {num}\n"
            response += "\n**Dica:** Para obter dados completos, acesse o link do edital e faça upload do PDF manualmente.\n\n"

        if duplicados:
            response += f"**⚠️ Já existentes (ignorados):** {len(duplicados)} edital(is)\n"
            for num in duplicados[:3]:
                response += f"- {num}\n"
            response += "\n"

        if erros:
            response += f"**❌ Erros:** {len(erros)}\n"
            for err in erros[:3]:
                response += f"- {err}\n"
            response += "\n"

        response += "Use **liste meus editais** para ver todos os editais salvos."

        return response, resultado_salvar
    else:
        return f"Erro ao salvar editais: {resultado_salvar.get('error')}", resultado_salvar


def formatar_resposta_tabular(resposta: str) -> str:
    """
    Melhora a formatação de respostas que contêm dados tabulares.
    Converte tabelas mal formatadas para markdown correto.
    """
    import re

    # Se já tem formato markdown de tabela bem formatada, preservar
    if re.search(r'\|[^|]+\|[^|]+\|', resposta) and '---' in resposta:
        return resposta

    linhas = resposta.strip().split("\n")

    # Detectar padrões de tabela (linha com múltiplas colunas separadas)
    # Padrão comum: "ID    Número    Órgão    Status"
    palavras_header = ["id", "número", "numero", "órgão", "orgao", "status", "valor", "data",
                       "nome", "objeto", "modalidade", "fonte", "url", "tipo"]

    for i, linha in enumerate(linhas):
        linha_lower = linha.lower()
        # Verifica se a linha parece um header de tabela
        matches = sum(1 for p in palavras_header if p in linha_lower)

        if matches >= 3:  # Pelo menos 3 palavras de header
            # Encontrou header, tentar converter para markdown
            # Detectar separador (múltiplos espaços ou tab)

            # Tentar separar por tabs primeiro
            if "\t" in linha:
                colunas_header = [c.strip() for c in linha.split("\t") if c.strip()]
            else:
                # Separar por múltiplos espaços (4+)
                colunas_header = [c.strip() for c in re.split(r'\s{4,}', linha) if c.strip()]

            if len(colunas_header) >= 3:
                # Montar tabela markdown
                md_linhas = []

                # Header
                md_linhas.append("| " + " | ".join(colunas_header) + " |")
                md_linhas.append("|" + "|".join("---" for _ in colunas_header) + "|")

                # Dados (linhas seguintes)
                for j in range(i + 1, len(linhas)):
                    data_linha = linhas[j].strip()
                    if not data_linha:
                        continue

                    if "\t" in data_linha:
                        colunas_data = [c.strip() for c in data_linha.split("\t")]
                    else:
                        colunas_data = [c.strip() for c in re.split(r'\s{4,}', data_linha)]

                    # Ajustar número de colunas
                    while len(colunas_data) < len(colunas_header):
                        colunas_data.append("")
                    colunas_data = colunas_data[:len(colunas_header)]

                    # Truncar valores muito longos
                    colunas_data = [c[:80] + "..." if len(c) > 80 else c for c in colunas_data]

                    md_linhas.append("| " + " | ".join(colunas_data) + " |")

                # Texto antes da tabela + tabela
                texto_antes = "\n".join(linhas[:i]).strip()
                tabela = "\n".join(md_linhas)

                if texto_antes:
                    return texto_antes + "\n\n" + tabela
                return tabela

    return resposta


def processar_consulta_mindsdb(message: str, user_id: str):
    """
    Processa consultas analíticas via MindsDB.
    Envia a pergunta em linguagem natural para o agente editais_database_searcher.
    """
    from tools import tool_consulta_mindsdb

    resultado = tool_consulta_mindsdb(message, user_id)

    if resultado.get("success"):
        resposta_mindsdb = resultado.get("resposta", "")

        # Melhorar formatação de tabelas
        resposta_formatada = formatar_resposta_tabular(resposta_mindsdb)

        response = f"""## 📊 Consulta Analítica

**Pergunta:** {message}

---

{resposta_formatada}

---
*Consulta realizada via MindsDB (GPT-4o)*"""
    else:
        error = resultado.get("error", "Erro desconhecido")
        response = f"""## ❌ Erro na Consulta

Não foi possível processar a consulta analítica.

**Erro:** {error}

**Dica:** Tente reformular a pergunta ou use comandos diretos como:
- "liste meus editais"
- "liste meus produtos"
- "calcule aderência do produto X ao edital Y"
"""

    return response, resultado


def processar_registrar_resultado(message: str, user_id: str):
    """
    Processa registro de resultado de certame (vitória/derrota).
    Alimenta a base de preços históricos e concorrentes.
    """
    from tools import tool_registrar_resultado

    resultado = tool_registrar_resultado(message, user_id)

    if not resultado.get("success"):
        error = resultado.get("error", "Erro desconhecido")
        response = f"""❌ **Erro ao registrar resultado**

{error}

**Exemplos de como registrar:**
- "Perdemos o PE-001/2026 para MedLab com R$ 365.000, nosso preço foi R$ 380.000"
- "Ganhamos o edital PE-002/2026 com R$ 290.000"
- "Perdemos o PE-003 por documentação"
- "O edital PE-004 foi cancelado"
"""
        return response, None

    # Montar resposta de sucesso
    emoji_resultado = "🏆" if resultado["resultado"] == "vitoria" else "📊"
    status_texto = {
        "vitoria": "VITÓRIA",
        "derrota": "DERROTA",
        "cancelado": "CANCELADO",
        "deserto": "DESERTO",
        "revogado": "REVOGADO"
    }.get(resultado["resultado"], resultado["resultado"].upper())

    response = f"""{emoji_resultado} **Resultado Registrado - {resultado['edital_numero']}**

**Órgão:** {resultado.get('orgao', 'N/A')}
**Resultado:** {status_texto}
"""

    # Tabela de preços se disponível
    if resultado.get("preco_vencedor") or resultado.get("nosso_preco"):
        response += "\n| Posição | Empresa | Preço |\n"
        response += "|---------|---------|-------|\n"

        if resultado.get("empresa_vencedora") and resultado["resultado"] != "vitoria":
            preco_venc = resultado["preco_vencedor"]
            preco_fmt = f"R$ {preco_venc:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") if preco_venc else "N/A"
            response += f"| 1º | {resultado['empresa_vencedora']} | {preco_fmt} |\n"

        if resultado.get("nosso_preco"):
            pos = "1º" if resultado["resultado"] == "vitoria" else "2º"
            nosso_preco = resultado["nosso_preco"]
            preco_fmt = f"R$ {nosso_preco:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") if nosso_preco else "N/A"
            response += f"| {pos} | Sua Empresa | {preco_fmt} |\n"

        response += "\n"

    # Análise se foi derrota por preço
    if resultado.get("diferenca") and resultado["resultado"] == "derrota":
        diferenca = resultado["diferenca"]
        diferenca_pct = resultado.get("diferenca_pct", 0)
        desconto = resultado.get("desconto_percentual")

        response += f"""**Análise:**
- Diferença: R$ {diferenca:,.2f} ({diferenca_pct:.1f}%)
"""
        if desconto:
            response += f"- Desconto do vencedor: {desconto:.1f}% sobre referência\n"

        if resultado.get("motivo"):
            motivo_texto = {
                "preco": "Preço",
                "tecnica": "Questão técnica",
                "documentacao": "Documentação",
                "prazo": "Prazo",
                "outro": "Outro"
            }.get(resultado["motivo"], resultado["motivo"])
            response += f"- Motivo principal: {motivo_texto}\n"

        response += f"""
💡 **Insight:** Para editais similares, considere preços ~{diferenca_pct:.0f}% menores.
"""

    # Mensagem de sucesso final
    response += """
✅ Dados salvos no histórico de preços e concorrentes!
"""

    return response, resultado


def processar_consultar_resultado(message: str, user_id: str):
    """
    Consulta resultado de um certame já registrado.
    Suporta consulta de um edital específico ou de todos os editais.
    """
    from models import get_db, Edital, PrecoHistorico, Concorrente, ParticipacaoEdital
    import re

    db = get_db()
    try:
        # Verificar se é consulta de TODOS os editais
        msg_lower = message.lower()
        consulta_todos = any(p in msg_lower for p in [
            "todos os editais", "todos editais", "resultados dos editais",
            "resultado dos editais", "todos os resultados", "listar resultados"
        ])

        if consulta_todos:
            return processar_consultar_todos_resultados(user_id, db)

        # Extrair número do edital da mensagem
        # Padrões: PE-001/2026, 90186, PE001, etc
        padrao = r'(?:PE[-\s]?)?(\d{2,6})(?:/\d{4})?'
        match = re.search(padrao, message, re.IGNORECASE)

        if not match:
            return "❌ Não identifiquei o número do edital. Informe o número (ex: PE-041/2026 ou 90186)\n\nPara ver todos os resultados, use: \"mostre os resultados de todos os editais\"", None

        numero_edital = match.group(0)

        # Buscar edital
        edital = db.query(Edital).filter(
            Edital.numero.ilike(f"%{numero_edital}%"),
            Edital.user_id == user_id
        ).first()

        if not edital:
            # Tentar busca mais flexível
            numero_limpo = match.group(1)
            edital = db.query(Edital).filter(
                Edital.numero.ilike(f"%{numero_limpo}%"),
                Edital.user_id == user_id
            ).first()

        if not edital:
            return f"❌ Edital '{numero_edital}' não encontrado no seu cadastro.", None

        # Buscar resultado registrado
        preco_hist = db.query(PrecoHistorico).filter(
            PrecoHistorico.edital_id == edital.id
        ).order_by(PrecoHistorico.data_registro.desc()).first()

        if not preco_hist:
            response = f"""📋 **Edital {edital.numero}**

**Órgão:** {edital.orgao}
**Status:** {edital.status or 'Não definido'}
**Valor Referência:** R$ {float(edital.valor_referencia):,.2f}

⚠️ **Nenhum resultado registrado ainda.**

Para registrar o resultado, use:
- "Perdemos o {edital.numero} para [EMPRESA] com R$ [VALOR]"
- "Ganhamos o {edital.numero} com R$ [VALOR]"
"""
            return response, None

        # Buscar participações
        participacoes = db.query(ParticipacaoEdital).filter(
            ParticipacaoEdital.edital_id == edital.id
        ).order_by(ParticipacaoEdital.posicao_final).all()

        # Montar resposta
        emoji_resultado = "🏆" if preco_hist.resultado == "vitoria" else "📊"
        status_texto = {
            "vitoria": "VITÓRIA",
            "derrota": "DERROTA",
            "cancelado": "CANCELADO",
            "deserto": "DESERTO",
            "revogado": "REVOGADO"
        }.get(preco_hist.resultado, preco_hist.resultado.upper() if preco_hist.resultado else "N/A")

        response = f"""{emoji_resultado} **Resultado do Edital {edital.numero}**

**Órgão:** {edital.orgao}
**Resultado:** {status_texto}
**Data:** {preco_hist.data_homologacao.strftime('%d/%m/%Y') if preco_hist.data_homologacao else 'N/A'}
"""

        # Tabela de participantes
        if participacoes:
            response += "\n**Participantes:**\n"
            response += "| Pos | Empresa | Preço |\n"
            response += "|-----|---------|-------|\n"

            for part in participacoes:
                if part.concorrente_id:
                    conc = db.query(Concorrente).get(part.concorrente_id)
                    nome = conc.nome if conc else "Desconhecido"
                else:
                    nome = "Sua Empresa"

                preco_fmt = f"R$ {float(part.preco_proposto):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") if part.preco_proposto else "N/A"
                pos = f"{part.posicao_final}º" if part.posicao_final else "-"
                response += f"| {pos} | {nome} | {preco_fmt} |\n"

            response += "\n"

        # Análise
        if preco_hist.resultado == "derrota" and preco_hist.nosso_preco and preco_hist.preco_vencedor:
            diferenca = float(preco_hist.nosso_preco) - float(preco_hist.preco_vencedor)
            diferenca_pct = (diferenca / float(preco_hist.nosso_preco)) * 100

            response += f"""**Análise:**
- Nosso preço: R$ {float(preco_hist.nosso_preco):,.2f}
- Preço vencedor: R$ {float(preco_hist.preco_vencedor):,.2f}
- Diferença: R$ {diferenca:,.2f} ({diferenca_pct:.1f}%)
"""
            if preco_hist.motivo_perda:
                motivo_texto = {
                    "preco": "Preço",
                    "tecnica": "Questão técnica",
                    "documentacao": "Documentação",
                    "prazo": "Prazo",
                    "outro": "Outro"
                }.get(preco_hist.motivo_perda, preco_hist.motivo_perda)
                response += f"- Motivo: {motivo_texto}\n"

        return response, {"edital_id": edital.id, "resultado": preco_hist.resultado}

    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"❌ Erro ao consultar resultado: {str(e)}", None
    finally:
        db.close()


def processar_consultar_todos_resultados(user_id: str, db):
    """
    Consulta resultados de TODOS os editais do usuário.
    Retorna uma tabela markdown com os resultados.
    """
    from models import Edital, PrecoHistorico

    try:
        # Buscar editais com resultado registrado (status diferente de 'novo', 'aberto', 'analisando')
        status_com_resultado = ['vencedor', 'perdedor', 'cancelado', 'deserto', 'revogado']

        editais = db.query(Edital).filter(
            Edital.user_id == user_id,
            Edital.status.in_(status_com_resultado)
        ).order_by(Edital.data_abertura.desc()).all()

        if not editais:
            # Verificar se tem editais sem resultado
            total_editais = db.query(Edital).filter(Edital.user_id == user_id).count()
            if total_editais > 0:
                return f"""📊 **Resultados de Certames**

⚠️ Nenhum edital com resultado registrado.

Você tem **{total_editais} editais** cadastrados, mas nenhum com resultado definido.

Para registrar um resultado, use:
- "Perdemos o edital [NUMERO] para [EMPRESA] com R$ [VALOR]"
- "Ganhamos o edital [NUMERO] com R$ [VALOR]"
- "O edital [NUMERO] foi cancelado"
""", None
            else:
                return "❌ Você não tem editais cadastrados.", None

        # Contar por status
        contagem = {}
        for e in editais:
            status = e.status or "indefinido"
            contagem[status] = contagem.get(status, 0) + 1

        # Montar tabela markdown
        response = f"""## 📊 Resultados dos Certames

**Total com resultado:** {len(editais)} editais

**Resumo:**
"""
        # Adicionar resumo com emojis
        emoji_status = {
            'vencedor': '🏆',
            'perdedor': '📉',
            'cancelado': '⛔',
            'deserto': '🚫',
            'revogado': '❌'
        }
        for status, qtd in sorted(contagem.items(), key=lambda x: -x[1]):
            emoji = emoji_status.get(status, '📋')
            response += f"- {emoji} **{status.capitalize()}:** {qtd}\n"

        response += "\n---\n\n"

        # Tabela de editais
        response += "| Número | Órgão | Status | Valor Ref. | Data |\n"
        response += "|--------|-------|--------|------------|------|\n"

        for edital in editais[:20]:  # Limitar a 20 para não ficar muito grande
            numero = edital.numero or "N/A"
            orgao = (edital.orgao[:30] + "...") if edital.orgao and len(edital.orgao) > 30 else (edital.orgao or "N/A")
            status = edital.status.capitalize() if edital.status else "N/A"
            valor = f"R$ {float(edital.valor_referencia):,.0f}".replace(",", ".") if edital.valor_referencia else "N/A"
            data = edital.data_abertura.strftime('%d/%m/%Y') if edital.data_abertura else "N/A"

            # Adicionar emoji ao status
            emoji = emoji_status.get(edital.status, '')
            response += f"| {numero} | {orgao} | {emoji} {status} | {valor} | {data} |\n"

        if len(editais) > 20:
            response += f"\n*... e mais {len(editais) - 20} editais*\n"

        response += "\n---\n*Para detalhes de um edital específico, use: \"Qual o resultado do edital [NUMERO]?\"*"

        return response, {"total": len(editais), "contagem": contagem}

    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"❌ Erro ao consultar resultados: {str(e)}", None


def processar_extrair_ata(texto_pdf: str, filepath: str, user_id: str, filename: str):
    """
    Processa extração de resultados de uma ata de sessão de pregão.

    Args:
        texto_pdf: Texto extraído do PDF
        filepath: Caminho do arquivo
        user_id: ID do usuário
        filename: Nome do arquivo original

    Returns:
        Tuple (response_text, resultado)
    """
    from tools import tool_extrair_ata_pdf

    resultado = tool_extrair_ata_pdf(texto_pdf, user_id)

    if not resultado.get("success"):
        response = f"""## ❌ Erro ao Extrair Ata

**Arquivo:** {filename}

**Erro:** {resultado.get('error', 'Erro desconhecido')}

**Dica:** Certifique-se de que o arquivo é uma ata de sessão de pregão eletrônico.
"""
        return response, resultado

    # Montar resposta formatada
    response = f"""## 📄 Resultados Extraídos da Ata

**Arquivo:** {filename}
**Edital:** {resultado.get('edital', 'Não identificado')}
**Órgão:** {resultado.get('orgao', 'Não identificado')}
**Data da Sessão:** {resultado.get('data_sessao', 'N/A')}
**Objeto:** {resultado.get('objeto', 'N/A')[:200]}{'...' if resultado.get('objeto') and len(resultado.get('objeto', '')) > 200 else ''}

---

### 📊 Itens/Lotes Extraídos

"""

    # Verificar se há aviso (sem itens extraídos)
    if resultado.get("aviso"):
        response += f"""⚠️ **{resultado['aviso']}**

O documento foi processado mas não foram encontrados itens com vencedores/preços estruturados.

**Possíveis causas:**
- A ata pode estar incompleta ou ser de outra fase do pregão
- O formato do PDF pode ser diferente do esperado
- O texto pode estar escaneado (imagem) e não selecionável

**Alternativa:** Registre o resultado manualmente:
- "Perdemos o edital {resultado.get('edital', '[NUMERO]')} para [EMPRESA] com R$ [VALOR]"
- "Ganhamos o edital {resultado.get('edital', '[NUMERO]')} com R$ [VALOR]"

"""
        return response, resultado

    for item in resultado.get("itens", []):
        emoji = "🏆" if item.get("vencedor") else "❓"
        response += f"""**Item {item.get('item', '?')}:** {item.get('descricao', 'N/A')[:100]}...
- {emoji} **Vencedor:** {item.get('vencedor', 'Não identificado')}
- 💰 **Preço:** R$ {item.get('preco_vencedor', 0):,.2f}
- 👥 **Participantes:** {item.get('participantes_count', 0)}

"""

    # Desclassificados
    if resultado.get("desclassificados"):
        response += "### ⚠️ Empresas Desclassificadas\n\n"
        for desc in resultado["desclassificados"]:
            response += f"- **{desc.get('empresa', 'N/A')}:** {desc.get('motivo', 'Motivo não informado')}\n"
        response += "\n"

    # Concorrentes registrados
    response += "---\n\n### 📁 Dados Registrados\n\n"

    if resultado.get("concorrentes_novos"):
        response += f"**Novos concorrentes:** {', '.join(resultado['concorrentes_novos'][:5])}"
        if len(resultado.get('concorrentes_novos', [])) > 5:
            response += f" (+{len(resultado['concorrentes_novos']) - 5})"
        response += "\n"

    if resultado.get("concorrentes_atualizados"):
        response += f"**Concorrentes atualizados:** {len(resultado['concorrentes_atualizados'])}\n"

    if resultado.get("edital_encontrado"):
        response += f"\n✅ **Edital {resultado['edital_encontrado']} encontrado no sistema - dados salvos no histórico!**\n"
    else:
        edital_num = resultado.get('edital', '[NUMERO]')
        itens = resultado.get('itens', [])

        # Obter dados do primeiro item de forma segura
        primeiro_item = itens[0] if itens else {}
        vencedor = primeiro_item.get('vencedor', 'EMPRESA')
        preco = primeiro_item.get('preco_vencedor', 0)
        objeto = resultado.get('objeto', 'equipamentos') or 'equipamentos'

        response += f"""
⚠️ **Edital não encontrado no sistema.**

Para salvar os dados no histórico, primeiro cadastre o edital:
- Busque editais: "busque editais de {objeto[:30]}"
- Ou registre manualmente o resultado: "Perdemos o edital {edital_num} para {vencedor} com R$ {preco:,.0f}"
"""

    return response, resultado


def processar_buscar_atas_pncp(message: str, user_id: str):
    """
    Processa busca de atas de sessão/registro de preço no PNCP.

    Args:
        message: Mensagem do usuário
        user_id: ID do usuário

    Returns:
        Tuple (response_text, resultado)
    """
    from tools import tool_buscar_atas_pncp

    # Extrair termo de busca usando helper
    palavras = ["busque", "buscar", "encontre", "encontrar", "baixe", "baixar",
                "atas", "ata", "de", "do", "da", "no", "na", "pncp", "registro",
                "preço", "preco", "sessão", "sessao"]
    termo = extrair_termo(message, palavras)

    if not termo or len(termo) < 3:
        return """## ❓ Termo de Busca Necessário

Por favor, especifique o que você está buscando. Exemplos:
- "Busque atas de **hematologia**"
- "Encontre atas de **equipamentos hospitalares**"
- "Baixe atas de **material de laboratório**"
""", None

    resultado = tool_buscar_atas_pncp(termo, user_id)

    if not resultado.get("success"):
        response = f"""## ❌ Erro na Busca de Atas

**Termo:** {termo}
**Erro:** {resultado.get('error', 'Erro desconhecido')}

**Dica:** Tente termos mais específicos como:
- "hematologia"
- "equipamento médico"
- "reagentes laboratoriais"
"""
        return response, resultado

    atas = resultado.get("atas", [])

    response = f"""## 📄 Atas Encontradas no PNCP

**Termo:** {termo}
**Total:** {resultado.get('total', len(atas))} atas encontradas
**Fonte:** {resultado.get('fonte', 'PNCP')}

---

"""

    for i, ata in enumerate(atas[:10], 1):
        titulo = ata.get('titulo', 'Sem título')
        orgao = ata.get('orgao', 'N/A')
        descricao = ata.get('descricao', '')[:150]
        data = ata.get('data_assinatura') or ata.get('data_publicacao', 'N/A')
        url = ata.get('url_pncp', ata.get('url', '#'))

        response += f"""### {i}. {titulo}

**Órgão:** {orgao}
**Data:** {data}
**Descrição:** {descricao}{'...' if len(ata.get('descricao', '')) > 150 else ''}

🔗 [Acessar no PNCP]({url})

---

"""

    response += """
### 💡 Como usar as atas:

1. Clique no link para acessar a ata no PNCP
2. Baixe o PDF da ata de sessão
3. Envie o PDF aqui com a mensagem: **"Extraia os resultados desta ata"**

O sistema irá extrair automaticamente os vencedores, preços e participantes!
"""

    return response, resultado


def processar_buscar_precos_pncp(message: str, user_id: str):
    """
    Processa busca de preços de contratos no PNCP.
    Funcionalidade 4 da Sprint 1.

    Args:
        message: Mensagem do usuário
        user_id: ID do usuário

    Returns:
        Tuple (response_text, resultado)
    """
    from tools import tool_buscar_precos_pncp

    # Extrair termo de busca usando helper
    palavras = ["busque", "buscar", "encontre", "encontrar", "preços", "precos",
                "de", "do", "da", "no", "na", "pncp", "mercado", "médio", "medio",
                "quanto", "custa", "valor", "valores", "contrato", "contratos",
                "praticado", "praticados", "histórico", "historico", "qual", "o"]
    termo = extrair_termo(message, palavras)

    if not termo or len(termo) < 3:
        return """## ❓ Termo de Busca Necessário

Por favor, especifique o produto/equipamento que deseja pesquisar. Exemplos:
- "Busque preços de **analisador hematológico**"
- "Qual o preço de mercado para **centrífuga**?"
- "Preços de **reagentes de bioquímica** no PNCP"
""", None

    resultado = tool_buscar_precos_pncp(termo, meses=12, user_id=user_id)

    if not resultado.get("success"):
        response = f"""## ❌ Nenhum Preço Encontrado

**Termo:** {termo}
**Erro:** {resultado.get('error', 'Nenhum contrato encontrado')}

**Dica:** Tente termos mais específicos como:
- "analisador hematológico"
- "equipamento laboratório"
- "reagentes diagnóstico"
"""
        return response, resultado

    # Formatar resposta
    stats = resultado.get("estatisticas", {})
    contratos = resultado.get("contratos", [])
    top_fornecedores = resultado.get("top_fornecedores", [])

    response = f"""## 💰 Preços de Mercado - PNCP

**Termo pesquisado:** {termo}
**Período:** Últimos {resultado.get('periodo_meses', 12)} meses
**Contratos encontrados:** {resultado.get('total_contratos', 0)}
**Fonte:** {resultado.get('fonte', 'PNCP')}

---

### 📊 Estatísticas de Preços

| Métrica | Valor |
|---------|-------|
| **Mínimo** | R$ {stats.get('preco_minimo', 0):,.2f} |
| **Médio** | R$ {stats.get('preco_medio', 0):,.2f} |
| **Mediano** | R$ {stats.get('preco_mediano', 0):,.2f} |
| **Máximo** | R$ {stats.get('preco_maximo', 0):,.2f} |

---

### 🏢 Principais Fornecedores

"""
    for i, forn in enumerate(top_fornecedores[:5], 1):
        response += f"{i}. **{forn.get('nome', 'N/A')[:40]}** - {forn.get('contratos', 0)} contratos (média: R$ {forn.get('preco_medio', 0):,.2f})\n"

    response += """

---

### 📋 Últimos Contratos

"""
    for i, contrato in enumerate(contratos[:10], 1):
        objeto = contrato.get('objeto', contrato.get('titulo', 'N/A'))[:80]
        fornecedor = contrato.get('fornecedor', 'N/A')[:30]
        valor = contrato.get('valor', 0)
        orgao = contrato.get('orgao', 'N/A')[:30]
        data = contrato.get('data_assinatura', contrato.get('data_publicacao', 'N/A'))
        url = contrato.get('url_pncp', '#')

        response += f"""**{i}. {objeto}...**
- 🏢 Órgão: {orgao}
- 🏭 Fornecedor: {fornecedor}
- 💵 Valor: **R$ {valor:,.2f}**
- 📅 Data: {data}
"""
        if url and url != '#':
            response += f"- 🔗 [Ver no PNCP]({url})\n"
        response += "\n"

    response += """
---

### 💡 Como usar esses dados:

1. **Para definir preço de proposta:** Use o preço médio como referência
2. **Para análise de concorrentes:** Veja os principais fornecedores
3. **Para justificativa de preços:** Cite os contratos como referência

📌 **Dica:** Para salvar esses preços no histórico, registre um resultado de edital!
"""

    return response, resultado


# ==================== HELPER: EXTRAÇÃO DE TERMOS ====================

def extrair_termo(message: str, palavras_remover: list) -> str:
    """
    Extrai termo de busca removendo palavras-chave de comando.
    Usa regex com word boundaries para não cortar partes de palavras.

    Args:
        message: Mensagem do usuário
        palavras_remover: Lista de palavras a remover

    Returns:
        Termo extraído limpo
    """
    import re

    texto = message.lower()

    # Remover palavras usando word boundaries para não cortar partes de palavras
    for palavra in palavras_remover:
        # \b = word boundary - só casa com palavra completa
        pattern = r'\b' + re.escape(palavra) + r'\b'
        texto = re.sub(pattern, ' ', texto, flags=re.IGNORECASE)

    # Limpar espaços extras e pontuação no início/fim
    texto = re.sub(r'\s+', ' ', texto).strip()
    texto = re.sub(r'^[.,!?:;\s]+|[.,!?:;\s]+$', '', texto)

    return texto


# ==================== SPRINT 1 - FUNCIONALIDADE 5: HISTÓRICO DE PREÇOS ====================

def processar_historico_precos(message: str, user_id: str):
    """Processa consulta de histórico de preços."""
    from tools import tool_historico_precos

    # Extrair termo usando helper
    palavras = ["histórico", "historico", "preços", "precos", "de", "do", "da",
                "registrados", "salvos", "mostre", "mostrar", "ver", "consultar",
                "quais", "já", "ja"]
    termo = extrair_termo(message, palavras)

    resultado = tool_historico_precos(termo=termo if termo else None, user_id=user_id)

    if not resultado.get("success"):
        return f"""## ❌ Histórico de Preços

**Erro:** {resultado.get('error', 'Nenhum registro encontrado')}

**Dica:** Registre resultados de editais para criar histórico de preços.
""", resultado

    stats = resultado.get("estatisticas", {})
    historico = resultado.get("historico", [])

    response = f"""## 📈 Histórico de Preços

**Termo:** {termo or 'Todos'}
**Total de registros:** {resultado.get('total', 0)}

---

### 📊 Estatísticas

| Métrica | Valor |
|---------|-------|
| **Mínimo** | R$ {stats.get('preco_minimo', 0):,.2f} |
| **Médio** | R$ {stats.get('preco_medio', 0):,.2f} |
| **Mediano** | R$ {stats.get('preco_mediano', 0):,.2f} |
| **Máximo** | R$ {stats.get('preco_maximo', 0):,.2f} |

---

### 📋 Últimos Registros

"""
    for i, h in enumerate(historico[:10], 1):
        resultado_emoji = "🏆" if h.get('resultado') == 'vitoria' else "📊"
        response += f"{i}. {resultado_emoji} **R$ {h.get('preco_vencedor', 0):,.2f}** - {h.get('empresa_vencedora', 'N/A')} ({h.get('data', 'N/A')})\n"

    return response, resultado


# ==================== SPRINT 1 - FUNCIONALIDADE 6: ANÁLISE DE CONCORRENTES ====================

def processar_listar_concorrentes(user_id: str):
    """Processa listagem de concorrentes."""
    from tools import tool_listar_concorrentes

    resultado = tool_listar_concorrentes(user_id=user_id)

    if not resultado.get("success"):
        return f"""## ❌ Concorrentes

**Erro:** {resultado.get('error', 'Nenhum concorrente cadastrado')}

**Dica:** {resultado.get('dica', 'Registre resultados de editais para cadastrar concorrentes automaticamente.')}
""", resultado

    concorrentes = resultado.get("concorrentes", [])

    response = f"""## 👥 Concorrentes Conhecidos

**Total:** {resultado.get('total', 0)} concorrentes

---

| # | Empresa | Participações | Vitórias | Taxa |
|---|---------|---------------|----------|------|
"""
    for i, c in enumerate(concorrentes[:15], 1):
        response += f"| {i} | {c.get('nome', 'N/A')[:25]} | {c.get('editais_participados', 0)} | {c.get('editais_ganhos', 0)} | {c.get('taxa_vitoria', 0):.1f}% |\n"

    response += """

---

💡 **Dica:** Use "analise o concorrente [NOME]" para ver detalhes.
"""
    return response, resultado


def processar_analisar_concorrente(message: str, user_id: str):
    """Processa análise de concorrente específico."""
    from tools import tool_analisar_concorrente

    # Extrair nome do concorrente usando helper
    palavras = ["analise", "analisar", "análise", "concorrente", "o", "do", "da",
                "empresa", "histórico", "historico", "como", "está", "esta"]
    nome = extrair_termo(message, palavras)

    if not nome:
        return """## ❓ Nome do Concorrente

Por favor, especifique o concorrente. Exemplo:
- "Analise o concorrente **MedLab**"
- "Histórico do concorrente **TechSaúde**"
""", None

    resultado = tool_analisar_concorrente(nome, user_id=user_id)

    if not resultado.get("success"):
        return f"""## ❌ Concorrente Não Encontrado

**Buscado:** {nome}
**Erro:** {resultado.get('error', 'Não encontrado')}

**Dica:** {resultado.get('dica', 'Use "liste concorrentes" para ver os cadastrados.')}
""", resultado

    conc = resultado.get("concorrente", {})
    stats = resultado.get("estatisticas_precos", {})
    historico = resultado.get("historico_participacoes", [])

    response = f"""## 🔍 Análise do Concorrente

### {conc.get('nome', 'N/A')}
**CNPJ:** {conc.get('cnpj', 'Não informado')}

---

### 📊 Estatísticas

| Métrica | Valor |
|---------|-------|
| **Editais Participados** | {conc.get('editais_participados', 0)} |
| **Editais Ganhos** | {conc.get('editais_ganhos', 0)} |
| **Taxa de Vitória** | {conc.get('taxa_vitoria', 0):.1f}% |

### 💰 Preços Praticados

| Métrica | Valor |
|---------|-------|
| **Mínimo** | R$ {stats.get('preco_minimo', 0):,.2f} |
| **Médio** | R$ {stats.get('preco_medio', 0):,.2f} |
| **Máximo** | R$ {stats.get('preco_maximo', 0):,.2f} |

---

### 📋 Últimas Participações

"""
    for i, h in enumerate(historico[:10], 1):
        emoji = "🏆" if h.get('venceu') else "📊"
        response += f"{i}. {emoji} {h.get('edital', 'N/A')} - R$ {h.get('preco', 0):,.2f} (#{h.get('posicao', '?')}º)\n"

    return response, resultado


# ==================== SPRINT 1 - FUNCIONALIDADE 7: RECOMENDAÇÃO DE PREÇOS ====================

def processar_recomendar_preco(message: str, user_id: str):
    """Processa recomendação de preço."""
    from tools import tool_recomendar_preco

    # Extrair termo usando helper
    palavras = ["recomendar", "recomende", "sugerir", "sugira", "preço", "preco",
                "que", "qual", "colocar", "para", "de", "do", "da"]
    termo = extrair_termo(message, palavras)

    if not termo:
        return """## ❓ Produto/Termo Necessário

Por favor, especifique o produto. Exemplo:
- "Recomende preço para **analisador hematológico**"
- "Qual preço sugerir para **reagentes bioquímica**?"
""", None

    resultado = tool_recomendar_preco(termo, user_id=user_id)

    if not resultado.get("success"):
        return f"""## ❌ Recomendação de Preço

**Termo:** {termo}
**Erro:** {resultado.get('error', 'Dados insuficientes')}

**Dica:** {resultado.get('dica', 'Registre mais resultados de editais ou busque preços no PNCP.')}
""", resultado

    rec = resultado.get("recomendacao", {})
    stats = resultado.get("estatisticas_historico", resultado.get("estatisticas_mercado", {}))
    fonte = resultado.get("fonte", "")

    response = f"""## 💡 Recomendação de Preço

**Termo:** {termo}
**Fonte:** {fonte.replace('_', ' ').title()}
**Registros analisados:** {stats.get('total_registros', 0)}

---

### 🎯 Preços Sugeridos

| Estratégia | Preço Sugerido |
|------------|----------------|
| 🔥 **Agressivo** | R$ {rec.get('preco_agressivo', rec.get('preco_minimo_sugerido', 0)):,.2f} |
| ✅ **Ideal** | R$ {rec.get('preco_ideal', 0):,.2f} |
| 🛡️ **Conservador** | R$ {rec.get('preco_conservador', rec.get('preco_maximo_sugerido', 0)):,.2f} |

---

### 📊 Referência de Mercado

| Métrica | Valor |
|---------|-------|
| **Preço Médio Vencedor** | R$ {stats.get('preco_medio_vencedor', stats.get('preco_medio', 0)):,.2f} |
| **Preço Mínimo** | R$ {stats.get('preco_minimo_vencedor', stats.get('preco_minimo', 0)):,.2f} |

---

**Justificativa:** {resultado.get('justificativa', 'N/A')}

💡 **Dica:** O preço **ideal** oferece boa margem de vitória com lucro razoável.
"""
    return response, resultado


# ==================== SPRINT 1 - FUNCIONALIDADE 8: CLASSIFICAÇÃO DE EDITAIS ====================

def processar_classificar_edital(message: str, user_id: str):
    """Processa classificação de edital."""
    from tools import tool_classificar_edital

    # Extrair texto do edital ou ID
    msg_lower = message.lower()

    # Verificar se tem ID de edital
    import re
    match_id = re.search(r'edital\s*(\d+)', msg_lower)
    edital_id = int(match_id.group(1)) if match_id else None

    # Usar mensagem como texto se não tem ID
    texto = message if not edital_id else None

    resultado = tool_classificar_edital(edital_id=edital_id, texto_edital=texto, user_id=user_id)

    if not resultado.get("success"):
        return f"""## ❌ Classificação de Edital

**Erro:** {resultado.get('error', 'Não foi possível classificar')}

**Dica:** Forneça o texto do objeto do edital ou o ID de um edital cadastrado.
""", resultado

    categoria = resultado.get("categoria", "outros")
    confianca = resultado.get("confianca", 0)

    # Mapeamento de categorias
    categorias_nome = {
        "comodato": "🤝 Comodato de Equipamentos",
        "aluguel_reagentes": "📦 Aluguel com Reagentes",
        "aluguel_simples": "🏷️ Aluguel Simples",
        "venda": "💰 Venda/Aquisição",
        "consumo_reagentes": "🧪 Consumo de Reagentes",
        "insumos_hospitalares": "🏥 Insumos Hospitalares",
        "insumos_laboratoriais": "🔬 Insumos Laboratoriais",
        "outros": "❓ Outros"
    }

    response = f"""## 🏷️ Classificação do Edital

**Categoria Identificada:** {categorias_nome.get(categoria, categoria)}
**Confiança:** {confianca}%

---

### 📊 Todas as Categorias Detectadas

"""
    for cat, score in resultado.get("todas_categorias", {}).items():
        emoji = "✅" if cat == categoria else "⬜"
        response += f"{emoji} **{cat}**: {score} matches\n"

    response += f"""

---

**Justificativa:** {resultado.get('justificativa', 'N/A')}
"""
    return response, resultado


# ==================== SPRINT 1 - FUNCIONALIDADE 9: VERIFICAR COMPLETUDE ====================

def processar_verificar_completude(message: str, user_id: str):
    """Processa verificação de completude de produto — usa mesmo tool do endpoint REST."""
    from tools import tool_verificar_completude_produto

    # Extrair nome do produto usando helper
    palavras = ["verificar", "verifique", "completude", "produto", "está", "esta",
                "completo", "falta", "informação", "informacao", "faltando", "o", "do", "da"]
    nome = extrair_termo(message, palavras)

    resultado = tool_verificar_completude_produto(nome_produto=nome if nome else None, user_id=user_id)

    if not resultado.get("success"):
        return f"""## ❌ Verificação de Completude

**Erro:** {resultado.get('error', 'Produto não encontrado')}

**Dica:** Informe o nome do produto. Exemplo: "Verifique completude do **Analisador XYZ**"
""", resultado

    produto = resultado.get("produto", {})
    completude = resultado.get("completude", {})
    campos_basicos = resultado.get("campos_basicos", [])
    mascara_check = resultado.get("mascara_check", [])
    subclasse_nome = resultado.get("subclasse_nome")

    status_emoji = {
        "completo": "✅",
        "quase_completo": "🟡",
        "incompleto": "🟠",
        "muito_incompleto": "🔴"
    }

    response = f"""## 📋 Verificação de Completude

### Produto: {produto.get('nome', 'N/A')}

**Status:** {status_emoji.get(completude.get('status'), '❓')} {completude.get('status', 'N/A').replace('_', ' ').title()}

| Métrica | Percentual |
|---------|-----------|
| **Geral** | {completude.get('percentual_geral', 0)}% |
| **Dados Básicos** | {completude.get('percentual_basicos', 0)}% |
| **Especificações** | {completude.get('percentual_mascara', 0)}% |

---

### Dados Básicos

"""
    for c in campos_basicos:
        emoji = "✅" if c["preenchido"] else "❌"
        valor = c["valor"] if c["preenchido"] else "Não preenchido"
        response += f"- {emoji} **{c['campo']}**: {valor}\n"

    if subclasse_nome and mascara_check:
        preenchidos = sum(1 for c in mascara_check if c["preenchido"])
        response += f"""
---

### Especificações — {subclasse_nome} ({preenchidos}/{len(mascara_check)})

"""
        faltantes = [c for c in mascara_check if not c["preenchido"]]
        if faltantes:
            response += "**Faltantes:**\n"
            for c in faltantes[:20]:
                un = f" ({c['unidade']})" if c.get("unidade") else ""
                response += f"- ❌ {c['campo']}{un}\n"
            if len(faltantes) > 20:
                response += f"- ... e mais {len(faltantes) - 20} campos\n"
    elif not subclasse_nome:
        response += "\n⚠️ Produto sem subclasse atribuída — não é possível verificar especificações contra a máscara.\n"

    recomendacoes = resultado.get("recomendacoes", [])
    if recomendacoes:
        response += "\n---\n\n### 💡 Recomendações\n\n"
        for rec in recomendacoes:
            response += f"- {rec}\n"

    return response, resultado


def processar_cadastrar_edital(message: str, user_id: str, intencao_resultado: dict = None):
    """
    Processa ação: Cadastrar edital manualmente no sistema.

    Extrai dados da mensagem do usuário:
    - Número do edital
    - Órgão
    - Objeto (descrição)
    - Modalidade (opcional)
    - Data de abertura (opcional)
    - UF (opcional)
    """
    from models import Edital
    from database import SessionLocal
    import re
    from datetime import datetime

    # Usar LLM para extrair dados estruturados da mensagem
    prompt_extracao = f"""Extraia os dados do edital da mensagem abaixo e retorne APENAS um JSON:

MENSAGEM: "{message}"

Extraia:
- numero: número/identificador do edital (ex: PE-001/2026, Pregão 15/2026)
- orgao: nome do órgão licitante
- objeto: descrição/objeto da licitação
- modalidade: uma de [pregao_eletronico, pregao_presencial, concorrencia, tomada_precos, convite, leilao, dispensa, inexigibilidade] (default: pregao_eletronico)
- data_abertura: data no formato YYYY-MM-DD (se mencionada)
- uf: sigla do estado (se mencionado)
- cidade: nome da cidade (se mencionado)
- valor_referencia: valor estimado (se mencionado, apenas número)

Retorne APENAS o JSON, sem explicações:
{{"numero": "...", "orgao": "...", "objeto": "...", "modalidade": "...", "data_abertura": null, "uf": null, "cidade": null, "valor_referencia": null}}"""

    try:
        resposta_llm = call_deepseek(
            [{"role": "user", "content": prompt_extracao}],
            max_tokens=500,
            model_override="deepseek-chat"
        )

        # Extrair JSON da resposta
        import json
        json_match = re.search(r'\{[\s\S]*?\}', resposta_llm)
        if not json_match:
            return """❌ **Não consegui extrair os dados do edital.**

Por favor, forneça pelo menos:
- **Número do edital** (ex: PE-001/2026)
- **Órgão** (ex: Hospital das Clínicas)
- **Objeto** (ex: Aquisição de equipamentos)

**Exemplo:**
```
Cadastre o edital PE-001/2026, órgão Hospital das Clínicas UFMG, objeto: Aquisição de analisadores hematológicos
```""", None

        dados = json.loads(json_match.group())

        # Validar campos obrigatórios
        if not dados.get("numero"):
            return "❌ **Número do edital é obrigatório.** Informe o número (ex: PE-001/2026)", None

        if not dados.get("orgao"):
            return "❌ **Órgão é obrigatório.** Informe o órgão licitante.", None

        if not dados.get("objeto"):
            return "❌ **Objeto é obrigatório.** Informe a descrição/objeto da licitação.", None

        # Criar edital no banco
        db = SessionLocal()
        try:
            # Verificar se já existe
            edital_existente = db.query(Edital).filter(
                Edital.numero == dados["numero"],
                Edital.user_id == user_id
            ).first()

            if edital_existente:
                return f"""⚠️ **Edital já cadastrado!**

**Número:** {edital_existente.numero}
**Órgão:** {edital_existente.orgao}
**Status:** {edital_existente.status}

Se deseja atualizar, use: "Atualize o edital {dados['numero']} com..." """, None

            # Criar novo edital
            novo_edital = Edital(
                user_id=user_id,
                numero=dados["numero"],
                orgao=dados["orgao"],
                objeto=dados["objeto"],
                modalidade=dados.get("modalidade", "pregao_eletronico"),
                status="novo",
                fonte="manual",
                uf=dados.get("uf"),
                cidade=dados.get("cidade"),
                valor_referencia=float(dados["valor_referencia"]) if dados.get("valor_referencia") else None
            )

            # Converter data_abertura se existir
            if dados.get("data_abertura"):
                try:
                    novo_edital.data_abertura = datetime.strptime(dados["data_abertura"], "%Y-%m-%d")
                except:
                    pass

            db.add(novo_edital)
            db.commit()
            db.refresh(novo_edital)

            response = f"""✅ **Edital cadastrado com sucesso!**

📋 **Dados do Edital:**
| Campo | Valor |
|-------|-------|
| **Número** | {novo_edital.numero} |
| **Órgão** | {novo_edital.orgao} |
| **Objeto** | {novo_edital.objeto[:100]}{'...' if len(novo_edital.objeto) > 100 else ''} |
| **Modalidade** | {novo_edital.modalidade} |
| **Status** | {novo_edital.status} |
| **UF** | {novo_edital.uf or '-'} |
| **Cidade** | {novo_edital.cidade or '-'} |

---
**Próximos passos:**
- Calcule a aderência: "Calcule aderência do produto X ao edital {novo_edital.numero}"
- Gere uma proposta: "Gere proposta para o edital {novo_edital.numero}"
- Liste seus editais: "Liste meus editais"
"""
            return response, {"edital_id": str(novo_edital.id), "numero": novo_edital.numero}

        finally:
            db.close()

    except json.JSONDecodeError as e:
        return f"❌ Erro ao interpretar dados: {str(e)}", None
    except Exception as e:
        print(f"[ERRO] processar_cadastrar_edital: {e}")
        return f"❌ Erro ao cadastrar edital: {str(e)}", None


# =============================================================================
# PROCESSADORES SPRINT 2: ALERTAS E AUTOMAÇÃO
# =============================================================================

def processar_configurar_alertas(message: str, user_id: str):
    """Processa configuração de alertas de prazo para editais."""
    import re
    from tools import tool_configurar_alertas

    msg = message.lower()

    # Extrair número do edital - aceita formatos: PE-123/2026, PE-TESTE/2026, PE 123, Pregão 123/2026
    match_edital = re.search(r'(PE[-\s]?[\w]+[-/]?\d*|[Pp]reg[aã]o\s*n?[ºo°]?\s*[\w/]+|\d{1,5}[/]\d{4})', message, re.IGNORECASE)
    edital_numero = match_edital.group(1).strip() if match_edital else None

    if not edital_numero:
        return "⚠️ Para configurar alertas, preciso saber qual edital. Informe o número do edital, por exemplo:\n\n*\"Configure alertas para o PE 123/2024\"*"

    # Extrair tempos (horas/minutos antes)
    tempos_minutos = []

    # Padrões de tempo
    match_horas = re.findall(r'(\d+)\s*(?:hora|h)', msg)
    match_dias = re.findall(r'(\d+)\s*(?:dia|d)', msg)
    match_minutos = re.findall(r'(\d+)\s*(?:minuto|min|m\b)', msg)

    for h in match_horas:
        tempos_minutos.append(int(h) * 60)
    for d in match_dias:
        tempos_minutos.append(int(d) * 1440)
    for m in match_minutos:
        tempos_minutos.append(int(m))

    # Se não especificou tempo, usar padrões
    if not tempos_minutos:
        tempos_minutos = [1440, 60, 15]  # 1 dia, 1 hora, 15 min

    # Detectar tipo de alerta
    tipo = "abertura"
    if "impugna" in msg:
        tipo = "impugnacao"
    elif "recurso" in msg:
        tipo = "recursos"
    elif "proposta" in msg:
        tipo = "proposta"

    # Detectar canais
    canais = {"email": True, "push": True}
    if "apenas email" in msg or "só email" in msg:
        canais = {"email": True, "push": False}
    elif "apenas push" in msg or "só push" in msg:
        canais = {"email": False, "push": True}

    resultado = tool_configurar_alertas(
        user_id=user_id,
        edital_numero=edital_numero,
        tempos_minutos=tempos_minutos,
        tipo=tipo,
        canais=canais
    )

    if resultado.get("success"):
        alertas = resultado.get("alertas_criados", [])
        msg_resp = f"✅ **Alertas configurados para {edital_numero}**\n\n"

        if alertas:
            msg_resp += "📋 **Alertas agendados:**\n"
            for a in alertas:
                msg_resp += f"- ⏰ {a['tempo_antes']} antes → {a['data_disparo']}\n"
        else:
            msg_resp += "ℹ️ Os alertas foram configurados com os tempos padrão.\n"

        msg_resp += f"\n🔔 **Canais:** Email: {'✅' if canais['email'] else '❌'} | Push: {'✅' if canais['push'] else '❌'}"

        return msg_resp
    else:
        return f"❌ {resultado.get('error', 'Erro ao configurar alertas')}"


def processar_listar_alertas(message: str, user_id: str):
    """Processa listagem de alertas configurados."""
    from tools import tool_listar_alertas
    msg = message.lower()

    apenas_agendados = "todos" not in msg and "histórico" not in msg and "historico" not in msg
    periodo_dias = 30

    if "semana" in msg:
        periodo_dias = 7
    elif "mês" in msg or "mes" in msg:
        periodo_dias = 30
    elif "ano" in msg:
        periodo_dias = 365

    resultado = tool_listar_alertas(
        user_id=user_id,
        apenas_agendados=apenas_agendados,
        periodo_dias=periodo_dias
    )

    if resultado.get("success"):
        editais_com_alertas = resultado.get("editais", [])
        total_alertas = resultado.get("total_alertas", 0)

        if total_alertas == 0:
            return "📭 Você não tem alertas configurados.\n\nPara criar alertas, diga algo como:\n*\"Configure alertas para o PE 123/2024 com 1 dia e 1 hora de antecedência\"*"

        msg_resp = f"🔔 **Seus Alertas** ({total_alertas} encontrados)\n\n"

        for ed in editais_com_alertas:
            edital_info = ed.get('edital', {})
            alertas = ed.get('alertas', [])
            numero = edital_info.get('numero', 'N/A')
            orgao = edital_info.get('orgao', '')[:40]

            msg_resp += f"📋 **{numero}** - {orgao}\n"

            for a in alertas:
                status_icon = {"agendado": "⏳", "disparado": "✅", "lido": "👁️", "cancelado": "❌"}.get(a.get('status', ''), "📌")
                tipo = a.get('tipo', 'abertura').title()
                data_disparo = a.get('data_disparo', 'N/A')
                # Formatar data ISO para legível
                if data_disparo and data_disparo != 'N/A':
                    try:
                        from datetime import datetime
                        dt = datetime.fromisoformat(data_disparo)
                        data_disparo = dt.strftime("%d/%m/%Y %H:%M")
                    except:
                        pass
                tempo_antes = a.get('tempo_antes', '')

                msg_resp += f"   {status_icon} {tipo} - 📅 {data_disparo}\n"
                if tempo_antes:
                    msg_resp += f"      ⏰ {tempo_antes} antes da abertura\n"

            msg_resp += "\n"

        return msg_resp
    else:
        return f"❌ {resultado.get('error', 'Erro ao listar alertas')}"


def processar_dashboard_prazos(message: str, user_id: str):
    """Processa exibição do dashboard de prazos."""
    from tools import tool_dashboard_prazos
    msg = message.lower()

    dias = 7  # Padrão: próximos 7 dias
    if "mês" in msg or "mes" in msg or "30" in msg:
        dias = 30
    elif "15" in msg:
        dias = 15
    elif "semana" in msg or "7" in msg:
        dias = 7

    resultado = tool_dashboard_prazos(user_id=user_id, dias=dias)

    if resultado.get("success"):
        editais = resultado.get("editais", [])
        stats = resultado.get("estatisticas", {})

        msg_resp = f"📊 **Dashboard de Prazos** (próximos {dias} dias)\n\n"

        # Resumo (estatísticas: total, criticos, altos, medios, normais)
        msg_resp += "### 📈 Resumo\n"
        msg_resp += f"- Total: **{stats.get('total', 0)}** editais\n"
        msg_resp += f"- 🔴 Críticos (< 1 dia): **{stats.get('criticos', 0)}**\n"
        msg_resp += f"- 🟠 Altos (1-3 dias): **{stats.get('altos', 0)}**\n"
        msg_resp += f"- 🟡 Médios (3-7 dias): **{stats.get('medios', 0)}**\n"
        msg_resp += f"- 🟢 Normais (> 7 dias): **{stats.get('normais', 0)}**\n\n"

        if not editais:
            msg_resp += "ℹ️ Nenhum edital com prazo neste período.\n"
        else:
            msg_resp += "### 📋 Editais por Prazo\n\n"

            for e in editais[:10]:  # Limitar a 10
                # Estrutura: edital{id,numero,orgao,...}, datas{...}, tempo_restante{texto,dias,horas}, urgencia, emoji, alertas_configurados
                edital_info = e.get('edital', {})
                datas = e.get('datas', {})
                tempo = e.get('tempo_restante', {})
                icon = e.get('emoji', '🟢')

                numero = edital_info.get('numero', 'N/A')
                orgao = edital_info.get('orgao', '')[:40]
                data_abertura = datas.get('abertura_formatada', 'N/A')
                tempo_texto = tempo.get('texto', 'N/A')
                alertas_qtd = e.get('alertas_configurados', 0)

                msg_resp += f"{icon} **{numero}** - {orgao}\n"
                msg_resp += f"   📅 Abertura: {data_abertura}\n"
                msg_resp += f"   ⏱️ **{tempo_texto}**\n"
                if alertas_qtd:
                    msg_resp += f"   🔔 Alertas: {alertas_qtd}\n"
                msg_resp += "\n"

        return msg_resp
    else:
        return f"❌ {resultado.get('error', 'Erro ao carregar dashboard')}"


def processar_calendario_editais(message: str, user_id: str):
    """Processa exibição do calendário de editais."""
    from datetime import datetime
    from tools import tool_calendario_editais
    import re

    msg = message.lower()
    hoje = datetime.now()

    # Detectar mês/ano
    mes = hoje.month
    ano = hoje.year

    # Padrões para mês
    meses_pt = {
        "janeiro": 1, "fevereiro": 2, "março": 3, "marco": 3, "abril": 4,
        "maio": 5, "junho": 6, "julho": 7, "agosto": 8, "setembro": 9,
        "outubro": 10, "novembro": 11, "dezembro": 12
    }

    for nome, num in meses_pt.items():
        if nome in msg:
            mes = num
            break

    # Detectar ano
    match_ano = re.search(r'20\d{2}', msg)
    if match_ano:
        ano = int(match_ano.group())

    resultado = tool_calendario_editais(user_id=user_id, mes=mes, ano=ano)

    if resultado.get("success"):
        calendario = resultado.get("calendario", {})
        mes_nome = resultado.get("mes_nome", "")
        total = resultado.get("total_editais", 0)

        msg_resp = f"📅 **Calendário de Editais - {mes_nome} {ano}**\n\n"
        msg_resp += f"Total: **{total}** editais no mês\n\n"

        if not calendario:
            msg_resp += "ℹ️ Nenhum edital com data neste mês.\n"
        else:
            # Ordenar por dia
            for dia in sorted(calendario.keys(), key=int):
                editais_dia = calendario[dia]
                msg_resp += f"### 📆 Dia {dia}\n"

                for e in editais_dia:
                    status_icon = {"novo": "🆕", "analisando": "🔍", "participar": "✅", "proposta_enviada": "📤"}.get(e['status'], "📌")
                    msg_resp += f"{status_icon} **{e['numero']}** - {e['orgao'][:35]}\n"
                    if e.get('horario'):
                        msg_resp += f"   ⏰ {e['horario']}\n"
                msg_resp += "\n"

        return msg_resp
    else:
        return f"❌ {resultado.get('error', 'Erro ao carregar calendário')}"


def processar_configurar_monitoramento(message: str, user_id: str):
    """Processa configuração de monitoramento automático de editais."""
    import re
    from tools import tool_configurar_monitoramento

    msg = message.lower()

    # Extrair termo de busca - geralmente vem após "monitore" ou "monitorar"
    match_termo = re.search(r'monitor[ea]?\s+(?:editais\s+(?:de|para|sobre)\s+)?(.+?)(?:\s+(?:no|na|em|com|para)|$)', msg)
    termo = match_termo.group(1).strip() if match_termo else None

    if not termo:
        # Tentar extrair de outra forma
        match_termo2 = re.search(r'(?:busca automática|acompanhar)\s+(?:de\s+)?(.+?)(?:\s+(?:no|na|em)|$)', msg)
        termo = match_termo2.group(1).strip() if match_termo2 else None

    if not termo:
        return "⚠️ Para configurar um monitoramento, preciso saber o que monitorar.\n\nExemplos:\n- *\"Monitore editais de equipamentos laboratoriais\"*\n- *\"Configure monitoramento para reagentes em SP e MG\"*"

    # Detectar fontes
    fontes = []
    if "pncp" in msg:
        fontes.append("pncp")
    if "comprasnet" in msg:
        fontes.append("comprasnet")
    if "bec" in msg:
        fontes.append("bec")
    if not fontes:
        fontes = ["pncp"]  # Padrão

    # Detectar UFs
    ufs = []
    ufs_validas = ["AC", "AL", "AP", "AM", "BA", "CE", "DF", "ES", "GO", "MA", "MT", "MS",
                   "MG", "PA", "PB", "PR", "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC",
                   "SP", "SE", "TO"]
    for uf in ufs_validas:
        if uf.lower() in msg or uf in message:
            ufs.append(uf)

    # Detectar frequência
    frequencia_horas = 4  # Padrão
    if "hora em hora" in msg or "1 hora" in msg:
        frequencia_horas = 1
    elif "2 hora" in msg:
        frequencia_horas = 2
    elif "6 hora" in msg:
        frequencia_horas = 6
    elif "12 hora" in msg:
        frequencia_horas = 12
    elif "diário" in msg or "diario" in msg or "24 hora" in msg:
        frequencia_horas = 24

    # Detectar score mínimo
    score_minimo = 70
    match_score = re.search(r'score\s*(?:mínimo|minimo)?\s*(?:de\s+)?(\d+)', msg)
    if match_score:
        score_minimo = int(match_score.group(1))

    resultado = tool_configurar_monitoramento(
        user_id=user_id,
        termo=termo,
        fontes=fontes,
        ufs=ufs if ufs else None,
        frequencia_horas=frequencia_horas,
        score_minimo=score_minimo
    )

    if resultado.get("success"):
        mon = resultado.get("monitoramento", {})
        msg_resp = f"✅ **Monitoramento Configurado**\n\n"
        msg_resp += f"🔍 **Termo:** {mon.get('termo', termo)}\n"
        msg_resp += f"📡 **Fontes:** {', '.join(mon.get('fontes', fontes))}\n"
        msg_resp += f"📍 **UFs:** {', '.join(mon.get('ufs', ufs)) if mon.get('ufs') else 'Todas'}\n"
        msg_resp += f"⏱️ **Frequência:** A cada {mon.get('frequencia_horas', frequencia_horas)} hora(s)\n"
        msg_resp += f"📊 **Score mínimo para alerta:** {mon.get('score_minimo', score_minimo)}%\n"
        msg_resp += f"\n🆔 ID: `{mon.get('id', 'N/A')}`"

        return msg_resp
    else:
        return f"❌ {resultado.get('error', 'Erro ao configurar monitoramento')}"


def processar_listar_monitoramentos(message: str, user_id: str):
    """Processa listagem de monitoramentos configurados."""
    from tools import tool_listar_monitoramentos
    msg = message.lower()
    apenas_ativos = "todos" not in msg and "inativos" not in msg

    resultado = tool_listar_monitoramentos(user_id=user_id, apenas_ativos=apenas_ativos)

    if resultado.get("success"):
        monitoramentos = resultado.get("monitoramentos", [])

        if not monitoramentos:
            return "📭 Você não tem monitoramentos configurados.\n\nPara criar um monitoramento, diga algo como:\n*\"Monitore editais de equipamentos laboratoriais no PNCP\"*"

        msg_resp = f"🔍 **Seus Monitoramentos** ({len(monitoramentos)} encontrados)\n\n"

        for m in monitoramentos:
            status_icon = "✅" if m.get('ativo') else "⏸️"
            msg_resp += f"{status_icon} **{m['termo']}**\n"
            msg_resp += f"   📡 Fontes: {', '.join(m.get('fontes', []))}\n"
            msg_resp += f"   📍 UFs: {', '.join(m.get('ufs', [])) if m.get('ufs') else 'Todas'}\n"
            msg_resp += f"   ⏱️ A cada {m.get('frequencia_horas', 4)}h\n"
            msg_resp += f"   📊 Score mínimo: {m.get('score_minimo', 70)}%\n"
            if m.get('ultima_execucao'):
                msg_resp += f"   🕐 Última execução: {m['ultima_execucao']}\n"
            if m.get('editais_encontrados'):
                msg_resp += f"   📋 Editais encontrados: {m['editais_encontrados']}\n"
            msg_resp += "\n"

        return msg_resp
    else:
        return f"❌ {resultado.get('error', 'Erro ao listar monitoramentos')}"


def processar_desativar_monitoramento(message: str, user_id: str):
    """Processa desativação de monitoramento."""
    import re
    from tools import tool_desativar_monitoramento

    msg = message.lower()

    # Tentar extrair termo do monitoramento
    match_termo = re.search(r'(?:desativ|par|cancel|remov)[ea]?\s+(?:o\s+)?monitoramento\s+(?:de\s+)?(.+)', msg)
    termo = match_termo.group(1).strip() if match_termo else None

    # Tentar extrair ID
    match_id = re.search(r'id[:\s]+([a-f0-9-]+)', msg, re.IGNORECASE)
    monitoramento_id = match_id.group(1) if match_id else None

    if not termo and not monitoramento_id:
        return "⚠️ Para desativar um monitoramento, informe o termo ou ID.\n\nExemplos:\n- *\"Desative o monitoramento de equipamentos laboratoriais\"*\n- *\"Pare de monitorar reagentes\"*"

    resultado = tool_desativar_monitoramento(
        user_id=user_id,
        termo=termo,
        monitoramento_id=monitoramento_id
    )

    if resultado.get("success"):
        return f"✅ Monitoramento desativado com sucesso!\n\n🔍 **Termo:** {resultado.get('termo', termo or 'N/A')}"
    else:
        return f"❌ {resultado.get('error', 'Erro ao desativar monitoramento')}"


def processar_configurar_notificacoes(message: str, user_id: str):
    """Processa configuração de preferências de notificação."""
    import re
    from tools import tool_configurar_preferencias_notificacao

    msg = message.lower()

    # Extrair email
    match_email = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', message)
    email = match_email.group() if match_email else None

    # Detectar habilitação
    email_habilitado = True
    push_habilitado = True

    if "desativar email" in msg or "sem email" in msg:
        email_habilitado = False
    if "desativar push" in msg or "sem push" in msg:
        push_habilitado = False

    # Detectar horários
    horario_inicio = None
    horario_fim = None

    match_horario = re.search(r'(?:das|entre)\s+(\d{1,2})(?:h|:00)?\s+(?:às|e|até)\s+(\d{1,2})(?:h|:00)?', msg)
    if match_horario:
        horario_inicio = f"{int(match_horario.group(1)):02d}:00"
        horario_fim = f"{int(match_horario.group(2)):02d}:00"

    resultado = tool_configurar_preferencias_notificacao(
        user_id=user_id,
        email_habilitado=email_habilitado,
        push_habilitado=push_habilitado,
        email_notificacao=email,
        horario_inicio=horario_inicio,
        horario_fim=horario_fim
    )

    if resultado.get("success"):
        prefs = resultado.get("preferencias", {})
        msg_resp = "✅ **Preferências de Notificação Atualizadas**\n\n"
        msg_resp += f"📧 **Email:** {'✅ Habilitado' if prefs.get('email_habilitado') else '❌ Desabilitado'}\n"
        if prefs.get('email_notificacao'):
            msg_resp += f"   Enviar para: {prefs['email_notificacao']}\n"
        msg_resp += f"🔔 **Push:** {'✅ Habilitado' if prefs.get('push_habilitado') else '❌ Desabilitado'}\n"

        if prefs.get('horario_inicio') and prefs.get('horario_fim'):
            msg_resp += f"⏰ **Horário:** {prefs['horario_inicio']} às {prefs['horario_fim']}\n"

        return msg_resp
    else:
        return f"❌ {resultado.get('error', 'Erro ao configurar notificações')}"


def processar_historico_notificacoes(message: str, user_id: str):
    """Processa listagem do histórico de notificações."""
    from tools import tool_historico_notificacoes
    msg = message.lower()

    apenas_nao_lidas = "não lida" in msg or "nao lida" in msg or "pendente" in msg

    limite = 20
    if "últimas 10" in msg or "ultimas 10" in msg:
        limite = 10
    elif "últimas 50" in msg or "ultimas 50" in msg:
        limite = 50

    resultado = tool_historico_notificacoes(
        user_id=user_id,
        limite=limite,
        apenas_nao_lidas=apenas_nao_lidas
    )

    if resultado.get("success"):
        notificacoes = resultado.get("notificacoes", [])
        nao_lidas = resultado.get("nao_lidas", 0)

        if not notificacoes:
            return "📭 Você não tem notificações.\n\nAs notificações aparecem quando:\n- Alertas de prazo são disparados\n- Novos editais são encontrados pelo monitoramento\n- O sistema precisa informar algo importante"

        msg_resp = f"📬 **Suas Notificações** ({len(notificacoes)} exibidas"
        if nao_lidas > 0:
            msg_resp += f", {nao_lidas} não lidas"
        msg_resp += ")\n\n"

        for n in notificacoes:
            tipo_icon = {
                "alerta_prazo": "⏰",
                "novo_edital": "📋",
                "alta_aderencia": "⭐",
                "resultado": "📊",
                "sistema": "🔧"
            }.get(n.get('tipo'), "📌")

            lida_icon = "👁️" if n.get('lida') else "🔵"

            msg_resp += f"{lida_icon} {tipo_icon} **{n['titulo']}**\n"
            msg_resp += f"   {n['mensagem'][:100]}{'...' if len(n.get('mensagem', '')) > 100 else ''}\n"
            msg_resp += f"   🕐 {n['created_at']}\n\n"

        return msg_resp
    else:
        return f"❌ {resultado.get('error', 'Erro ao carregar notificações')}"


def processar_extrair_datas_edital(message: str, user_id: str, texto_pdf: str = None):
    """Processa extração de datas importantes de um edital."""
    import re
    from tools import tool_extrair_datas_edital

    # Se não tem texto PDF, informar como usar
    if not texto_pdf:
        # Verificar se há número de edital para buscar
        match_edital = re.search(r'(PE[-]?\d+[-/]?\d*|[Pp]reg[aã]o\s*n?[ºo°]?\s*[\d/]+|\d{1,5}[/]\d{4})', message, re.IGNORECASE)

        if match_edital:
            edital_numero = match_edital.group(1)
            return f"⚠️ Para extrair as datas do edital **{edital_numero}**, faça upload do PDF do edital.\n\nApós o upload, direi:\n*\"Extraia as datas do edital {edital_numero}\"*"
        else:
            return "⚠️ Para extrair datas de um edital, faça upload do PDF primeiro.\n\nApós o upload, diga:\n*\"Extraia as datas do edital PE 123/2024\"*"

    # Se temos texto PDF, extrair datas
    resultado = tool_extrair_datas_edital(
        user_id=user_id,
        texto_edital=texto_pdf
    )

    if resultado.get("success"):
        datas = resultado.get("datas", {})
        msg_resp = "📅 **Datas Extraídas do Edital**\n\n"

        if datas.get("data_abertura"):
            msg_resp += f"📆 **Data de Abertura:** {datas['data_abertura']}\n"
        if datas.get("horario_abertura"):
            msg_resp += f"⏰ **Horário:** {datas['horario_abertura']}\n"
        if datas.get("data_limite_propostas"):
            msg_resp += f"📝 **Limite para Propostas:** {datas['data_limite_propostas']}\n"
        if datas.get("data_impugnacao"):
            msg_resp += f"⚠️ **Prazo Impugnação:** {datas['data_impugnacao']}\n"
        if datas.get("data_recursos"):
            msg_resp += f"📑 **Prazo Recursos:** {datas['data_recursos']}\n"
        if datas.get("data_publicacao"):
            msg_resp += f"📰 **Data Publicação:** {datas['data_publicacao']}\n"

        msg_resp += "\n💡 *Deseja configurar alertas para estas datas?*"

        return msg_resp
    else:
        return f"❌ {resultado.get('error', 'Erro ao extrair datas')}"


def processar_cancelar_alerta(message: str, user_id: str):
    """Processa cancelamento de alertas."""
    import re
    from tools import tool_cancelar_alerta

    msg = message.lower()

    # Extrair número do edital
    match_edital = re.search(r'(PE[-]?\d+[-/]?\d*|[Pp]reg[aã]o\s*n?[ºo°]?\s*[\d/]+|\d{1,5}[/]\d{4})', message, re.IGNORECASE)
    edital_numero = match_edital.group(1) if match_edital else None

    # Extrair ID do alerta
    match_id = re.search(r'alerta[:\s]+([a-f0-9-]+)', msg, re.IGNORECASE)
    alerta_id = match_id.group(1) if match_id else None

    # Cancelar todos?
    cancelar_todos = "todos" in msg

    if not edital_numero and not alerta_id and not cancelar_todos:
        return "⚠️ Para cancelar alertas, informe o edital ou o ID do alerta.\n\nExemplos:\n- *\"Cancele os alertas do PE 123/2024\"*\n- *\"Cancele todos os alertas\"*"

    resultado = tool_cancelar_alerta(
        user_id=user_id,
        alerta_id=alerta_id,
        edital_numero=edital_numero,
        cancelar_todos=cancelar_todos
    )

    if resultado.get("success"):
        qtd = resultado.get("cancelados", 0)
        msg_resp = f"✅ **{qtd} alerta(s) cancelado(s)**\n\n"

        if edital_numero:
            msg_resp += f"📋 Edital: {edital_numero}"

        return msg_resp
    else:
        return f"❌ {resultado.get('error', 'Erro ao cancelar alertas')}"


# =============================================================================
# ANÁLISE DE EDITAIS (Resumir e Perguntar)
# =============================================================================

def processar_resumir_edital(message: str, user_id: str, intencao_resultado: dict = None):
    """
    Processa ação: Resumir um edital cadastrado no sistema.
    Gera um resumo executivo com principais informações.

    Returns: (response_text, resultado_dict)
    """
    import re
    db = get_db()

    try:
        # Extrair número do edital da mensagem
        edital_numero = None
        if intencao_resultado and intencao_resultado.get("edital"):
            edital_numero = intencao_resultado.get("edital")
        else:
            # Tentar extrair padrões como PE-001/2026, PE001/2026, 001/2026
            padrao = re.search(r'(PE[-\s]?\d+[/\-]\d{4}|\d+[/\-]\d{4})', message, re.IGNORECASE)
            if padrao:
                edital_numero = padrao.group(1)

        if not edital_numero:
            return (
                "## ❌ Número do Edital Não Informado\n\n"
                "Por favor, informe o número do edital que deseja resumir.\n\n"
                "**Exemplos:**\n"
                "- \"Resuma o edital PE-001/2026\"\n"
                "- \"Resumo do edital 123/2025\"\n"
                "- \"Faça um resumo do edital PE-041/2026\"",
                {"error": "Número do edital não informado"}
            )

        # Buscar edital no banco
        edital = db.query(Edital).filter(
            Edital.user_id == user_id,
            Edital.numero.ilike(f"%{edital_numero.replace('-', '%').replace('/', '%')}%")
        ).first()

        if not edital:
            return (
                f"## ❌ Edital Não Encontrado\n\n"
                f"O edital **{edital_numero}** não foi encontrado no seu cadastro.\n\n"
                "**Dica:** Use \"Liste meus editais\" para ver os editais cadastrados.",
                {"error": f"Edital {edital_numero} não encontrado"}
            )

        # Buscar requisitos do edital
        from models import EditalRequisito
        requisitos = db.query(EditalRequisito).filter(
            EditalRequisito.edital_id == edital.id
        ).all()

        # Montar contexto para o resumo
        contexto = f"""
EDITAL: {edital.numero}
ÓRGÃO: {edital.orgao or 'Não informado'}
OBJETO: {edital.objeto or 'Não informado'}
MODALIDADE: {edital.modalidade or 'Não informada'}
VALOR DE REFERÊNCIA: {f'R$ {edital.valor_referencia:,.2f}' if edital.valor_referencia else 'Não informado'}
DATA DE ABERTURA: {edital.data_abertura.strftime('%d/%m/%Y %H:%M') if edital.data_abertura else 'Não informada'}
DATA DE PUBLICAÇÃO: {edital.data_publicacao.strftime('%d/%m/%Y') if edital.data_publicacao else 'Não informada'}
UF: {edital.uf or 'Não informada'}
CIDADE: {edital.cidade or 'Não informada'}
STATUS: {edital.status or 'novo'}

REQUISITOS ({len(requisitos)} encontrados):
"""
        for req in requisitos[:20]:  # Limitar a 20 requisitos
            obrig = "OBRIGATÓRIO" if req.obrigatorio else "Desejável"
            contexto += f"- [{obrig}] {req.descricao[:200]}\n"

        if len(requisitos) > 20:
            contexto += f"\n... e mais {len(requisitos) - 20} requisitos"

        # Chamar LLM para gerar resumo
        prompt_resumo = f"""Faça um RESUMO EXECUTIVO do seguinte edital de licitação.

O resumo deve ser objetivo e incluir:
1. **Objeto**: O que está sendo licitado (em 1-2 frases)
2. **Valor**: Valor de referência e observações
3. **Prazos**: Data de abertura e prazos importantes
4. **Principais Requisitos**: Os 3-5 requisitos mais importantes/restritivos
5. **Alertas**: Pontos de atenção para participação

DADOS DO EDITAL:
{contexto}

Formate a resposta em Markdown com emojis para facilitar a leitura."""

        messages = [{"role": "user", "content": prompt_resumo}]

        resumo = call_deepseek(messages, max_tokens=2000)

        response_text = f"""## 📋 Resumo do Edital {edital.numero}

{resumo}

---
📊 **Dados do Sistema:**
- Status: {edital.status or 'novo'}
- Requisitos cadastrados: {len(requisitos)}
- URL: {edital.url or 'Não disponível'}
"""

        return response_text, {"success": True, "edital": edital.numero, "requisitos": len(requisitos)}

    except Exception as e:
        return f"## ❌ Erro ao Resumir Edital\n\n{str(e)}", {"error": str(e)}
    finally:
        db.close()


def processar_perguntar_edital(message: str, user_id: str, intencao_resultado: dict = None):
    """
    Processa ação: Responder dúvidas sobre um edital específico.

    Fluxo:
    1. Tenta responder com dados cadastrados no banco
    2. Se LLM indicar que não encontrou a informação, verifica se tem PDF
    3. Se tiver PDF, lê e responde
    4. Se não tiver PDF, pede para o usuário fazer upload

    Returns: (response_text, resultado_dict)
    """
    import re
    import os
    db = get_db()

    try:
        # Extrair número do edital da mensagem
        edital_numero = None
        if intencao_resultado and intencao_resultado.get("edital"):
            edital_numero = intencao_resultado.get("edital")
        else:
            # Tentar extrair padrões como PE-001/2026, PE001/2026, 001/2026
            padrao = re.search(r'(PE[-\s]?\d+[/\-]\d{4}|\d+[/\-]\d{4})', message, re.IGNORECASE)
            if padrao:
                edital_numero = padrao.group(1)

        if not edital_numero:
            return (
                "## ❌ Número do Edital Não Informado\n\n"
                "Por favor, informe o número do edital sobre o qual deseja perguntar.\n\n"
                "**Exemplos:**\n"
                "- \"Qual o prazo de entrega do edital PE-001/2026?\"\n"
                "- \"O edital PE-001/2026 exige garantia?\"\n"
                "- \"Quais documentos são exigidos no PE-041/2026?\"",
                {"error": "Número do edital não informado"}
            )

        # Buscar edital no banco
        edital = db.query(Edital).filter(
            Edital.user_id == user_id,
            Edital.numero.ilike(f"%{edital_numero.replace('-', '%').replace('/', '%')}%")
        ).first()

        if not edital:
            return (
                f"## ❌ Edital Não Encontrado\n\n"
                f"O edital **{edital_numero}** não foi encontrado no seu cadastro.\n\n"
                "**Dica:** Use \"Liste meus editais\" para ver os editais cadastrados.\n"
                f"Ou busque o edital: \"Busque o edital {edital_numero} no PNCP\"",
                {"error": f"Edital {edital_numero} não encontrado"}
            )

        # Buscar requisitos do edital
        from models import EditalRequisito, EditalDocumento
        requisitos = db.query(EditalRequisito).filter(
            EditalRequisito.edital_id == edital.id
        ).all()

        # Buscar documentos PDF do edital
        documentos = db.query(EditalDocumento).filter(
            EditalDocumento.edital_id == edital.id
        ).all()

        # Montar contexto do edital com dados cadastrados
        contexto_banco = f"""
EDITAL: {edital.numero}
ÓRGÃO: {edital.orgao or 'Não informado'}
OBJETO: {edital.objeto or 'Não informado'}
MODALIDADE: {edital.modalidade or 'Não informada'}
VALOR DE REFERÊNCIA: {f'R$ {edital.valor_referencia:,.2f}' if edital.valor_referencia else 'Não informado'}
DATA DE ABERTURA: {edital.data_abertura.strftime('%d/%m/%Y %H:%M') if edital.data_abertura else 'Não informada'}
DATA DE PUBLICAÇÃO: {edital.data_publicacao.strftime('%d/%m/%Y') if edital.data_publicacao else 'Não informada'}
UF: {edital.uf or 'Não informada'}
CIDADE: {edital.cidade or 'Não informada'}
STATUS: {edital.status or 'novo'}

REQUISITOS DO EDITAL ({len(requisitos)} requisitos):
"""
        for req in requisitos:
            obrig = "[OBRIGATÓRIO]" if req.obrigatorio else "[Desejável]"
            tipo = f"({req.tipo})" if req.tipo else ""
            contexto_banco += f"- {obrig} {tipo} {req.descricao}\n"

        if not requisitos:
            contexto_banco += "- Nenhum requisito cadastrado.\n"

        # Detectar se é pergunta específica que provavelmente está no PDF
        msg_lower = message.lower()
        perguntas_especificas = [
            'prazo de entrega', 'prazo entrega', 'dias para entrega', 'entregar em',
            'garantia', 'garantir',
            'local de entrega', 'onde entregar', 'endereço de entrega',
            'forma de pagamento', 'pagamento', 'pagar',
            'penalidade', 'multa', 'sanção', 'sanções',
            'documentos de habilitação', 'habilitação', 'documentos exigidos',
            'qualificação técnica', 'atestado', 'certidão',
            'termo de referência', 'anexo', 'especificação técnica',
            'cláusula', 'item', 'subitem',
            'critério de julgamento', 'desempate', 'lance',
            'contrato', 'vigência', 'aditivo',
            'reajuste', 'índice', 'igpm', 'ipca'
        ]
        eh_pergunta_especifica = any(p in msg_lower for p in perguntas_especificas)

        # Verificar se tem PDF disponível ANTES de tentar banco
        doc_com_texto = None
        for doc in documentos:
            if doc.texto_extraido and len(doc.texto_extraido) > 100:
                doc_com_texto = doc
                break
            elif doc.path_arquivo and os.path.exists(doc.path_arquivo):
                try:
                    from tools import tool_processar_upload
                    texto = tool_processar_upload(doc.path_arquivo)
                    if texto and len(texto) > 100:
                        doc.texto_extraido = texto[:200000]  # Aumentar limite
                        doc.processado = True
                        db.commit()
                        doc_com_texto = doc
                        break
                except Exception as e:
                    print(f"[PERGUNTAR_EDITAL] Erro ao ler PDF {doc.nome_arquivo}: {e}")

        # Se é pergunta específica E tem PDF, ir direto pro PDF
        if eh_pergunta_especifica and doc_com_texto:
            print(f"[PERGUNTAR_EDITAL] Pergunta específica detectada. Indo direto pro PDF...")
        else:
            # PASSO 1: Tentar responder com dados do banco (perguntas gerais)
            prompt_banco = f"""Você é um assistente especializado em licitações públicas.

DADOS CADASTRADOS DO EDITAL:
{contexto_banco}

PERGUNTA DO USUÁRIO:
{message}

INSTRUÇÕES IMPORTANTES:
1. Responda a pergunta usando os dados acima
2. Para perguntas amplas como "tudo sobre", "informações", "detalhes", "resumo": apresente TODOS os dados disponíveis de forma organizada
3. SOMENTE responda "INFORMACAO_NAO_ENCONTRADA_NO_CADASTRO" se a pergunta for sobre algo MUITO ESPECÍFICO que realmente não está nos dados (ex: cláusula específica, anexo, item de planilha)
4. Se há dados relevantes para responder, mesmo que parcialmente, responda com o que tem
5. Seja objetivo e organize a resposta em seções quando apropriado

Responda em Markdown (ou a frase especial APENAS se realmente não houver dados relevantes)."""

            messages = [{"role": "user", "content": prompt_banco}]
            resposta_banco = call_deepseek(messages, max_tokens=2000)

            # Verificar se encontrou a informação no banco
            if "INFORMACAO_NAO_ENCONTRADA" not in resposta_banco.upper() and "NÃO ENCONTRADA" not in resposta_banco.upper() and "NÃO CONSTA" not in resposta_banco.upper():
                # Encontrou no banco - retornar resposta
                response_text = f"""## 💬 Resposta sobre o Edital {edital.numero}

{resposta_banco}

---
📋 **Edital:** {edital.numero}
🏢 **Órgão:** {edital.orgao or 'N/I'}
📊 **Fonte:** Dados cadastrados no sistema
"""
                return response_text, {"success": True, "edital": edital.numero, "fonte": "banco"}

            print(f"[PERGUNTAR_EDITAL] Informação não encontrada no banco. Verificando PDFs...")

        # doc_com_texto já foi verificado antes
        if doc_com_texto and doc_com_texto.texto_extraido:
            # PASSO 3: Tem PDF - ler e responder
            print(f"[PERGUNTAR_EDITAL] Lendo PDF: {doc_com_texto.nome_arquivo}")

            texto_completo = doc_com_texto.texto_extraido
            print(f"[PERGUNTAR_EDITAL] Texto total: {len(texto_completo)} caracteres")

            # Extrair palavras-chave da pergunta para busca inteligente
            msg_lower = message.lower()
            palavras_busca = []

            # Mapeamento de termos comuns em editais
            if any(p in msg_lower for p in ['prazo', 'entrega', 'entregar']):
                palavras_busca.extend(['prazo', 'entrega', 'dias', 'úteis', 'corridos', 'fornecimento'])
            if any(p in msg_lower for p in ['garantia', 'garantir']):
                palavras_busca.extend(['garantia', 'garantir', 'meses', 'anos', 'cobertura'])
            if any(p in msg_lower for p in ['pagamento', 'pagar', 'valor']):
                palavras_busca.extend(['pagamento', 'pagar', 'fatura', 'nota fiscal', 'dias'])
            if any(p in msg_lower for p in ['local', 'onde', 'endereço']):
                palavras_busca.extend(['local', 'entrega', 'endereço', 'sede', 'almoxarifado'])
            if any(p in msg_lower for p in ['documento', 'habilitação', 'exigência', 'exige']):
                palavras_busca.extend(['habilitação', 'documento', 'certidão', 'atestado', 'declaração'])
            if any(p in msg_lower for p in ['penalidade', 'multa', 'sanção']):
                palavras_busca.extend(['penalidade', 'multa', 'sanção', 'advertência', 'suspensão'])

            # Chunkar o documento (chunks de 2000 caracteres com overlap de 300)
            chunk_size = 2000
            overlap = 300
            chunks = []
            for i in range(0, len(texto_completo), chunk_size - overlap):
                chunk = texto_completo[i:i + chunk_size]
                chunks.append((i, chunk))  # (posição, texto)

            print(f"[PERGUNTAR_EDITAL] Documento dividido em {len(chunks)} chunks")

            # Criar frases de busca combinadas (mais específicas)
            frases_exatas = []
            if any(p in msg_lower for p in ['prazo', 'entrega']):
                frases_exatas.extend(['prazo de entrega', 'prazo para entrega', 'entrega do objeto',
                                      'prazo de fornecimento', 'dias para entrega', 'dias após'])
            if any(p in msg_lower for p in ['garantia']):
                frases_exatas.extend(['prazo de garantia', 'garantia de', 'meses de garantia', 'anos de garantia'])
            if any(p in msg_lower for p in ['pagamento']):
                frases_exatas.extend(['prazo de pagamento', 'pagamento será', 'dias após', 'nota fiscal'])
            if any(p in msg_lower for p in ['local']):
                frases_exatas.extend(['local de entrega', 'endereço de entrega', 'entregar no', 'entregue no'])

            # Buscar nos chunks - priorizar frases exatas
            chunks_relevantes = []
            for pos, chunk in chunks:
                chunk_lower = chunk.lower()

                # Score alto para frases exatas
                score_exato = sum(3 for f in frases_exatas if f in chunk_lower)
                # Score menor para palavras individuais
                score_palavras = sum(1 for p in palavras_busca if p in chunk_lower)

                score_total = score_exato + score_palavras
                if score_total > 0:
                    chunks_relevantes.append((score_total, pos, chunk))

            # Ordenar por relevância (maior score primeiro)
            chunks_relevantes.sort(key=lambda x: -x[0])

            # Pegar os 12 chunks mais relevantes
            chunks_relevantes = chunks_relevantes[:12]
            print(f"[PERGUNTAR_EDITAL] {len(chunks_relevantes)} chunks relevantes encontrados")
            if chunks_relevantes:
                print(f"[PERGUNTAR_EDITAL] Top scores: {[c[0] for c in chunks_relevantes[:5]]}")

            # Montar contexto com chunks relevantes
            if chunks_relevantes:
                texto_pdf = f"=== TRECHOS RELEVANTES DO EDITAL ===\n\n"
                for i, (score, pos, chunk) in enumerate(chunks_relevantes, 1):
                    texto_pdf += f"--- Trecho {i} (posição {pos}) ---\n{chunk}\n\n"
            else:
                # Sem chunks relevantes, pegar documento inteiro até limite
                texto_pdf = texto_completo[:60000]
                print(f"[PERGUNTAR_EDITAL] Sem chunks relevantes, usando primeiros 60K caracteres")

            prompt_pdf = f"""Você é um assistente especializado em licitações públicas brasileiras.

CONTEÚDO DO EDITAL (extraído do PDF "{doc_com_texto.nome_arquivo}"):
{texto_pdf}

PERGUNTA DO USUÁRIO:
{message}

INSTRUÇÕES:
1. Responda a pergunta com base no conteúdo do edital acima
2. CITE O TRECHO EXATO do edital que contém a resposta (entre aspas)
3. Se a informação estiver em um Anexo ou Termo de Referência, indique qual
4. Se não encontrar a informação específica, diga claramente e sugira onde pode estar
5. Seja objetivo e direto

Responda em Markdown."""

            messages_pdf = [{"role": "user", "content": prompt_pdf}]
            resposta_pdf = call_deepseek(messages_pdf, max_tokens=3000)

            response_text = f"""## 💬 Resposta sobre o Edital {edital.numero}

{resposta_pdf}

---
📋 **Edital:** {edital.numero}
🏢 **Órgão:** {edital.orgao or 'N/I'}
📄 **Fonte:** PDF do edital ({doc_com_texto.nome_arquivo})
"""
            return response_text, {"success": True, "edital": edital.numero, "fonte": "pdf", "arquivo": doc_com_texto.nome_arquivo}

        # PASSO 4: Não tem PDF - pedir upload
        print(f"[PERGUNTAR_EDITAL] Nenhum PDF encontrado para o edital {edital.numero}")

        response_text = f"""## ⚠️ Informação Não Disponível

A informação solicitada **não foi encontrada** nos dados cadastrados do edital **{edital.numero}**.

Para responder sua pergunta, preciso do **PDF do edital**.

### 📤 Como fazer:
1. Faça upload do PDF do edital (arraste ou clique no 📎)
2. Junto com o arquivo, envie sua pergunta novamente

**Exemplo:** Envie o PDF e escreva:
> "Qual o prazo de entrega exigido neste edital?"

---
📋 **Edital:** {edital.numero}
🏢 **Órgão:** {edital.orgao or 'N/I'}
📊 **Dados no sistema:** {len(requisitos)} requisitos cadastrados
📄 **PDFs salvos:** Nenhum
"""
        return response_text, {
            "success": False,
            "edital": edital.numero,
            "precisa_upload": True,
            "mensagem": "PDF do edital necessário para responder esta pergunta"
        }

    except Exception as e:
        return f"## ❌ Erro ao Processar Pergunta\n\n{str(e)}", {"error": str(e)}
    finally:
        db.close()


def processar_baixar_pdf_edital(message: str, user_id: str, intencao_resultado: dict = None):
    """
    Processa ação: Baixar o PDF de um edital cadastrado.

    Fluxo:
    1. Identifica o edital pelo número
    2. Verifica se já tem PDF salvo
    3. Se não tem, usa a URL cadastrada para baixar
    4. Extrai texto do PDF e salva na tabela editais_documentos

    Returns: (response_text, resultado_dict)
    """
    import re
    import os
    import requests
    from datetime import datetime
    db = get_db()

    try:
        # Extrair número do edital da mensagem
        edital_numero = None
        if intencao_resultado and intencao_resultado.get("edital"):
            edital_numero = intencao_resultado.get("edital")
        else:
            # Tentar extrair padrões como PE-001/2026, PE001/2026, 001/2026
            padrao = re.search(r'(PE[-\s]?\d+[/\-]\d{4}|\d+[/\-]\d{4})', message, re.IGNORECASE)
            if padrao:
                edital_numero = padrao.group(1)

        if not edital_numero:
            return (
                "## ❌ Número do Edital Não Informado\n\n"
                "Por favor, informe o número do edital que deseja baixar.\n\n"
                "**Exemplos:**\n"
                "- \"Baixe o PDF do edital PE-001/2026\"\n"
                "- \"Faça download do edital 90006/2026\"\n"
                "- \"Baixar edital PE-041/2026\"",
                {"error": "Número do edital não informado"}
            )

        # Buscar edital no banco
        edital = db.query(Edital).filter(
            Edital.user_id == user_id,
            Edital.numero.ilike(f"%{edital_numero.replace('-', '%').replace('/', '%')}%")
        ).first()

        if not edital:
            return (
                f"## ❌ Edital Não Encontrado\n\n"
                f"O edital **{edital_numero}** não foi encontrado no seu cadastro.\n\n"
                "**Dica:** Use \"Liste meus editais\" para ver os editais cadastrados.",
                {"error": f"Edital {edital_numero} não encontrado"}
            )

        # Verificar se já tem documento salvo
        from models import EditalDocumento
        doc_existente = db.query(EditalDocumento).filter(
            EditalDocumento.edital_id == edital.id
        ).first()

        # Verificar se existe registro E se o arquivo físico ainda existe
        arquivo_existe = doc_existente and doc_existente.path_arquivo and os.path.exists(doc_existente.path_arquivo)

        if doc_existente and doc_existente.texto_extraido and len(doc_existente.texto_extraido) > 100 and arquivo_existe:
            return (
                f"## ✅ PDF Já Disponível\n\n"
                f"O edital **{edital.numero}** já possui PDF baixado e processado.\n\n"
                f"📄 **Arquivo:** {doc_existente.nome_arquivo}\n"
                f"📝 **Texto extraído:** {len(doc_existente.texto_extraido):,} caracteres\n\n"
                "Você pode fazer perguntas sobre o edital:\n"
                f"- \"Quais itens o edital {edital.numero} comporta?\"\n"
                f"- \"Qual o prazo de entrega do edital {edital.numero}?\"",
                {"success": True, "edital": edital.numero, "ja_existia": True, "arquivo": doc_existente.nome_arquivo, "edital_id": edital.id}
            )

        # Se o registro existe mas arquivo foi apagado, deletar o registro para recriar
        if doc_existente and not arquivo_existe:
            db.delete(doc_existente)
            db.commit()
            doc_existente = None  # Permitir re-download

        # ========== MÉTODO 1: Tentar baixar via API do PNCP (se tiver dados) ==========
        if edital.cnpj_orgao and edital.ano_compra and edital.seq_compra:
            from tools import tool_buscar_arquivos_edital_pncp, tool_baixar_pdf_pncp

            response_text = f"## ⏳ Baixando PDF do Edital {edital.numero}...\n\n"
            response_text += f"🔍 **Fonte:** API do PNCP\n"
            response_text += f"📌 **CNPJ:** {edital.cnpj_orgao} | **Ano:** {edital.ano_compra} | **Seq:** {edital.seq_compra}\n\n"

            # Buscar lista de arquivos
            arquivos_result = tool_buscar_arquivos_edital_pncp(edital_id=edital.id, user_id=user_id)

            if arquivos_result.get('success') and arquivos_result.get('arquivos'):
                arquivos = arquivos_result['arquivos']
                arquivo_edital = arquivos_result.get('arquivo_edital') or arquivos[0]

                response_text += f"📁 **Arquivos encontrados:** {len(arquivos)}\n"
                for arq in arquivos:
                    response_text += f"   - {arq['titulo']}\n"
                response_text += f"\n📥 **Baixando:** {arquivo_edital['titulo']}...\n\n"

                # Baixar o arquivo principal
                download_result = tool_baixar_pdf_pncp(
                    cnpj=edital.cnpj_orgao,
                    ano=edital.ano_compra,
                    seq=edital.seq_compra,
                    sequencial_arquivo=arquivo_edital['sequencial'],
                    user_id=user_id,
                    edital_id=edital.id
                )

                if download_result.get('success'):
                    filepath = download_result['filepath']
                    filename = download_result['filename']
                    filesize = download_result['filesize']

                    response_text += f"✅ **Download concluído:** {filename} ({filesize/1024:.1f} KB)\n\n"

                    # Extrair texto do PDF
                    texto_extraido = ""
                    try:
                        from PyPDF2 import PdfReader
                        reader = PdfReader(filepath)
                        for page in reader.pages:
                            texto_extraido += page.extract_text() or ""
                        response_text += f"📝 **Texto extraído:** {len(texto_extraido):,} caracteres\n\n"
                    except Exception as e:
                        response_text += f"⚠️ **Aviso:** Não foi possível extrair texto: {str(e)}\n\n"

                    # Salvar no banco
                    novo_doc = EditalDocumento(
                        id=str(uuid.uuid4()),
                        edital_id=edital.id,
                        tipo='edital_principal',
                        nome_arquivo=filename,
                        path_arquivo=filepath,
                        texto_extraido=texto_extraido[:100000] if texto_extraido else None,
                        processado=True,
                        created_at=datetime.now()
                    )
                    db.add(novo_doc)
                    db.commit()

                    response_text += "## ✅ PDF Salvo com Sucesso!\n\n"

                    # Extrair requisitos automaticamente (incluindo documentais/certidões)
                    requisitos_extraidos = 0
                    if texto_extraido and len(texto_extraido) > 200:
                        try:
                            resultado_req = tool_extrair_requisitos(edital.id, texto_extraido, user_id)
                            requisitos_extraidos = resultado_req.get("requisitos_extraidos", 0)
                            if requisitos_extraidos > 0:
                                response_text += f"📋 **Requisitos extraídos:** {requisitos_extraidos} (técnicos, documentais e comerciais)\n\n"
                        except Exception as e_req:
                            response_text += f"⚠️ Não foi possível extrair requisitos automaticamente: {str(e_req)}\n\n"

                    response_text += "Agora você pode fazer perguntas sobre o edital:\n"
                    response_text += f"- \"Quais itens o edital {edital.numero} comporta?\"\n"
                    response_text += f"- \"Qual o local de entrega do edital {edital.numero}?\"\n"

                    return response_text, {
                        "success": True,
                        "edital": edital.numero,
                        "arquivo": filename,
                        "tamanho": filesize,
                        "texto_extraido": len(texto_extraido),
                        "requisitos_extraidos": requisitos_extraidos,
                        "fonte": "PNCP"
                    }
                else:
                    response_text += f"⚠️ **Erro no download via PNCP:** {download_result.get('error', 'Erro desconhecido')}\n\n"
                    response_text += "Tentando método alternativo...\n\n"
            else:
                response_text = f"## ⏳ Baixando PDF do Edital {edital.numero}...\n\n"
                response_text += f"⚠️ **PNCP:** Nenhum arquivo encontrado via API\n"
                response_text += "Tentando método alternativo...\n\n"

        # ========== MÉTODO 2: Tentar baixar da URL cadastrada ==========

        # Verificar se tem URL do edital
        if not edital.url:
            return (
                f"## ⚠️ URL do Edital Não Cadastrada\n\n"
                f"O edital **{edital.numero}** não possui URL cadastrada para download.\n\n"
                "**Opções:**\n"
                f"1. Atualize o edital com a URL: \"Atualize o edital {edital.numero} com URL: [URL_DO_PDF]\"\n"
                "2. Faça upload manual do PDF (arraste o arquivo para o chat)",
                {"error": "URL não cadastrada", "edital": edital.numero}
            )

        # Tentar baixar o PDF da URL
        if 'response_text' not in locals():
            response_text = f"## ⏳ Baixando PDF do Edital {edital.numero}...\n\n"
        response_text += f"🔗 **URL:** {edital.url}\n\n"

        try:
            # Baixar o arquivo
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(edital.url, headers=headers, timeout=60, allow_redirects=True)
            response.raise_for_status()

            # Determinar nome do arquivo
            content_type = response.headers.get('Content-Type', '')
            filename = f"edital_{edital.numero.replace('/', '_').replace('-', '_')}"

            if 'pdf' in content_type.lower() or edital.url.lower().endswith('.pdf'):
                filename += ".pdf"
            elif 'html' in content_type.lower():
                # É uma página HTML, não PDF direto - tentar encontrar link do PDF
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(response.content, 'html.parser')

                # Procurar links de PDF na página
                pdf_links = []
                for link in soup.find_all('a', href=True):
                    href = link['href']
                    if '.pdf' in href.lower():
                        # Converter para URL absoluta se necessário
                        if href.startswith('/'):
                            from urllib.parse import urlparse
                            parsed = urlparse(edital.url)
                            href = f"{parsed.scheme}://{parsed.netloc}{href}"
                        elif not href.startswith('http'):
                            href = edital.url.rsplit('/', 1)[0] + '/' + href
                        pdf_links.append(href)

                if pdf_links:
                    # Verificar se a URL cadastrada parece ser genérica (página inicial de portal)
                    url_generica = any(x in edital.url.lower() for x in [
                        '/pt-br', '/home', 'compras.gov', 'pncp.gov', 'bec.sp.gov'
                    ]) and edital.url.count('/') <= 4

                    if url_generica:
                        # URL é genérica demais - pedir URL específica
                        return (
                            f"## ⚠️ URL Genérica Cadastrada\n\n"
                            f"A URL cadastrada para o edital **{edital.numero}** parece ser a página inicial do portal:\n"
                            f"`{edital.url}`\n\n"
                            "Essa URL não aponta diretamente para o edital.\n\n"
                            "**O que fazer:**\n"
                            f"1. Acesse o portal e busque pelo edital {edital.numero}\n"
                            "2. Copie a URL da página específica do edital (ou do PDF)\n"
                            "3. Atualize com a URL correta:\n"
                            f"   `Atualize o edital {edital.numero} com URL: [URL_DO_EDITAL]`\n\n"
                            "Ou faça upload manual do PDF.",
                            {"error": "URL genérica", "edital": edital.numero, "url_atual": edital.url}
                        )

                    # Filtrar links relevantes ao edital
                    # Extrair apenas números do edital para comparação
                    numero_limpo = re.sub(r'[^\d]', '', edital.numero)

                    # Prioridade 1: Links que contenham o número do edital
                    links_com_numero = [l for l in pdf_links if numero_limpo in l]

                    # Prioridade 2: Links com palavras-chave do edital
                    palavras_edital = ['edital', 'pregao', 'pregão', 'pe-', 'pe_', 'licitacao', 'licitação',
                                       'termo_referencia', 'termo-referencia', 'tr_', 'tr-']
                    links_com_palavra = [l for l in pdf_links if any(p in l.lower() for p in palavras_edital)]

                    # Prioridade 3: Excluir links claramente não relacionados
                    palavras_excluir = ['politica', 'policy', 'manual', 'regulamento', 'instrucao', 'normativa',
                                        'template', 'modelo', 'anexo_unico', 'formulario', 'cadastro']
                    links_filtrados = [l for l in pdf_links if not any(p in l.lower() for p in palavras_excluir)]

                    # Escolher o melhor link
                    pdf_url = None
                    if links_com_numero:
                        pdf_url = links_com_numero[0]
                        response_text += f"🔍 **URL original era HTML.** Encontrado PDF com número do edital:\n`{pdf_url}`\n\n"
                    elif links_com_palavra:
                        pdf_url = links_com_palavra[0]
                        response_text += f"🔍 **URL original era HTML.** Encontrado PDF de edital:\n`{pdf_url}`\n\n"
                    elif links_filtrados:
                        pdf_url = links_filtrados[0]
                        response_text += f"🔍 **URL original era HTML.** Encontrado possível PDF:\n`{pdf_url}`\n\n"
                    else:
                        # Todos os links parecem não relacionados - listar para o usuário
                        links_lista = "\n".join([f"- `{l}`" for l in pdf_links[:5]])
                        return (
                            f"## ⚠️ Nenhum PDF do Edital Encontrado\n\n"
                            f"A página do edital **{edital.numero}** contém PDFs, mas nenhum parece ser o edital:\n\n"
                            f"{links_lista}\n\n"
                            "**O que fazer:**\n"
                            "1. Acesse a URL no navegador e encontre o PDF correto\n"
                            "2. Atualize com a URL direta do PDF:\n"
                            f"   `Atualize o edital {edital.numero} com URL: [URL_DO_PDF]`\n\n"
                            "Ou faça upload manual do PDF.",
                            {"error": "PDF não identificado", "edital": edital.numero, "links_encontrados": pdf_links[:5]}
                        )

                    # Baixar o PDF encontrado
                    response_pdf = requests.get(pdf_url, headers=headers, timeout=60, allow_redirects=True)
                    response_pdf.raise_for_status()
                    response = response_pdf  # Substituir o response pelo PDF
                    filename += ".pdf"
                else:
                    # Não encontrou PDF - informar usuário
                    return (
                        f"## ⚠️ URL Não Contém PDF Direto\n\n"
                        f"A URL cadastrada para o edital **{edital.numero}** aponta para uma página HTML, "
                        f"não para o arquivo PDF:\n`{edital.url}`\n\n"
                        "**O que fazer:**\n"
                        "1. Acesse a URL acima no navegador\n"
                        "2. Encontre o link do PDF do edital\n"
                        "3. Atualize com a URL correta:\n"
                        f"   `Atualize o edital {edital.numero} com URL: [URL_DO_PDF]`\n\n"
                        "Ou faça upload manual do PDF (arraste o arquivo para o chat).",
                        {"error": "URL não é PDF direto", "edital": edital.numero, "url_atual": edital.url}
                    )
            else:
                filename += ".pdf"

            # Salvar arquivo
            upload_dir = os.path.join(os.path.dirname(__file__), 'uploads', user_id, 'editais')
            os.makedirs(upload_dir, exist_ok=True)
            filepath = os.path.join(upload_dir, filename)

            with open(filepath, 'wb') as f:
                f.write(response.content)

            filesize = len(response.content)
            response_text += f"✅ **Download concluído:** {filename} ({filesize/1024:.1f} KB)\n\n"

            # Extrair texto do PDF
            texto_extraido = ""
            try:
                if filename.endswith('.pdf'):
                    from PyPDF2 import PdfReader
                    reader = PdfReader(filepath)
                    for page in reader.pages:
                        texto_extraido += page.extract_text() or ""
                elif filename.endswith('.html'):
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(response.content, 'html.parser')
                    texto_extraido = soup.get_text(separator='\n', strip=True)

                response_text += f"📝 **Texto extraído:** {len(texto_extraido):,} caracteres\n\n"
            except Exception as e:
                response_text += f"⚠️ **Aviso:** Não foi possível extrair texto: {str(e)}\n\n"

            # Salvar no banco
            if doc_existente:
                doc_existente.path_arquivo = filepath
                doc_existente.nome_arquivo = filename
                doc_existente.texto_extraido = texto_extraido[:100000] if texto_extraido else None
                doc_existente.processado = True
            else:
                novo_doc = EditalDocumento(
                    id=str(uuid.uuid4()),
                    edital_id=edital.id,
                    tipo='edital_principal',  # Valores: edital_principal, termo_referencia, anexo, planilha, outro
                    nome_arquivo=filename,
                    path_arquivo=filepath,
                    texto_extraido=texto_extraido[:100000] if texto_extraido else None,
                    processado=True,
                    created_at=datetime.now()
                )
                db.add(novo_doc)

            db.commit()

            response_text += "## ✅ PDF Salvo com Sucesso!\n\n"

            # Extrair requisitos automaticamente (incluindo documentais/certidões)
            requisitos_extraidos = 0
            if texto_extraido and len(texto_extraido) > 200:
                try:
                    resultado_req = tool_extrair_requisitos(edital.id, texto_extraido, user_id)
                    requisitos_extraidos = resultado_req.get("requisitos_extraidos", 0)
                    if requisitos_extraidos > 0:
                        response_text += f"📋 **Requisitos extraídos:** {requisitos_extraidos} (técnicos, documentais e comerciais)\n\n"
                except Exception as e_req:
                    response_text += f"⚠️ Não foi possível extrair requisitos automaticamente: {str(e_req)}\n\n"

            response_text += "Agora você pode fazer perguntas sobre o edital:\n"
            response_text += f"- \"Quais itens o edital {edital.numero} comporta?\"\n"
            response_text += f"- \"Qual o local de entrega do edital {edital.numero}?\"\n"
            response_text += f"- \"Me conte tudo sobre o edital {edital.numero}\"\n"

            return response_text, {
                "success": True,
                "edital": edital.numero,
                "arquivo": filename,
                "tamanho": filesize,
                "texto_extraido": len(texto_extraido),
                "requisitos_extraidos": requisitos_extraidos
            }

        except requests.exceptions.RequestException as e:
            return (
                f"## ❌ Erro ao Baixar PDF\n\n"
                f"Não foi possível baixar o arquivo da URL:\n`{edital.url}`\n\n"
                f"**Erro:** {str(e)}\n\n"
                "**Opções:**\n"
                "1. Verifique se a URL está correta\n"
                "2. Tente acessar a URL manualmente e baixar o PDF\n"
                "3. Faça upload manual do PDF (arraste o arquivo para o chat)",
                {"error": str(e), "edital": edital.numero}
            )

    except Exception as e:
        return f"## ❌ Erro ao Processar Download\n\n{str(e)}", {"error": str(e)}
    finally:
        db.close()


def processar_atualizar_url_edital(message: str, user_id: str, intencao_resultado: dict = None):
    """Atualiza a URL de um edital cadastrado"""
    db = get_db()
    try:
        import re

        # Extrair número do edital da mensagem
        # Padrões: PE-001/2026, 02223/2025, PE001, etc.
        edital_numero = None
        if intencao_resultado and intencao_resultado.get("edital"):
            edital_numero = intencao_resultado["edital"]
        else:
            # Tentar extrair da mensagem
            patterns = [
                r'edital\s+([A-Za-z]{0,3}[-]?\d+[/]\d{4})',  # edital PE-001/2026 ou 02223/2025
                r'edital\s+([A-Za-z]{2,3}[-]?\d+)',  # edital PE-001 ou PE001
                r'([A-Za-z]{2,3}[-]\d+[/]\d{4})',  # PE-001/2026
                r'(\d{4,}[/]\d{4})',  # 02223/2025
            ]
            for pattern in patterns:
                match = re.search(pattern, message, re.IGNORECASE)
                if match:
                    edital_numero = match.group(1)
                    break

        if not edital_numero:
            return (
                "## ⚠️ Número do Edital Não Informado\n\n"
                "Por favor, informe o número do edital que deseja atualizar.\n\n"
                "**Formato:**\n"
                "`Atualize o edital PE-001/2026 com URL: https://exemplo.com/edital.pdf`",
                {"error": "Número do edital não informado"}
            )

        # Extrair a nova URL da mensagem
        url_match = re.search(r'(https?://[^\s<>"]+)', message)
        if not url_match:
            return (
                f"## ⚠️ URL Não Informada\n\n"
                f"Por favor, informe a nova URL para o edital **{edital_numero}**.\n\n"
                "**Formato:**\n"
                f"`Atualize o edital {edital_numero} com URL: https://exemplo.com/edital.pdf`",
                {"error": "URL não informada", "edital": edital_numero}
            )

        nova_url = url_match.group(1)

        # Buscar edital no banco
        editais = db.query(Edital).filter(Edital.user_id == user_id).all()

        # Normalizar número para comparação
        numero_busca = edital_numero.replace('-', '').replace('/', '').upper()

        edital = None
        for e in editais:
            num_edital = e.numero.replace('-', '').replace('/', '').upper()
            if num_edital == numero_busca or numero_busca in num_edital or num_edital in numero_busca:
                edital = e
                break

        if not edital:
            return (
                f"## ❌ Edital Não Encontrado\n\n"
                f"O edital **{edital_numero}** não foi encontrado no seu cadastro.\n\n"
                "**Dica:** Use \"Liste meus editais\" para ver os editais cadastrados.",
                {"error": f"Edital {edital_numero} não encontrado"}
            )

        # Atualizar a URL
        url_antiga = edital.url
        edital.url = nova_url
        db.commit()

        response_text = f"## ✅ URL Atualizada com Sucesso!\n\n"
        response_text += f"**Edital:** {edital.numero}\n"
        response_text += f"**Órgão:** {edital.orgao or 'N/A'}\n\n"

        if url_antiga:
            response_text += f"**URL anterior:** `{url_antiga}`\n"
        response_text += f"**Nova URL:** `{nova_url}`\n\n"

        response_text += "Agora você pode baixar o PDF:\n"
        response_text += f"- `Baixe o PDF do edital {edital.numero}`"

        return response_text, {
            "success": True,
            "edital": edital.numero,
            "edital_id": edital.id,
            "url_antiga": url_antiga,
            "url_nova": nova_url
        }

    except Exception as e:
        return f"## ❌ Erro ao Atualizar URL\n\n{str(e)}", {"error": str(e)}
    finally:
        db.close()


# =============================================================================
# FASE 1: PRECIFICAÇÃO — Handlers (UC-P01 a UC-P10)
# =============================================================================

def _extrair_id_da_mensagem(message: str, prefixo: str = "edital") -> str:
    """Extrai UUID ou referência de ID de uma mensagem de chat."""
    import re
    # UUID
    m = re.search(r'[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}', message, re.I)
    if m:
        return m.group(0)
    return None


def processar_precif_organizar_lotes(message: str, user_id: str, empresa_id: str = None):
    """UC-P01: Organizar edital por lotes"""
    from tools import tool_organizar_lotes
    from models import get_db, Edital
    db = get_db()
    try:
        edital_id = _extrair_id_da_mensagem(message)
        if not edital_id:
            # Pegar último edital do usuário
            edital = db.query(Edital).filter(Edital.user_id == user_id).order_by(Edital.created_at.desc()).first()
            edital_id = edital.id if edital else None
        if not edital_id:
            return "❌ Nenhum edital encontrado. Salve um edital primeiro.", None

        resultado = tool_organizar_lotes(edital_id=edital_id, user_id=user_id, empresa_id=empresa_id, importar_pncp=True)
        if not resultado.get("success"):
            return f"❌ {resultado.get('error')}", resultado

        lotes = resultado.get("lotes", [])
        total = resultado.get("total_itens", 0)
        txt = f"## ✅ Lotes Organizados\n\n{resultado.get('mensagem')}\n\n"
        txt += f"**Total de itens:** {total}\n\n"
        for l in lotes:
            txt += f"- **Lote {l['numero_lote']}** — {l.get('nome', 'N/A')} | Status: {l.get('status')} | Valor est.: R$ {l.get('valor_estimado', 0):,.2f}\n"

        return txt, resultado
    finally:
        db.close()


def processar_precif_selecao_portfolio(message: str, user_id: str, empresa_id: str = None):
    """UC-P02: Seleção inteligente de portfolio"""
    from tools import tool_selecao_portfolio
    edital_item_id = _extrair_id_da_mensagem(message)
    if not edital_item_id:
        return "❌ Informe o ID do item do edital para seleção de portfolio.", None

    resultado = tool_selecao_portfolio(edital_item_id=edital_item_id, user_id=user_id, empresa_id=empresa_id)
    if not resultado.get("success"):
        return f"❌ {resultado.get('error')}", resultado

    sugestoes = resultado.get("sugestoes", [])
    txt = f"## 🎯 Seleção de Portfolio\n\n{resultado.get('mensagem')}\n\n"
    if sugestoes:
        txt += "| # | Produto | Fabricante | Match | Preço Ref. |\n|---|---------|-----------|-------|------------|\n"
        for i, s in enumerate(sugestoes[:10], 1):
            txt += f"| {i} | {s['nome'][:30]} | {s.get('fabricante', '-')[:20]} | {s['match_score']}% | R$ {s.get('preco_referencia', 0) or 0:,.2f} |\n"
    else:
        txt += "Nenhum produto com match encontrado. Cadastre produtos no portfolio.\n"

    return txt, resultado


def processar_precif_calcular_volumetria(message: str, user_id: str):
    """UC-P03: Cálculo técnico de volumetria"""
    from tools import tool_calcular_volumetria
    import re

    vinculo_id = _extrair_id_da_mensagem(message)
    if not vinculo_id:
        return "❌ Informe o ID do vínculo item-produto para calcular volumetria.", None

    # Extrair números da mensagem
    nums = re.findall(r'[\d.]+', message)
    volume = float(nums[0]) if len(nums) > 0 else None
    rendimento = float(nums[1]) if len(nums) > 1 else None

    resultado = tool_calcular_volumetria(
        edital_item_produto_id=vinculo_id, user_id=user_id,
        volume_edital=volume, rendimento_produto=rendimento
    )
    if not resultado.get("success"):
        return f"❌ {resultado.get('error')}", resultado

    txt = f"## 📊 Volumetria Calculada\n\n"
    txt += f"**Fórmula:** {resultado.get('formula')}\n\n"
    txt += f"- Volume do edital: {resultado.get('volume_edital'):,.0f}\n"
    txt += f"- Rendimento: {resultado.get('rendimento_produto'):,.0f} por kit\n"
    txt += f"- Volume real ajustado: {resultado.get('volume_real_ajustado'):,.0f}\n"
    txt += f"- **Quantidade de kits: {resultado.get('quantidade_kits')}**\n"

    return txt, resultado


def processar_precif_configurar_custos(message: str, user_id: str, empresa_id: str = None):
    """UC-P04: Configurar base de custos (Camada A)"""
    from tools import tool_configurar_custos
    import re

    vinculo_id = _extrair_id_da_mensagem(message)
    if not vinculo_id:
        return "❌ Informe o ID do vínculo item-produto para configurar custos.", None

    # Extrair custo da mensagem
    m = re.search(r'R\$\s*([\d.,]+)', message)
    custo = float(m.group(1).replace('.', '').replace(',', '.')) if m else None

    resultado = tool_configurar_custos(
        edital_item_produto_id=vinculo_id, user_id=user_id,
        custo_unitario=custo, empresa_id=empresa_id
    )
    if not resultado.get("success"):
        return f"❌ {resultado.get('error')}", resultado

    txt = f"## 💰 Base de Custos (Camada A)\n\n{resultado.get('mensagem')}\n\n"
    txt += f"- Custo unitário: R$ {resultado.get('custo_unitario', 0):,.2f}\n"
    txt += f"- NCM: {resultado.get('ncm', 'N/A')}\n"
    txt += f"- ICMS: {resultado.get('icms', 0)}%" + (" (ISENTO)" if resultado.get('isencao_icms') else "") + "\n"
    txt += f"- IPI: {resultado.get('ipi', 0)}% | PIS+COFINS: {resultado.get('pis_cofins', 0)}%\n"
    txt += f"- **Custo base final: R$ {resultado.get('custo_base_final', 0):,.2f}**\n"

    return txt, resultado


def processar_precif_preco_base(message: str, user_id: str):
    """UC-P05: Montar preço base (Camada B)"""
    from tools import tool_montar_preco_base
    import re

    vinculo_id = _extrair_id_da_mensagem(message)
    if not vinculo_id:
        return "❌ Informe o ID do vínculo item-produto.", None

    # Detectar modo
    modo = 'manual'
    if 'markup' in message.lower():
        modo = 'markup'
    elif 'upload' in message.lower():
        modo = 'upload'

    preco = None
    markup = None
    m = re.search(r'R\$\s*([\d.,]+)', message)
    if m:
        preco = float(m.group(1).replace('.', '').replace(',', '.'))
    m2 = re.search(r'(\d+[.,]?\d*)\s*%', message)
    if m2 and modo == 'markup':
        markup = float(m2.group(1).replace(',', '.'))

    resultado = tool_montar_preco_base(
        edital_item_produto_id=vinculo_id, user_id=user_id,
        modo=modo, preco_base=preco, markup_percentual=markup
    )
    if not resultado.get("success"):
        return f"❌ {resultado.get('error')}", resultado

    txt = f"## 📈 Preço Base (Camada B)\n\n{resultado.get('mensagem')}\n\n"
    txt += f"- Custo base (A): R$ {resultado.get('custo_base', 0):,.2f}\n"
    txt += f"- **Preço base (B): R$ {resultado.get('preco_base', 0):,.2f}**\n"
    if resultado.get('markup_percentual'):
        txt += f"- Markup: {resultado.get('markup_percentual'):.1f}%\n"

    return txt, resultado


def processar_precif_valor_referencia(message: str, user_id: str):
    """UC-P06: Definir valor de referência (Camada C)"""
    from tools import tool_definir_referencia
    import re

    vinculo_id = _extrair_id_da_mensagem(message)
    if not vinculo_id:
        return "❌ Informe o ID do vínculo item-produto.", None

    valor_ref = None
    pct = None
    m = re.search(r'R\$\s*([\d.,]+)', message)
    if m:
        valor_ref = float(m.group(1).replace('.', '').replace(',', '.'))
    m2 = re.search(r'(\d+[.,]?\d*)\s*%', message)
    if m2:
        pct = float(m2.group(1).replace(',', '.'))

    resultado = tool_definir_referencia(
        edital_item_produto_id=vinculo_id, user_id=user_id,
        valor_referencia=valor_ref, percentual_sobre_base=pct
    )
    if not resultado.get("success"):
        return f"❌ {resultado.get('error')}", resultado

    comp = resultado.get("comparativo", {})
    txt = f"## 🎯 Valor de Referência (Camada C)\n\n{resultado.get('mensagem')}\n\n"
    txt += f"**Comparativo:**\n"
    txt += f"- Custo (A): R$ {comp.get('custo_base_A', 0):,.2f}\n"
    txt += f"- Preço Base (B): R$ {comp.get('preco_base_B', 0):,.2f}\n"
    txt += f"- **Target (C): R$ {comp.get('target_C', 0):,.2f}**\n"
    if resultado.get('margem_sobre_custo'):
        txt += f"- Margem sobre custo: {resultado['margem_sobre_custo']:.1f}%\n"

    return txt, resultado


def processar_precif_estruturar_lances(message: str, user_id: str):
    """UC-P07: Estruturar lances (Camadas D e E)"""
    from tools import tool_estruturar_lances
    import re

    vinculo_id = _extrair_id_da_mensagem(message)
    if not vinculo_id:
        return "❌ Informe o ID do vínculo item-produto.", None

    # Extrair valores
    valores = re.findall(r'R\$\s*([\d.,]+)', message)
    lance_ini = float(valores[0].replace('.', '').replace(',', '.')) if len(valores) > 0 else None
    lance_min = float(valores[1].replace('.', '').replace(',', '.')) if len(valores) > 1 else None

    resultado = tool_estruturar_lances(
        edital_item_produto_id=vinculo_id, user_id=user_id,
        lance_inicial=lance_ini, lance_minimo=lance_min
    )
    if not resultado.get("success"):
        return f"❌ {resultado.get('error')}", resultado

    txt = f"## ⚡ Estrutura de Lances (Camadas D+E)\n\n{resultado.get('mensagem')}\n\n"
    txt += f"- Custo base: R$ {resultado.get('custo_base', 0):,.2f}\n"
    txt += f"- Lance inicial (D): R$ {resultado.get('lance_inicial', 0):,.2f}\n"
    txt += f"- Lance mínimo (E): R$ {resultado.get('lance_minimo', 0):,.2f}\n"
    txt += f"- **Faixa de disputa:** {resultado.get('faixa_disputa')}\n"
    if resultado.get('margem_minima') is not None:
        txt += f"- Margem mínima: {resultado['margem_minima']:.1f}%\n"
    for w in resultado.get('warnings', []):
        txt += f"\n{w}\n"

    return txt, resultado


def processar_precif_estrategia(message: str, user_id: str, empresa_id: str = None):
    """UC-P08: Estratégia competitiva"""
    from tools import tool_estrategia_competitiva
    from models import get_db, Edital

    edital_id = _extrair_id_da_mensagem(message)
    if not edital_id:
        db = get_db()
        try:
            edital = db.query(Edital).filter(Edital.user_id == user_id).order_by(Edital.created_at.desc()).first()
            edital_id = edital.id if edital else None
        finally:
            db.close()
    if not edital_id:
        return "❌ Nenhum edital encontrado.", None

    perfil = 'quero_ganhar'
    msg_lower = message.lower()
    if 'não ganhei' in msg_lower or 'nao ganhei' in msg_lower or 'mínimo' in msg_lower or 'minimo' in msg_lower:
        perfil = 'nao_ganhei_minimo'

    resultado = tool_estrategia_competitiva(edital_id=edital_id, user_id=user_id, perfil=perfil, empresa_id=empresa_id)
    if not resultado.get("success"):
        return f"❌ {resultado.get('error')}", resultado

    txt = f"## 🏆 Estratégia Competitiva\n\n{resultado.get('mensagem')}\n\n"
    txt += f"**Perfil:** {'QUERO GANHAR' if perfil == 'quero_ganhar' else 'NÃO GANHEI NO MÍNIMO'}\n\n"
    for c in resultado.get("cenarios", []):
        txt += f"**Item {c.get('item', '?')}:**\n"
        for key in ['cenario_1', 'cenario_2', 'cenario_3']:
            cen = c.get(key, {})
            txt += f"  - {cen.get('label')}: R$ {cen.get('valor', 0):,.2f} (margem: {cen.get('margem', 0):.1f}%)\n"
        txt += "\n"

    return txt, resultado


def processar_precif_historico_camada_f(message: str, user_id: str):
    """UC-P09: Histórico de preços (Camada F)"""
    from tools import tool_historico_precos_camada_f
    import re

    vinculo_id = _extrair_id_da_mensagem(message)
    # Extrair termo de busca
    termo = None
    m = re.search(r'(?:de|para|termo:?)\s+["\']?([^"\']+)["\']?', message, re.I)
    if m:
        termo = m.group(1).strip()

    if not vinculo_id and not termo:
        return "❌ Informe o ID do vínculo item-produto ou um termo de busca.", None

    resultado = tool_historico_precos_camada_f(
        edital_item_produto_id=vinculo_id or "", user_id=user_id, termo=termo
    )
    if not resultado.get("success"):
        return f"❌ {resultado.get('error')}", resultado

    txt = f"## 📜 Histórico de Preços (Camada F)\n\n{resultado.get('mensagem')}\n\n"
    if resultado.get('preco_medio'):
        txt += f"- Preço médio: R$ {resultado['preco_medio']:,.2f}\n"
        txt += f"- Faixa: R$ {resultado.get('preco_min', 0):,.2f} — R$ {resultado.get('preco_max', 0):,.2f}\n"
        txt += f"- Registros: {resultado.get('qtd_registros', 0)}\n"
    else:
        txt += "Nenhum registro de preço encontrado.\n"

    return txt, resultado


def processar_precif_comodato(message: str, user_id: str, empresa_id: str = None):
    """UC-P10: Gestão de comodato"""
    from tools import tool_gestao_comodato
    from models import get_db, Edital
    import re

    edital_id = _extrair_id_da_mensagem(message)
    if not edital_id:
        db = get_db()
        try:
            edital = db.query(Edital).filter(Edital.user_id == user_id).order_by(Edital.created_at.desc()).first()
            edital_id = edital.id if edital else None
        finally:
            db.close()
    if not edital_id:
        return "❌ Nenhum edital encontrado.", None

    # Extrair nome e valor da mensagem
    nome = None
    valor = None
    meses = None

    m_nome = re.search(r'equipamento[:\s]+["\']?([^"\',$]+)', message, re.I)
    if m_nome:
        nome = m_nome.group(1).strip()

    m_valor = re.search(r'R\$\s*([\d.,]+)', message)
    if m_valor:
        valor = float(m_valor.group(1).replace('.', '').replace(',', '.'))

    m_meses = re.search(r'(\d+)\s*meses', message, re.I)
    if m_meses:
        meses = int(m_meses.group(1))

    resultado = tool_gestao_comodato(
        edital_id=edital_id, user_id=user_id, empresa_id=empresa_id,
        nome_equipamento=nome or "Equipamento",
        valor_equipamento=valor,
        duracao_meses=meses,
    )
    if not resultado.get("success"):
        return f"❌ {resultado.get('error')}", resultado

    com = resultado.get('comodato', {})
    txt = f"## 🔧 Comodato\n\n{resultado.get('mensagem')}\n\n"
    txt += f"- Equipamento: {com.get('nome_equipamento')}\n"
    if com.get('valor_equipamento'):
        txt += f"- Valor: R$ {com['valor_equipamento']:,.2f}\n"
    if com.get('duracao_meses'):
        txt += f"- Prazo: {com['duracao_meses']} meses\n"
    if resultado.get('amortizacao_mensal'):
        txt += f"- **Amortização mensal: R$ {resultado['amortizacao_mensal']:,.2f}**\n"

    return txt, resultado


def processar_chat_livre(message: str, user_id: str, session_id: str, db):
    """Processa chat livre sobre licitações"""
    # Buscar histórico
    historico = db.query(Message).filter(
        Message.session_id == session_id
    ).order_by(Message.created_at.desc()).limit(MAX_HISTORY_MESSAGES).all()

    historico = list(reversed(historico))

    # Montar mensagens
    system_prompt = """Você é um especialista em licitações públicas brasileiras.
Seu conhecimento inclui:
- Lei 14.133/2021 (Nova Lei de Licitações)
- Pregão eletrônico e presencial
- Elaboração de propostas técnicas
- Análise de editais
- Impugnações e recursos
- Comodato de equipamentos
- Contratos administrativos

Responda de forma clara, objetiva e fundamentada na legislação quando aplicável."""

    messages = [{"role": "system", "content": system_prompt}]

    for msg in historico:
        messages.append({"role": msg.role, "content": msg.content})

    messages.append({"role": "user", "content": message})

    try:
        response = call_deepseek(messages, max_tokens=4000)
        return response
    except Exception as e:
        return f"Erro ao processar: {str(e)}"


# =============================================================================
# Upload Routes
# =============================================================================

@app.route("/api/upload", methods=["POST"])
@require_auth
def upload_manual():
    """
    Upload de manual PDF para extração de especificações.

    Form data:
    - file: arquivo PDF
    - nome_produto: nome do produto
    - categoria: equipamento, reagente, insumo_hospitalar, insumo_laboratorial
    - fabricante: (opcional)
    - modelo: (opcional)
    """
    user_id = get_current_user_id()

    if 'file' not in request.files:
        return jsonify({"error": "Arquivo não enviado"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Nenhum arquivo selecionado"}), 400

    nome_produto = request.form.get('nome_produto')
    categoria = request.form.get('categoria', 'equipamento')
    fabricante = request.form.get('fabricante')
    modelo = request.form.get('modelo')

    if not nome_produto:
        return jsonify({"error": "nome_produto é obrigatório"}), 400

    # Salvar arquivo
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    filename = f"{user_id}_{uuid.uuid4().hex[:8]}_{file.filename}"
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # Processar
    resultado = tool_processar_upload(
        filepath=filepath,
        user_id=user_id,
        nome_produto=nome_produto,
        categoria=categoria,
        fabricante=fabricante,
        modelo=modelo
    )

    if resultado.get("success"):
        return jsonify(resultado), 201
    else:
        return jsonify(resultado), 400


@app.route("/api/chat-upload", methods=["POST"])
@require_auth
def chat_upload():
    """
    Envia mensagem com arquivo anexo.
    O agente classificador interpreta a intenção do usuário.

    Form data:
    - file: arquivo PDF
    - session_id: ID da sessão de chat
    - message: mensagem do usuário (opcional)
    """
    user_id = get_current_user_id()

    if 'file' not in request.files:
        return jsonify({"error": "Arquivo não enviado"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Nenhum arquivo selecionado"}), 400

    session_id = request.form.get('session_id')
    message = request.form.get('message', '').strip()
    subclasse_id = request.form.get('subclasse_id')  # Para guiar extração com campos_mascara

    if not session_id:
        return jsonify({"error": "session_id é obrigatório"}), 400

    # ========== USAR AGENTE CLASSIFICADOR ==========
    print(f"[CHAT-UPLOAD] Classificando intenção: '{message}' (tem_arquivo=True)")
    intencao_resultado = detectar_intencao_ia(message, tem_arquivo=True)
    intencao_arquivo = intencao_resultado.get("intencao", "arquivo_cadastrar")
    nome_produto = intencao_resultado.get("nome_produto")
    print(f"[CHAT-UPLOAD] Intenção detectada: {intencao_arquivo}")

    # Mapear intenções do classificador para ações internas
    mapa_intencoes = {
        "arquivo_cadastrar": "cadastrar",
        "arquivo_mostrar": "mostrar_conteudo",
        "arquivo_specs": "extrair_specs",
        "arquivo_resumir": "resumir",
        "arquivo_analisar": "analisar",
        "extrair_ata": "extrair_ata",  # Nova ação: extrair resultados de ata de pregão
        # Fallbacks para compatibilidade
        "upload_manual": "cadastrar",
        "chat_livre": "cadastrar"  # Se não entendeu, cadastra
    }
    intencao_arquivo = mapa_intencoes.get(intencao_arquivo, "cadastrar")

    db = get_db()
    try:
        # Verificar sessão
        session = db.query(Session).filter(
            Session.id == session_id,
            Session.user_id == user_id
        ).first()

        if not session:
            return jsonify({"error": "Sessão não encontrada"}), 404

        # Salvar arquivo
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        filename = f"{user_id}_{uuid.uuid4().hex[:8]}_{file.filename}"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)

        # Extrair texto do PDF
        import fitz
        texto_pdf = ""
        try:
            doc = fitz.open(filepath)
            for page in doc:
                texto_pdf += page.get_text()
            doc.close()
        except Exception as e:
            texto_pdf = f"Erro ao extrair texto: {e}"

        # Salvar mensagem do usuário
        acoes_desc = {
            "cadastrar": "Cadastrar como produto",
            "mostrar_conteudo": "Mostrar conteúdo",
            "extrair_specs": "Extrair especificações",
            "resumir": "Resumir documento",
            "analisar": "Analisar documento",
            "extrair_ata": "Extrair resultados da ata"
        }
        user_msg_content = f"📎 **{file.filename}**\n*{acoes_desc.get(intencao_arquivo, 'Processar')}*"
        user_msg = Message(
            session_id=session_id,
            role='user',
            content=user_msg_content,
            action_type='upload_manual'
        )
        db.add(user_msg)

        resultado = {"success": True}
        response_text = ""

        # ========== AÇÃO: MOSTRAR CONTEÚDO ==========
        if intencao_arquivo == "mostrar_conteudo":
            response_text = f"""## 📄 Conteúdo do Documento

**Arquivo:** {file.filename}
**Tamanho:** {len(texto_pdf)} caracteres

---

{texto_pdf[:5000]}

{"..." if len(texto_pdf) > 5000 else ""}

---
*Para cadastrar como produto, envie: "cadastre"*"""

        # ========== AÇÃO: EXTRAIR SPECS (sem cadastrar) ==========
        elif intencao_arquivo == "extrair_specs":
            # info e specs já importados no topo
            info = _extrair_info_produto(texto_pdf[:8000])

            # Extrair specs via IA
            prompt = PROMPT_EXTRAIR_SPECS.format(texto=texto_pdf[:15000])
            resposta_ia = call_deepseek([{"role": "user", "content": prompt}], max_tokens=8000)

            response_text = f"""## 📊 Especificações Extraídas

**Produto identificado:** {info.get('nome', 'N/A')}
**Fabricante:** {info.get('fabricante', 'N/A')}
**Modelo:** {info.get('modelo', 'N/A')}

### Especificações:

{resposta_ia[:4000]}

---
*Para cadastrar como produto, envie: "cadastre"*"""

        # ========== AÇÃO: RESUMIR ==========
        elif intencao_arquivo == "resumir":
            prompt_resumo = f"""Resuma o documento abaixo em português, destacando:
1. Tipo de documento (manual, datasheet, etc)
2. Produto/equipamento descrito
3. Principais características
4. Pontos importantes

DOCUMENTO:
{texto_pdf[:10000]}

RESUMO:"""
            resumo = call_deepseek([{"role": "user", "content": prompt_resumo}], max_tokens=2000, model_override="deepseek-chat")

            response_text = f"""## 📝 Resumo do Documento

**Arquivo:** {file.filename}

{resumo}

---
*Para cadastrar como produto, envie: "cadastre"*"""

        # ========== AÇÃO: ANALISAR ==========
        elif intencao_arquivo == "analisar":
            prompt_analise = f"""Analise o documento técnico abaixo e forneça:
1. Tipo de documento
2. Produto/equipamento descrito
3. Fabricante
4. Principais especificações técnicas
5. Aplicações/uso indicado
6. Pontos fortes e fracos (se identificáveis)

DOCUMENTO:
{texto_pdf[:12000]}

ANÁLISE:"""
            analise = call_deepseek([{"role": "user", "content": prompt_analise}], max_tokens=3000, model_override="deepseek-chat")

            response_text = f"""## 🔍 Análise do Documento

**Arquivo:** {file.filename}

{analise}

---
*Para cadastrar como produto, envie: "cadastre"*"""

        # ========== AÇÃO: EXTRAIR ATA DE SESSÃO ==========
        elif intencao_arquivo == "extrair_ata":
            response_text, resultado = processar_extrair_ata(texto_pdf, filepath, user_id, file.filename)

        # ========== AÇÃO: CADASTRAR (padrão) ==========
        else:
            resultado = tool_processar_upload(
                filepath=filepath,
                user_id=user_id,
                nome_produto=nome_produto,
                categoria=None,
                fabricante=None,
                modelo=None,
                subclasse_id=subclasse_id
            )

            if resultado.get("success"):
                produto = resultado.get("produto", {})
                specs = resultado.get("especificacoes", [])

                response_text = f"""## ✅ Produto Cadastrado com Sucesso!

**Nome:** {produto.get('nome', 'N/A')}
**Fabricante:** {produto.get('fabricante', 'Não identificado')}
**Modelo:** {produto.get('modelo', 'Não identificado')}
**Categoria:** {produto.get('categoria', 'equipamento')}
**ID:** {produto.get('id', 'N/A')}

### Especificações Extraídas ({len(specs)} encontradas):
"""
                for spec in specs[:15]:
                    response_text += f"- **{spec.get('nome', 'N/A')}:** {spec.get('valor', 'N/A')}\n"

                if len(specs) > 15:
                    response_text += f"\n... e mais {len(specs) - 15} especificações.\n"

                response_text += "\n---\n✅ Produto pronto para calcular aderência com editais!"
            elif resultado.get("duplicado"):
                prod_exist = resultado.get("produto_existente", {})
                response_text = f"""## ⚠️ Produto já cadastrado!

**Nome:** {prod_exist.get('nome', 'N/A')}
**Modelo:** {prod_exist.get('modelo', 'N/A')}
**ID:** {prod_exist.get('id', 'N/A')}

Use **reprocesse o produto {prod_exist.get('nome')}** para atualizar as especificações."""
            else:
                response_text = f"❌ Erro ao processar arquivo: {resultado.get('error', 'Erro desconhecido')}"

        # Salvar resposta do assistente
        assistant_msg = Message(
            session_id=session_id,
            role='assistant',
            content=response_text,
            action_type='upload_manual'
        )
        db.add(assistant_msg)

        # Atualizar sessão
        session.updated_at = datetime.now()
        db.commit()

        return jsonify({
            "success": resultado.get("success", False),
            "response": response_text,
            "session_id": session_id,
            "action_type": "upload_manual"
        })

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


@app.route("/api/upload-chat", methods=["POST"])
@require_auth
def upload_chat():
    """
    DEPRECATED - Use /api/chat-upload instead.
    Mantido para compatibilidade.
    """
    user_id = get_current_user_id()

    if 'file' not in request.files:
        return jsonify({"error": "Arquivo não enviado"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Nenhum arquivo selecionado"}), 400

    session_id = request.form.get('session_id')
    nome_produto = request.form.get('nome_produto', '').strip()

    if not session_id:
        return jsonify({"error": "session_id é obrigatório"}), 400
    if not nome_produto:
        return jsonify({"error": "nome_produto é obrigatório"}), 400

    db = get_db()
    try:
        # Verificar sessão
        session = db.query(Session).filter(
            Session.id == session_id,
            Session.user_id == user_id
        ).first()

        if not session:
            return jsonify({"error": "Sessão não encontrada"}), 404

        # Salvar arquivo
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        filename = f"{user_id}_{uuid.uuid4().hex[:8]}_{file.filename}"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)

        # Salvar mensagem do usuário
        user_msg_content = f"📎 Upload: **{file.filename}**\nCadastrar como: **{nome_produto}**"
        user_msg = Message(
            session_id=session_id,
            role='user',
            content=user_msg_content,
            action_type='upload_manual'
        )
        db.add(user_msg)

        # Determinar categoria automaticamente
        categoria = "equipamento"
        nome_lower = nome_produto.lower()
        if any(t in nome_lower for t in ["analisador", "bioquímic", "laborat"]):
            categoria = "equipamento"
        elif any(t in nome_lower for t in ["centrifuga", "microscop"]):
            categoria = "equipamento"
        elif any(t in nome_lower for t in ["cama", "maca", "cadeira"]):
            categoria = "mobiliario"
        elif any(t in nome_lower for t in ["monitor", "desfibrilador", "eletrocard"]):
            categoria = "equipamento"

        # Processar arquivo
        resultado = tool_processar_upload(
            filepath=filepath,
            user_id=user_id,
            nome_produto=nome_produto,
            categoria=categoria,
            fabricante=None,
            modelo=None
        )

        # Montar resposta
        if resultado.get("success"):
            produto = resultado.get("produto", {})
            specs = resultado.get("especificacoes", [])

            response_text = f"""## ✅ Produto Cadastrado com Sucesso!

**Nome:** {produto.get('nome', nome_produto)}
**Categoria:** {categoria}
**ID:** {produto.get('id', 'N/A')}

### Especificações Extraídas ({len(specs)} encontradas):
"""
            for spec in specs[:15]:
                response_text += f"- **{spec.get('nome', 'N/A')}:** {spec.get('valor', 'N/A')}\n"

            if len(specs) > 15:
                response_text += f"\n... e mais {len(specs) - 15} especificações.\n"

            response_text += "\n---\n✅ Produto pronto para calcular aderência com editais!"
        else:
            response_text = f"❌ Erro ao processar arquivo: {resultado.get('error', 'Erro desconhecido')}"

        # Salvar resposta do assistente
        assistant_msg = Message(
            session_id=session_id,
            role='assistant',
            content=response_text,
            action_type='upload_manual'
        )
        db.add(assistant_msg)

        # Atualizar sessão
        session.updated_at = datetime.now()
        db.commit()

        return jsonify({
            "success": resultado.get("success", False),
            "response": response_text,
            "session_id": session_id,
            "produto": resultado.get("produto"),
            "especificacoes_extraidas": len(resultado.get("especificacoes", []))
        })

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


# =============================================================================
# Session Routes
# =============================================================================

@app.route("/api/sessions", methods=["GET"])
@require_auth
def get_sessions():
    user_id = get_current_user_id()
    db = get_db()
    try:
        sessions = db.query(Session).filter(
            Session.user_id == user_id
        ).order_by(Session.updated_at.desc()).all()

        return jsonify({"sessions": [s.to_dict() for s in sessions]})
    finally:
        db.close()


@app.route("/api/sessions", methods=["POST"])
@require_auth
def new_session():
    user_id = get_current_user_id()
    data = request.json or {}
    name = data.get("name", "Nova conversa")

    db = get_db()
    try:
        session = Session(user_id=user_id, name=name)
        db.add(session)
        db.commit()
        return jsonify(session.to_dict()), 201
    finally:
        db.close()


@app.route("/api/sessions/<session_id>", methods=["GET"])
@require_auth
def get_session_detail(session_id):
    user_id = get_current_user_id()
    db = get_db()
    try:
        session = db.query(Session).filter(
            Session.id == session_id,
            Session.user_id == user_id
        ).first()

        if not session:
            return jsonify({"error": "Sessão não encontrada"}), 404

        return jsonify(session.to_dict(include_messages=True))
    finally:
        db.close()


@app.route("/api/sessions/<session_id>", methods=["DELETE"])
@require_auth
def delete_session(session_id):
    user_id = get_current_user_id()
    db = get_db()
    try:
        session = db.query(Session).filter(
            Session.id == session_id,
            Session.user_id == user_id
        ).first()

        if not session:
            return jsonify({"error": "Sessão não encontrada"}), 404

        db.delete(session)
        db.commit()
        return jsonify({"message": "Sessão excluída"})
    finally:
        db.close()


@app.route("/api/sessions/<session_id>", methods=["PATCH"])
@require_auth
def update_session(session_id):
    user_id = get_current_user_id()
    data = request.json or {}
    new_name = data.get("name")

    if not new_name:
        return jsonify({"error": "name é obrigatório"}), 400

    db = get_db()
    try:
        session = db.query(Session).filter(
            Session.id == session_id,
            Session.user_id == user_id
        ).first()

        if not session:
            return jsonify({"error": "Sessão não encontrada"}), 404

        session.name = new_name
        db.commit()
        return jsonify({"message": "Sessão renomeada", "name": new_name})
    finally:
        db.close()


# =============================================================================
# Produtos Routes
# =============================================================================

@app.route("/api/produtos", methods=["GET"])
@require_auth
def listar_produtos_api():
    user_id = get_current_user_id()
    categoria = request.args.get("categoria")
    nome = request.args.get("nome")

    resultado = tool_listar_produtos(user_id, categoria=categoria, nome=nome)
    return jsonify(resultado)


@app.route("/api/produtos/<produto_id>", methods=["GET"])
@require_auth
def get_produto(produto_id):
    user_id = get_current_user_id()
    db = get_db()
    try:
        produto = db.query(Produto).filter(
            Produto.id == produto_id,
            Produto.user_id == user_id
        ).first()

        if not produto:
            return jsonify({"error": "Produto não encontrado"}), 404

        return jsonify(produto.to_dict(include_specs=True))
    finally:
        db.close()


@app.route("/api/produtos/<produto_id>/completude", methods=["GET"])
@require_auth
def get_produto_completude(produto_id):
    """Verifica completude do produto: dados básicos + specs vs máscara da subclasse."""
    from tools import tool_verificar_completude_produto
    user_id = get_current_user_id()
    resultado = tool_verificar_completude_produto(produto_id=produto_id, user_id=user_id)
    if not resultado.get("success"):
        return jsonify({"error": resultado.get("error", "Erro")}), 404
    return jsonify(resultado)


# =============================================================================
# Editais Routes
# =============================================================================

@app.route("/api/editais", methods=["GET"])
@require_auth
def listar_editais_api():
    user_id = get_current_user_id()
    status = request.args.get("status")
    uf = request.args.get("uf")
    categoria = request.args.get("categoria")

    resultado = tool_listar_editais(user_id, status=status, uf=uf, categoria=categoria)
    return jsonify(resultado)


@app.route("/api/editais/<edital_id>", methods=["GET"])
@require_auth
def get_edital(edital_id):
    user_id = get_current_user_id()
    db = get_db()
    try:
        edital = db.query(Edital).filter(
            Edital.id == edital_id,
            Edital.user_id == user_id
        ).first()

        if not edital:
            return jsonify({"error": "Edital não encontrado"}), 404

        return jsonify(edital.to_dict(include_requisitos=True))
    finally:
        db.close()


@app.route("/api/editais/<edital_id>/pdf", methods=["GET"])
@require_auth
def download_edital_pdf(edital_id):
    """Download ou visualização do PDF do edital.
    Reutiliza a mesma lógica do chat: baixa, salva localmente, cria EditalDocumento.
    Na segunda vez, serve do disco instantaneamente.
    """
    import re
    import requests as req
    from models import EditalDocumento
    from datetime import datetime
    user_id = get_current_user_id()
    db = get_db()
    try:
        edital = db.query(Edital).filter(
            Edital.id == edital_id,
            Edital.user_id == user_id
        ).first()

        if not edital:
            return jsonify({"error": "Edital não encontrado"}), 404

        download_flag = request.args.get('download', 'false').lower() == 'true'

        # === 1. Já tem arquivo local? Servir direto ===
        doc = db.query(EditalDocumento).filter(
            EditalDocumento.edital_id == edital_id,
            EditalDocumento.tipo == 'edital_principal'
        ).first()

        # Helper: se arquivo é ZIP, extrair PDF do edital de dentro
        import zipfile as _zf
        def _extrair_pdf_de_zip(zip_path):
            if not _zf.is_zipfile(zip_path):
                return zip_path  # Não é ZIP, retorna como está
            print(f"[PDF] Arquivo é ZIP, extraindo PDF: {zip_path}")
            upload_dir = os.path.dirname(zip_path)
            all_pdfs = []
            with _zf.ZipFile(zip_path, 'r') as zf:
                for name in zf.namelist():
                    if name.lower().endswith('.pdf'):
                        zf.extract(name, upload_dir)
                        fpath = os.path.join(upload_dir, name)
                        all_pdfs.append((fpath, os.path.getsize(fpath), os.path.basename(name).lower()))
                for name in zf.namelist():
                    if name.lower().endswith('.zip'):
                        zf.extract(name, upload_dir)
                        nested = os.path.join(upload_dir, name)
                        if _zf.is_zipfile(nested):
                            with _zf.ZipFile(nested, 'r') as zf2:
                                for n2 in zf2.namelist():
                                    if n2.lower().endswith('.pdf'):
                                        zf2.extract(n2, upload_dir)
                                        fp2 = os.path.join(upload_dir, n2)
                                        all_pdfs.append((fp2, os.path.getsize(fp2), os.path.basename(n2).lower()))
                            try: os.remove(nested)
                            except OSError: pass
            if not all_pdfs:
                return zip_path
            edital_pdfs = [p for p in all_pdfs if 'edital' in p[2] and 'anexo' not in p[2] and 'termo' not in p[2] and 'minuta' not in p[2]]
            best = max(edital_pdfs, key=lambda x: x[1]) if edital_pdfs else max(all_pdfs, key=lambda x: x[1])
            for p in all_pdfs:
                if p[0] != best[0]:
                    try: os.remove(p[0])
                    except OSError: pass
            pdf_filename = re.sub(r'[^\w\-_\.]', '_', os.path.basename(best[0]))
            final_path = os.path.join(upload_dir, pdf_filename)
            if best[0] != final_path:
                if os.path.exists(final_path): os.remove(final_path)
                os.rename(best[0], final_path)
            try: os.remove(zip_path)
            except OSError: pass
            # Limpar subpastas
            for p in all_pdfs:
                parent = os.path.dirname(p[0])
                if parent != upload_dir:
                    try: os.removedirs(parent)
                    except OSError: pass
            print(f"[PDF] PDF extraído de ZIP: {final_path}")
            return final_path

        if doc and doc.path_arquivo and os.path.exists(doc.path_arquivo):
            filepath = _extrair_pdf_de_zip(doc.path_arquivo)
            if filepath != doc.path_arquivo:
                doc.path_arquivo = filepath
                db.commit()
            return send_file(
                filepath,
                mimetype='application/pdf',
                as_attachment=download_flag,
                download_name=doc.nome_arquivo or f"edital_{edital.numero}.pdf"
            )

        # Se doc existe mas sem arquivo no disco, deletar para permitir re-download
        if doc and (not doc.path_arquivo or not os.path.exists(doc.path_arquivo or "")):
            db.delete(doc)
            db.commit()
            doc = None

        if edital.pdf_path and os.path.exists(edital.pdf_path):
            filepath = _extrair_pdf_de_zip(edital.pdf_path)
            if filepath != edital.pdf_path:
                edital.pdf_path = filepath
                db.commit()
            return send_file(
                filepath,
                mimetype='application/pdf',
                as_attachment=download_flag,
                download_name=edital.pdf_titulo or f"edital_{edital.numero}.pdf"
            )

        # === 2. Extrair dados PNCP da URL se necessário ===
        if not (edital.cnpj_orgao and edital.ano_compra and edital.seq_compra):
            if edital.url and 'pncp.gov.br' in edital.url:
                # Formato /compras/{cnpj}/{ano}/{seq}
                match = re.search(r'/compras/(\d{14}\d{0,2})/(\d{4})/(\d+)', edital.url)
                if match:
                    edital.cnpj_orgao = match.group(1)
                    edital.ano_compra = int(match.group(2))
                    edital.seq_compra = int(match.group(3))
                else:
                    # Formato /editais/{cnpj}-{id}-{seq}/{ano}
                    match = re.search(r'/editais/(\d{14}\d{0,2})-(\d+)-(\d+)/(\d{4})', edital.url)
                    if match:
                        edital.cnpj_orgao = match.group(1)
                        edital.ano_compra = int(match.group(4))
                        edital.seq_compra = int(match.group(3))
                if match:
                    db.commit()
                    print(f"[PDF] Dados PNCP extraídos da URL: cnpj={edital.cnpj_orgao}, ano={edital.ano_compra}, seq={edital.seq_compra}")

        # === 3. Baixar via PNCP e SALVAR localmente (igual ao chat) ===
        if edital.cnpj_orgao and edital.ano_compra and edital.seq_compra:
            try:
                from tools import tool_buscar_arquivos_edital_pncp, tool_baixar_pdf_pncp

                arquivos_result = tool_buscar_arquivos_edital_pncp(
                    cnpj=edital.cnpj_orgao, ano=edital.ano_compra, seq=edital.seq_compra
                )

                if arquivos_result.get('success') and arquivos_result.get('arquivos'):
                    arquivo_edital = arquivos_result.get('arquivo_edital') or arquivos_result['arquivos'][0]

                    # Baixar e salvar no disco (tool_baixar_pdf_pncp salva em uploads/{user_id}/editais/)
                    download_result = tool_baixar_pdf_pncp(
                        cnpj=edital.cnpj_orgao,
                        ano=edital.ano_compra,
                        seq=edital.seq_compra,
                        sequencial_arquivo=arquivo_edital['sequencial'],
                        user_id=user_id,
                        edital_id=edital.id
                    )

                    if download_result.get('success'):
                        filepath = download_result['filepath']
                        filename = download_result['filename']

                        # Extrair texto do PDF
                        texto_extraido = ""
                        try:
                            from PyPDF2 import PdfReader
                            reader = PdfReader(filepath)
                            for page in reader.pages:
                                texto_extraido += page.extract_text() or ""
                        except Exception:
                            pass

                        # Salvar no banco (EditalDocumento) — igual processar_baixar_pdf_edital
                        novo_doc = EditalDocumento(
                            id=str(uuid.uuid4()),
                            edital_id=edital.id,
                            tipo='edital_principal',
                            nome_arquivo=filename,
                            path_arquivo=filepath,
                            texto_extraido=texto_extraido[:100000] if texto_extraido else None,
                            processado=True,
                            created_at=datetime.now()
                        )
                        db.add(novo_doc)

                        # Atualizar edital com pdf_path, pdf_url, pdf_titulo
                        edital.pdf_path = filepath
                        edital.pdf_url = download_result.get('url') or edital.pdf_url
                        edital.pdf_titulo = arquivo_edital.get('titulo')
                        db.commit()

                        print(f"[PDF] Salvo localmente: {filepath} ({download_result['filesize']/1024:.1f} KB)")

                        return send_file(
                            filepath,
                            mimetype='application/pdf',
                            as_attachment=download_flag,
                            download_name=filename
                        )
            except Exception as e:
                print(f"[PDF] Erro ao baixar via PNCP: {e}")
                import traceback
                traceback.print_exc()

        # === 4. Fallback: proxy da pdf_url ou url genérica ===
        fallback_url = edital.pdf_url or edital.url
        if fallback_url:
            try:
                resp = req.get(fallback_url, timeout=60, headers={
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
                }, allow_redirects=True)
                content_type = resp.headers.get('Content-Type', '')
                if resp.status_code == 200 and ('pdf' in content_type.lower() or fallback_url.lower().endswith('.pdf')):
                    # Salvar localmente
                    upload_dir = os.path.join(UPLOAD_FOLDER, user_id, "editais")
                    os.makedirs(upload_dir, exist_ok=True)
                    safe_numero = re.sub(r'[^\w\-.]', '_', edital.numero or edital_id)
                    filepath = os.path.join(upload_dir, f"edital_{safe_numero}.pdf")
                    with open(filepath, 'wb') as f:
                        f.write(resp.content)

                    novo_doc = EditalDocumento(
                        id=str(uuid.uuid4()),
                        edital_id=edital.id,
                        tipo='edital_principal',
                        nome_arquivo=f"edital_{safe_numero}.pdf",
                        path_arquivo=filepath,
                        processado=False,
                        created_at=datetime.now()
                    )
                    db.add(novo_doc)
                    edital.pdf_path = filepath
                    db.commit()

                    return send_file(
                        filepath,
                        mimetype='application/pdf',
                        as_attachment=download_flag,
                        download_name=f"edital_{safe_numero}.pdf"
                    )
            except Exception as e:
                print(f"[PDF] Erro fallback URL: {e}")

        return jsonify({"error": "PDF não disponível para este edital"}), 404
    finally:
        db.close()


@app.route("/api/editais/<edital_id>/pdf/view", methods=["GET"])
def view_edital_pdf(edital_id):
    """Serve PDF já baixado para visualização inline (iframe/object).
    Usa token via query param para permitir acesso direto pelo browser.
    """
    from models import EditalDocumento
    token = request.args.get('token')
    if not token:
        return jsonify({"error": "Token não fornecido"}), 401
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=['HS256'])
        user_id = payload.get('user_id')
    except Exception:
        return jsonify({"error": "Token inválido"}), 401

    db = get_db()
    try:
        edital = db.query(Edital).filter(Edital.id == edital_id, Edital.user_id == user_id).first()
        if not edital:
            return jsonify({"error": "Edital não encontrado"}), 404

        doc = db.query(EditalDocumento).filter(
            EditalDocumento.edital_id == edital_id,
            EditalDocumento.tipo == 'edital_principal'
        ).first()

        pdf_path = None
        if doc and doc.path_arquivo and os.path.exists(doc.path_arquivo):
            pdf_path = doc.path_arquivo
        elif edital.pdf_path and os.path.exists(edital.pdf_path):
            pdf_path = edital.pdf_path

        if not pdf_path:
            return jsonify({"error": "PDF não encontrado. Clique em 'Baixar PDF do Edital' na aba IA primeiro."}), 404

        # Ler o arquivo e retornar como response puro — sem Content-Disposition
        # para evitar que o browser trate como download dentro de iframe
        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()
        resp = app.response_class(pdf_bytes, mimetype='application/pdf')
        resp.headers['Content-Length'] = str(len(pdf_bytes))
        resp.headers['Cache-Control'] = 'public, max-age=3600'
        return resp
    finally:
        db.close()


# ─── Endpoint: Extrair/Listar Lotes do Edital ────────────────────────────────

@app.route("/api/editais/<edital_id>/lotes", methods=["GET"])
@require_auth
def listar_lotes_edital(edital_id):
    """Lista lotes do edital. Se não existem, retorna lista vazia."""
    from models import Lote, LoteItem, EditalItem
    user_id = get_current_user_id()
    db = get_db()
    try:
        lotes = db.query(Lote).filter(
            Lote.edital_id == edital_id,
            Lote.user_id == user_id
        ).order_by(Lote.numero_lote).all()

        lotes_resp = []
        for l in lotes:
            ld = l.to_dict()
            lote_itens = db.query(LoteItem).filter(LoteItem.lote_id == l.id).order_by(LoteItem.ordem).all()
            ld["itens"] = []
            for li in lote_itens:
                ei = db.query(EditalItem).filter(EditalItem.id == li.edital_item_id).first()
                if ei:
                    ld["itens"].append(ei.to_dict())
            lotes_resp.append(ld)

        return jsonify({"success": True, "lotes": lotes_resp, "total": len(lotes_resp)})
    finally:
        db.close()


@app.route("/api/editais/<edital_id>/buscar-itens-pncp", methods=["POST"])
@require_auth
def buscar_itens_pncp(edital_id):
    """Busca itens do edital diretamente na API PNCP e salva no banco."""
    user_id = get_current_user_id()
    try:
        from tools import tool_buscar_itens_edital_pncp
        resultado = tool_buscar_itens_edital_pncp(edital_id=edital_id, user_id=user_id)
        return jsonify(resultado)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/editais/<edital_id>/lotes/extrair", methods=["POST"])
@require_auth
def extrair_lotes_edital(edital_id):
    """Extrai lotes do edital via IA (lê PDF) e cria no banco."""
    from tools import tool_organizar_lotes
    user_id = get_current_user_id()
    empresa_id = getattr(request, 'empresa_id', None)
    forcar = request.json.get("forcar", False) if request.is_json else False

    result = tool_organizar_lotes(
        edital_id=edital_id,
        user_id=user_id,
        empresa_id=empresa_id,
        importar_pncp=True,
        forcar=forcar
    )

    status_code = 200 if result.get("success") else 400
    return jsonify(result), status_code


@app.route("/api/editais/numero/<numero>/pdf", methods=["GET"])
@require_auth
def download_edital_pdf_by_numero(numero):
    """Download ou visualização do PDF do edital pelo número"""
    from models import EditalDocumento
    user_id = get_current_user_id()
    db = get_db()
    try:
        # Normalizar número para busca
        numero_busca = numero.replace('_', '/').replace('-', '/').upper()

        # Buscar edital pelo número
        edital = db.query(Edital).filter(
            Edital.user_id == user_id
        ).all()

        # Encontrar edital com número similar
        edital_encontrado = None
        for e in edital:
            num_edital = e.numero.replace('-', '/').upper()
            if num_edital == numero_busca or numero_busca in num_edital or num_edital in numero_busca:
                edital_encontrado = e
                break

        if not edital_encontrado:
            return jsonify({"error": f"Edital {numero} não encontrado"}), 404

        # Buscar documento do edital
        doc = db.query(EditalDocumento).filter(
            EditalDocumento.edital_id == edital_encontrado.id,
            EditalDocumento.tipo == 'edital_principal'
        ).first()

        if not doc or not doc.path_arquivo:
            return jsonify({"error": "PDF não disponível para este edital"}), 404

        if not os.path.exists(doc.path_arquivo):
            return jsonify({"error": "Arquivo não encontrado no servidor"}), 404

        # Parâmetro para forçar download
        download = request.args.get('download', 'false').lower() == 'true'

        return send_file(
            doc.path_arquivo,
            mimetype='application/pdf',
            as_attachment=download,
            download_name=doc.nome_arquivo or f"edital_{edital_encontrado.numero}.pdf"
        )
    finally:
        db.close()


@app.route("/api/proxy/pdf/pncp/<cnpj>/<int:ano>/<int:seq>", methods=["GET"])
@app.route("/api/proxy/pdf/pncp/<cnpj>/<int:ano>/<int:seq>/<int:arquivo_id>", methods=["GET"])
def proxy_pdf_pncp(cnpj, ano, seq, arquivo_id=1):
    """
    Proxy para visualizar PDFs do PNCP no navegador.
    Busca o arquivo da API do PNCP, extrai PDF do ZIP se necessário,
    e retorna com headers corretos para visualização inline.
    """
    import requests as req
    from io import BytesIO
    import zipfile

    # Construir URL do PNCP
    pdf_url = f"https://pncp.gov.br/api/pncp/v1/orgaos/{cnpj}/compras/{ano}/{seq}/arquivos/{arquivo_id}"

    try:
        print(f"[PDF PROXY] Buscando: {pdf_url}")
        resp = req.get(pdf_url, timeout=60)

        if resp.status_code == 200:
            content = resp.content
            content_type = resp.headers.get('Content-Type', '')
            content_disp = resp.headers.get('Content-Disposition', '')

            # Detectar nome do arquivo
            filename = f"edital_pncp_{ano}_{seq}.pdf"
            if 'filename=' in content_disp:
                import re
                match = re.search(r'filename[^;=\n]*=(([\'"]).*?\2|[^;\n]*)', content_disp)
                if match:
                    filename = match.group(1).strip('"\'')

            # Verificar se é um ZIP (PNCP frequentemente retorna ZIP com PDFs dentro)
            is_zip = (
                filename.lower().endswith('.zip') or
                'zip' in content_type.lower() or
                content[:4] == b'PK\x03\x04'  # Magic bytes do ZIP
            )

            if is_zip:
                print(f"[PDF PROXY] Arquivo é ZIP, extraindo PDF...")
                try:
                    zip_buffer = BytesIO(content)
                    with zipfile.ZipFile(zip_buffer, 'r') as zf:
                        # Listar arquivos no ZIP
                        pdf_files = [f for f in zf.namelist() if f.lower().endswith('.pdf')]
                        print(f"[PDF PROXY] PDFs no ZIP: {pdf_files}")

                        if pdf_files:
                            # Priorizar arquivo com "Edital" no nome, senão pega o primeiro
                            pdf_to_extract = pdf_files[0]
                            for pf in pdf_files:
                                if 'edital' in pf.lower():
                                    pdf_to_extract = pf
                                    break

                            print(f"[PDF PROXY] Extraindo: {pdf_to_extract}")
                            pdf_content = zf.read(pdf_to_extract)
                            filename = pdf_to_extract
                            content = pdf_content
                        else:
                            print(f"[PDF PROXY] Nenhum PDF encontrado no ZIP")
                            return jsonify({"error": "ZIP não contém arquivos PDF"}), 404
                except zipfile.BadZipFile:
                    print(f"[PDF PROXY] Arquivo não é ZIP válido, tentando como PDF")

            # Parâmetro para forçar download
            download = request.args.get('download', 'false').lower() == 'true'

            pdf_buffer = BytesIO(content)
            return send_file(
                pdf_buffer,
                mimetype='application/pdf',
                as_attachment=download,
                download_name=filename if filename.lower().endswith('.pdf') else f"{filename}.pdf"
            )
        else:
            print(f"[PDF PROXY] Erro PNCP: {resp.status_code}")
            return jsonify({"error": f"PDF não encontrado no PNCP (status {resp.status_code})"}), 404

    except Exception as e:
        print(f"[PDF PROXY] Erro: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Erro ao buscar PDF: {str(e)}"}), 500


# =============================================================================
# Fontes Routes
# =============================================================================

@app.route("/api/fontes", methods=["GET"])
def listar_fontes_api():
    resultado = tool_listar_fontes()
    return jsonify(resultado)


# =============================================================================
# Análises Routes
# =============================================================================

@app.route("/api/analises", methods=["GET"])
@require_auth
def listar_analises():
    user_id = get_current_user_id()
    db = get_db()
    try:
        analises = db.query(Analise).filter(
            Analise.user_id == user_id
        ).order_by(Analise.created_at.desc()).limit(50).all()

        return jsonify({"analises": [a.to_dict() for a in analises]})
    finally:
        db.close()


@app.route("/api/analises/<analise_id>", methods=["GET"])
@require_auth
def get_analise(analise_id):
    user_id = get_current_user_id()
    db = get_db()
    try:
        analise = db.query(Analise).filter(
            Analise.id == analise_id,
            Analise.user_id == user_id
        ).first()

        if not analise:
            return jsonify({"error": "Análise não encontrada"}), 404

        return jsonify(analise.to_dict(include_detalhes=True))
    finally:
        db.close()


# =============================================================================
# Propostas Routes
# =============================================================================

@app.route("/api/propostas", methods=["GET"])
@require_auth
def listar_propostas():
    user_id = get_current_user_id()
    db = get_db()
    try:
        propostas = db.query(Proposta).filter(
            Proposta.user_id == user_id
        ).order_by(Proposta.created_at.desc()).limit(50).all()

        return jsonify({"propostas": [p.to_dict() for p in propostas]})
    finally:
        db.close()


@app.route("/api/propostas/<proposta_id>", methods=["GET"])
@require_auth
def get_proposta(proposta_id):
    user_id = get_current_user_id()
    db = get_db()
    try:
        proposta = db.query(Proposta).filter(
            Proposta.id == proposta_id,
            Proposta.user_id == user_id
        ).first()

        if not proposta:
            return jsonify({"error": "Proposta não encontrada"}), 404

        return jsonify(proposta.to_dict())
    finally:
        db.close()


# =============================================================================
# Dashboard Stats Route
# =============================================================================

@app.route("/api/dashboard/stats", methods=["GET"])
@require_auth
def get_dashboard_stats():
    """Retorna estatísticas reais do banco para o Dashboard, filtradas pelo user_id."""
    from sqlalchemy import func, case
    from datetime import date, timedelta

    user_id = get_current_user_id()
    db = get_db()
    try:
        # Total de editais
        total_editais = db.query(Edital).filter(Edital.user_id == user_id).count()

        # Editais por status
        status_rows = (
            db.query(Edital.status, func.count(Edital.id))
            .filter(Edital.user_id == user_id)
            .group_by(Edital.status)
            .all()
        )
        editais_por_status = {row[0]: row[1] for row in status_rows}

        # Total de propostas
        total_propostas = db.query(Proposta).filter(Proposta.user_id == user_id).count()

        # Propostas por status
        prop_status_rows = (
            db.query(Proposta.status, func.count(Proposta.id))
            .filter(Proposta.user_id == user_id)
            .group_by(Proposta.status)
            .all()
        )
        propostas_por_status = {row[0]: row[1] for row in prop_status_rows}

        # Taxa de sucesso: editais ganhos / (ganhos + perdidos)
        ganhos = editais_por_status.get('ganho', 0) + editais_por_status.get('vencedor', 0)
        perdidos = editais_por_status.get('perdido', 0) + editais_por_status.get('perdedor', 0)
        total_encerrados = ganhos + perdidos
        taxa_sucesso = round(ganhos / total_encerrados, 4) if total_encerrados > 0 else 0.0

        # Valor total contratado
        valor_row = (
            db.query(func.coalesce(func.sum(Contrato.valor_total), 0))
            .filter(Contrato.user_id == user_id)
            .scalar()
        )
        valor_total_contratado = float(valor_row) if valor_row else 0.0

        # Editais por mês (últimos 6 meses)
        seis_meses_atras = date.today().replace(day=1)
        # Recalcular como 6 meses atrás
        mes = seis_meses_atras.month - 6
        ano = seis_meses_atras.year
        if mes <= 0:
            mes += 12
            ano -= 1
        from datetime import datetime as dt
        data_inicio = dt(ano, mes, 1)

        mes_rows = (
            db.query(
                func.date_format(Edital.created_at, '%Y-%m').label('mes'),
                func.count(Edital.id).label('total')
            )
            .filter(
                Edital.user_id == user_id,
                Edital.created_at >= data_inicio
            )
            .group_by('mes')
            .order_by('mes')
            .all()
        )
        editais_por_mes = [{"mes": row[0], "total": row[1]} for row in mes_rows]

        # Próximos prazos (editais com data_abertura nos próximos 30 dias)
        hoje = date.today()
        em_30_dias = hoje + timedelta(days=30)
        prazo_rows = (
            db.query(Edital.numero, Edital.data_abertura)
            .filter(
                Edital.user_id == user_id,
                Edital.data_abertura >= hoje,
                Edital.data_abertura <= em_30_dias,
                Edital.status.in_(['novo', 'analisando', 'participando'])
            )
            .order_by(Edital.data_abertura)
            .limit(10)
            .all()
        )
        proximos_prazos = [
            {
                "edital": row[0],
                "prazo": row[1].date().isoformat() if row[1] else None,
                "dias_restantes": (row[1].date() - hoje).days if row[1] else None
            }
            for row in prazo_rows
        ]

        return jsonify({
            "total_editais": total_editais,
            "editais_por_status": editais_por_status,
            "total_propostas": total_propostas,
            "propostas_por_status": propostas_por_status,
            "taxa_sucesso": taxa_sucesso,
            "valor_total_contratado": valor_total_contratado,
            "editais_por_mes": editais_por_mes,
            "proximos_prazos": proximos_prazos
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


# =============================================================================
# Endpoints: Modalidades, Origens, Áreas de Produto (filtros parametrizáveis)
# =============================================================================

def inferir_tipo_produto(objeto: str) -> str:
    """Infere o tipo de produto a partir do texto do objeto do edital."""
    if not objeto:
        return ""
    obj = objeto.lower()
    # Reagentes
    if any(w in obj for w in ['reagente', 'kit diagnóstico', 'kit diagnostico', 'teste rápido', 'teste rapido',
                                'hematologi', 'bioquímic', 'bioquimic', 'imunologi', 'microbiologi', 'sorologi']):
        return "Reagentes"
    # Equipamentos
    if any(w in obj for w in ['equipamento', 'analisador', 'microscópio', 'microscopio', 'centrífuga', 'centrifuga',
                                'autoclave', 'espectrofotômetro', 'espectrofotometro', 'aparelho', 'máquina', 'maquina']):
        return "Equipamentos"
    # Comodato
    if any(w in obj for w in ['comodato', 'cessão de uso', 'cessao de uso']):
        return "Comodato"
    # Aluguel
    if any(w in obj for w in ['aluguel', 'locação', 'locacao', 'locaçao']):
        return "Aluguel"
    # Insumos Hospitalares
    if any(w in obj for w in ['insumo', 'material hospitalar', 'luva', 'seringa', 'agulha', 'gaze',
                                'cateter', 'sonda', 'material médico', 'material medico']):
        return "Insumos Hospitalares"
    # Informática
    if any(w in obj for w in ['computador', 'servidor', 'notebook', 'software', 'licença software',
                                'impressora', 'monitor', 'informática', 'informatica', 'desktop']):
        return "Informática"
    # Redes
    if any(w in obj for w in ['switch', 'roteador', 'firewall', 'cabeamento', 'fibra óptica', 'fibra optica', 'rede']):
        return "Redes"
    # Mobiliário
    if any(w in obj for w in ['mobiliário', 'mobiliario', 'mesa', 'cadeira', 'armário', 'armario', 'estante']):
        return "Mobiliário"
    return ""


def inferir_origem_orgao(orgao: str, orgao_tipo: str = "") -> str:
    """Infere a origem/tipo do órgão licitante."""
    if not orgao:
        return ""
    org = orgao.lower()
    tipo = (orgao_tipo or "").lower()

    if any(w in org for w in ['hospital', 'hc ', 'hcpa', 'hu ', 'santa casa']):
        return "Hospital"
    if any(w in org for w in ['universidade', 'ufrj', 'usp', 'unicamp', 'ufmg', 'ufrgs', 'ufsc', 'ufba',
                                'ufscar', 'ufpe', 'unb', 'ufpr', 'unifesp', 'unesp', 'uerj']):
        return "Universidade"
    if any(w in org for w in ['lacen', 'laboratório central', 'laboratorio central']):
        return "LACEN"
    if any(w in org for w in ['exército', 'exercito', 'marinha', 'aeronáutica', 'aeronautica', 'força aérea',
                                'forca aerea', 'comando militar', 'arsenal']):
        return "Força Armada"
    if any(w in org for w in ['fapesp', 'faperj', 'fapemig', 'cnpq', 'capes', 'fundação de amparo',
                                'fundacao de amparo']):
        return "Fundações de Pesquisa"
    if any(w in org for w in ['fiocruz', 'embrapa', 'inpe', 'inpa', 'impa', 'lncc', 'cbpf',
                                'centro de pesquisa', 'instituto de pesquisa']):
        return "Centros de Pesquisas"
    if any(w in org for w in ['fundação', 'fundacao']) and 'pesquisa' not in org:
        return "Fundações Diversas"
    if any(w in org for w in ['autarquia', 'agência', 'agencia', 'anatel', 'anvisa', 'anp', 'ans', 'ana',
                                'ibama', 'inmetro']):
        return "Autarquia"
    if tipo == 'federal' or any(w in org for w in ['ministério', 'ministerio', 'governo federal', 'união', 'uniao']):
        return "Federal"
    if tipo == 'estadual' or any(w in org for w in ['governo do estado', 'secretaria de estado', 'estado de',
                                                      'governo estadual']):
        return "Estadual"
    if tipo == 'municipal' or any(w in org for w in ['prefeitura', 'município', 'municipio', 'câmara municipal',
                                                       'camara municipal']):
        return "Municipal"
    return ""


@app.route("/api/modalidades", methods=["GET"])
@require_auth
def listar_modalidades():
    """Lista todas as modalidades de licitação ativas."""
    db = get_db()
    try:
        items = db.query(ModalidadeLicitacao).filter(ModalidadeLicitacao.ativo == True).order_by(ModalidadeLicitacao.ordem).all()
        return jsonify({"success": True, "items": [m.to_dict() for m in items]})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        db.close()


@app.route("/api/origens", methods=["GET"])
@require_auth
def listar_origens():
    """Lista todas as origens de órgão ativas."""
    db = get_db()
    try:
        items = db.query(OrigemOrgao).filter(OrigemOrgao.ativo == True).order_by(OrigemOrgao.ordem).all()
        return jsonify({"success": True, "items": [o.to_dict() for o in items]})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        db.close()


@app.route("/api/areas-produto", methods=["GET"])
@require_auth
def listar_areas_produto():
    """Lista áreas com classes e subclasses aninhadas (filtradas por empresa)."""
    db = get_db()
    try:
        empresa_id = get_current_empresa_id()
        query = db.query(AreaProduto).filter(AreaProduto.ativo == True)
        if empresa_id:
            query = query.filter(AreaProduto.empresa_id == empresa_id)
        items = query.order_by(AreaProduto.ordem).all()
        return jsonify({"success": True, "items": [a.to_dict() for a in items]})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        db.close()


# =============================================================================
# Score Profundo: 6 dimensões via Reasoner para editais da busca
# =============================================================================

def _salvar_edital_temp_para_score(edital_data, user_id, empresa_id):
    """Salva edital da busca no banco para poder chamar tool_calcular_scores_validacao.
    Retorna edital_id se criado/encontrado, None se erro."""
    db = get_db()
    try:
        numero = edital_data.get("numero", "")
        orgao = edital_data.get("orgao", "")

        # Verificar se já existe
        existe = db.query(Edital).filter(
            Edital.user_id == user_id,
            Edital.numero == numero,
            Edital.orgao == orgao
        ).first()

        if existe:
            return str(existe.id)

        # Criar novo
        from datetime import datetime as _dt
        novo = Edital(
            user_id=user_id,
            empresa_id=empresa_id,
            numero=numero or f"TEMP-{str(uuid.uuid4())[:8]}",
            orgao=orgao or "Não identificado",
            objeto=edital_data.get("objeto") or edital_data.get("descricao") or "—",
            modalidade="pregao_eletronico",
            status="temp_score",
            fonte=edital_data.get("fonte"),
            uf=edital_data.get("uf"),
            cidade=edital_data.get("cidade"),
            url=edital_data.get("url") or edital_data.get("link"),
            cnpj_orgao=edital_data.get("cnpj_orgao"),
            ano_compra=edital_data.get("ano_compra"),
            seq_compra=edital_data.get("seq_compra"),
        )
        if edital_data.get("valor_referencia") or edital_data.get("valor_estimado"):
            try:
                novo.valor_referencia = float(edital_data.get("valor_referencia") or edital_data.get("valor_estimado"))
            except (ValueError, TypeError):
                pass
        if edital_data.get("data_abertura"):
            try:
                d = edital_data["data_abertura"]
                if isinstance(d, str):
                    for fmt in ("%d/%m/%Y %H:%M", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d"):
                        try:
                            novo.data_abertura = _dt.strptime(d[:19], fmt)
                            break
                        except (ValueError, TypeError):
                            continue
            except Exception:
                pass

        db.add(novo)
        db.commit()
        edital_id = str(novo.id)
        print(f"[SCORE_PROFUNDO] Edital salvo: {numero} -> {edital_id}")
        return edital_id
    except Exception as e:
        db.rollback()
        print(f"[SCORE_PROFUNDO] Erro ao salvar edital: {e}")
        return None
    finally:
        db.close()


def _calcular_score_profundo(editais, user_id, empresa_id):
    """Score profundo: 6 dimensões via DeepSeek para cada edital."""
    import time
    from concurrent.futures import ThreadPoolExecutor, as_completed

    if not editais:
        return editais

    t0 = time.time()

    def _score_one(edital):
        try:
            edital_id = _salvar_edital_temp_para_score(edital, user_id, empresa_id)
            if not edital_id:
                edital["score_tecnico"] = edital.get("score_tecnico", 0)
                edital["score_profundo"] = None
                return edital

            resultado = tool_calcular_scores_validacao(
                edital_id=edital_id, user_id=user_id, produto_id=None
            )

            if resultado.get("success"):
                scores = resultado.get("scores", {})
                score_geral = resultado.get("score_geral", 0)
                # Só sobrescreve score_tecnico se não havia (modo profundo puro)
                # No modo hibrido, preserva o score rápido para ordenação
                if not edital.get("score_tecnico"):
                    edital["score_tecnico"] = score_geral
                edital["score_profundo"] = {
                    "scores": scores,
                    "score_geral": score_geral,
                    "decisao": resultado.get("decisao", "AVALIAR"),
                    "justificativa": resultado.get("justificativa", ""),
                    "pontos_positivos": resultado.get("pontos_positivos", []),
                    "pontos_atencao": resultado.get("pontos_atencao", []),
                }
                # Mapear decisao GO/NO-GO para recomendacao PARTICIPAR/NÃO PARTICIPAR
                decisao = resultado.get("decisao", "AVALIAR")
                if decisao == "GO":
                    edital["recomendacao"] = "PARTICIPAR"
                elif decisao == "NO-GO":
                    edital["recomendacao"] = "NÃO PARTICIPAR"
                else:
                    edital["recomendacao"] = "AVALIAR"
                edital["justificativa"] = resultado.get("justificativa", "")
            else:
                edital["score_profundo"] = None
        except Exception as e:
            print(f"[SCORE_PROFUNDO] Erro: {e}")
            edital["score_profundo"] = None
        return edital

    MAX_WORKERS = min(len(editais), 2)
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = {executor.submit(_score_one, e): e for e in editais}
        for future in as_completed(futures, timeout=600):
            try:
                future.result()
            except Exception as e:
                print(f"[SCORE_PROFUNDO] Future erro: {e}")

    editais.sort(key=lambda x: x.get("score_tecnico", 0), reverse=True)

    # Limpar editais temporários criados para o cálculo de score
    try:
        db = get_db()
        temp_deletados = db.query(Edital).filter(
            Edital.user_id == user_id,
            Edital.status == "temp_score"
        ).delete()
        db.commit()
        if temp_deletados:
            print(f"[SCORE_PROFUNDO] {temp_deletados} editais temporarios removidos")
        db.close()
    except Exception as e:
        print(f"[SCORE_PROFUNDO] Erro ao limpar temporarios: {e}")

    elapsed = time.time() - t0
    print(f"[SCORE_PROFUNDO] {len(editais)} editais em {elapsed:.1f}s")
    return editais


# =============================================================================
# Onda 2 - T8: Endpoint REST GET /api/editais/buscar
# =============================================================================

@app.route("/api/editais/buscar", methods=["GET"])
@require_auth
def buscar_editais_rest():
    """
    Endpoint REST para busca de editais — mesma lógica do chat.
    Usa _buscar_editais_multifonte() (PNCP API + scraper web + deduplicação).
    Query params:
      - termo (obrigatório): termo de busca
      - uf (opcional): filtrar por UF (ex: SP, RJ)
      - calcularScore (opcional, default true): calcular score aderência
      - incluirEncerrados (opcional, default false): incluir editais encerrados
      - limite / limit (opcional, default 20): máximo de resultados
      - diasBusca (opcional, default 90): janela de publicação em dias (0 = sem limite)
      - fonte (opcional): fonte específica (PNCP, ComprasNet, etc). Se omitido, busca em todas.
      - modalidade (opcional): filtrar por modalidade (ex: Pregão Eletrônico)
      - tipoProduto / tipo_produto (opcional): filtrar por tipo de produto inferido do objeto
      - origem (opcional): filtrar por origem do órgão inferida
    """
    user_id = get_current_user_id()
    termo = request.args.get("termo", "").strip()
    uf = request.args.get("uf", "").strip().upper() or None
    fonte = request.args.get("fonte", request.args.get("fontes", "")).strip() or None
    modalidade = request.args.get("modalidade", "").strip() or None
    tipo_produto = request.args.get("tipoProduto", request.args.get("tipo_produto", "")).strip() or None
    origem = request.args.get("origem", "").strip() or None
    # Score: tipoScore (nenhum/rapido/hibrido/profundo) com retrocompatibilidade
    tipo_score = request.args.get("tipoScore", request.args.get("tipo_score", "")).strip()
    limite_score_profundo = int(request.args.get("limiteScoreProfundo", request.args.get("limite_score_profundo", 10)))
    if limite_score_profundo == 0:
        limite_score_profundo = 9999  # "Todos"

    # Retrocompatibilidade: calcularScore=true sem tipoScore → "rapido"
    if not tipo_score:
        calcular_score_str = request.args.get("calcularScore", request.args.get("calcular_score", "true"))
        tipo_score = "rapido" if calcular_score_str.lower() == "true" else "nenhum"

    calcular_score = tipo_score != "nenhum"
    incluir_encerrados_str = request.args.get("incluirEncerrados", request.args.get("incluir_encerrados", "false"))
    incluir_encerrados = incluir_encerrados_str.lower() == "true"
    limite = int(request.args.get("limite", request.args.get("limit", 20)))
    dias_busca = int(request.args.get("diasBusca", request.args.get("dias_busca", 90)))
    if dias_busca != 0:
        dias_busca = max(7, min(730, dias_busca))

    if not termo:
        return jsonify({"success": False, "error": "Parâmetro 'termo' é obrigatório"}), 400

    try:
        # Busca multi-fonte: mesma lógica do chat (PNCP + scraper web + deduplicação)
        resultado_busca = _buscar_editais_multifonte(
            termo=termo,
            user_id=user_id,
            uf=uf,
            incluir_encerrados=incluir_encerrados,
            buscar_detalhes=True,
            dias_busca=dias_busca,
            fonte=fonte,
            modalidade=modalidade,
            tipo_produto=tipo_produto,
            origem=origem,
        )

        editais = resultado_busca.get("editais", [])

        # Buscas paralelas extras com termos CATMAT/semânticos dos produtos
        if tipo_score != "nenhum":
            try:
                from concurrent.futures import ThreadPoolExecutor, as_completed
                import unicodedata as _ud

                def _norm_termo(txt):
                    if not txt:
                        return ""
                    txt = _ud.normalize('NFKD', txt).encode('ascii', 'ignore').decode('ascii')
                    return txt.lower().strip()

                # Normalizar termo de busca para matching
                termo_norm = _norm_termo(termo) if termo else ""
                termo_words = {w for w in termo_norm.split() if len(w) > 3}

                db_termos = get_db()
                produtos_user = db_termos.query(Produto).filter(Produto.user_id == user_id).all()
                termos_busca_extras = set()

                # Só extrair termos de produtos relevantes ao termo buscado
                for p in produtos_user:
                    p_nome_norm = _norm_termo(p.nome or "")
                    p_cat_norm = _norm_termo(p.categoria or "")
                    p_words = set(p_nome_norm.split()) | set(p_cat_norm.split())
                    overlap = termo_words & {w for w in p_words if len(w) > 3}
                    if not overlap:
                        continue
                    for desc in (p.catmat_descricoes or [])[:2]:
                        palavras = [w for w in desc.split() if len(w) > 3][:4]
                        if palavras:
                            termos_busca_extras.add(" ".join(palavras))
                    for t in (p.termos_busca or [])[:3]:
                        termos_busca_extras.add(t)
                db_termos.close()

                # Remover termos genéricos demais que trazem ruído
                TERMOS_GENERICOS_NORM = {
                    "material hospitalar", "material medico", "equipamento hospitalar",
                    "insumo hospitalar", "insumo laboratorial", "material laboratorio",
                    "reagente bioquimico", "equipamento medico", "material odontologico",
                    "servico geral", "material consumo", "material permanente",
                    "material medico hospitalar", "equipamento laboratorio",
                }
                termos_busca_extras = {
                    t for t in termos_busca_extras
                    if _norm_termo(t) not in TERMOS_GENERICOS_NORM and len(t.split()) >= 2
                }

                termos_extras = list(termos_busca_extras)[:5]
                if termos_extras:
                    print(f"[BUSCA] Buscas extras com {len(termos_extras)} termos CATMAT/semânticos: {termos_extras}")
                    with ThreadPoolExecutor(max_workers=3) as executor:
                        futures = {}
                        for t_extra in termos_extras:
                            f = executor.submit(
                                _buscar_editais_multifonte,
                                termo=t_extra, user_id=user_id, uf=uf,
                                incluir_encerrados=incluir_encerrados,
                                buscar_detalhes=True, dias_busca=dias_busca,
                                fonte=fonte,
                            )
                            futures[f] = t_extra

                        for f in as_completed(futures, timeout=30):
                            try:
                                res = f.result()
                                extras = res.get("editais", [])
                                if extras:
                                    # Filtrar: objeto deve ter pelo menos 1 palavra do termo original
                                    filtered = []
                                    for e in extras:
                                        obj_norm = _norm_termo(e.get("objeto", ""))
                                        if any(w in obj_norm for w in termo_words):
                                            filtered.append(e)
                                    editais.extend(filtered)
                                    print(f"[BUSCA] Termo extra '{futures[f]}': {len(filtered)}/{len(extras)} editais após filtro")
                            except Exception:
                                pass

                    # Deduplicar por número+órgão e fallback órgão+valor
                    seen = set()
                    seen2 = set()
                    dedup = []
                    for e in editais:
                        chave = f"{e.get('numero', '')}-{e.get('orgao', '')}"
                        orgao = (e.get('orgao') or '').strip().lower()
                        valor = e.get('valor_estimado') or e.get('valor_referencia') or 0
                        chave2 = f"{orgao}|{valor}" if orgao and valor else ""
                        if chave in seen:
                            continue
                        if chave2 and chave2 in seen2:
                            continue
                        seen.add(chave)
                        if chave2:
                            seen2.add(chave2)
                        dedup.append(e)
                    editais = dedup
                    print(f"[BUSCA] Total após dedup: {len(editais)} editais")
            except Exception as e:
                print(f"[BUSCA] Erro buscas extras: {e}")

        # Calcular score ANTES de aplicar limite (para pegar os melhores)
        empresa_id = get_current_empresa_id()
        if tipo_score == "rapido" and editais:
            resultado_score = tool_calcular_score_aderencia(editais=editais, user_id=user_id, empresa_id=empresa_id)
            if resultado_score.get("success"):
                editais = resultado_score.get("editais_com_score", editais)

        elif tipo_score == "hibrido" and editais:
            # Passo 1: Score rápido em TODOS
            resultado_score = tool_calcular_score_aderencia(editais=editais, user_id=user_id, empresa_id=empresa_id)
            if resultado_score.get("success"):
                editais = resultado_score.get("editais_com_score", editais)
            # Passo 2: Ordenar por score desc, pegar os N melhores para profundo
            editais.sort(key=lambda x: x.get("score_tecnico", 0), reverse=True)
            top_n = editais[:limite_score_profundo]
            rest = editais[limite_score_profundo:]
            # Passo 3: Score profundo nos N melhores
            top_n = _calcular_score_profundo(top_n, user_id, empresa_id)
            # Retornar TODOS: top N enriquecidos + restantes com score rápido
            editais = top_n + rest

        elif tipo_score == "profundo" and editais:
            # Score profundo direto nos N primeiros (sem triagem rápida)
            target = editais[:limite_score_profundo]
            rest = editais[limite_score_profundo:]
            target = _calcular_score_profundo(target, user_id, empresa_id)
            editais = target + rest

        # Ordenar: se tem score, por score desc; senão abertos primeiro por data
        from datetime import datetime as _dt
        def _sort_key_rest(e):
            encerrado = e.get("encerrado", False)
            score = e.get("score_tecnico", -1)
            d = e.get("data_encerramento") or e.get("data_abertura") or ""
            dt = _dt(2099, 12, 31)
            for fmt in ("%d/%m/%Y %H:%M", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d"):
                try:
                    dt = _dt.strptime(d[:19], fmt)
                    break
                except (ValueError, TypeError):
                    continue
            if calcular_score:
                # Score desc (negativo para reverse), depois abertos primeiro, depois data
                return (-score, 1 if encerrado else 0, dt)
            return (1 if encerrado else 0, dt)
        editais.sort(key=_sort_key_rest)

        # Aplicar limite DEPOIS da ordenação
        editais = editais[:limite]

        # Normalizar campos de retorno para o frontend
        editais_normalizados = []
        for i, e in enumerate(editais):
            # Gerar ID único: usar ID do banco se existir, senão criar hash do número+órgão
            edital_id = e.get("id")
            if not edital_id:
                import hashlib
                chave = f"{e.get('numero', '')}-{e.get('orgao', '')}-{i}"
                edital_id = hashlib.md5(chave.encode()).hexdigest()
            _objeto = e.get("objeto") or e.get("descricao") or ""
            _orgao = e.get("orgao") or ""
            _orgao_tipo = e.get("orgao_tipo") or "municipal"
            editais_normalizados.append({
                "id": edital_id,
                "numero": e.get("numero"),
                "orgao": _orgao,
                "orgao_tipo": _orgao_tipo,
                "objeto": _objeto,
                "modalidade": e.get("modalidade"),
                "valor_estimado": e.get("valor_referencia") or e.get("valor_estimado"),
                "data_publicacao": e.get("data_publicacao"),
                "data_abertura": e.get("data_abertura"),
                "data_encerramento": e.get("data_encerramento"),
                "uf": e.get("uf") or uf,
                "cidade": e.get("cidade"),
                "url": e.get("url") or e.get("link"),
                "fonte": e.get("fonte"),
                "fonte_tipo": e.get("fonte_tipo"),
                "score_tecnico": e.get("score_tecnico"),
                "recomendacao": e.get("recomendacao"),
                "justificativa": e.get("justificativa"),
                "produtos_aderentes": e.get("produtos_aderentes", []),
                "cnpj_orgao": e.get("cnpj_orgao"),
                "ano_compra": e.get("ano_compra"),
                "seq_compra": e.get("seq_compra"),
                "total_itens": e.get("total_itens", 0),
                "itens": e.get("itens", []),
                "dados_completos": e.get("dados_completos", False),
                "encerrado": e.get("encerrado", False),
                "pdf_url": e.get("pdf_url"),
                "tipo_produto_inferido": inferir_tipo_produto(_objeto),
                "origem_inferida": inferir_origem_orgao(_orgao, _orgao_tipo),
                "score_profundo": e.get("score_profundo"),
            })

        return jsonify({
            "success": True,
            "editais": editais_normalizados,
            "total": len(editais_normalizados),
            "termo": termo,
            "uf": uf,
            "fontes_usadas": resultado_busca.get("fontes_consultadas", []),
            "erros_fontes": resultado_busca.get("erros"),
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# =============================================================================
# Onda 2 - T9: Endpoint REST GET /api/editais/salvos
# =============================================================================

@app.route("/api/editais/salvos", methods=["GET"])
@require_auth
def listar_editais_salvos():
    """
    Lista editais salvos do usuário com scores e estratégias.
    Query params: status, uf, categoria, com_score (bool), com_estrategia (bool)
    """
    user_id = get_current_user_id()
    status = request.args.get("status")
    uf = request.args.get("uf")
    categoria = request.args.get("categoria")
    com_score = request.args.get("com_score", "false").lower() == "true"
    com_estrategia = request.args.get("com_estrategia", "false").lower() == "true"
    page = int(request.args.get("page", 1))
    per_page = int(request.args.get("per_page", 20))

    db = get_db()
    try:
        query = db.query(Edital).filter(
            Edital.user_id == user_id,
            Edital.status != "temp_score",  # Excluir editais temporários criados pelo score profundo
        )

        if status:
            query = query.filter(Edital.status == status)
        if uf:
            query = query.filter(Edital.uf == uf.upper())
        if categoria:
            query = query.filter(Edital.categoria == categoria)

        total = query.count()
        editais = query.order_by(Edital.created_at.desc()).offset((page - 1) * per_page).limit(per_page).all()

        resultado = []
        for edital in editais:
            edital_dict = edital.to_dict()

            # Incluir melhor score/análise se solicitado
            if com_score:
                melhor_analise = db.query(Analise).filter(
                    Analise.edital_id == edital.id,
                    Analise.user_id == user_id
                ).order_by(Analise.score_final.desc()).first()
                if melhor_analise:
                    edital_dict["score_tecnico"] = float(melhor_analise.score_tecnico) if melhor_analise.score_tecnico else None
                    edital_dict["score_final"] = float(melhor_analise.score_final) if melhor_analise.score_final else None
                    edital_dict["score_geral"] = float(melhor_analise.score_final) if melhor_analise.score_final else None
                    edital_dict["analise_id"] = melhor_analise.id
                    # Parsear recomendacao: pode ser JSON (novo) ou texto puro (antigo)
                    scores_parsed = False
                    try:
                        rec_data = json.loads(melhor_analise.recomendacao) if melhor_analise.recomendacao else {}
                        if isinstance(rec_data, dict) and "scores" in rec_data:
                            edital_dict["scores"] = rec_data["scores"]
                            edital_dict["decisao_ia"] = rec_data.get("decisao")
                            edital_dict["justificativa_ia"] = rec_data.get("justificativa")
                            edital_dict["pontos_positivos"] = rec_data.get("pontos_positivos", [])
                            edital_dict["pontos_atencao"] = rec_data.get("pontos_atencao", [])
                            scores_parsed = True
                    except (json.JSONDecodeError, TypeError):
                        pass

                    if not scores_parsed:
                        # Recomendacao é texto puro (score profundo antigo) — reconstruir scores dos campos
                        edital_dict["scores"] = {
                            "tecnico": float(melhor_analise.score_tecnico) if melhor_analise.score_tecnico else 0,
                            "documental": 0,
                            "complexidade": 0,
                            "juridico": 0,
                            "logistico": 0,
                            "comercial": float(melhor_analise.score_comercial) if melhor_analise.score_comercial else 0,
                        }
                        rec_text = melhor_analise.recomendacao or ""
                        if rec_text.startswith("GO"):
                            edital_dict["decisao_ia"] = "GO"
                        elif rec_text.startswith("NO-GO"):
                            edital_dict["decisao_ia"] = "NO-GO"
                        else:
                            edital_dict["decisao_ia"] = "AVALIAR"
                        edital_dict["justificativa_ia"] = rec_text.split(": ", 1)[1] if ": " in rec_text else rec_text
                        edital_dict["recomendacao"] = rec_text

            # Incluir estratégia/decisão go-nogo se solicitado
            if com_estrategia:
                estrategia = db.query(EstrategiaEdital).filter(
                    EstrategiaEdital.edital_id == edital.id,
                    EstrategiaEdital.user_id == user_id
                ).first()
                if estrategia:
                    edital_dict["decisao"] = estrategia.decisao
                    edital_dict["prioridade"] = estrategia.prioridade
                    edital_dict["margem_desejada"] = float(estrategia.margem_desejada) if estrategia.margem_desejada else None
                    edital_dict["justificativa_estrategia"] = estrategia.justificativa
                    edital_dict["estrategia_id"] = estrategia.id

            resultado.append(edital_dict)

        return jsonify({
            "success": True,
            "editais": resultado,
            "total": total,
            "page": page,
            "per_page": per_page,
            "pages": (total + per_page - 1) // per_page
        })

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        db.close()


# =============================================================================
# Onda 2 - T10: Endpoint POST /api/editais/<id>/scores-validacao
# =============================================================================

@app.route("/api/editais/<edital_id>/scores-validacao", methods=["POST"])
@require_auth
def calcular_scores_validacao_rest(edital_id):
    """
    Calcula 6 dimensões de score de validação para um edital.
    Body (opcional): {"produto_id": "<uuid>"}
    """
    user_id = get_current_user_id()
    data = request.json or {}
    produto_id = data.get("produto_id")

    try:
        resultado = tool_calcular_scores_validacao(
            edital_id=edital_id,
            user_id=user_id,
            produto_id=produto_id
        )
        return jsonify(resultado)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# Endpoint para calcular score profundo sob demanda (edital não salvo)
@app.route("/api/editais/salvar-scores-captacao", methods=["POST"])
@require_auth
def salvar_scores_captacao():
    """Salva scores calculados na Captação para que a Validação não precise recalcular."""
    user_id = get_current_user_id()
    data = request.get_json()
    edital_id = data.get("edital_id")
    if not edital_id:
        return jsonify({"error": "edital_id obrigatório"}), 400

    db = get_db()
    try:
        edital = db.query(Edital).filter(Edital.id == edital_id, Edital.user_id == user_id).first()
        if not edital:
            return jsonify({"error": "Edital não encontrado"}), 404

        empresa = db.query(Empresa).filter(Empresa.user_id == user_id).first()
        empresa_id = empresa.id if empresa else None

        # Buscar produto correspondente (necessário para Analise — produto_id é NOT NULL)
        from models import Produto
        produto = db.query(Produto).filter(Produto.user_id == user_id).first()
        if not produto:
            return jsonify({"error": "Nenhum produto cadastrado"}), 400
        produto_id = produto.id

        # Verificar se já existe análise para este edital
        analise_existente = db.query(Analise).filter(
            Analise.edital_id == edital_id,
            Analise.user_id == user_id
        ).first()

        score_tecnico = data.get("score_tecnico", 0)
        score_comercial = data.get("score_comercial", 0)
        score_final = data.get("score_final", 0)
        decisao = data.get("decisao", "AVALIAR")
        justificativa = data.get("justificativa", "")

        recomendacao_text = json.dumps({
            "decisao": decisao,
            "justificativa": justificativa,
            "scores": {
                "tecnico": score_tecnico,
                "documental": data.get("score_documental", 0),
                "complexidade": data.get("score_complexidade", 0),
                "juridico": data.get("score_juridico", 0),
                "logistico": data.get("score_logistico", 0),
                "comercial": score_comercial,
            },
            "pontos_positivos": data.get("pontos_positivos", []),
            "pontos_atencao": data.get("pontos_atencao", []),
        }, ensure_ascii=False)

        if analise_existente:
            analise_existente.score_tecnico = score_tecnico
            analise_existente.score_comercial = score_comercial
            analise_existente.score_final = score_final
            analise_existente.recomendacao = recomendacao_text
        else:
            nova_analise = Analise(
                id=str(uuid.uuid4()),
                edital_id=edital_id,
                produto_id=produto_id,
                user_id=user_id,
                empresa_id=empresa_id,
                score_tecnico=score_tecnico,
                score_comercial=score_comercial,
                score_final=score_final,
                recomendacao=recomendacao_text,
                created_at=datetime.now(),
            )
            db.add(nova_analise)

        db.commit()
        return jsonify({"success": True})
    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


@app.route("/api/editais/score-profundo-sob-demanda", methods=["POST"])
@require_auth
def score_profundo_sob_demanda():
    """
    Recebe dados de um edital não salvo, salva temporariamente,
    calcula score profundo e retorna resultado.
    """
    user_id = get_current_user_id()
    data = request.json or {}

    if not data.get("objeto"):
        return jsonify({"success": False, "error": "Campo 'objeto' obrigatório"}), 400

    try:
        # Buscar empresa_id do usuário
        db = get_db()
        empresa = db.query(Empresa).filter(Empresa.user_id == user_id).first()
        empresa_id = str(empresa.id) if empresa else user_id
        db.close()

        # Salvar edital temporário
        edital_id = _salvar_edital_temp_para_score(data, user_id, empresa_id)
        if not edital_id:
            return jsonify({"success": False, "error": "Erro ao salvar edital temporário"}), 500

        # Calcular score profundo
        resultado = tool_calcular_scores_validacao(
            edital_id=edital_id,
            user_id=user_id,
            produto_id=data.get("produto_id")
        )

        # Retornar o edital_id para que o frontend possa usá-lo depois
        if resultado.get("success"):
            resultado["edital_id_temp"] = edital_id

        return jsonify(resultado)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# =============================================================================
# Onda 2 - T11: Endpoint PUT /api/propostas/<id>/status
# =============================================================================

@app.route("/api/propostas/<proposta_id>/status", methods=["PUT"])
@require_auth
def atualizar_status_proposta(proposta_id):
    """
    Atualiza o status de uma proposta seguindo o workflow:
    rascunho -> revisao -> aprovada -> enviada -> aceita/rejeitada
    """
    user_id = get_current_user_id()
    data = request.json or {}
    novo_status = data.get("status", "").strip()

    # Status do modelo: rascunho, revisao, aprovada, enviada
    TRANSICOES_VALIDAS = {
        "rascunho": ["revisao"],
        "revisao": ["rascunho", "aprovada"],
        "aprovada": ["revisao", "enviada"],
        "enviada": ["aprovada", "rascunho"]
    }

    STATUS_VALIDOS = list(TRANSICOES_VALIDAS.keys())

    if not novo_status:
        return jsonify({"success": False, "error": "Campo 'status' é obrigatório"}), 400

    if novo_status not in STATUS_VALIDOS:
        return jsonify({
            "success": False,
            "error": f"Status inválido. Válidos: {', '.join(STATUS_VALIDOS)}"
        }), 400

    db = get_db()
    try:
        proposta = db.query(Proposta).filter(
            Proposta.id == proposta_id,
            Proposta.user_id == user_id
        ).first()

        if not proposta:
            return jsonify({"success": False, "error": "Proposta não encontrada"}), 404

        status_atual = proposta.status
        transicoes = TRANSICOES_VALIDAS.get(status_atual, [])

        if novo_status not in transicoes:
            return jsonify({
                "success": False,
                "error": f"Transição inválida: '{status_atual}' -> '{novo_status}'. Permitidas: {transicoes}"
            }), 400

        proposta.status = novo_status
        proposta.updated_at = datetime.now()
        db.commit()

        return jsonify({
            "success": True,
            "proposta_id": proposta_id,
            "status_anterior": status_atual,
            "status_atual": novo_status,
            "proposta": proposta.to_dict()
        })

    except Exception as e:
        db.rollback()
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        db.close()


# =============================================================================
# Onda 2 - T12: Endpoint export PDF/DOCX propostas
# =============================================================================

@app.route("/api/propostas/<proposta_id>/export", methods=["GET"])
@require_auth
def exportar_proposta(proposta_id):
    """
    Exporta proposta em PDF ou DOCX.
    Query param: formato (pdf | docx), padrão = pdf
    """
    user_id = get_current_user_id()
    formato = request.args.get("formato", "pdf").lower()

    if formato not in ("pdf", "docx"):
        return jsonify({"success": False, "error": "Formato inválido. Use 'pdf' ou 'docx'"}), 400

    db = get_db()
    try:
        proposta = db.query(Proposta).filter(
            Proposta.id == proposta_id,
            Proposta.user_id == user_id
        ).first()

        if not proposta:
            return jsonify({"success": False, "error": "Proposta não encontrada"}), 404

        edital = db.query(Edital).filter(Edital.id == proposta.edital_id).first()
        produto = db.query(Produto).filter(Produto.id == proposta.produto_id).first()

        # Montar conteúdo do documento
        titulo = f"Proposta Técnica - {edital.numero if edital else proposta_id}"
        conteudo_linhas = [
            f"PROPOSTA TÉCNICA E COMERCIAL",
            f"",
            f"Edital: {edital.numero if edital else 'N/A'}",
            f"Órgão: {edital.orgao if edital else 'N/A'}",
            f"Objeto: {edital.objeto[:300] if edital and edital.objeto else 'N/A'}",
            f"",
            f"PRODUTO OFERTADO",
            f"Produto: {produto.nome if produto else 'N/A'}",
            f"Fabricante: {produto.fabricante if produto else 'N/A'}",
            f"Modelo: {produto.modelo if produto else 'N/A'}",
            f"",
            f"PROPOSTA COMERCIAL",
            f"Preço Unitário: R$ {proposta.preco_unitario:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') if proposta.preco_unitario else "Preço Unitário: N/A",
            f"Quantidade: {proposta.quantidade or 1}",
            f"Preço Total: R$ {proposta.preco_total:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.') if proposta.preco_total else "Preço Total: N/A",
            f"",
            f"DESCRIÇÃO TÉCNICA",
            proposta.texto_tecnico or "Sem descrição técnica.",
            f"",
            f"Status: {proposta.status}",
            f"Data: {proposta.created_at.strftime('%d/%m/%Y') if proposta.created_at else 'N/A'}",
        ]
        conteudo_texto = "\n".join(conteudo_linhas)

        import tempfile, os as _os

        if formato == "docx":
            try:
                from docx import Document
                doc = Document()
                doc.add_heading(titulo, 0)
                for linha in conteudo_linhas:
                    if linha:
                        doc.add_paragraph(linha)
                    else:
                        doc.add_paragraph("")
                tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                doc.save(tmp.name)
                tmp.close()
                return send_file(
                    tmp.name,
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    as_attachment=True,
                    download_name=f"proposta_{proposta_id[:8]}.docx"
                )
            except ImportError:
                return jsonify({"success": False, "error": "python-docx não instalado. Execute: pip install python-docx"}), 500

        else:  # pdf
            try:
                from weasyprint import HTML
                html_content = f"""
                <html><head><meta charset="utf-8">
                <style>
                  body {{ font-family: Arial, sans-serif; margin: 40px; font-size: 12px; }}
                  h1 {{ color: #1e3a5f; border-bottom: 2px solid #1e3a5f; }}
                  h2 {{ color: #1e3a5f; margin-top: 20px; }}
                  .label {{ font-weight: bold; }}
                  table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
                  td {{ padding: 4px 8px; border: 1px solid #ddd; }}
                </style>
                </head><body>
                <h1>{titulo}</h1>
                <pre style="white-space: pre-wrap;">{conteudo_texto}</pre>
                </body></html>
                """
                tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
                HTML(string=html_content).write_pdf(tmp.name)
                tmp.close()
                return send_file(
                    tmp.name,
                    mimetype="application/pdf",
                    as_attachment=True,
                    download_name=f"proposta_{proposta_id[:8]}.pdf"
                )
            except ImportError:
                # Fallback: gerar PDF simples com reportlab
                try:
                    from reportlab.lib.pagesizes import A4
                    from reportlab.pdfgen import canvas
                    from reportlab.lib.units import cm

                    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
                    c = canvas.Canvas(tmp.name, pagesize=A4)
                    largura, altura = A4
                    y = altura - 2 * cm

                    c.setFont("Helvetica-Bold", 16)
                    c.drawString(2 * cm, y, titulo)
                    y -= 1 * cm

                    c.setFont("Helvetica", 10)
                    for linha in conteudo_linhas:
                        if y < 2 * cm:
                            c.showPage()
                            y = altura - 2 * cm
                            c.setFont("Helvetica", 10)
                        c.drawString(2 * cm, y, linha[:100])
                        y -= 0.5 * cm

                    c.save()
                    tmp.close()
                    return send_file(
                        tmp.name,
                        mimetype="application/pdf",
                        as_attachment=True,
                        download_name=f"proposta_{proposta_id[:8]}.pdf"
                    )
                except ImportError:
                    return jsonify({"success": False, "error": "Nenhuma lib de PDF disponível. Instale weasyprint ou reportlab"}), 500

    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    finally:
        db.close()


# =============================================================================
# Documentos Necessários — lista de tipos para dropdown
# =============================================================================

@app.route("/api/documentos-necessarios", methods=["GET"])
@require_auth
def listar_documentos_necessarios():
    """Retorna todos os tipos de documentos necessários (ativos), agrupados por categoria."""
    from models import DocumentoNecessario, CategoriaDocumento
    db = get_db()
    try:
        docs = db.query(DocumentoNecessario).outerjoin(CategoriaDocumento).filter(
            DocumentoNecessario.ativo == True
        ).order_by(CategoriaDocumento.ordem, DocumentoNecessario.ordem).all()
        return jsonify([d.to_dict() for d in docs])
    finally:
        db.close()


# =============================================================================
# B1: Upload real de arquivo na EmpresaPage (persistir no disco)
# =============================================================================

@app.route("/api/empresa-documentos/upload", methods=["POST"])
@require_auth
def upload_empresa_documento():
    """
    Upload de documento da empresa (contrato social, atestados, etc).
    Salva arquivo no disco e cria/atualiza registro EmpresaDocumento.

    Form data:
    - file: arquivo (PDF, DOCX, etc.)
    - empresa_id: ID da empresa
    - tipo: tipo do documento (contrato_social, atestado_capacidade, etc.)
    - data_emissao: (opcional) data de emissão
    - data_vencimento: (opcional) data de vencimento
    """
    from models import Empresa, EmpresaDocumento, DocumentoNecessario
    user_id = get_current_user_id()

    if 'file' not in request.files:
        return jsonify({"error": "Arquivo não enviado"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Nenhum arquivo selecionado"}), 400

    empresa_id = request.form.get('empresa_id')
    documento_necessario_id = request.form.get('documento_necessario_id')

    _tipos_validos = {
        'contrato_social', 'atestado_capacidade', 'balanco', 'alvara',
        'registro_conselho', 'procuracao', 'certidao_negativa',
        'habilitacao_fiscal', 'habilitacao_economica', 'qualificacao_tecnica',
        'afe', 'cbpad', 'cbpp', 'bombeiros', 'outro'
    }

    tipo = 'outro'
    doc_nec = None
    if documento_necessario_id:
        # Resolver tipo a partir do DocumentoNecessario
        db_tmp = get_db()
        try:
            doc_nec = db_tmp.query(DocumentoNecessario).filter(
                DocumentoNecessario.id == documento_necessario_id
            ).first()
            if doc_nec and doc_nec.tipo_chave in _tipos_validos:
                tipo = doc_nec.tipo_chave
        finally:
            db_tmp.close()

    if tipo == 'outro' and not documento_necessario_id:
        # Fallback: tentar resolver pelo campo tipo legado
        tipo_raw = request.form.get('tipo', 'outro')
        tipo_norm = tipo_raw.lower().replace(' ', '_').replace('de_', '')
        _tipo_map = {
            'contrato_social': 'contrato_social', 'procuracao': 'procuracao',
            'certidao_negativa': 'certidao_negativa', 'habilitacao_fiscal': 'habilitacao_fiscal',
            'habilitacao_economica': 'habilitacao_economica', 'balanco_patrimonial': 'balanco',
            'qualificacao_tecnica': 'qualificacao_tecnica', 'atestado_capacidade': 'atestado_capacidade',
            'afe': 'afe', 'cbpad': 'cbpad', 'cbpp': 'cbpp',
            'corpo_bombeiros': 'bombeiros', 'bombeiros': 'bombeiros',
            'alvara': 'alvara', 'registro_conselho': 'registro_conselho',
            'outro': 'outro',
        }
        tipo = _tipo_map.get(tipo_norm, tipo_norm)
        if tipo not in _tipos_validos:
            tipo = 'outro'

    if not empresa_id:
        return jsonify({"error": "empresa_id é obrigatório"}), 400

    db = get_db()
    try:
        # Verificar que a empresa pertence ao usuário
        empresa = db.query(Empresa).filter(
            Empresa.id == empresa_id,
            Empresa.user_id == user_id
        ).first()
        if not empresa:
            return jsonify({"error": "Empresa não encontrada"}), 404

        # Criar diretório para a empresa
        upload_dir = os.path.join(UPLOAD_FOLDER, 'empresa', empresa_id, tipo)
        os.makedirs(upload_dir, exist_ok=True)

        # Salvar arquivo com nome seguro
        from werkzeug.utils import secure_filename
        safe_name = secure_filename(file.filename) or f"doc_{uuid.uuid4().hex[:8]}.pdf"
        filepath = os.path.join(upload_dir, f"{uuid.uuid4().hex[:8]}_{safe_name}")
        file.save(filepath)

        # Criar registro EmpresaDocumento
        doc = EmpresaDocumento()
        doc.id = str(uuid.uuid4())
        doc.empresa_id = empresa_id
        doc.tipo = tipo
        doc.nome_arquivo = file.filename
        doc.path_arquivo = filepath
        if documento_necessario_id:
            doc.documento_necessario_id = documento_necessario_id

        # Datas opcionais
        data_emissao = request.form.get('data_emissao')
        data_vencimento = request.form.get('data_vencimento')
        if data_emissao:
            from datetime import date as date_type
            try:
                doc.data_emissao = date_type.fromisoformat(data_emissao)
            except ValueError:
                pass
        if data_vencimento:
            from datetime import date as date_type
            try:
                doc.data_vencimento = date_type.fromisoformat(data_vencimento)
            except ValueError:
                pass

        db.add(doc)
        db.commit()
        db.refresh(doc)

        return jsonify({
            "success": True,
            "message": f"Documento '{file.filename}' enviado com sucesso",
            "documento": doc.to_dict(),
            "path": filepath
        }), 201

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


@app.route("/api/empresa-documentos/<doc_id>/download", methods=["GET"])
@require_auth
def download_empresa_documento(doc_id):
    """Download/visualização de documento da empresa."""
    from models import Empresa, EmpresaDocumento
    user_id = get_current_user_id()
    download = request.args.get('download', 'false').lower() == 'true'

    db = get_db()
    try:
        doc = db.query(EmpresaDocumento).filter(EmpresaDocumento.id == doc_id).first()
        if not doc:
            return jsonify({"error": "Documento não encontrado"}), 404

        # Verificar propriedade via empresa
        empresa = db.query(Empresa).filter(
            Empresa.id == doc.empresa_id,
            Empresa.user_id == user_id
        ).first()
        if not empresa:
            return jsonify({"error": "Acesso negado"}), 403

        if not doc.path_arquivo or not os.path.exists(doc.path_arquivo):
            return jsonify({"error": "Arquivo não encontrado no disco"}), 404

        # Detectar mimetype
        mimetype = 'application/pdf'
        nome = doc.nome_arquivo or 'documento'
        if nome.lower().endswith('.docx'):
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif nome.lower().endswith('.doc'):
            mimetype = 'application/msword'
        elif nome.lower().endswith('.txt'):
            mimetype = 'text/plain'

        return send_file(
            doc.path_arquivo,
            mimetype=mimetype,
            as_attachment=download,
            download_name=nome
        )
    finally:
        db.close()


@app.route("/api/empresa-certidoes/<certidao_id>/download", methods=["GET"])
@require_auth
def download_empresa_certidao(certidao_id):
    """Download/visualização de certidão da empresa."""
    from models import Empresa, EmpresaCertidao
    user_id = get_current_user_id()
    download = request.args.get('download', 'false').lower() == 'true'

    db = get_db()
    try:
        cert = db.query(EmpresaCertidao).filter(EmpresaCertidao.id == certidao_id).first()
        if not cert:
            return jsonify({"error": "Certidão não encontrada"}), 404

        empresa = db.query(Empresa).filter(
            Empresa.id == cert.empresa_id,
            Empresa.user_id == user_id
        ).first()
        if not empresa:
            return jsonify({"error": "Acesso negado"}), 403

        if not cert.path_arquivo or not os.path.exists(cert.path_arquivo):
            return jsonify({"error": "Arquivo não encontrado no disco"}), 404

        return send_file(
            cert.path_arquivo,
            mimetype='application/pdf',
            as_attachment=download,
            download_name=f"certidao_{cert.tipo}_{cert.id[:8]}.pdf"
        )
    finally:
        db.close()


def _extrair_dados_certidao_pdf(filepath):
    """
    Extrai texto do PDF de certidão e usa LLM para identificar dados estruturados:
    numero, data_emissao, data_vencimento, orgao_emissor, situacao, resumo.
    Retorna dict ou None se falhar.
    """
    try:
        import fitz
        texto = ""
        doc = fitz.open(filepath)
        for page in doc:
            texto += page.get_text()
        doc.close()

        if not texto or len(texto.strip()) < 30:
            print(f"[CERTIDAO-IA] PDF sem texto suficiente: {len(texto)} chars")
            return None

        # Limitar texto para não estourar contexto
        texto_truncado = texto[:6000]

        from llm import call_deepseek
        prompt = f"""Analise o texto abaixo extraído de uma certidão e retorne APENAS um JSON com os campos:

- "numero": número/código da certidão (string ou null)
- "data_emissao": data de emissão no formato YYYY-MM-DD (string ou null)
- "data_vencimento": data de validade/vencimento no formato YYYY-MM-DD (string ou null)
- "orgao_emissor": órgão que emitiu (ex: "Receita Federal", "FGTS/Caixa", "TST") (string ou null)
- "situacao": uma das opções: "valida", "vencida", "invalida", "negativa", "positiva", "positiva com efeito de negativa", "nada consta", "regular", "irregular" (string)
- "resumo": resumo em 1-2 frases do conteúdo da certidão (string)

Regras:
- Se o texto indicar "NADA CONSTA", "NEGATIVA", "REGULAR" → situacao = correspondente
- Se indicar "POSITIVA" (sem "efeito de negativa") → situacao = "positiva" (débitos existem)
- Datas devem ser convertidas para YYYY-MM-DD
- Se não conseguir identificar um campo, use null
- Retorne APENAS o JSON, sem explicações

TEXTO DA CERTIDÃO:
{texto_truncado}"""

        resposta = call_deepseek(
            [{"role": "user", "content": prompt}],
            max_tokens=1000,
            model_override="deepseek-chat"
        )

        if not resposta:
            return None

        # Extrair JSON da resposta
        import json as json_mod
        import re
        # Tentar encontrar JSON na resposta
        json_match = re.search(r'\{[^{}]*\}', resposta, re.DOTALL)
        if json_match:
            dados = json_mod.loads(json_match.group())
            print(f"[CERTIDAO-IA] Dados extraídos: {dados}")
            return dados

        print(f"[CERTIDAO-IA] Resposta sem JSON válido: {resposta[:200]}")
        return None

    except Exception as e:
        print(f"[CERTIDAO-IA] Erro na extração: {e}")
        return None


@app.route("/api/empresa-certidoes/<certidao_id>/upload", methods=["POST"])
@require_auth
def upload_empresa_certidao(certidao_id):
    """
    Upload de arquivo PDF para uma certidão existente.
    Atualiza path_arquivo, status para 'valida', e data_vencimento se fornecida.

    Form data:
    - file: arquivo PDF
    - data_vencimento: (opcional) data de vencimento no formato YYYY-MM-DD
    - numero: (opcional) número da certidão
    """
    from models import Empresa, EmpresaCertidao
    user_id = get_current_user_id()

    if 'file' not in request.files:
        return jsonify({"error": "Arquivo não enviado"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Nenhum arquivo selecionado"}), 400

    db = get_db()
    try:
        cert = db.query(EmpresaCertidao).filter(EmpresaCertidao.id == certidao_id).first()
        if not cert:
            return jsonify({"error": "Certidão não encontrada"}), 404

        empresa = db.query(Empresa).filter(
            Empresa.id == cert.empresa_id,
            Empresa.user_id == user_id
        ).first()
        if not empresa:
            return jsonify({"error": "Acesso negado"}), 403

        # Criar diretório para certidões
        upload_dir = os.path.join(UPLOAD_FOLDER, 'empresa', cert.empresa_id, 'certidoes')
        os.makedirs(upload_dir, exist_ok=True)

        # Salvar arquivo
        from werkzeug.utils import secure_filename
        safe_name = secure_filename(file.filename) or f"certidao_{uuid.uuid4().hex[:8]}.pdf"
        filepath = os.path.join(upload_dir, f"{uuid.uuid4().hex[:8]}_{safe_name}")
        file.save(filepath)

        # Extrair texto do PDF e usar IA para identificar dados da certidão
        from datetime import date as _date, datetime as _dt
        dados_ia = _extrair_dados_certidao_pdf(filepath)

        cert.path_arquivo = filepath
        cert.data_emissao = _date.today()

        if dados_ia:
            # IA extraiu dados — aplicar automaticamente
            if dados_ia.get("data_vencimento"):
                try:
                    cert.data_vencimento = _dt.strptime(dados_ia["data_vencimento"], '%Y-%m-%d').date()
                except (ValueError, TypeError):
                    pass
            if dados_ia.get("data_emissao"):
                try:
                    cert.data_emissao = _dt.strptime(dados_ia["data_emissao"], '%Y-%m-%d').date()
                except (ValueError, TypeError):
                    pass
            if dados_ia.get("numero"):
                cert.numero = dados_ia["numero"]
            if dados_ia.get("orgao_emissor"):
                cert.orgao_emissor = dados_ia["orgao_emissor"]

            situacao = (dados_ia.get("situacao") or "").lower()
            if situacao in ("valida", "válida", "regular", "positiva com efeito de negativa", "nada consta", "negativa"):
                cert.status = 'valida'
            elif situacao in ("vencida", "expirada"):
                cert.status = 'vencida'
            elif situacao in ("invalida", "inválida", "positiva", "irregular"):
                cert.status = 'pendente'
            else:
                cert.status = 'valida'  # Upload manual = assume válida

            cert.mensagem = dados_ia.get("resumo") or "Certidão enviada por upload manual"
            cert.dados_extras = {k: v for k, v in dados_ia.items() if k not in ("resumo",) and v}
        else:
            # Falha na extração IA — salvar com dados mínimos
            cert.status = 'valida'
            cert.mensagem = "Certidão enviada por upload manual (extração automática indisponível)"

        # Fallback: dados do formulário sobrescrevem IA se enviados
        data_venc = request.form.get('data_vencimento')
        if data_venc:
            try:
                cert.data_vencimento = _dt.strptime(data_venc, '%Y-%m-%d').date()
            except ValueError:
                pass
        numero = request.form.get('numero')
        if numero:
            cert.numero = numero

        db.commit()
        return jsonify({
            "success": True,
            "message": f"Certidão atualizada com sucesso" + (" (dados extraídos por IA)" if dados_ia else ""),
            "certidao": cert.to_dict(),
            "dados_extraidos": dados_ia,
        })

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


# =============================================================================
# B4: Usar porte/regime da empresa no cálculo de validação
# (Implementado em tools.py via PROMPT_SCORES_VALIDACAO)
# =============================================================================


# =============================================================================
# B5: Gerar classes/subclasses automaticamente via IA
# =============================================================================

@app.route("/api/parametrizacoes/gerar-classes", methods=["POST"])
@require_auth
def gerar_classes_ia():
    """
    Gera classes e subclasses de produtos automaticamente via IA.
    Lista todos os produtos do usuário, envia para DeepSeek e retorna
    uma estrutura de classes/subclasses sugeridas.
    """
    from models import Produto
    user_id = get_current_user_id()

    db = get_db()
    try:
        produtos = db.query(Produto).filter(Produto.user_id == user_id).all()

        if not produtos:
            return jsonify({"error": "Nenhum produto cadastrado. Cadastre produtos primeiro."}), 400

        # Montar lista de produtos para o prompt
        lista_produtos = []
        for p in produtos:
            lista_produtos.append(
                f"- {p.nome} (categoria: {p.categoria}, NCM: {p.ncm or 'N/A'}, "
                f"fabricante: {p.fabricante or 'N/A'})"
            )

        prompt = f"""Você é um especialista em classificação de produtos para licitações públicas.

Dado os seguintes produtos cadastrados:

{chr(10).join(lista_produtos)}

Agrupe esses produtos em CLASSES e SUBCLASSES lógicas para facilitar a gestão em processos licitatórios.

Para cada classe e subclasse, sugira o NCM (Nomenclatura Comum do Mercosul) mais adequado.

RESPONDA APENAS EM JSON com a seguinte estrutura:
{{
  "classes": [
    {{
      "nome": "Nome da Classe",
      "descricao": "Descrição breve",
      "ncm_principal": "XXXX.XX.XX",
      "subclasses": [
        {{
          "nome": "Nome da Subclasse",
          "descricao": "Descrição breve",
          "ncm": "XXXX.XX.XX",
          "produtos_sugeridos": ["Nome Produto 1", "Nome Produto 2"]
        }}
      ]
    }}
  ]
}}
"""

        resposta = call_deepseek(
            [{"role": "user", "content": prompt}],
            max_tokens=3000,
            model_override="deepseek-chat"
        )

        # Extrair JSON da resposta
        import json, re
        texto = resposta if isinstance(resposta, str) else str(resposta)
        json_match = re.search(r'\{[\s\S]*\}', texto)
        if json_match:
            try:
                resultado = json.loads(json_match.group(0))
                return jsonify({
                    "success": True,
                    "classes": resultado.get("classes", []),
                    "total_produtos": len(produtos)
                })
            except json.JSONDecodeError:
                pass

        return jsonify({
            "success": True,
            "classes": [],
            "raw_response": texto[:2000],
            "message": "IA retornou resposta mas não no formato JSON esperado"
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


# =============================================================================
# B6: Vincular documento da empresa a item/requisito do edital
# =============================================================================

@app.route("/api/editais/<edital_id>/vincular-documento", methods=["POST"])
@require_auth
def vincular_documento_edital(edital_id):
    """
    Vincula um documento da empresa a um requisito do edital.

    JSON body:
    - documento_id: ID do EmpresaDocumento ou EmpresaCertidao
    - requisito_id: ID do EditalRequisito
    - tipo_documento: "documento" ou "certidao"
    """
    from models import Empresa, EmpresaDocumento, EmpresaCertidao, EditalRequisito, Edital
    user_id = get_current_user_id()
    data = request.json or {}

    documento_id = data.get('documento_id')
    requisito_id = data.get('requisito_id')
    tipo_documento = data.get('tipo_documento', 'documento')

    if not documento_id or not requisito_id:
        return jsonify({"error": "documento_id e requisito_id são obrigatórios"}), 400

    db = get_db()
    try:
        # Verificar edital pertence ao usuário
        edital = db.query(Edital).filter(
            Edital.id == edital_id,
            Edital.user_id == user_id
        ).first()
        if not edital:
            return jsonify({"error": "Edital não encontrado"}), 404

        # Verificar requisito pertence ao edital
        requisito = db.query(EditalRequisito).filter(
            EditalRequisito.id == requisito_id,
            EditalRequisito.edital_id == edital_id
        ).first()
        if not requisito:
            return jsonify({"error": "Requisito não encontrado"}), 404

        # Vincular (via campo edital_requisito_id no documento)
        if tipo_documento == 'certidao':
            cert = db.query(EmpresaCertidao).filter(EmpresaCertidao.id == documento_id).first()
            if not cert:
                return jsonify({"error": "Certidão não encontrada"}), 404
            cert.edital_requisito_id = requisito_id
        else:
            doc = db.query(EmpresaDocumento).filter(EmpresaDocumento.id == documento_id).first()
            if not doc:
                return jsonify({"error": "Documento não encontrado"}), 404
            doc.edital_requisito_id = requisito_id

        db.commit()

        return jsonify({
            "success": True,
            "message": f"Documento vinculado ao requisito '{requisito.descricao[:60]}'"
        })

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


@app.route("/api/editais/<edital_id>/documentacao-necessaria", methods=["GET"])
@require_auth
def documentacao_necessaria(edital_id):
    """
    Retorna documentação necessária para o edital, cruzando requisitos extraídos
    do PDF com documentos/certidões reais da empresa.
    Se o edital não tem requisitos documentais extraídos, usa lista padrão de licitações.
    Formato: { documentos: [{ nome, categoria, status, validade, exigido }] }
    """
    from models import Empresa, EmpresaDocumento, EmpresaCertidao, EditalRequisito, Edital
    from datetime import date
    user_id = get_current_user_id()

    db = get_db()
    try:
        edital = db.query(Edital).filter(
            Edital.id == edital_id,
            Edital.user_id == user_id
        ).first()
        if not edital:
            return jsonify({"error": "Edital não encontrado"}), 404

        empresa = db.query(Empresa).filter(Empresa.user_id == user_id).first()

        # Requisitos documentais extraídos do edital (se houver)
        requisitos_doc = db.query(EditalRequisito).filter(
            EditalRequisito.edital_id == edital_id,
            EditalRequisito.tipo == 'documental'
        ).all()

        # Documentos e certidões da empresa (indexados por tipo)
        from models import DocumentoNecessario, CategoriaDocumento, FonteCertidao
        docs_empresa_por_tipo = {}
        certs_por_tipo = {}
        certs_por_fonte = {}
        if empresa:
            for d in db.query(EmpresaDocumento).filter(EmpresaDocumento.empresa_id == empresa.id).all():
                docs_empresa_por_tipo.setdefault(d.tipo, []).append(d)
            for c in db.query(EmpresaCertidao).filter(EmpresaCertidao.empresa_id == empresa.id).all():
                certs_por_tipo.setdefault(c.tipo, []).append(c)
                if c.fonte_certidao_id:
                    certs_por_fonte[c.fonte_certidao_id] = c

        hoje = date.today()
        documentos = []

        # === DOCUMENTOS: carregar do banco (documentos_necessarios) ===
        docs_necessarios_db = db.query(DocumentoNecessario).filter(
            DocumentoNecessario.ativo == True
        ).order_by(DocumentoNecessario.ordem).all()

        # === CERTIDÕES: carregar do banco (fontes_certidoes) ===
        fontes_certidoes_db = db.query(FonteCertidao).filter(
            FonteCertidao.ativo == True
        ).order_by(FonteCertidao.nome).all()

        # Tipos de certidão conhecidos (para checar_status)
        TIPOS_CERTIDAO = {fc.tipo_certidao for fc in fontes_certidoes_db if fc.tipo_certidao}

        def checar_status_documento(tipo_ref):
            """Cruza tipo_ref com documentos da empresa."""
            encontrados = docs_empresa_por_tipo.get(tipo_ref, [])
            if encontrados:
                doc = encontrados[0]
                venc = doc.data_vencimento
                if venc and venc < hoje:
                    return "vencido", venc.strftime("%d/%m/%Y") if venc else None
                return "ok", venc.strftime("%d/%m/%Y") if venc else None
            return "faltando", None

        def checar_status_certidao_por_fonte(fonte_id):
            """Cruza fonte_certidao_id com certidões da empresa."""
            cert = certs_por_fonte.get(fonte_id)
            if cert:
                venc = cert.data_vencimento
                if cert.status == "vencida" or (venc and venc < hoje):
                    return "vencido", venc.strftime("%d/%m/%Y") if venc else None
                elif cert.status in ("valida",):
                    return "ok", venc.strftime("%d/%m/%Y") if venc else None
                else:
                    return "faltando", None
            return "faltando", None

        # Tipos exigidos pelo edital (se requisitos foram extraídos do PDF)
        tipos_exigidos_edital = set()
        if requisitos_doc:
            for req in requisitos_doc:
                tipo_ref = (req.nome_especificacao or "").strip().lower()
                tipos_exigidos_edital.add(tipo_ref)

        # === PARTE 1: Documentos da empresa (do banco documentos_necessarios) ===
        for dn in docs_necessarios_db:
            categoria = dn.categoria_rel.nome if dn.categoria_rel else "Outros"

            # Não incluir tipos de certidão aqui (vão na parte 2)
            if dn.tipo_chave in TIPOS_CERTIDAO:
                continue

            status, validade = checar_status_documento(dn.tipo_chave)
            exigido = dn.tipo_chave in tipos_exigidos_edital if requisitos_doc else dn.obrigatorio
            documentos.append({
                "nome": dn.nome,
                "categoria": categoria,
                "status": status,
                "validade": validade,
                "exigido": exigido,
            })

        # === PARTE 2: Certidões (do banco fontes_certidoes) ===
        for fc in fontes_certidoes_db:
            status, validade = checar_status_certidao_por_fonte(fc.id)
            exigido = fc.tipo_certidao in tipos_exigidos_edital if requisitos_doc else True
            documentos.append({
                "nome": fc.nome,
                "categoria": "Certidões e Fiscal",
                "status": status,
                "validade": validade,
                "exigido": exigido,
            })

        # === PARTE 3: Documentos extras do edital — inserir no banco se não existem ===
        if requisitos_doc:
            tipos_ja_incluidos = {dn.tipo_chave for dn in docs_necessarios_db} | {fc.tipo_certidao for fc in fontes_certidoes_db}
            # Buscar categoria "Outros" para inserir novos tipos
            cat_outros = db.query(CategoriaDocumento).filter(CategoriaDocumento.nome.ilike("%outro%")).first()
            cat_outros_id = cat_outros.id if cat_outros else None
            nomes_ja_inseridos = set()

            for req in requisitos_doc:
                tipo_ref = (req.nome_especificacao or "").strip().lower()
                if tipo_ref not in tipos_ja_incluidos and tipo_ref != "":
                    nome = req.descricao[:80] if req.descricao else "Documento exigido pelo edital"

                    # Evitar duplicatas dentro do mesmo lote
                    if nome in nomes_ja_inseridos:
                        continue
                    nomes_ja_inseridos.add(nome)

                    # Gerar tipo_chave único a partir do nome
                    import re as _re
                    tipo_chave_novo = _re.sub(r'[^a-z0-9]', '_', nome.lower().strip())[:50]

                    # Verificar se tipo_chave já existe no banco
                    existe = db.query(DocumentoNecessario).filter(
                        DocumentoNecessario.tipo_chave == tipo_chave_novo
                    ).first()

                    if not existe:
                        # Inserir novo tipo de documento na tabela documentos_necessarios
                        novo_doc_nec = DocumentoNecessario(
                            id=str(uuid.uuid4()),
                            categoria_id=cat_outros_id,
                            nome=nome,
                            tipo_chave=tipo_chave_novo,
                            obrigatorio=True,
                            ativo=True,
                            ordem=200,
                        )
                        db.add(novo_doc_nec)
                        print(f"[DOCS] Novo tipo inserido: {nome} (chave={tipo_chave_novo})")

                    documentos.append({
                        "nome": nome,
                        "categoria": "Outros",
                        "status": "faltando",
                        "validade": None,
                        "exigido": True,
                        "descricao_edital": req.descricao,
                    })

            if nomes_ja_inseridos:
                db.commit()

        return jsonify({
            "success": True,
            "edital_id": edital_id,
            "edital_numero": edital.numero,
            "documentos": documentos,
            "total": len(documentos),
            "fonte": "requisitos_edital" if requisitos_doc else "padrao_licitacao",
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


@app.route("/api/editais/<edital_id>/analisar-mercado", methods=["POST"])
@require_auth
def analisar_mercado_edital(edital_id):
    """Análise de mercado do órgão: dados PNCP + histórico interno + IA."""
    from models import Edital, OrgaoPerfil, EstrategiaEdital
    from tools import _extrair_json_object
    import requests as req
    from datetime import datetime, timedelta

    user_id = get_current_user_id()
    db = get_db()
    try:
        edital = db.query(Edital).filter(Edital.id == edital_id, Edital.user_id == user_id).first()
        if not edital:
            return jsonify({"error": "Edital não encontrado"}), 404

        cnpj_orgao = edital.cnpj_orgao
        nome_orgao = edital.orgao or ""
        uf_orgao = edital.uf or ""

        if not cnpj_orgao and not nome_orgao:
            return jsonify({"error": "Edital sem dados do órgão (CNPJ ou nome)"}), 400

        # Verificar cache no banco (< 30 dias)
        forcar = (request.get_json(silent=True) or {}).get("forcar", False)
        perfil_db = None
        if cnpj_orgao:
            perfil_db = db.query(OrgaoPerfil).filter(OrgaoPerfil.cnpj == cnpj_orgao).first()

        if perfil_db and not forcar:
            if perfil_db.ultima_atualizacao and perfil_db.ultima_atualizacao > datetime.now() - timedelta(days=30):
                # Retornar do cache
                hist_interno = _calcular_historico_interno(db, nome_orgao, user_id)
                return jsonify({
                    "success": True,
                    "cache": True,
                    "orgao": {"nome": perfil_db.nome, "cnpj": perfil_db.cnpj, "uf": perfil_db.uf},
                    "estatisticas": {
                        "total_compras": perfil_db.total_compras,
                        "valor_total": float(perfil_db.valor_total_compras) if perfil_db.valor_total_compras else 0,
                        "valor_medio": float(perfil_db.valor_medio_compras) if perfil_db.valor_medio_compras else 0,
                        "modalidades": perfil_db.modalidades_frequentes or {},
                    },
                    "compras_similares": perfil_db.compras_similares or [],
                    "historico_interno": hist_interno,
                    "analise_ia": perfil_db.analise_ia or "",
                })

        # === Buscar no PNCP ===
        print(f"[MERCADO] Buscando compras do órgão: {nome_orgao}")
        # Buscar pelo nome do órgão (primeiras 3 palavras significativas)
        import re as _re
        palavras_orgao = [w for w in _re.sub(r'[^\w\s]', '', nome_orgao).split() if len(w) > 2][:4]
        termo_orgao = " ".join(palavras_orgao)

        compras = []
        total_pncp = 0
        try:
            url_search = f"https://pncp.gov.br/api/search/?q={termo_orgao}&tipos_documento=edital&tam_pagina=50&pagina=1&ordenacao=-data"
            resp = req.get(url_search, timeout=20, headers={"Accept": "application/json"})
            if resp.status_code == 200:
                data = resp.json()
                total_pncp = data.get("total", 0)
                for item in data.get("items", []):
                    # Filtrar só do mesmo órgão (CNPJ ou nome)
                    item_cnpj = item.get("orgao_cnpj", "")
                    item_orgao = item.get("orgao_nome", "")
                    if cnpj_orgao and item_cnpj != cnpj_orgao:
                        continue
                    if not cnpj_orgao and nome_orgao[:15].lower() not in item_orgao.lower():
                        continue
                    compras.append({
                        "objeto": (item.get("description") or "")[:80],
                        "valor": item.get("valor_global"),
                        "data": str(item.get("data_publicacao_pncp", ""))[:10],
                        "modalidade": item.get("modalidade_licitacao_nome", ""),
                        "uf": item.get("uf", ""),
                    })
        except Exception as e:
            print(f"[MERCADO] Erro ao buscar PNCP: {e}")

        # Calcular estatísticas
        valores = [c["valor"] for c in compras if c.get("valor") and c["valor"] > 0]
        valor_total = sum(valores) if valores else 0
        valor_medio = valor_total / len(valores) if valores else 0
        modalidades = {}
        for c in compras:
            mod = c.get("modalidade", "Outro")
            modalidades[mod] = modalidades.get(mod, 0) + 1

        # Compras similares (mesmo segmento que o edital)
        objeto_edital = (edital.objeto or "").lower()
        palavras_obj = [w for w in _re.sub(r'[^\w\s]', '', objeto_edital).split() if len(w) > 4][:3]
        similares = []
        for c in compras:
            obj_c = (c.get("objeto") or "").lower()
            if any(p in obj_c for p in palavras_obj):
                similares.append(c)

        # Histórico interno
        hist_interno = _calcular_historico_interno(db, nome_orgao, user_id)

        # === Análise IA ===
        analise_ia_texto = ""
        try:
            prompt_mercado = f"""Analise o perfil deste órgão contratante para licitações:

ÓRGÃO: {nome_orgao}
CNPJ: {cnpj_orgao or 'N/A'}
UF: {uf_orgao}

DADOS DO PNCP:
- Total de compras encontradas: {len(compras)}
- Valor total: R$ {valor_total:,.2f}
- Valor médio por compra: R$ {valor_medio:,.2f}
- Modalidades: {json.dumps(modalidades, ensure_ascii=False)}
- Compras similares ao objeto "{edital.objeto[:50]}": {len(similares)}

HISTÓRICO INTERNO DO SISTEMA:
- Editais deste órgão no sistema: {hist_interno.get('total', 0)}
- Decisões GO: {hist_interno.get('go', 0)} | NO-GO: {hist_interno.get('nogo', 0)}

Gere uma análise concisa (3-5 frases) sobre:
1. Perfil do órgão (volume, tipo de compras, porte)
2. Risco de pagamento (baseado no tipo de órgão — federal/estadual/municipal)
3. Oportunidades (recorrência, volume, segmento)
4. Recomendação (vale a pena investir neste órgão?)

Responda APENAS o texto da análise, sem JSON."""

            analise_ia_texto = call_deepseek([{"role": "user", "content": prompt_mercado}], max_tokens=1000, model_override="deepseek-chat")
            analise_ia_texto = analise_ia_texto.strip()
        except Exception as e:
            analise_ia_texto = f"Erro na análise IA: {e}"

        # === Salvar no banco ===
        if perfil_db:
            perfil_db.nome = nome_orgao
            perfil_db.uf = uf_orgao
            perfil_db.total_compras = len(compras)
            perfil_db.valor_total_compras = valor_total
            perfil_db.valor_medio_compras = valor_medio
            perfil_db.modalidades_frequentes = modalidades
            perfil_db.compras_similares = similares[:10]
            perfil_db.analise_ia = analise_ia_texto
            perfil_db.ultima_atualizacao = datetime.now()
        else:
            perfil_db = OrgaoPerfil(
                id=str(uuid.uuid4()),
                cnpj=cnpj_orgao or f"sem_cnpj_{nome_orgao[:20]}",
                nome=nome_orgao,
                uf=uf_orgao,
                total_compras=len(compras),
                valor_total_compras=valor_total,
                valor_medio_compras=valor_medio,
                modalidades_frequentes=modalidades,
                compras_similares=similares[:10],
                analise_ia=analise_ia_texto,
                ultima_atualizacao=datetime.now(),
                user_id=user_id,
            )
            db.add(perfil_db)

        db.commit()
        print(f"[MERCADO] Perfil salvo: {nome_orgao}, {len(compras)} compras, {len(similares)} similares")

        return jsonify({
            "success": True,
            "cache": False,
            "orgao": {"nome": nome_orgao, "cnpj": cnpj_orgao, "uf": uf_orgao},
            "estatisticas": {
                "total_compras": len(compras),
                "valor_total": valor_total,
                "valor_medio": valor_medio,
                "modalidades": modalidades,
            },
            "compras_similares": similares[:10],
            "historico_interno": hist_interno,
            "analise_ia": analise_ia_texto,
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


def _calcular_historico_interno(db, nome_orgao, user_id):
    """Calcula decisões internas GO/NO-GO para um órgão."""
    try:
        editais_orgao = db.query(Edital).filter(
            Edital.user_id == user_id,
            Edital.orgao.ilike(f"%{nome_orgao[:20]}%"),
            Edital.status != "temp_score",
        ).all()
        total = len(editais_orgao)
        go = sum(1 for e in editais_orgao if e.status == "go")
        nogo = sum(1 for e in editais_orgao if e.status == "nogo")
        avaliando = sum(1 for e in editais_orgao if e.status == "avaliando")
        return {"total": total, "go": go, "nogo": nogo, "avaliando": avaliando}
    except Exception:
        return {"total": 0, "go": 0, "nogo": 0, "avaliando": 0}


@app.route("/api/editais/<edital_id>/analisar-riscos", methods=["POST"])
@require_auth
def analisar_riscos_edital(edital_id):
    """Análise completa de riscos: PDF via IA + atas + vencedores + concorrentes."""
    from models import Edital, EditalDocumento, Concorrente, PrecoHistorico, Produto, AtaConsultada
    from tools import _extrair_json_object, tool_buscar_atas_pncp
    import requests as req
    import re as _re
    from datetime import datetime

    user_id = get_current_user_id()
    db = get_db()
    try:
        edital = db.query(Edital).filter(Edital.id == edital_id, Edital.user_id == user_id).first()
        if not edital:
            return jsonify({"error": "Edital não encontrado"}), 404

        # === PASSO 1: Identificar produto match ===
        produto_match = None
        analise_existente = db.query(Analise).filter(Analise.edital_id == edital_id, Analise.user_id == user_id).first()
        if analise_existente and analise_existente.produto_id:
            produto_match = db.query(Produto).filter(Produto.id == analise_existente.produto_id).first()
        if not produto_match:
            produto_match = db.query(Produto).filter(Produto.user_id == user_id).first()

        empresa = db.query(Empresa).filter(Empresa.user_id == user_id).first()
        empresa_id = empresa.id if empresa else None

        # === PASSO 2: Buscar atas no PNCP ===
        # Usar palavras-chave do objeto (mesmo critério do endpoint historico-vencedores)
        import re as _re2
        objeto = edital.objeto or ""
        palavras_stop = {"aquisicao", "aquisição", "de", "do", "da", "dos", "das", "para", "com", "registro", "preco", "preço", "futura", "eventual", "fornecimento"}
        palavras = [w for w in _re2.sub(r'[^\w\s]', ' ', objeto.lower()).split() if w not in palavras_stop and len(w) > 3][:3]
        termo_ata = " ".join(palavras) if palavras else (produto_match.nome if produto_match else objeto[:30])
        print(f"[RISCOS] Buscando atas para: {termo_ata}")
        atas_resultado = tool_buscar_atas_pncp(termo_ata, user_id)
        atas = atas_resultado.get("atas", []) if atas_resultado.get("success") else []
        print(f"[RISCOS] {len(atas)} atas encontradas")

        # === PASSO 3: Buscar vencedores e preços das atas ===
        vencedores_todos = []
        cnpjs_processados = {}
        salvos_concorrentes = 0
        salvos_precos = 0

        salvos_atas = 0
        for ata in atas[:5]:
            cnpj = ata.get("cnpj_orgao") or ""
            ano = ata.get("ano")
            url_pncp = ata.get("url_pncp") or ""
            match_url = _re.search(r'/atas/\d+/\d+/(\d+)/\d+', url_pncp)
            seq_compra = int(match_url.group(1)) if match_url else None
            if not cnpj or not ano or not seq_compra:
                continue

            # Salvar ata no banco (evitar duplicata por numero_controle)
            nc = ata.get("numero_controle") or ""
            ata_db = db.query(AtaConsultada).filter(AtaConsultada.numero_controle_pncp == nc).first() if nc else None
            if not ata_db:
                url_edital_origem = f"https://pncp.gov.br/app/editais/{cnpj}/{ano}/{seq_compra}"
                dp = ata.get("data_publicacao")
                dt_pub = None
                if dp:
                    try:
                        dt_pub = datetime.fromisoformat(str(dp).replace("Z", "")[:19])
                    except (ValueError, TypeError):
                        pass
                ata_db = AtaConsultada(
                    id=str(uuid.uuid4()),
                    edital_id=edital_id,
                    user_id=user_id,
                    numero_controle_pncp=nc or None,
                    titulo=ata.get("titulo"),
                    orgao=ata.get("orgao"),
                    cnpj_orgao=cnpj,
                    uf=ata.get("uf"),
                    ano=ano,
                    seq_compra=seq_compra,
                    url_pncp=url_pncp,
                    url_edital_origem=url_edital_origem,
                    data_publicacao=dt_pub,
                )
                db.add(ata_db)
                db.flush()
                salvos_atas += 1
            try:
                url_itens = f"https://pncp.gov.br/api/pncp/v1/orgaos/{cnpj}/compras/{ano}/{seq_compra}/itens"
                resp = req.get(url_itens, timeout=15, headers={"Accept": "application/json"})
                if resp.status_code != 200:
                    continue
                itens = resp.json()
                if not isinstance(itens, list):
                    continue
                vencedores_ata = []
                for item in itens[:10]:
                    num_item = item.get("numeroItem")
                    if not num_item:
                        continue
                    try:
                        url_r = f"https://pncp.gov.br/api/pncp/v1/orgaos/{cnpj}/compras/{ano}/{seq_compra}/itens/{num_item}/resultados"
                        resp_r = req.get(url_r, timeout=10, headers={"Accept": "application/json"})
                        if resp_r.status_code == 200:
                            res_item = resp_r.json()
                            if isinstance(res_item, list) and res_item:
                                r = res_item[0]
                                cnpj_venc = (r.get("niFornecedor") or "").replace(".", "").replace("/", "").replace("-", "")
                                nome_venc = r.get("nomeRazaoSocialFornecedor")
                                vencedores_ata.append({
                                    "item": num_item,
                                    "descricao": str(item.get("descricao", ""))[:80],
                                    "quantidade": item.get("quantidade"),
                                    "valor_estimado": item.get("valorUnitarioEstimado"),
                                    "vencedor": nome_venc,
                                    "cnpj_vencedor": cnpj_venc,
                                    "valor_homologado": r.get("valorUnitarioHomologado"),
                                    "qtd_homologada": r.get("quantidadeHomologada"),
                                    "valor_total_homologado": r.get("valorTotalHomologado"),
                                    "porte": r.get("porteFornecedorNome"),
                                })
                                # Salvar concorrente
                                if nome_venc and cnpj_venc:
                                    conc = cnpjs_processados.get(cnpj_venc)
                                    if not conc:
                                        conc = db.query(Concorrente).filter(Concorrente.cnpj == cnpj_venc).first()
                                    if not conc:
                                        conc = Concorrente(id=str(uuid.uuid4()), nome=nome_venc, cnpj=cnpj_venc, editais_participados=1, editais_ganhos=1)
                                        db.add(conc)
                                        db.flush()
                                        salvos_concorrentes += 1
                                    else:
                                        conc.editais_participados = (conc.editais_participados or 0) + 1
                                        conc.editais_ganhos = (conc.editais_ganhos or 0) + 1
                                    cnpjs_processados[cnpj_venc] = conc
                                    # Salvar preço histórico
                                    try:
                                        ph = PrecoHistorico(
                                            id=str(uuid.uuid4()), edital_id=edital_id, user_id=user_id,
                                            empresa_id=empresa_id, concorrente_id=conc.id,
                                            produto_id=produto_match.id if produto_match else None,
                                            ata_consultada_id=ata_db.id,
                                            preco_referencia=item.get("valorUnitarioEstimado"),
                                            preco_vencedor=r.get("valorUnitarioHomologado"),
                                            empresa_vencedora=nome_venc, cnpj_vencedor=cnpj_venc,
                                            resultado="derrota", fonte="pncp", data_registro=datetime.now(),
                                        )
                                        if item.get("valorUnitarioEstimado") and r.get("valorUnitarioHomologado") and item["valorUnitarioEstimado"] > 0:
                                            ph.desconto_percentual = round((1 - r["valorUnitarioHomologado"] / item["valorUnitarioEstimado"]) * 100, 2)
                                        db.add(ph)
                                        salvos_precos += 1
                                    except Exception:
                                        pass
                    except Exception:
                        pass
                if vencedores_ata:
                    vencedores_todos.append({
                        "ata_titulo": ata.get("titulo") or f"Ata {ano}",
                        "orgao": ata.get("orgao"), "uf": ata.get("uf"),
                        "vencedores": vencedores_ata,
                    })
            except Exception:
                pass

        # Commit concorrentes e preços
        try:
            db.commit()
            print(f"[RISCOS] Salvos: {salvos_atas} atas, {salvos_concorrentes} concorrentes, {salvos_precos} preços")
        except Exception:
            db.rollback()

        # === PASSO 4: Calcular recorrência ===
        datas = []
        for ata in atas:
            dp = ata.get("data_publicacao") or ata.get("data_assinatura")
            if dp:
                try:
                    dt = datetime.fromisoformat(str(dp).replace("Z", "")[:19])
                    datas.append(dt)
                except (ValueError, TypeError):
                    pass
        frequencia = None
        if len(datas) >= 2:
            datas.sort()
            diffs = [(datas[i+1] - datas[i]).days for i in range(len(datas)-1)]
            media_dias = sum(diffs) / len(diffs) if diffs else 0
            frequencia = "semestral" if media_dias <= 200 else "anual" if media_dias <= 400 else "esporadica"

        # === PASSO 5: Montar contexto competitivo para IA ===
        contexto_competitivo = ""
        if vencedores_todos:
            nomes_venc = {}
            for va in vencedores_todos:
                for v in va["vencedores"]:
                    n = v.get("vencedor", "")
                    if n:
                        nomes_venc[n] = nomes_venc.get(n, 0) + 1
            top_conc = sorted(nomes_venc.items(), key=lambda x: -x[1])[:3]
            contexto_competitivo = f"\n\nCONTEXTO COMPETITIVO:\n"
            contexto_competitivo += f"- {len(atas)} atas encontradas no PNCP para produto similar\n"
            if frequencia:
                contexto_competitivo += f"- Recorrência: {frequencia}\n"
            for nome, cnt in top_conc:
                contexto_competitivo += f"- Concorrente: {nome} ({cnt} vitória(s))\n"

        # === PASSO 6: Análise via IA do PDF ===
        doc = db.query(EditalDocumento).filter(
            EditalDocumento.edital_id == edital_id,
            EditalDocumento.texto_extraido != None
        ).first()

        analise_ia = {}
        if doc and doc.texto_extraido and len(doc.texto_extraido) >= 200:
            prompt = f"""Você é um especialista em análise de riscos em editais de licitação pública brasileira.

Analise o texto do edital abaixo e extraia uma análise completa de riscos.
{contexto_competitivo}
Retorne APENAS um JSON válido (sem texto adicional) com esta estrutura:
{{
  "modalidade": "tipo da modalidade",
  "prazo_pagamento": "prazo de pagamento ou 'Não informado'",
  "sinais_mercado": ["sinais relevantes detectados"],
  "riscos": [{{"tipo": "juridico|tecnico|financeiro|logistico", "descricao": "...", "severidade": "alto|medio|baixo", "mitigacao": "..."}}],
  "fatal_flaws": ["problemas CRÍTICOS que impedem participação"],
  "flags_juridicos": ["cláusulas restritivas"],
  "trechos_relevantes": [{{"trecho": "trecho do edital (máx 150 chars)", "tipo": "risco|oportunidade|alerta", "comentario": "comentário analítico"}}]
}}

REGRAS:
- 3-5 riscos de diferentes categorias. Se há dados competitivos, inclua riscos de concorrência.
- Fatal flaws são RAROS
- 3-8 trechos relevantes
- Mitigações objetivas e práticas

TEXTO DO EDITAL:
{doc.texto_extraido[:20000]}

ANÁLISE DE RISCOS (JSON):"""

            resposta = call_deepseek([{"role": "user", "content": prompt}], max_tokens=8000, model_override="deepseek-chat")
            analise_ia = _extrair_json_object(resposta)

        # === PASSO 7: Carregar concorrentes do banco ===
        concorrentes_db = db.query(Concorrente).order_by(Concorrente.editais_ganhos.desc()).limit(20).all()
        concorrentes_lista = [{
            "nome": c.nome, "cnpj": c.cnpj,
            "editais_participados": c.editais_participados or 0,
            "editais_ganhos": c.editais_ganhos or 0,
            "taxa_vitoria": round((c.editais_ganhos or 0) / max(c.editais_participados or 1, 1) * 100, 1),
        } for c in concorrentes_db]

        # === RETORNAR TUDO ===
        return jsonify({
            "success": True,
            # Análise IA
            **analise_ia,
            # Atas
            "atas": atas[:10],
            "frequencia": frequencia,
            "termo_ata": termo_ata,
            # Vencedores
            "vencedores_resultados": vencedores_todos,
            # Concorrentes
            "concorrentes": concorrentes_lista,
            # Salvos
            "salvos": {"atas": salvos_atas, "concorrentes": salvos_concorrentes, "precos": salvos_precos},
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


@app.route("/api/editais/<edital_id>/historico-vencedores", methods=["POST"])
@require_auth
def historico_vencedores_edital(edital_id):
    """Busca atas de registro de preço no PNCP para o mesmo objeto do edital, identificando vencedores anteriores."""
    from tools import tool_buscar_atas_pncp
    user_id = get_current_user_id()
    db = get_db()
    try:
        edital = db.query(Edital).filter(Edital.id == edital_id, Edital.user_id == user_id).first()
        if not edital:
            return jsonify({"error": "Edital não encontrado"}), 404

        # Extrair termo de busca do objeto do edital (primeiras 3 palavras significativas)
        import re
        objeto = edital.objeto or ""
        palavras_stop = {"aquisicao", "aquisição", "de", "do", "da", "dos", "das", "para", "com", "registro", "preco", "preço", "futura", "eventual", "fornecimento"}
        palavras = [w for w in re.sub(r'[^\w\s]', ' ', objeto.lower()).split() if w not in palavras_stop and len(w) > 3][:3]
        termo = " ".join(palavras) if palavras else objeto[:30]

        resultado = tool_buscar_atas_pncp(termo, user_id)

        if resultado.get("success"):
            atas = resultado.get("atas", [])
            # Enriquecer: tentar identificar frequência
            datas = []
            for ata in atas:
                dp = ata.get("data_publicacao") or ata.get("data_assinatura")
                if dp:
                    try:
                        from datetime import datetime
                        dt = datetime.fromisoformat(str(dp).replace("Z", "")[:19])
                        datas.append(dt)
                    except (ValueError, TypeError):
                        pass

            frequencia = None
            if len(datas) >= 2:
                datas.sort()
                diffs = [(datas[i+1] - datas[i]).days for i in range(len(datas)-1)]
                media_dias = sum(diffs) / len(diffs) if diffs else 0
                if media_dias <= 200:
                    frequencia = "semestral"
                elif media_dias <= 400:
                    frequencia = "anual"
                else:
                    frequencia = "esporadica"

            return jsonify({
                "success": True,
                "termo": termo,
                "total": resultado.get("total", len(atas)),
                "atas": atas[:10],
                "frequencia": frequencia,
            })
        else:
            return jsonify({"success": False, "error": resultado.get("error", "Nenhuma ata encontrada"), "termo": termo})

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


@app.route("/api/concorrentes/listar", methods=["GET"])
@require_auth
def listar_concorrentes_api():
    """Lista concorrentes cadastrados com estatísticas."""
    from models import Concorrente
    db = get_db()
    try:
        concorrentes = db.query(Concorrente).order_by(Concorrente.editais_ganhos.desc()).limit(20).all()
        return jsonify({
            "success": True,
            "total": len(concorrentes),
            "concorrentes": [{
                "id": c.id,
                "nome": c.nome,
                "cnpj": c.cnpj,
                "editais_participados": c.editais_participados or 0,
                "editais_ganhos": c.editais_ganhos or 0,
                "taxa_vitoria": round((c.editais_ganhos or 0) / max(c.editais_participados or 1, 1) * 100, 1),
                "preco_medio": float(c.preco_medio) if c.preco_medio else None,
            } for c in concorrentes],
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


@app.route("/api/editais/<edital_id>/vencedores-atas", methods=["POST"])
@require_auth
def vencedores_atas_edital(edital_id):
    """Busca vencedores e preços das atas encontradas para o edital."""
    user_id = get_current_user_id()
    db = get_db()
    try:
        edital = db.query(Edital).filter(Edital.id == edital_id, Edital.user_id == user_id).first()
        if not edital:
            return jsonify({"error": "Edital não encontrado"}), 404

        data = request.get_json(silent=True) or {}
        atas = data.get("atas", [])
        if not atas:
            return jsonify({"error": "Nenhuma ata fornecida"}), 400

        import requests as req
        import re as _re
        resultados = []
        for ata in atas[:5]:  # Limitar a 5 atas
            cnpj = ata.get("cnpj_orgao") or ""
            ano = ata.get("ano")

            # Extrair sequencial da COMPRA (não da ata) da URL ou numero_controle
            seq_compra = None
            url_pncp = ata.get("url_pncp") or ""
            # URL formato: /atas/{cnpj}/{ano}/{seq_compra}/{seq_ata}
            match_url = _re.search(r'/atas/\d+/\d+/(\d+)/\d+', url_pncp)
            if match_url:
                seq_compra = int(match_url.group(1))
            else:
                # numero_controle formato: cnpj-1-seq_compra/ano-seq_ata
                nc = ata.get("numero_controle") or ""
                match_nc = _re.search(r'\d{14}-\d+-(\d+)/\d+-\d+', nc)
                if match_nc:
                    seq_compra = int(match_nc.group(1))

            if not cnpj or not ano or not seq_compra:
                continue

            # Buscar itens da compra
            try:
                url_itens = f"https://pncp.gov.br/api/pncp/v1/orgaos/{cnpj}/compras/{ano}/{seq_compra}/itens"
                resp = req.get(url_itens, timeout=15, headers={"Accept": "application/json"})
                if resp.status_code != 200:
                    continue
                itens = resp.json()
                if not isinstance(itens, list):
                    continue

                # Para cada item, buscar resultado (vencedor)
                vencedores_ata = []
                for item in itens[:10]:  # Limitar a 10 itens por ata
                    num_item = item.get("numeroItem")
                    if not num_item:
                        continue
                    try:
                        url_resultado = f"https://pncp.gov.br/api/pncp/v1/orgaos/{cnpj}/compras/{ano}/{seq_compra}/itens/{num_item}/resultados"
                        resp_r = req.get(url_resultado, timeout=10, headers={"Accept": "application/json"})
                        if resp_r.status_code == 200:
                            resultados_item = resp_r.json()
                            if isinstance(resultados_item, list) and resultados_item:
                                r = resultados_item[0]  # Primeiro resultado = vencedor
                                vencedores_ata.append({
                                    "item": num_item,
                                    "descricao": str(item.get("descricao", ""))[:80],
                                    "quantidade": item.get("quantidade"),
                                    "valor_estimado": item.get("valorUnitarioEstimado"),
                                    "vencedor": r.get("nomeRazaoSocialFornecedor"),
                                    "cnpj_vencedor": r.get("niFornecedor"),
                                    "valor_homologado": r.get("valorUnitarioHomologado"),
                                    "qtd_homologada": r.get("quantidadeHomologada"),
                                    "valor_total_homologado": r.get("valorTotalHomologado"),
                                    "porte": r.get("porteFornecedorNome"),
                                })
                    except Exception:
                        pass

                if vencedores_ata:
                    resultados.append({
                        "ata_titulo": ata.get("titulo") or f"Ata {ano}/{seq}",
                        "orgao": ata.get("orgao"),
                        "uf": ata.get("uf"),
                        "vencedores": vencedores_ata,
                    })
            except Exception:
                pass

        # Salvar vencedores no banco (concorrentes + precos_historicos)
        from models import Concorrente, PrecoHistorico
        from datetime import datetime, date as date_type
        salvos_concorrentes = 0
        salvos_precos = 0
        cnpjs_processados = {}  # cache para evitar duplicatas no mesmo batch

        empresa = db.query(Empresa).filter(Empresa.user_id == user_id).first()
        empresa_id = empresa.id if empresa else None

        for res_ata in resultados:
            for v in res_ata.get("vencedores", []):
                cnpj_venc = (v.get("cnpj_vencedor") or "").replace(".", "").replace("/", "").replace("-", "")
                nome_venc = v.get("vencedor")
                if not nome_venc:
                    continue

                # Criar/atualizar concorrente (evitar duplicatas)
                concorrente = cnpjs_processados.get(cnpj_venc) if cnpj_venc else None
                if not concorrente and cnpj_venc:
                    concorrente = db.query(Concorrente).filter(Concorrente.cnpj == cnpj_venc).first()
                if not concorrente and nome_venc:
                    concorrente = db.query(Concorrente).filter(Concorrente.nome.ilike(f"%{nome_venc[:20]}%")).first()
                if not concorrente:
                    concorrente = Concorrente(
                        id=str(uuid.uuid4()),
                        nome=nome_venc,
                        cnpj=cnpj_venc or None,
                        editais_participados=1,
                        editais_ganhos=1,
                    )
                    db.add(concorrente)
                    db.flush()  # Garante ID gerado antes de usar
                    salvos_concorrentes += 1
                else:
                    concorrente.editais_participados = (concorrente.editais_participados or 0) + 1
                    concorrente.editais_ganhos = (concorrente.editais_ganhos or 0) + 1

                if cnpj_venc:
                    cnpjs_processados[cnpj_venc] = concorrente

                # Salvar preço histórico
                try:
                    preco_hist = PrecoHistorico(
                        id=str(uuid.uuid4()),
                        edital_id=edital_id,
                        user_id=user_id,
                        empresa_id=empresa_id,
                        concorrente_id=concorrente.id,
                        preco_referencia=v.get("valor_estimado"),
                        preco_vencedor=v.get("valor_homologado"),
                        empresa_vencedora=nome_venc,
                        cnpj_vencedor=cnpj_venc or None,
                        resultado="derrota",
                        fonte="pncp",
                        data_registro=datetime.now(),
                    )
                    if v.get("valor_estimado") and v.get("valor_homologado") and v["valor_estimado"] > 0:
                        preco_hist.desconto_percentual = round((1 - v["valor_homologado"] / v["valor_estimado"]) * 100, 2)
                    db.add(preco_hist)
                    salvos_precos += 1
                except Exception:
                    pass  # Ignora duplicatas

        try:
            db.commit()
            print(f"[VENCEDORES] Salvos: {salvos_concorrentes} concorrentes, {salvos_precos} preços históricos")
        except Exception:
            db.rollback()
            print("[VENCEDORES] Erro ao salvar, rollback")

        return jsonify({
            "success": True,
            "resultados": resultados,
            "total_atas_analisadas": len(resultados),
            "salvos": {"concorrentes": salvos_concorrentes, "precos_historicos": salvos_precos},
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


@app.route("/api/editais/<edital_id>/extrair-requisitos", methods=["POST"])
@require_auth
def extrair_requisitos_edital(edital_id):
    """
    Extrai requisitos (técnicos, documentais, comerciais) do PDF já baixado do edital.
    Chamado pela aba Documentos quando o edital tem PDF mas não tem requisitos extraídos.
    """
    from models import Edital, EditalDocumento, EditalRequisito
    user_id = get_current_user_id()

    db = get_db()
    try:
        edital = db.query(Edital).filter(
            Edital.id == edital_id,
            Edital.user_id == user_id
        ).first()
        if not edital:
            return jsonify({"error": "Edital não encontrado"}), 404

        # Verificar se já tem requisitos extraídos
        forcar = request.get_json(silent=True) or {}
        forcar = forcar.get("forcar", False)
        requisitos_existentes = db.query(EditalRequisito).filter(
            EditalRequisito.edital_id == edital_id
        ).count()
        if requisitos_existentes > 0 and not forcar:
            return jsonify({
                "success": True,
                "message": f"Edital já possui {requisitos_existentes} requisitos extraídos",
                "requisitos_extraidos": requisitos_existentes,
                "ja_existia": True,
            })

        # Se forçar, deletar requisitos antigos
        if forcar and requisitos_existentes > 0:
            db.query(EditalRequisito).filter(EditalRequisito.edital_id == edital_id).delete()
            db.commit()
            print(f"[EXTRAIR-REQUISITOS] Removidos {requisitos_existentes} requisitos antigos (forcar=True)")

        # Buscar texto do PDF do edital
        doc = db.query(EditalDocumento).filter(
            EditalDocumento.edital_id == edital_id,
            EditalDocumento.texto_extraido != None
        ).first()

        if not doc or not doc.texto_extraido or len(doc.texto_extraido) < 200:
            return jsonify({
                "error": "Edital não possui PDF com texto extraído. Baixe o PDF primeiro.",
                "sem_pdf": True,
            }), 400

        # Extrair requisitos via IA
        resultado = tool_extrair_requisitos(edital_id, doc.texto_extraido, user_id)

        if resultado.get("success"):
            return jsonify({
                "success": True,
                "requisitos_extraidos": resultado.get("requisitos_extraidos", 0),
                "categoria": resultado.get("categoria"),
            })
        else:
            return jsonify({"error": resultado.get("error", "Erro na extração")}), 500

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


# =============================================================================
# B7: SICONV como fonte de busca de editais
# (Implementado em tools.py — parser + integração via scraper web)
# =============================================================================


# =============================================================================
# B8: Busca automática de certidões (portais oficiais)
# =============================================================================


def _tentar_obter_certidao(fonte, cnpj_limpo, empresa, upload_dir):
    """
    Tenta obter a certidão automaticamente via HTTP nos portais públicos.
    Retorna dict com {sucesso, path_arquivo, numero, data_vencimento, mensagem, dados_extras}.

    Testado em 10/03/2026 — resultados reais:

    BUSCA AUTOMÁTICA REAL (sem captcha):
    1. CGU Certidão Correcional — API REST, consolida CEIS+CNEP+CEPIM+CGU-PJ+ePAD
    2. CRF/FGTS (Caixa)        — JSF sem captcha, retorna Regular/Irregular
    3. BrasilAPI CNPJ           — REST gratuita, dados cadastrais + Simples Nacional
    4. API Portal Transparência — REST oficial (requer chave gratuita)

    CAPTCHA (upload manual obrigatório):
    5. CND Federal (Receita)    — URL antiga MORTA (404), migrou para e-CAC com login gov.br
    6. CNDT/TST                 — Captcha próprio (imagem+áudio)
    7. CND Estadual (SEFAZ)     — reCAPTCHA v2 (SP) ou hCaptcha (varia)
    8. CND Municipal            — Varia por município, maioria com captcha
    """
    import requests
    from datetime import date, datetime, timedelta
    import json as json_lib
    import re
    import time

    nome_lower = fonte.nome.lower()
    tipo = fonte.tipo_certidao or ''
    resultado = {
        "sucesso": False, "mensagem": "", "path_arquivo": None,
        "numero": None, "data_vencimento": None, "dados_extras": None
    }

    HEADERS = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "application/json, text/html, */*",
    }

    try:
        os.makedirs(upload_dir, exist_ok=True)

        # ═══════════════════════════════════════════════════════════════════
        # 1. CGU CERTIDÃO CORRECIONAL — Consolida CEIS+CNEP+CEPIM+CGU-PJ+ePAD
        #    API pública, SEM captcha, SEM chave. Testado: FUNCIONA.
        # ═══════════════════════════════════════════════════════════════════
        if any(k in nome_lower for k in ['cgu', 'correcional', 'ceis', 'cnep', 'cepim']):
            try:
                # Passo 1: Solicitar consulta
                resp1 = requests.post(
                    "https://certidoes.cgu.gov.br/api/publico/consulta-responsabilizacao",
                    json={
                        "cpfCnpj": [cnpj_limpo],
                        "certidoes": [8],  # ID=8 = Entes Privados (CNPJ)
                        "captcha": "",
                        "idTipoNatureza": 0,
                        "isConsultaReceita": False,
                    },
                    headers=HEADERS,
                    timeout=20,
                )
                if resp1.status_code != 200:
                    resultado["mensagem"] = f"CGU: Erro ao solicitar consulta (HTTP {resp1.status_code})"
                    return resultado

                ids = resp1.json()
                if not ids or not isinstance(ids, list):
                    resultado["mensagem"] = "CGU: Resposta inesperada ao solicitar consulta"
                    return resultado

                result_id = ids[0]

                # Passo 2: Obter resultado (pode levar alguns segundos)
                for tentativa in range(5):
                    time.sleep(1)
                    resp2 = requests.get(
                        f"https://certidoes.cgu.gov.br/api/publico/resultado-consulta-responsabilizacao/{result_id}",
                        headers=HEADERS,
                        timeout=20,
                    )
                    if resp2.status_code == 200:
                        dados = resp2.json()
                        if dados and isinstance(dados, list) and len(dados) > 0:
                            break
                else:
                    resultado["mensagem"] = "CGU: Timeout aguardando resultado da consulta"
                    return resultado

                consulta = dados[0]
                certidoes_result = consulta.get("consultasCertidoes", [])

                # Analisar resultados
                todas_nada_consta = True
                detalhes = []
                for cert_r in certidoes_result:
                    for ws in cert_r.get("consultasWebservices", []):
                        ws_nome = ws.get("nomeWebservice", "")
                        ws_tipo = ws.get("tipoRetornoWebservice", "")
                        detalhes.append(f"{ws_nome}: {ws_tipo}")
                        if ws_tipo != "NADA_CONSTA":
                            todas_nada_consta = False

                nome_encontrado = consulta.get("nomeRequerente", "")
                data_consulta = consulta.get("dataHoraConsulta", "")

                if todas_nada_consta:
                    resultado["sucesso"] = True
                    resultado["data_vencimento"] = (date.today() + timedelta(days=180)).isoformat()
                    resultado["numero"] = f"CGU-{result_id[:8]}"
                    resultado["mensagem"] = f"CGU Certidão Correcional: NADA CONSTA ({', '.join(detalhes)})"
                else:
                    resultado["mensagem"] = f"CGU: ATENÇÃO — Registros encontrados: {', '.join(detalhes)}"

                resultado["dados_extras"] = json_lib.dumps({
                    "nome_encontrado": nome_encontrado,
                    "data_consulta": data_consulta,
                    "webservices": detalhes,
                    "certidao_negativa": todas_nada_consta,
                }, ensure_ascii=False)
                return resultado

            except requests.exceptions.Timeout:
                resultado["mensagem"] = "CGU: Timeout — tente novamente"
            except Exception as e:
                resultado["mensagem"] = f"CGU: Erro — {str(e)[:100]}"
            return resultado

        # ═══════════════════════════════════════════════════════════════════
        # 2. CRF/FGTS — Caixa Econômica — Automação via Playwright
        #    ShieldSquare anti-bot: Playwright executa o JS challenge.
        #    Fallback: mensagem de upload manual.
        # ═══════════════════════════════════════════════════════════════════
        if 'fgts' in nome_lower or 'crf' in nome_lower or 'caixa' in nome_lower or tipo == 'fgts':
            try:
                from certidao_browser import buscar_fgts_browser
                print(f"[CERTIDAO] FGTS: Tentando automação via Playwright...")
                resultado_browser = buscar_fgts_browser(cnpj_limpo, upload_dir)
                if resultado_browser.get("sucesso") or resultado_browser.get("dados_extras"):
                    return resultado_browser
                # Falhou — usar mensagem do browser como fallback
                resultado["mensagem"] = resultado_browser.get("mensagem", "")
            except ImportError:
                pass
            except Exception as e:
                print(f"[CERTIDAO] FGTS browser erro: {e}")

            if not resultado["mensagem"]:
                resultado["mensagem"] = (
                    "FGTS/CRF: Automação indisponível. "
                    "Acesse https://consulta-crf.caixa.gov.br/consultacrf/pages/consultaEmpregador.jsf, "
                    "informe o CNPJ e faça upload do PDF. Validade: 30 dias."
                )
            return resultado

        # ═══════════════════════════════════════════════════════════════════
        # 3. BrasilAPI CNPJ — Dados cadastrais + Simples Nacional (GRATUITA)
        # ═══════════════════════════════════════════════════════════════════
        if 'brasilapi' in nome_lower or 'cnpj' in nome_lower or 'situação cadastral' in nome_lower or 'situacao cadastral' in nome_lower:
            try:
                resp = requests.get(
                    f"https://brasilapi.com.br/api/cnpj/v1/{cnpj_limpo}",
                    headers=HEADERS,
                    timeout=15,
                )
                if resp.status_code == 200:
                    dados = resp.json()
                    situacao = dados.get("descricao_situacao_cadastral", "")
                    simples = dados.get("opcao_pelo_simples", False)
                    razao = dados.get("razao_social", "")

                    if situacao.upper() == "ATIVA":
                        resultado["sucesso"] = True
                        resultado["mensagem"] = f"CNPJ: Situação ATIVA — {razao}. Simples: {'Sim' if simples else 'Não'}"
                        resultado["data_vencimento"] = (date.today() + timedelta(days=30)).isoformat()
                    else:
                        resultado["mensagem"] = f"CNPJ: Situação '{situacao}' — {razao}"

                    resultado["dados_extras"] = json_lib.dumps({
                        "razao_social": razao,
                        "situacao": situacao,
                        "simples_nacional": simples,
                        "uf": dados.get("uf"),
                        "municipio": dados.get("municipio"),
                        "cnae_fiscal": dados.get("cnae_fiscal"),
                        "cnae_descricao": dados.get("cnae_fiscal_descricao"),
                        "porte": dados.get("porte"),
                        "capital_social": dados.get("capital_social"),
                        "socios": [s.get("nome_socio") for s in dados.get("qsa", [])],
                    }, ensure_ascii=False)
                    return resultado
                else:
                    resultado["mensagem"] = f"BrasilAPI: HTTP {resp.status_code}"
            except requests.exceptions.Timeout:
                resultado["mensagem"] = "BrasilAPI: Timeout"
            except Exception as e:
                resultado["mensagem"] = f"BrasilAPI: Erro — {str(e)[:100]}"
            return resultado

        # ═══════════════════════════════════════════════════════════════════
        # 4. API OFICIAL Portal da Transparência (requer chave gratuita)
        #    Cadastro em: portaldatransparencia.gov.br/api-de-dados/cadastrar-email
        #    Header: chave-api-dados
        # ═══════════════════════════════════════════════════════════════════
        if any(k in nome_lower for k in ['transparência', 'transparencia', 'portal transparencia']):
            chave_api = fonte.api_key or ""
            if not chave_api:
                resultado["mensagem"] = (
                    "Portal da Transparência API: Chave API não configurada. "
                    "Cadastre-se gratuitamente em https://portaldatransparencia.gov.br/api-de-dados/cadastrar-email "
                    "e configure a chave na fonte de certidão."
                )
                return resultado

            try:
                api_base = "https://api.portaldatransparencia.gov.br/api-de-dados"
                api_headers = {**HEADERS, "chave-api-dados": chave_api}
                resultados_api = {}

                for cadastro, param in [("ceis", "codigoSancionado"), ("cnep", "codigoSancionado"), ("cepim", "cnpjSancionado")]:
                    resp = requests.get(
                        f"{api_base}/{cadastro}",
                        params={param: cnpj_limpo, "pagina": 1},
                        headers=api_headers,
                        timeout=15,
                    )
                    if resp.status_code == 200:
                        dados = resp.json()
                        registros = dados if isinstance(dados, list) else dados.get("data", [])
                        resultados_api[cadastro.upper()] = len(registros)
                    elif resp.status_code == 401:
                        resultado["mensagem"] = "Portal da Transparência: Chave API inválida"
                        return resultado
                    else:
                        resultados_api[cadastro.upper()] = f"HTTP {resp.status_code}"

                todas_limpo = all(v == 0 for v in resultados_api.values() if isinstance(v, int))
                if todas_limpo:
                    resultado["sucesso"] = True
                    resultado["data_vencimento"] = (date.today() + timedelta(days=180)).isoformat()
                    resultado["mensagem"] = f"Portal Transparência: NADA CONSTA ({', '.join(f'{k}:0' for k in resultados_api)})"
                else:
                    resultado["mensagem"] = f"Portal Transparência: {', '.join(f'{k}:{v}' for k, v in resultados_api.items())}"

                resultado["dados_extras"] = json_lib.dumps(resultados_api, ensure_ascii=False)
                return resultado

            except Exception as e:
                resultado["mensagem"] = f"Portal Transparência API: Erro — {str(e)[:100]}"
            return resultado

        # ═══════════════════════════════════════════════════════════════════
        # 5. CND FEDERAL (Receita Federal / PGFN) — MANUAL
        #    Portal exige hCaptcha invisível (domínios gov.br bloqueados por solvers)
        # ═══════════════════════════════════════════════════════════════════
        if tipo == 'cnd_federal' or 'receita federal' in nome_lower or 'cnd federal' in nome_lower or 'pgfn' in nome_lower:
            resultado["mensagem"] = (
                "CND Federal: Busca automática não disponível — portal da Receita exige hCaptcha "
                "que não pode ser resolvido automaticamente. "
                "Acesse https://servicos.receitafederal.gov.br/servico/certidoes/, "
                "selecione Pessoa Jurídica, informe o CNPJ e clique Emitir Certidão. "
                "Faça upload do PDF — a IA extrai os dados automaticamente."
            )
            resultado["dados_extras"] = {
                "Método": "MANUAL — hCaptcha gov.br não suportado por solvers",
                "Portal": "https://servicos.receitafederal.gov.br/servico/certidoes/",
                "Procedimento": "Selecionar Pessoa Jurídica → Informar CNPJ → Emitir Certidão → Upload PDF",
                "Validade": "180 dias",
            }
            return resultado

        # ═══════════════════════════════════════════════════════════════════
        # 6. CNDT/TST — Automação via Playwright + CapSolver (captcha imagem)
        # ═══════════════════════════════════════════════════════════════════
        if 'cndt' in nome_lower or 'tst' in nome_lower or tipo == 'trabalhista':
            try:
                from certidao_browser import buscar_tst_browser
                from config import CAPSOLVER_API_KEY as _cap_key
                if _cap_key:
                    print(f"[CERTIDAO] TST: Tentando automação via Playwright + CapSolver...")
                    resultado_browser = buscar_tst_browser(cnpj_limpo, upload_dir)
                    if resultado_browser.get("sucesso") or resultado_browser.get("dados_extras"):
                        return resultado_browser
                    resultado["mensagem"] = resultado_browser.get("mensagem", "")
            except ImportError:
                pass
            except Exception as e:
                print(f"[CERTIDAO] TST browser erro: {e}")

            if not resultado["mensagem"]:
                resultado["mensagem"] = (
                    "CNDT/TST: Automação indisponível (CapSolver não configurado ou falha). "
                    "Acesse https://cndt-certidao.tst.jus.br/gerarCertidao.faces, "
                    "informe o CNPJ formatado, resolva o captcha e faça upload do PDF."
                )
            return resultado

        # ═══════════════════════════════════════════════════════════════════
        # 7. CND ESTADUAL (SEFAZ) — Automação via Playwright + CapSolver
        # ═══════════════════════════════════════════════════════════════════
        if tipo == 'cnd_estadual' or 'sefaz' in nome_lower or 'cnd estadual' in nome_lower:
            uf = fonte.uf or (empresa.uf if hasattr(empresa, 'uf') else '')
            url_sefaz = fonte.url_portal or (f"https://www.sefaz.{uf.lower()}.gov.br" if uf else "")

            try:
                from certidao_browser import buscar_sefaz_browser
                from config import CAPSOLVER_API_KEY as _cap_key
                if _cap_key and url_sefaz:
                    ie = fonte.cnpj_consulta or ""  # IE pode estar no cnpj_consulta
                    print(f"[CERTIDAO] SEFAZ-{uf}: Tentando automação via Playwright + CapSolver...")
                    resultado_browser = buscar_sefaz_browser(cnpj_limpo, uf, url_sefaz, upload_dir, ie)
                    if resultado_browser.get("sucesso") or resultado_browser.get("dados_extras"):
                        return resultado_browser
                    resultado["mensagem"] = resultado_browser.get("mensagem", "")
            except ImportError:
                pass
            except Exception as e:
                print(f"[CERTIDAO] SEFAZ browser erro: {e}")

            if not resultado["mensagem"]:
                resultado["mensagem"] = (
                    f"CND Estadual ({uf or 'UF?'}): Automação indisponível. "
                    f"Acesse {url_sefaz}, informe CNPJ + Inscrição Estadual, "
                    "resolva o captcha e faça upload do PDF."
                )
            return resultado

        # ═══════════════════════════════════════════════════════════════════
        # 8. CND MUNICIPAL — Varia por município
        # ═══════════════════════════════════════════════════════════════════
        if tipo == 'cnd_municipal' or 'municipal' in nome_lower or 'prefeitura' in nome_lower:
            cidade = fonte.cidade or (empresa.cidade if hasattr(empresa, 'cidade') else '')
            url_pref = fonte.url_portal or ""
            resultado["mensagem"] = (
                f"CND Municipal ({cidade or 'cidade?'}): Acesse o portal da prefeitura"
                f"{' em ' + url_pref if url_pref else ''}, "
                "informe CNPJ + Inscrição Municipal (CCM) e faça upload do PDF."
            )
            return resultado

        # ═══════════════════════════════════════════════════════════════════
        # FALLBACK — Orientação genérica
        # ═══════════════════════════════════════════════════════════════════
        url_portal = fonte.url_portal or ""
        resultado["mensagem"] = (
            f"Acesse o portal{' em ' + url_portal if url_portal else ''} manualmente, "
            "obtenha a certidão e faça upload do PDF."
        )
        return resultado

    except Exception as e:
        resultado["mensagem"] = f"Erro geral: {str(e)[:100]}"
        return resultado


@app.route("/api/empresa-certidoes/sincronizar-fontes", methods=["POST"])
@require_auth
def sincronizar_fontes_certidoes():
    """
    Sincroniza fontes de certidões com registros de certidões da empresa.
    Para cada fonte ativa que NÃO tem certidão vinculada, cria um registro
    com status 'pendente' ou 'nao_disponivel'.
    Chamado automaticamente ao carregar a página Empresa.
    """
    from models import Empresa, EmpresaCertidao, FonteCertidao
    user_id = get_current_user_id()
    data = request.json or {}
    empresa_id = data.get('empresa_id')
    if not empresa_id:
        return jsonify({"error": "empresa_id é obrigatório"}), 400

    db = get_db()
    try:
        empresa = db.query(Empresa).filter(
            Empresa.id == empresa_id,
            Empresa.user_id == user_id
        ).first()
        if not empresa:
            return jsonify({"error": "Empresa não encontrada"}), 404

        # Buscar fontes ativas
        fontes = db.query(FonteCertidao).filter(
            FonteCertidao.user_id == user_id,
            FonteCertidao.ativo == True
        ).all()

        # Buscar certidões existentes
        certs_existentes = db.query(EmpresaCertidao).filter(
            EmpresaCertidao.empresa_id == empresa_id
        ).all()
        fontes_com_cert = {c.fonte_certidao_id for c in certs_existentes if c.fonte_certidao_id}

        criadas = 0
        from datetime import date
        for fonte in fontes:
            if fonte.id in fontes_com_cert:
                continue  # Já tem certidão vinculada

            # Determinar status inicial
            if not fonte.permite_busca_automatica:
                status_novo = "nao_disponivel"
            else:
                status_novo = "pendente"

            nova_cert = EmpresaCertidao()
            nova_cert.id = str(uuid.uuid4())
            nova_cert.empresa_id = empresa_id
            nova_cert.tipo = fonte.tipo_certidao
            nova_cert.orgao_emissor = fonte.orgao_emissor or fonte.nome
            nova_cert.url_consulta = fonte.url_portal or ""
            nova_cert.fonte_certidao_id = fonte.id
            nova_cert.status = status_novo
            nova_cert.data_vencimento = date.today()
            db.add(nova_cert)
            criadas += 1

        if criadas > 0:
            db.commit()

        return jsonify({
            "message": f"Sincronização concluída: {criadas} certidões criadas, {len(certs_existentes)} já existentes",
            "criadas": criadas,
            "total": len(certs_existentes) + criadas,
        })
    finally:
        db.close()


@app.route("/api/empresa-certidoes/buscar-automatica", methods=["POST"])
@require_auth
def buscar_certidoes_automatica():
    """
    Busca certidões da empresa automaticamente nos portais oficiais.
    Itera TODAS as fontes cadastradas (não apenas 5 tipos fixos).
    Para cada fonte ativa:
    1. Verifica se já existe certidão vinculada a essa fonte
    2. Se busca automática permitida e pública: tenta web search
    3. Se requer autenticação: marca como pendente
    4. Se manual: marca como nao_disponivel

    JSON body:
    - empresa_id: ID da empresa
    """
    from models import Empresa, EmpresaCertidao, FonteCertidao
    user_id = get_current_user_id()
    data = request.json or {}

    empresa_id = data.get('empresa_id')
    if not empresa_id:
        return jsonify({"error": "empresa_id é obrigatório"}), 400

    db = get_db()
    try:
        empresa = db.query(Empresa).filter(
            Empresa.id == empresa_id,
            Empresa.user_id == user_id
        ).first()
        if not empresa:
            return jsonify({"error": "Empresa não encontrada"}), 404

        if not empresa.cnpj:
            return jsonify({"error": "Empresa sem CNPJ cadastrado"}), 400

        cnpj = empresa.cnpj.replace('.', '').replace('/', '').replace('-', '').strip()

        # Buscar TODAS as fontes ativas do usuário
        fontes = db.query(FonteCertidao).filter(
            FonteCertidao.user_id == user_id,
            FonteCertidao.ativo == True
        ).all()

        if not fontes:
            return jsonify({"error": "Nenhuma fonte de certidão cadastrada. Acesse Cadastros > Empresa > Fontes de Certidões para configurar."}), 400

        resultados = []
        from datetime import date, datetime
        stats = {"publicas": 0, "autenticadas": 0, "manuais": 0, "sem_api_key": 0}

        for fonte in fontes:
            # Verificar se já existe certidão vinculada a esta fonte
            cert_existente = db.query(EmpresaCertidao).filter(
                EmpresaCertidao.empresa_id == empresa_id,
                EmpresaCertidao.fonte_certidao_id == fonte.id
            ).first()

            url_consulta = fonte.url_portal or ""
            resultado_auto = None

            # Classificar a fonte
            if not fonte.permite_busca_automatica:
                # Fonte manual — usar observacoes da fonte como orientação detalhada
                stats["manuais"] += 1
                status_novo = "nao_disponivel"
                acao = fonte.observacoes or "Upload manual necessário"
                # Guardar dados estruturados da fonte para o modal de detalhes
                resultado_auto = {
                    "dados_extras": {
                        "Método de Acesso": fonte.metodo_acesso or "publico",
                        "Portal": fonte.url_portal or "Não informado",
                        "Requer Autenticação": "Sim" if fonte.requer_autenticacao else "Não",
                        "Procedimento": "Acesse o portal, emita a certidão e faça upload do PDF",
                    }
                }
            elif fonte.requer_autenticacao and not fonte.usuario and not fonte.api_key and not fonte.certificado_path:
                # Requer autenticação mas sem credenciais
                stats["autenticadas"] += 1
                status_novo = "pendente"
                acao = f"Configure credenciais ({fonte.metodo_acesso}) em Fontes de Certidões"
            else:
                # Fonte pública ou com credenciais — tentar busca automática real
                stats["publicas"] += 1
                status_novo = "buscando"
                acao = "Buscando automaticamente..."

                # Tentar obter a certidão via HTTP
                upload_dir = os.path.join(UPLOAD_FOLDER, 'empresa', empresa_id, 'certidoes')
                resultado_auto = _tentar_obter_certidao(fonte, cnpj, empresa, upload_dir)

                if resultado_auto and resultado_auto.get("sucesso"):
                    status_novo = "valida"
                    acao = resultado_auto.get("mensagem", "Certidão obtida automaticamente")
                    stats["obtidas_auto"] = stats.get("obtidas_auto", 0) + 1
                elif resultado_auto:
                    status_novo = "pendente"
                    acao = resultado_auto.get("mensagem", "Acesse o portal manualmente")
                    # Contar fontes que precisam de API key separadamente
                    if "Chave de API não configurada" in acao or "Chave de API inválida" in acao:
                        stats["sem_api_key"] += 1
                else:
                    status_novo = "pendente"
                    acao = "Acesse o portal para emitir a certidão"

            # Atualizar ultima_consulta na fonte
            fonte.ultima_consulta = datetime.now()
            fonte.resultado_ultima_consulta = 'sucesso' if status_novo == 'valida' else 'erro'

            if cert_existente:
                cert_existente.url_consulta = url_consulta
                cert_existente.orgao_emissor = fonte.orgao_emissor or fonte.nome
                cert_existente.fonte_certidao_id = fonte.id
                # Persistir mensagem e dados_extras da busca
                cert_existente.mensagem = acao
                if resultado_auto and resultado_auto.get("dados_extras"):
                    import json as json_mod
                    de = resultado_auto["dados_extras"]
                    cert_existente.dados_extras = json_mod.loads(de) if isinstance(de, str) else de
                if status_novo == 'valida':
                    # Nova busca retornou válida — atualizar dados PRIMEIRO
                    cert_existente.status = 'valida'
                    if resultado_auto and resultado_auto.get("path_arquivo"):
                        cert_existente.path_arquivo = resultado_auto["path_arquivo"]
                    if resultado_auto and resultado_auto.get("data_vencimento"):
                        try:
                            cert_existente.data_vencimento = datetime.strptime(resultado_auto["data_vencimento"], '%Y-%m-%d').date()
                        except (ValueError, TypeError):
                            pass
                    if resultado_auto and resultado_auto.get("numero"):
                        cert_existente.numero = resultado_auto["numero"]
                    cert_existente.data_emissao = date.today()
                    # Verificar se a nova validade já venceu
                    if cert_existente.data_vencimento and cert_existente.data_vencimento < date.today():
                        cert_existente.status = 'vencida'
                elif cert_existente.data_vencimento and cert_existente.data_vencimento < date.today():
                    cert_existente.status = 'vencida'
                elif cert_existente.status in ('buscando', 'erro'):
                    cert_existente.status = status_novo
                resultados.append({
                    "tipo": fonte.tipo_certidao,
                    "status": "atualizado",
                    "certidao": cert_existente.to_dict(),
                    "fonte_nome": fonte.nome,
                    "fonte_id": fonte.id,
                    "url_consulta": url_consulta,
                    "permite_busca_automatica": fonte.permite_busca_automatica,
                    "requer_autenticacao": fonte.requer_autenticacao,
                    "acao": acao,
                })
            else:
                nova_cert = EmpresaCertidao()
                nova_cert.id = str(uuid.uuid4())
                nova_cert.empresa_id = empresa_id
                nova_cert.tipo = fonte.tipo_certidao
                nova_cert.orgao_emissor = fonte.orgao_emissor or fonte.nome
                nova_cert.url_consulta = url_consulta
                nova_cert.fonte_certidao_id = fonte.id
                nova_cert.status = status_novo
                if resultado_auto and resultado_auto.get("path_arquivo"):
                    nova_cert.path_arquivo = resultado_auto["path_arquivo"]
                if resultado_auto and resultado_auto.get("data_vencimento"):
                    try:
                        nova_cert.data_vencimento = datetime.strptime(resultado_auto["data_vencimento"], '%Y-%m-%d').date()
                    except (ValueError, TypeError):
                        nova_cert.data_vencimento = date.today()
                else:
                    nova_cert.data_vencimento = date.today()
                if resultado_auto and resultado_auto.get("numero"):
                    nova_cert.numero = resultado_auto["numero"]
                # Persistir mensagem e dados_extras
                nova_cert.mensagem = acao
                if resultado_auto and resultado_auto.get("dados_extras"):
                    import json as json_mod
                    de = resultado_auto["dados_extras"]
                    nova_cert.dados_extras = json_mod.loads(de) if isinstance(de, str) else de
                if status_novo == 'valida':
                    nova_cert.data_emissao = date.today()
                db.add(nova_cert)
                resultados.append({
                    "tipo": fonte.tipo_certidao,
                    "status": "criado",
                    "certidao": nova_cert.to_dict(),
                    "fonte_nome": fonte.nome,
                    "fonte_id": fonte.id,
                    "url_consulta": url_consulta,
                    "permite_busca_automatica": fonte.permite_busca_automatica,
                    "requer_autenticacao": fonte.requer_autenticacao,
                    "acao": acao,
                })

        db.commit()

        obtidas = stats.get("obtidas_auto", 0)
        sem_key = stats.get("sem_api_key", 0)
        pendentes_captcha = stats['publicas'] - obtidas - sem_key

        # Montar mensagem detalhada
        partes = [f"Busca concluída: {len(resultados)} fontes"]
        detalhes = []
        if obtidas > 0:
            detalhes.append(f"{obtidas} obtidas automaticamente")
        if sem_key > 0:
            detalhes.append(f"{sem_key} aguardando chave de API")
        if pendentes_captcha > 0:
            detalhes.append(f"{pendentes_captcha} pendentes (portal manual)")
        if stats['autenticadas'] > 0:
            detalhes.append(f"{stats['autenticadas']} aguardando credenciais")
        if stats['manuais'] > 0:
            detalhes.append(f"{stats['manuais']} manuais")
        msg = partes[0] + " — " + ", ".join(detalhes) if detalhes else partes[0]

        return jsonify({
            "success": True,
            "message": msg,
            "resultados": resultados,
            "cnpj": empresa.cnpj,
            "stats": stats,
            "nota": "Certidões marcadas como 'Válida' foram obtidas automaticamente. "
                    "Para as demais, clique no ícone de olho para abrir o portal e faça upload do PDF."
        })

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


# =============================================================================
# Busca de Certidões com SSE (Server-Sent Events) — progresso em tempo real
# =============================================================================

@app.route("/api/empresa-certidoes/buscar-stream", methods=["POST"])
@require_auth
def buscar_certidoes_stream():
    """
    Mesmo que buscar_certidoes_automatica(), mas retorna SSE stream com progresso.
    Cada fonte emite 2 eventos: 'searching' (início) e 'result' (resultado).
    Evento final: 'complete' com stats.
    """
    import json as json_mod
    from models import Empresa, EmpresaCertidao, FonteCertidao
    from datetime import date, datetime

    user_id = get_current_user_id()
    data = request.json or {}
    empresa_id = data.get('empresa_id')

    if not empresa_id:
        return jsonify({"error": "empresa_id é obrigatório"}), 400

    def generate():
        db = get_db()
        try:
            empresa = db.query(Empresa).filter(
                Empresa.id == empresa_id,
                Empresa.user_id == user_id
            ).first()
            if not empresa:
                yield f"data: {json_mod.dumps({'type': 'error', 'message': 'Empresa não encontrada'})}\n\n"
                return
            if not empresa.cnpj:
                yield f"data: {json_mod.dumps({'type': 'error', 'message': 'Empresa sem CNPJ cadastrado'})}\n\n"
                return

            cnpj = empresa.cnpj.replace('.', '').replace('/', '').replace('-', '').strip()

            fontes = db.query(FonteCertidao).filter(
                FonteCertidao.user_id == user_id,
                FonteCertidao.ativo == True
            ).all()

            if not fontes:
                yield f"data: {json_mod.dumps({'type': 'error', 'message': 'Nenhuma fonte de certidão cadastrada.'})}\n\n"
                return

            total = len(fontes)
            nomes_fontes = [f.nome for f in fontes]
            yield f"data: {json_mod.dumps({'type': 'start', 'total': total, 'fontes': nomes_fontes})}\n\n"

            resultados = []
            stats = {"publicas": 0, "autenticadas": 0, "manuais": 0, "sem_api_key": 0, "obtidas_auto": 0}

            for idx, fonte in enumerate(fontes):
                # Emit searching event
                yield f"data: {json_mod.dumps({'type': 'searching', 'index': idx + 1, 'total': total, 'fonte': fonte.nome, 'fonte_id': fonte.id})}\n\n"

                cert_existente = db.query(EmpresaCertidao).filter(
                    EmpresaCertidao.empresa_id == empresa_id,
                    EmpresaCertidao.fonte_certidao_id == fonte.id
                ).first()

                url_consulta = fonte.url_portal or ""
                resultado_auto = None

                # Classificar a fonte
                if not fonte.permite_busca_automatica:
                    stats["manuais"] += 1
                    status_novo = "nao_disponivel"
                    acao = fonte.observacoes or "Upload manual necessário"
                    resultado_auto = {
                        "dados_extras": {
                            "Método de Acesso": fonte.metodo_acesso or "publico",
                            "Portal": fonte.url_portal or "Não informado",
                            "Requer Autenticação": "Sim" if fonte.requer_autenticacao else "Não",
                            "Procedimento": "Acesse o portal, emita a certidão e faça upload do PDF",
                        }
                    }
                    sucesso_flag = False
                elif fonte.requer_autenticacao and not fonte.usuario and not fonte.api_key and not fonte.certificado_path:
                    stats["autenticadas"] += 1
                    status_novo = "pendente"
                    acao = f"Configure credenciais ({fonte.metodo_acesso}) em Fontes de Certidões"
                    sucesso_flag = False
                else:
                    stats["publicas"] += 1
                    status_novo = "buscando"
                    acao = "Buscando automaticamente..."

                    upload_dir = os.path.join(UPLOAD_FOLDER, 'empresa', empresa_id, 'certidoes')
                    resultado_auto = _tentar_obter_certidao(fonte, cnpj, empresa, upload_dir)

                    if resultado_auto and resultado_auto.get("sucesso"):
                        status_novo = "valida"
                        acao = resultado_auto.get("mensagem", "Certidão obtida automaticamente")
                        stats["obtidas_auto"] += 1
                        sucesso_flag = True
                    elif resultado_auto:
                        status_novo = "pendente"
                        acao = resultado_auto.get("mensagem", "Acesse o portal manualmente")
                        if "Chave de API não configurada" in acao or "Chave de API inválida" in acao:
                            stats["sem_api_key"] += 1
                        sucesso_flag = False
                    else:
                        status_novo = "pendente"
                        acao = "Acesse o portal para emitir a certidão"
                        sucesso_flag = False

                # Atualizar fonte
                fonte.ultima_consulta = datetime.now()
                fonte.resultado_ultima_consulta = 'sucesso' if status_novo == 'valida' else 'erro'

                # Criar/atualizar certidão
                if cert_existente:
                    cert_existente.url_consulta = url_consulta
                    cert_existente.orgao_emissor = fonte.orgao_emissor or fonte.nome
                    cert_existente.fonte_certidao_id = fonte.id
                    cert_existente.mensagem = acao
                    if resultado_auto and resultado_auto.get("dados_extras"):
                        de = resultado_auto["dados_extras"]
                        cert_existente.dados_extras = json_mod.loads(de) if isinstance(de, str) else de
                    if status_novo == 'valida':
                        cert_existente.status = 'valida'
                        if resultado_auto and resultado_auto.get("path_arquivo"):
                            cert_existente.path_arquivo = resultado_auto["path_arquivo"]
                        if resultado_auto and resultado_auto.get("data_vencimento"):
                            try:
                                cert_existente.data_vencimento = datetime.strptime(resultado_auto["data_vencimento"], '%Y-%m-%d').date()
                            except (ValueError, TypeError):
                                pass
                        if resultado_auto and resultado_auto.get("numero"):
                            cert_existente.numero = resultado_auto["numero"]
                        cert_existente.data_emissao = date.today()
                        if cert_existente.data_vencimento and cert_existente.data_vencimento < date.today():
                            cert_existente.status = 'vencida'
                    elif cert_existente.data_vencimento and cert_existente.data_vencimento < date.today():
                        cert_existente.status = 'vencida'
                    elif cert_existente.status in ('buscando', 'erro'):
                        cert_existente.status = status_novo
                    resultados.append({"fonte_nome": fonte.nome, "acao": acao, "status": status_novo})
                else:
                    nova_cert = EmpresaCertidao()
                    nova_cert.id = str(uuid.uuid4())
                    nova_cert.empresa_id = empresa_id
                    nova_cert.tipo = fonte.tipo_certidao
                    nova_cert.orgao_emissor = fonte.orgao_emissor or fonte.nome
                    nova_cert.url_consulta = url_consulta
                    nova_cert.fonte_certidao_id = fonte.id
                    nova_cert.status = status_novo
                    if resultado_auto and resultado_auto.get("path_arquivo"):
                        nova_cert.path_arquivo = resultado_auto["path_arquivo"]
                    if resultado_auto and resultado_auto.get("data_vencimento"):
                        try:
                            nova_cert.data_vencimento = datetime.strptime(resultado_auto["data_vencimento"], '%Y-%m-%d').date()
                        except (ValueError, TypeError):
                            nova_cert.data_vencimento = date.today()
                    else:
                        nova_cert.data_vencimento = date.today()
                    if resultado_auto and resultado_auto.get("numero"):
                        nova_cert.numero = resultado_auto["numero"]
                    nova_cert.mensagem = acao
                    if resultado_auto and resultado_auto.get("dados_extras"):
                        de = resultado_auto["dados_extras"]
                        nova_cert.dados_extras = json_mod.loads(de) if isinstance(de, str) else de
                    if status_novo == 'valida':
                        nova_cert.data_emissao = date.today()
                    db.add(nova_cert)
                    resultados.append({"fonte_nome": fonte.nome, "acao": acao, "status": status_novo})

                # Emit result event
                yield f"data: {json_mod.dumps({'type': 'result', 'index': idx + 1, 'total': total, 'fonte': fonte.nome, 'fonte_id': fonte.id, 'sucesso': sucesso_flag, 'status': status_novo, 'mensagem': acao})}\n\n"

            db.commit()

            # Build summary
            obtidas = stats.get("obtidas_auto", 0)
            sem_key = stats.get("sem_api_key", 0)
            partes = [f"Busca concluída: {len(resultados)} fontes"]
            detalhes = []
            if obtidas > 0:
                detalhes.append(f"{obtidas} obtidas automaticamente")
            if sem_key > 0:
                detalhes.append(f"{sem_key} aguardando chave de API")
            if stats['autenticadas'] > 0:
                detalhes.append(f"{stats['autenticadas']} aguardando credenciais")
            if stats['manuais'] > 0:
                detalhes.append(f"{stats['manuais']} manuais")
            msg = partes[0] + " — " + ", ".join(detalhes) if detalhes else partes[0]

            yield f"data: {json_mod.dumps({'type': 'complete', 'message': msg, 'stats': stats})}\n\n"

        except Exception as e:
            db.rollback()
            yield f"data: {json_mod.dumps({'type': 'error', 'message': str(e)})}\n\n"
        finally:
            db.close()

    return Response(generate(), mimetype='text/event-stream', headers={
        'Cache-Control': 'no-cache',
        'X-Accel-Buffering': 'no',
        'Connection': 'keep-alive',
    })


# =============================================================================
# CapSolver - Status da automação de certidões
# =============================================================================

@app.route("/api/capsolver/status", methods=["GET"])
@require_auth
def capsolver_status():
    """Retorna status do CapSolver (configurado, saldo) para o frontend."""
    import requests
    from config import CAPSOLVER_API_KEY
    if not CAPSOLVER_API_KEY:
        return jsonify({"configurado": False, "mensagem": "CapSolver não configurado. FGTS, TST e SEFAZ requerem upload manual."})

    try:
        resp = requests.post("https://api.capsolver.com/getBalance", json={
            "clientKey": CAPSOLVER_API_KEY
        }, timeout=10)
        data = resp.json()
        if data.get("errorId", 0) != 0:
            return jsonify({"configurado": False, "mensagem": f"CapSolver: {data.get('errorDescription', 'Erro')}"})
        return jsonify({
            "configurado": True,
            "saldo": data.get("balance", 0),
            "mensagem": f"CapSolver ativo — saldo: ${data.get('balance', 0):.2f}"
        })
    except Exception as e:
        return jsonify({"configurado": False, "mensagem": f"CapSolver: erro de conexão — {str(e)[:80]}"})


# =============================================================================
# Fontes de Certidões - Inicializar fontes padrão
# =============================================================================

@app.route("/api/fontes-certidoes/inicializar", methods=["POST"])
@require_auth
def inicializar_fontes_certidoes():
    """
    Cria as 5 fontes padrão de certidões para o usuário, caso ainda não existam.
    Usa dados da empresa (UF, cidade) para personalizar SEFAZ e Prefeitura.
    """
    from models import FonteCertidao, Empresa
    user_id = get_current_user_id()
    db = get_db()
    try:
        # Verificar se já tem fontes
        existentes = db.query(FonteCertidao).filter(
            FonteCertidao.user_id == user_id
        ).count()
        if existentes > 0:
            return jsonify({"success": True, "message": f"Já existem {existentes} fontes cadastradas", "criadas": 0})

        # Buscar empresa para personalizar fontes estadual/municipal
        empresa = db.query(Empresa).filter(Empresa.user_id == user_id).first()
        uf = empresa.uf if empresa and empresa.uf else "SP"
        cidade = empresa.cidade if empresa and empresa.cidade else ""

        fontes_padrao = [
            # ── BUSCA AUTOMÁTICA REAL (testado 10/03/2026) ──
            {
                "tipo_certidao": "outro",
                "nome": "CGU - Certidão Correcional",
                "orgao_emissor": "Controladoria-Geral da União (CGU)",
                "url_portal": "https://certidoes.cgu.gov.br/",
                "metodo_acesso": "publico",
                "requer_autenticacao": False,
                "permite_busca_automatica": True,
                "observacoes": "API pública SEM captcha. Consolida CEIS, CNEP, CEPIM, CGU-PJ e ePAD em uma única consulta. Busca automática funcional.",
            },
            {
                "tipo_certidao": "fgts",
                "nome": "Caixa - CRF/FGTS",
                "orgao_emissor": "Caixa Econômica Federal",
                "url_portal": "https://consulta-crf.caixa.gov.br/consultacrf/pages/consultaEmpregador.jsf",
                "metodo_acesso": "publico",
                "requer_autenticacao": False,
                "permite_busca_automatica": True,
                "observacoes": "Portal com ShieldSquare (anti-bot). Automação via Playwright. Validade: 30 dias.",
            },
            {
                "tipo_certidao": "outro",
                "nome": "BrasilAPI - Situação Cadastral CNPJ",
                "orgao_emissor": "Receita Federal (via BrasilAPI)",
                "url_portal": "https://brasilapi.com.br/",
                "metodo_acesso": "publico",
                "requer_autenticacao": False,
                "permite_busca_automatica": True,
                "observacoes": "API gratuita, sem captcha. Retorna situação cadastral, Simples Nacional, sócios, CNAE. Consulta em tempo real.",
            },
            # ── AUTOMAÇÃO VIA PLAYWRIGHT + CAPSOLVER ──
            {
                "tipo_certidao": "cnd_federal",
                "nome": "Receita Federal - CND Federal",
                "orgao_emissor": "Receita Federal / PGFN",
                "url_portal": "https://cav.receita.fazenda.gov.br",
                "metodo_acesso": "login_senha",
                "requer_autenticacao": True,
                "permite_busca_automatica": True,
                "observacoes": "Automação via e-CAC com certificado digital A1 ou login gov.br. Configure credenciais na fonte para busca automática.",
            },
            {
                "tipo_certidao": "trabalhista",
                "nome": "TST - CNDT",
                "orgao_emissor": "Tribunal Superior do Trabalho (TST)",
                "url_portal": "https://cndt-certidao.tst.jus.br/gerarCertidao.faces",
                "metodo_acesso": "publico",
                "requer_autenticacao": False,
                "permite_busca_automatica": True,
                "observacoes": "Automação via Playwright + CapSolver (captcha imagem). Validade: 180 dias.",
            },
            {
                "tipo_certidao": "cnd_estadual",
                "nome": f"SEFAZ-{uf} - CND Estadual",
                "orgao_emissor": f"Secretaria da Fazenda - {uf}",
                "url_portal": f"https://www.sefaz.{uf.lower()}.gov.br",
                "uf": uf,
                "metodo_acesso": "publico",
                "requer_autenticacao": False,
                "permite_busca_automatica": True,
                "observacoes": f"Automação via Playwright + CapSolver (reCAPTCHA/hCaptcha). Estado: {uf}.",
            },
            {
                "tipo_certidao": "cnd_municipal",
                "nome": f"Prefeitura{' de ' + cidade if cidade else ''} - CND Municipal",
                "orgao_emissor": f"Prefeitura{' de ' + cidade if cidade else ''}",
                "url_portal": "",
                "cidade": cidade,
                "uf": uf,
                "metodo_acesso": "publico",
                "requer_autenticacao": False,
                "permite_busca_automatica": False,
                "observacoes": "A URL varia por município (+5.500 portais diferentes). Preencha a URL do portal da prefeitura. Maioria exige captcha. Upload manual.",
            },
        ]

        criadas = 0
        for fonte_data in fontes_padrao:
            fonte = FonteCertidao()
            fonte.id = str(uuid.uuid4())
            fonte.user_id = user_id
            for k, v in fonte_data.items():
                setattr(fonte, k, v)
            db.add(fonte)
            criadas += 1

        db.commit()
        return jsonify({
            "success": True,
            "message": f"{criadas} fontes de certidão criadas com sucesso",
            "criadas": criadas,
            "nota": "Revise as URLs (especialmente Estadual e Municipal) e adicione credenciais se necessário."
        })

    except Exception as e:
        db.rollback()
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


# =============================================================================
# B3: Notificação real por email (SMTP)
# =============================================================================

@app.route("/api/notificacoes/enviar-email", methods=["POST"])
@require_auth
def enviar_notificacao_email():
    """
    Envia uma notificação por email.
    Usa a configuração SMTP do scheduler.py.

    JSON body:
    - notificacao_id: (opcional) ID da notificação existente para reenviar
    - assunto: assunto do email
    - corpo: corpo do email (texto ou HTML)
    - destinatario: (opcional) email. Se não informado, usa o email do usuário.
    """
    from models import Notificacao, PreferenciasNotificacao, User
    user_id = get_current_user_id()
    data = request.json or {}

    db = get_db()
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            return jsonify({"error": "Usuário não encontrado"}), 404

        # Determinar destinatário
        prefs = db.query(PreferenciasNotificacao).filter(
            PreferenciasNotificacao.user_id == user_id
        ).first()
        destinatario = data.get('destinatario') or (prefs.email_notificacao if prefs else None) or user.email

        assunto = data.get('assunto', 'Notificação - Agente Editais')
        corpo = data.get('corpo', '')

        # Se notificacao_id fornecido, buscar dados da notificação
        notificacao_id = data.get('notificacao_id')
        if notificacao_id:
            notif = db.query(Notificacao).filter(
                Notificacao.id == notificacao_id,
                Notificacao.user_id == user_id
            ).first()
            if notif:
                assunto = notif.titulo
                corpo = notif.mensagem

        if not corpo:
            return jsonify({"error": "Corpo do email não pode ser vazio"}), 400

        # Enviar via scheduler.enviar_email_alerta
        from scheduler import enviar_email_alerta
        sucesso = enviar_email_alerta(destinatario, assunto, corpo)

        # Atualizar flag enviado_email na notificação
        if notificacao_id and sucesso:
            notif = db.query(Notificacao).filter(Notificacao.id == notificacao_id).first()
            if notif:
                notif.enviado_email = True
                db.commit()

        if sucesso:
            return jsonify({
                "success": True,
                "message": f"Email enviado para {destinatario}"
            })
        else:
            return jsonify({
                "success": False,
                "message": "SMTP não configurado ou erro no envio. "
                           "Configure SMTP_HOST, SMTP_USER e SMTP_PASSWORD nas variáveis de ambiente.",
                "email_salvo": True
            })

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        db.close()


@app.route("/api/notificacoes/config-smtp", methods=["GET"])
@require_auth
def verificar_config_smtp():
    """Verifica se o SMTP está configurado."""
    from scheduler import SMTP_HOST, SMTP_USER, SMTP_PASSWORD
    configurado = bool(SMTP_USER and SMTP_PASSWORD)
    return jsonify({
        "smtp_configurado": configurado,
        "smtp_host": SMTP_HOST if configurado else None,
        "smtp_user": SMTP_USER[:3] + "***" if configurado else None
    })


# =============================================================================
# Scheduler Status + Notificações não-lidas
# =============================================================================

@app.route("/api/scheduler/status", methods=["GET"])
@require_auth
def scheduler_status():
    """Retorna status do scheduler e monitoramentos ativos."""
    from models import Monitoramento as Mon, Notificacao as Notif
    import scheduler as sched_mod

    user_id = get_current_user_id()
    db = get_db()

    # Status do scheduler
    sched_available = getattr(sched_mod, 'SCHEDULER_AVAILABLE', False)
    sched_instance = getattr(sched_mod, 'scheduler', None)
    scheduler_ativo = bool(sched_instance and sched_instance.running) if sched_available else False
    check_interval = getattr(sched_mod, 'CHECK_MONITORAMENTOS_INTERVAL', 60)

    # Monitoramentos do usuário
    monitoramentos = db.query(Mon).filter(
        Mon.user_id == user_id
    ).all()

    ativos = [m for m in monitoramentos if m.ativo]
    inativos = [m for m in monitoramentos if not m.ativo]

    # Notificações não lidas
    nao_lidas = db.query(Notif).filter(
        Notif.user_id == user_id,
        Notif.lida == False
    ).count()

    return jsonify({
        "scheduler_disponivel": sched_available,
        "scheduler_ativo": scheduler_ativo,
        "intervalo_verificacao_min": check_interval,
        "monitoramentos_ativos": len(ativos),
        "monitoramentos_inativos": len(inativos),
        "notificacoes_nao_lidas": nao_lidas,
        "monitoramentos": [{
            "id": m.id,
            "termo": m.termo,
            "ncm": m.ncm,
            "ufs": m.ufs,
            "ativo": m.ativo,
            "frequencia_horas": m.frequencia_horas,
            "ultima_execucao": m.ultima_execucao.isoformat() if m.ultima_execucao else None,
            "proximo_check": m.proximo_check.isoformat() if m.proximo_check else None,
            "editais_encontrados": m.editais_encontrados or 0,
        } for m in monitoramentos],
    })


@app.route("/api/notificacoes/nao-lidas", methods=["GET"])
@require_auth
def notificacoes_nao_lidas():
    """Retorna notificações não lidas do usuário."""
    from models import Notificacao
    user_id = get_current_user_id()
    db = get_db()

    notifs = db.query(Notificacao).filter(
        Notificacao.user_id == user_id,
        Notificacao.lida == False
    ).order_by(Notificacao.created_at.desc()).limit(20).all()

    return jsonify({
        "total": len(notifs),
        "notificacoes": [{
            "id": n.id,
            "tipo": n.tipo,
            "titulo": n.titulo,
            "mensagem": n.mensagem,
            "created_at": n.created_at.isoformat() if n.created_at else None,
        } for n in notifs],
    })


@app.route("/api/notificacoes/<notif_id>/lida", methods=["PUT"])
@require_auth
def marcar_notificacao_lida(notif_id):
    """Marca uma notificação como lida."""
    from models import Notificacao
    user_id = get_current_user_id()
    db = get_db()

    notif = db.query(Notificacao).filter(
        Notificacao.id == notif_id,
        Notificacao.user_id == user_id
    ).first()

    if not notif:
        return jsonify({"error": "Notificação não encontrada"}), 404

    notif.lida = True
    notif.lida_em = datetime.now()
    db.commit()

    return jsonify({"success": True})


@app.route("/api/notificacoes/marcar-todas-lidas", methods=["PUT"])
@require_auth
def marcar_todas_notificacoes_lidas():
    """Marca todas as notificações do usuário como lidas."""
    from models import Notificacao
    user_id = get_current_user_id()
    db = get_db()

    atualizadas = db.query(Notificacao).filter(
        Notificacao.user_id == user_id,
        Notificacao.lida == False
    ).update({"lida": True, "lida_em": datetime.now()})

    db.commit()

    return jsonify({"success": True, "atualizadas": atualizadas})


# =============================================================================
# Endpoints: Reprocessar metadados de produto (CATMAT + termos_busca)
# =============================================================================

@app.route("/api/produtos/<produto_id>/reprocessar-metadados", methods=["POST"])
@require_auth
def reprocessar_metadados(produto_id):
    """Reprocessa CATMAT/CATSER e termos de busca para um produto."""
    from tools import processar_metadados_produto
    user_id = get_current_user_id()
    db = get_db()
    try:
        produto = db.query(Produto).filter(
            Produto.id == produto_id,
            Produto.user_id == user_id
        ).first()
        if not produto:
            return jsonify({"success": False, "error": "Produto não encontrado"}), 404
    finally:
        db.close()

    processar_metadados_produto(produto_id)
    return jsonify({"success": True, "message": f"Metadados reprocessados para {produto_id}"})


@app.route("/api/admin/reprocessar-todos-metadados", methods=["POST"])
@require_auth
def reprocessar_todos_metadados():
    """Reprocessa CATMAT/CATSER e termos de busca para todos os produtos do usuário."""
    from tools import processar_metadados_produto
    user_id = get_current_user_id()
    db = get_db()
    try:
        produtos = db.query(Produto).filter(Produto.user_id == user_id).all()
        ids = [p.id for p in produtos]
    finally:
        db.close()

    for pid in ids:
        processar_metadados_produto(pid)

    return jsonify({"success": True, "total": len(ids)})


# =============================================================================
# Main
# =============================================================================

if __name__ == "__main__":
    print("=" * 50)
    print("AGENTE DE EDITAIS - MVP")
    print("=" * 50)

    # Criar pasta de uploads
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)

    print("Inicializando banco de dados...")
    init_db()

    # Iniciar scheduler para alertas e monitoramentos (Sprint 2)
    # Proteger contra inicialização dupla em debug mode (Flask reloader cria 2 processos)
    if os.environ.get('WERKZEUG_RUN_MAIN') == 'true' or not app.debug:
        try:
            from scheduler import iniciar_scheduler
            print("Iniciando scheduler de alertas e monitoramentos...")
            iniciar_scheduler()
        except Exception as e:
            print(f"[AVISO] Scheduler não iniciado: {e}")

    print("Servidor pronto na porta 5007!")
    print("=" * 50)

    app.run(host="0.0.0.0", port=5007, debug=True)

